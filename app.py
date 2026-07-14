#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import sqlite3
import webbrowser
import logging
import re
import unicodedata
import zipfile
import tempfile
import threading
import concurrent.futures
import traceback
import shutil
import urllib.request
import requests
import urllib.error
import urllib.parse
import html
import time
import ssl
import base64
import mimetypes
import uuid
import hashlib
from datetime import datetime, timedelta, date
from io import BytesIO
from urllib.parse import urlparse, quote_plus
from pathlib import Path
from xml.etree import ElementTree as ET
from flask import Flask, jsonify, request, send_from_directory, send_file, Response, stream_with_context, redirect
from flask_cors import CORS
from werkzeug.utils import secure_filename
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.edge.options import Options as EdgeOptions
    from selenium.webdriver.edge.service import Service as EdgeService
    SELENIUM_AVAILABLE = True
except Exception:
    SELENIUM_AVAILABLE = False
from werkzeug.exceptions import HTTPException
from autotoca import AccountAddressService
from integrations.outlook_graph import (
    OutlookOAuthError,
    OutlookSyncError,
    build_authorize_url as outlook_graph_build_authorize_url,
    ensure_schema as outlook_graph_ensure_schema,
    exchange_code_and_store as outlook_graph_exchange_code_and_store,
    fetch_messages as outlook_graph_fetch_messages,
    get_valid_access_token as outlook_graph_get_valid_access_token,
    parse_state as outlook_graph_parse_state,
    send_mail as outlook_graph_send_mail,
)
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None
    PDFPLUMBER_AVAILABLE = False

# Suprime logs verbosos do pdfminer (usado internamente pelo pdfplumber) —
# por padrão, o pdfminer emite DEBUG para cada token/objeto do PDF, poluindo o terminal.
import logging as _logging
_logging.getLogger('pdfminer').setLevel(_logging.WARNING)
_logging.getLogger('pdfminer.psparser').setLevel(_logging.WARNING)
_logging.getLogger('pdfminer.pdfdocument').setLevel(_logging.WARNING)
_logging.getLogger('pdfminer.pdfinterp').setLevel(_logging.WARNING)

try:
    import docx as python_docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    python_docx = None
    PYTHON_DOCX_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None
    PYTESSERACT_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    convert_from_path = None
    PDF2IMAGE_AVAILABLE = False

try:
    import pypdfium2 as pdfium
    PYPDFIUM2_AVAILABLE = True
except ImportError:
    pdfium = None
    PYPDFIUM2_AVAILABLE = False

try:
    from PIL import Image as PILImage
    from PIL import ImageOps, ImageFilter
    PIL_AVAILABLE = True
except ImportError:
    PILImage = None
    ImageOps = None
    ImageFilter = None
    PIL_AVAILABLE = False

REPORTLAB_IMPORT_ERROR = None
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.platypus import Paragraph
    REPORTLAB_AVAILABLE = True
except Exception as e:
    colors = None
    TA_LEFT = TA_CENTER = None
    A4 = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    mm = None
    ImageReader = None
    stringWidth = None
    Paragraph = None
    REPORTLAB_AVAILABLE = False
    REPORTLAB_IMPORT_ERROR = e

WHISPER_IMPORT_ERROR = None
WHISPER_IMPORT_ATTEMPTED = False
WhisperModel = None
WHISPER_AVAILABLE = True

TRANSCRIPTION_BACKEND = 'faster-whisper'

try:
    import imageio_ffmpeg
    IMAGEIO_FFMPEG_AVAILABLE = True
except ImportError:
    imageio_ffmpeg = None
    IMAGEIO_FFMPEG_AVAILABLE = False

# Configuracao
app = Flask(__name__, static_folder='public', static_url_path='')
CORS(app)

# Diretorio de dados
if sys.platform == 'win32':
    DATA_DIR = Path.home() / 'AppData' / 'Roaming' / 'toca-do-coelho'
    LEGACY_DATA_DIR_V2 = Path('C:/toca-do-coelho-version2')
    LEGACY_DATA_DIR_V1 = Path('C:/toca-do-coelho')
else:
    DATA_DIR = Path.home() / '.toca-do-coelho'
    LEGACY_DATA_DIR_V2 = None
    LEGACY_DATA_DIR_V1 = None

DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / 'toca-do-coelho.db'
TEST_DB_TEMPLATE_PATH = Path(__file__).resolve().parent / 'BD_teste' / 'toca-do-coelho-ficticio-reduzido.db'
BACKUP_DIR = DATA_DIR / 'backups'
BACKUP_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR = DATA_DIR / 'logs'
LOG_DIR.mkdir(parents=True, exist_ok=True)
LOG_FILE = LOG_DIR / 'app.log'
UPLOAD_DIR = DATA_DIR / 'uploads'
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
ACCOUNT_UPLOAD_DIR = UPLOAD_DIR / 'accounts'
ACCOUNT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
WIKI_UPLOAD_DIR = UPLOAD_DIR / 'wikitoca'
WIKI_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
AUTOTOCA_UPLOAD_DIR = UPLOAD_DIR / 'autotoca'
AUTOTOCA_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
AUTOTOCA_SUPPORT_FILES_DIR = Path(app.static_folder) / 'assets' / 'autotoca' / 'chamado-juridico'
AUTOTOCA_SUPPORT_FILES_DIR.mkdir(parents=True, exist_ok=True)
CHAMADO_JURIDICO_UPLOAD_DIR = AUTOTOCA_UPLOAD_DIR / 'chamado-juridico'
CHAMADO_JURIDICO_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
REEMBOLSOS_UPLOAD_DIR = AUTOTOCA_UPLOAD_DIR / 'reembolsos'
REEMBOLSOS_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

WHISPER_MODEL = None
WHISPER_MODEL_LOCK = threading.Lock()
TRANSCRIPTION_DEBUG = os.environ.get('TRANSCRIPTION_DEBUG', '').lower() in {'1', 'true', 'yes', 'on'}


def setup_logging():
    formatter = logging.Formatter(
        '[%(asctime)s] [%(levelname)s] [%(name)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )

    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)

    if not root_logger.handlers:
        console_handler = logging.StreamHandler(sys.stdout)
        console_handler.setFormatter(formatter)
        root_logger.addHandler(console_handler)

    file_handler_exists = any(
        isinstance(handler, logging.FileHandler) and getattr(handler, 'baseFilename', '') == str(LOG_FILE)
        for handler in root_logger.handlers
    )
    if not file_handler_exists:
        file_handler = logging.FileHandler(LOG_FILE, encoding='utf-8')
        file_handler.setFormatter(formatter)
        root_logger.addHandler(file_handler)


setup_logging()
logger = logging.getLogger('toca-do-coelho')

def _resolve_app_version():
    default_version = '1.0.0'
    env_version = (os.environ.get('TOCA_APP_VERSION') or '').strip()
    candidate_dirs = [Path(__file__).resolve().parent]

    if getattr(sys, 'frozen', False):
        meipass = getattr(sys, '_MEIPASS', None)
        if meipass:
            candidate_dirs.append(Path(meipass))
        candidate_dirs.append(Path(sys.executable).resolve().parent)

    for base_dir in candidate_dirs:
        version_file = base_dir / 'version.txt'
        try:
            if version_file.exists():
                file_version = version_file.read_text(encoding='utf-8').strip()
                if file_version:
                    return file_version
        except Exception as error:
            logger.warning(
                '[Update] Não foi possível ler %s: %s',
                version_file,
                error
            )

    return env_version or default_version


APP_VERSION = _resolve_app_version()
DEFAULT_GITHUB_OWNER = os.environ.get('TOCA_UPDATE_GITHUB_OWNER', 'rochanets').strip()
DEFAULT_GITHUB_REPO = os.environ.get('TOCA_UPDATE_GITHUB_REPO', 'TocaDoCoelho').strip()
DEFAULT_GITHUB_BRANCH = os.environ.get('TOCA_UPDATE_GITHUB_BRANCH', 'version5').strip()

logger.info('[Transcription] Backend faster-whisper será carregado sob demanda (lazy).')

AUTOMAPPING_CANCELLED_REQUESTS = set()
AUTOMAPPING_CANCELLED_LOCK = threading.Lock()


def _mark_automapping_cancelled(request_id):
    if not request_id:
        return
    with AUTOMAPPING_CANCELLED_LOCK:
        AUTOMAPPING_CANCELLED_REQUESTS.add(request_id)


def _is_automapping_cancelled(request_id, consume=False):
    if not request_id:
        return False
    with AUTOMAPPING_CANCELLED_LOCK:
        if request_id not in AUTOMAPPING_CANCELLED_REQUESTS:
            return False
        if consume:
            AUTOMAPPING_CANCELLED_REQUESTS.discard(request_id)
        return True


def configure_ffmpeg_for_whisper():
    bundled_candidates = []

    if getattr(sys, 'frozen', False):
        bundled_candidates.append(Path(sys.executable).parent / ('ffmpeg.exe' if sys.platform == 'win32' else 'ffmpeg'))
        meipass = getattr(sys, '_MEIPASS', None)
        if meipass:
            bundled_candidates.append(Path(meipass) / ('ffmpeg.exe' if sys.platform == 'win32' else 'ffmpeg'))

    bundled_candidates.append(Path(__file__).resolve().parent / ('ffmpeg.exe' if sys.platform == 'win32' else 'ffmpeg'))

    for candidate in bundled_candidates:
        if candidate.exists():
            os.environ['FFMPEG_BINARY'] = str(candidate)
            return str(candidate)

    ffmpeg_binary = os.environ.get('FFMPEG_BINARY')
    if ffmpeg_binary and Path(ffmpeg_binary).exists():
        return ffmpeg_binary

    ffmpeg_on_path = shutil.which('ffmpeg')
    if ffmpeg_on_path:
        os.environ['FFMPEG_BINARY'] = ffmpeg_on_path
        return ffmpeg_on_path

    if IMAGEIO_FFMPEG_AVAILABLE:
        try:
            ffmpeg_imageio = imageio_ffmpeg.get_ffmpeg_exe()
            ffmpeg_dir = str(Path(ffmpeg_imageio).parent)
            os.environ['PATH'] = f"{ffmpeg_dir}{os.pathsep}{os.environ.get('PATH', '')}"
            os.environ['FFMPEG_BINARY'] = ffmpeg_imageio
            return ffmpeg_imageio
        except Exception as e:
            if TRANSCRIPTION_DEBUG:
                logger.debug(f"[Transcription] Falha ao obter ffmpeg via imageio_ffmpeg: {e}")

    return None


def get_ffmpeg_install_instructions():
    if sys.platform == 'win32':
        return [
            'Windows (winget): winget install -e --id Gyan.FFmpeg',
            'Ou Windows (choco): choco install ffmpeg -y',
            'Depois reinicie o app.'
        ]
    if sys.platform == 'darwin':
        return ['macOS: brew install ffmpeg', 'Depois reinicie o app.']
    return ['Linux (Debian/Ubuntu): sudo apt update && sudo apt install -y ffmpeg', 'Depois reinicie o app.']


def get_whisper_model():
    global WHISPER_MODEL, WhisperModel, WHISPER_AVAILABLE, WHISPER_IMPORT_ERROR, WHISPER_IMPORT_ATTEMPTED

    if not WHISPER_AVAILABLE:
        return None

    with WHISPER_MODEL_LOCK:
        if WHISPER_MODEL is not None:
            return WHISPER_MODEL

        if WhisperModel is None and not WHISPER_IMPORT_ATTEMPTED:
            WHISPER_IMPORT_ATTEMPTED = True
            try:
                from faster_whisper import WhisperModel as _WhisperModel
                WhisperModel = _WhisperModel
            except Exception as e:
                WHISPER_AVAILABLE = False
                WHISPER_IMPORT_ERROR = str(e)
                logger.warning(f'[Transcription] Backend faster-whisper indisponível: {WHISPER_IMPORT_ERROR}')
                return None

        if WhisperModel is None:
            return None

        try:
            WHISPER_MODEL = WhisperModel('base', device='cpu', compute_type='int8')
        except Exception as e:
            WHISPER_AVAILABLE = False
            WHISPER_IMPORT_ERROR = str(e)
            logger.warning(f'[Transcription] Falha ao inicializar faster-whisper: {WHISPER_IMPORT_ERROR}')
            return None

    return WHISPER_MODEL


TEST_DB_REQUIRED_SCHEMA = {
    'clients': {'id', 'name', 'company', 'position'},
    'accounts': {'id', 'name'},
    'activities': {'id', 'client_id'},
    'commitments': {'id', 'client_id', 'title', 'due_date'},
    'app_settings': {'key', 'value'},
    'itoca_chat_history': {'id', 'session_id', 'role', 'content'},
}


def _is_test_db_schema_valid(db_path):
    try:
        conn = sqlite3.connect(str(db_path))
        c = conn.cursor()
        for table_name, required_columns in TEST_DB_REQUIRED_SCHEMA.items():
            c.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name = ?",
                (table_name,)
            )
            if c.fetchone() is None:
                logger.warning(
                    f'[Database] Banco de teste inválido: tabela obrigatória ausente ({table_name}).'
                )
                conn.close()
                return False

            c.execute(f"PRAGMA table_info({table_name})")
            table_columns = {row[1] for row in c.fetchall()}
            missing_columns = required_columns - table_columns
            if missing_columns:
                logger.warning(
                    '[Database] Banco de teste inválido: colunas obrigatórias ausentes em %s: %s',
                    table_name,
                    ', '.join(sorted(missing_columns))
                )
                conn.close()
                return False

        conn.close()
        return True
    except Exception as e:
        logger.warning(f'[Database] Falha ao validar banco de teste: {e}')
        return False


def maybe_seed_test_db_fallback():
    if DB_PATH.exists():
        return

    should_use_test_db = os.environ.get('TOCA_ENABLE_TEST_DB_FALLBACK', '').strip().lower() in {
        '1', 'true', 'yes', 'on'
    }
    if not should_use_test_db:
        return

    if getattr(sys, 'frozen', False):
        logger.info('[Database] Fallback de banco de teste ignorado em build instalada.')
        return

    if not TEST_DB_TEMPLATE_PATH.exists():
        logger.warning(f'[Database] Banco de teste não encontrado em {TEST_DB_TEMPLATE_PATH}.')
        return

    if not _is_test_db_schema_valid(TEST_DB_TEMPLATE_PATH):
        logger.warning('[Database] Banco de teste inválido. O app seguirá com banco vazio padrão.')
        return

    shutil.copy2(str(TEST_DB_TEMPLATE_PATH), str(DB_PATH))
    logger.info(f'[Database] Banco de teste copiado para uso local: {DB_PATH}')

# Migração automática do banco de dados antigo
if sys.platform == 'win32' and not DB_PATH.exists():
    import shutil
    migrated = False

    legacy_sources = [
        ('C:/toca-do-coelho-version2', LEGACY_DATA_DIR_V2, 'toca-do-coelho-version2.db'),
        ('C:/toca-do-coelho', LEGACY_DATA_DIR_V1, 'toca-do-coelho.db'),
    ]

    for source_label, legacy_dir, legacy_db_name in legacy_sources:
        if migrated or not legacy_dir or not legacy_dir.exists():
            continue

        old_db = legacy_dir / legacy_db_name
        if old_db.exists():
            logger.info(f'[Database] Migrando banco de dados de {old_db} para {DB_PATH}')
            shutil.copy2(str(old_db), str(DB_PATH))

            old_uploads = legacy_dir / 'uploads'
            if old_uploads.exists():
                for item in old_uploads.iterdir():
                    dest = UPLOAD_DIR / item.name
                    if dest.exists():
                        continue
                    if item.is_file():
                        shutil.copy2(str(item), str(dest))
                    elif item.is_dir():
                        shutil.copytree(str(item), str(dest))

            logger.info(f'[Database] Migração de {source_label} concluída com sucesso!')
            migrated = True

maybe_seed_test_db_fallback()

logger.info(f'[Database] Caminho: {DB_PATH}')

# Inicializar banco de dados
def init_db():
    conn = sqlite3.connect(str(DB_PATH))
    c = conn.cursor()
    
    # Tabela de clientes
    c.execute('''CREATE TABLE IF NOT EXISTS clients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        company TEXT NOT NULL,
        position TEXT NOT NULL,
        area_of_activity TEXT,
        email TEXT,
        phone TEXT,
        linkedin TEXT,
        photo_url TEXT,
        is_target INTEGER DEFAULT 0,
        is_cold_contact INTEGER DEFAULT 0,
        is_archived INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS job_groupings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS job_grouping_positions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        grouping_id INTEGER NOT NULL,
        position TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(grouping_id) REFERENCES job_groupings(id) ON DELETE CASCADE,
        UNIQUE(grouping_id, position)
    )''')

    # Tabela de configuracoes gerais
    c.execute('''CREATE TABLE IF NOT EXISTS app_settings (
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # WikiToca (base de conhecimento + documentos)
    c.execute('''CREATE TABLE IF NOT EXISTS wiki_entries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        category TEXT,
        content TEXT NOT NULL,
        tags TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS wiki_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        file_name TEXT NOT NULL,
        original_name TEXT NOT NULL,
        file_url TEXT NOT NULL,
        file_ext TEXT,
        file_size INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # Tabela de regras de status por cargo
    c.execute('''CREATE TABLE IF NOT EXISTS status_rules (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        position TEXT NOT NULL UNIQUE,
        green_days INTEGER NOT NULL,
        yellow_days INTEGER NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # Tabela de perfil do usuario
    c.execute('''CREATE TABLE IF NOT EXISTS user_profile (
        id INTEGER PRIMARY KEY CHECK (id = 1),
        full_name TEXT NOT NULL,
        nickname TEXT NOT NULL,
        position TEXT NOT NULL,
        photo_url TEXT,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS account_sectors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS accounts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        logo_url TEXT,
        is_target INTEGER DEFAULT 0,
        sector TEXT,
        average_revenue_cents INTEGER,
        professionals_count INTEGER,
        global_presence TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS account_main_contacts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        account_id INTEGER NOT NULL,
        client_id INTEGER NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE CASCADE,
        FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE,
        UNIQUE(account_id, client_id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS account_presences (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        account_id INTEGER NOT NULL,
        delivery_name TEXT NOT NULL,
        stf_owner TEXT,
        delivery_cell TEXT,
        service_id TEXT,
        current_revenue_cents INTEGER,
        validity_month TEXT,
        focal_client_id INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE CASCADE,
        FOREIGN KEY(focal_client_id) REFERENCES clients(id) ON DELETE SET NULL
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS account_renewal_events (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        account_id INTEGER NOT NULL,
        presence_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        due_date TEXT NOT NULL,
        due_time TEXT DEFAULT '09:00',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE CASCADE,
        FOREIGN KEY(presence_id) REFERENCES account_presences(id) ON DELETE CASCADE,
        UNIQUE(presence_id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS account_activities (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        account_id INTEGER NOT NULL,
        description TEXT NOT NULL,
        activity_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE CASCADE
    )''')

    # Arquivo morto de contas excluídas (snapshot completo em JSON para restauração futura)
    c.execute('''CREATE TABLE IF NOT EXISTS account_archives (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        account_name TEXT NOT NULL,
        snapshot_json TEXT NOT NULL,
        contacts_count INTEGER DEFAULT 0,
        activities_count INTEGER DEFAULT 0,
        archived_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS message_templates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        description TEXT NOT NULL,
        available_whatsapp INTEGER DEFAULT 1,
        available_email INTEGER DEFAULT 1,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS portfolio_offers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        summary TEXT,
        raw_input TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS portfolio_offer_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        offer_id INTEGER NOT NULL,
        pain TEXT,
        solution TEXT,
        sort_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(offer_id) REFERENCES portfolio_offers(id) ON DELETE CASCADE
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS iata_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        meeting_date TEXT,
        meeting_time TEXT,
        topic TEXT,
        participants TEXT,
        ata_json TEXT,
        insights_json TEXT,
        raw_text TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Tabela de compromissos (agenda)
    c.execute('''CREATE TABLE IF NOT EXISTS commitments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        client_id INTEGER NOT NULL,
        activity_id INTEGER,
        title TEXT NOT NULL,
        notes TEXT,
        due_date TEXT NOT NULL,
        due_time TEXT,
        source_type TEXT DEFAULT 'activity',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE,
        FOREIGN KEY(activity_id) REFERENCES activities(id) ON DELETE SET NULL
    )''')

    # Tabela de atividades
    c.execute('''CREATE TABLE IF NOT EXISTS activities (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        client_id INTEGER NOT NULL,
        contact_type TEXT DEFAULT 'Outro',
        information TEXT,
        description TEXT,
        activity_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE
    )''')
    
    # Tabela de log de sincronizacao WhatsApp (deduplicacao)
    c.execute('''CREATE TABLE IF NOT EXISTS whatsapp_sync_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        client_id INTEGER NOT NULL,
        phone TEXT,
        period_days INTEGER,
        message_count INTEGER,
        content_hash TEXT NOT NULL,
        last_message_ts INTEGER DEFAULT 0,
        activity_id INTEGER,
        imported_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE,
        FOREIGN KEY(activity_id) REFERENCES activities(id) ON DELETE SET NULL
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_wa_sync_hash ON whatsapp_sync_log(client_id, content_hash)')
    # Migração: coluna last_message_ts para deduplicação incremental (só resume mensagens novas)
    c.execute("PRAGMA table_info(whatsapp_sync_log)")
    _wa_sync_cols = [col[1] for col in c.fetchall()]
    if 'last_message_ts' not in _wa_sync_cols:
        c.execute('ALTER TABLE whatsapp_sync_log ADD COLUMN last_message_ts INTEGER DEFAULT 0')

    # Tabela de cards de mapeamento de ambiente
    c.execute('''CREATE TABLE IF NOT EXISTS environment_cards (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        description TEXT,
        display_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    # Tabela de respostas de mapeamento por cliente
    c.execute('''CREATE TABLE IF NOT EXISTS environment_responses (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        card_id INTEGER NOT NULL,
        client_id INTEGER NOT NULL,
        response TEXT,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(card_id) REFERENCES environment_cards(id) ON DELETE CASCADE,
        FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE,
        UNIQUE(card_id, client_id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS kanban_columns (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        display_order INTEGER DEFAULT 0,
        is_system INTEGER DEFAULT 0,
        is_locked INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS kanban_cards (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        description TEXT NOT NULL,
        tag TEXT,
        account_id INTEGER,
        contact_id INTEGER,
        activity TEXT,
        urgency TEXT DEFAULT 'Média',
        column_id INTEGER NOT NULL,
        display_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(column_id) REFERENCES kanban_columns(id) ON DELETE CASCADE,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE SET NULL,
        FOREIGN KEY(contact_id) REFERENCES clients(id) ON DELETE SET NULL
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS kanban_card_activities (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        card_id INTEGER NOT NULL,
        content TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(card_id) REFERENCES kanban_cards(id) ON DELETE CASCADE
    )''')
    
    # Tabela de sugestões diárias
    c.execute('''CREATE TABLE IF NOT EXISTS daily_suggestions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        date TEXT NOT NULL,
        suggestion_type TEXT NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        target_id INTEGER,
        target_data TEXT,
        completed INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        completed_at TIMESTAMP
    )''')

    # Histórico de execuções de AutoMapping
    c.execute('''CREATE TABLE IF NOT EXISTS automapping_runs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company TEXT NOT NULL,
        country TEXT NOT NULL,
        industry TEXT NOT NULL,
        query_key TEXT NOT NULL,
        result_json TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')


    c.execute('CREATE INDEX IF NOT EXISTS idx_automapping_query_key ON automapping_runs(query_key)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_automapping_created_at ON automapping_runs(created_at)')

    # Histórico de preenchimentos do robô do Chamado Jurídico (permite reaproveitar
    # texto e arquivos de uma execução anterior num novo chamado)
    c.execute('''CREATE TABLE IF NOT EXISTS chamado_juridico_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        conta TEXT NOT NULL,
        proposta_original_name TEXT,
        payload_json TEXT NOT NULL,
        files_json TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_chamado_juridico_history_created_at ON chamado_juridico_history(created_at)')

    # Reembolsos — histórico de endereços de Origem digitados (dropdown de
    # reaproveitamento) e endereço de Destino salvo por conta.
    c.execute('''CREATE TABLE IF NOT EXISTS reembolso_origem_historico (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        texto TEXT NOT NULL UNIQUE,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS account_reembolso_enderecos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        account_id INTEGER NOT NULL,
        endereco TEXT NOT NULL,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE CASCADE,
        UNIQUE(account_id)
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS reembolsos_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tipo TEXT NOT NULL,
        payload_json TEXT NOT NULL,
        files_json TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_reembolsos_history_created_at ON reembolsos_history(created_at)')

    # Tokens de integrações OAuth por usuário (Outlook Graph e futuros conectores)
    outlook_graph_ensure_schema(conn)

    # Inserir cards pré-definidos se não existirem
    predefined_cards = [
        ('ERP', '', 1),
        ('Cloud', '', 2),
        ('Principal Concorrente', '', 3),
        ('Gestão de Ativos/HAAS', '', 4),
        ('Cyber/Observability', '', 5),
        ('DWS', '', 6)
    ]
    
    for title, description, order in predefined_cards:
        c.execute('SELECT id FROM environment_cards WHERE title = ?', (title,))
        if not c.fetchone():
            c.execute('INSERT INTO environment_cards (title, description, display_order) VALUES (?, ?, ?)', 
                      (title, description, order))

    c.execute("PRAGMA table_info(kanban_columns)")
    kanban_column_cols = [col[1] for col in c.fetchall()]
    if 'is_system' not in kanban_column_cols:
        c.execute('ALTER TABLE kanban_columns ADD COLUMN is_system INTEGER DEFAULT 0')
    if 'is_locked' not in kanban_column_cols:
        c.execute('ALTER TABLE kanban_columns ADD COLUMN is_locked INTEGER DEFAULT 0')

    c.execute('SELECT COUNT(*) FROM kanban_columns')
    has_columns = c.fetchone()[0] > 0
    if not has_columns:
        default_columns = [
            ('Backlog', 1, 1, 0),
            ('Em Andamento', 2, 1, 0),
            ('Hold', 3, 1, 0),
            ('Done', 4, 1, 1),
            ('Descartado', 5, 1, 1)
        ]
        c.executemany(
            'INSERT INTO kanban_columns (title, display_order, is_system, is_locked) VALUES (?, ?, ?, ?)',
            default_columns
        )

    c.execute("UPDATE kanban_columns SET title = 'Done', updated_at = CURRENT_TIMESTAMP WHERE lower(title) = 'finalizado'")

    c.execute("PRAGMA table_info(kanban_cards)")
    kanban_card_cols = [col[1] for col in c.fetchall()]
    if 'urgency' not in kanban_card_cols:
        c.execute('ALTER TABLE kanban_cards ADD COLUMN urgency TEXT DEFAULT "Média"')
    
    # Adicionar coluna last_activity_date à tabela clients se não existir
    c.execute("PRAGMA table_info(clients)")
    columns = [col[1] for col in c.fetchall()]
    if 'last_activity_date' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN last_activity_date TIMESTAMP')
    if 'is_target' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN is_target INTEGER DEFAULT 0')
    if 'area_of_activity' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN area_of_activity TEXT')
    if 'is_cold_contact' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN is_cold_contact INTEGER DEFAULT 0')
    if 'linkedin' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN linkedin TEXT')
    if 'is_archived' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN is_archived INTEGER DEFAULT 0')
    if 'relationship_stage' not in columns:
        c.execute('ALTER TABLE clients ADD COLUMN relationship_stage TEXT DEFAULT ""')

    c.execute("PRAGMA table_info(commitments)")
    commitment_columns = [col[1] for col in c.fetchall()]
    if 'due_time' not in commitment_columns:
        c.execute('ALTER TABLE commitments ADD COLUMN due_time TEXT')
    if 'source_type' not in commitment_columns:
        c.execute('ALTER TABLE commitments ADD COLUMN source_type TEXT DEFAULT "activity"')

    c.execute("PRAGMA table_info(account_presences)")
    account_presence_columns = [col[1] for col in c.fetchall()]
    if 'delivery_cell' not in account_presence_columns:
        c.execute('ALTER TABLE account_presences ADD COLUMN delivery_cell TEXT')
    if 'service_id' not in account_presence_columns:
        c.execute('ALTER TABLE account_presences ADD COLUMN service_id TEXT')
    if 'billing_type' not in account_presence_columns:
        c.execute("ALTER TABLE account_presences ADD COLUMN billing_type TEXT DEFAULT 'Mensal'")
    if 'contract_end_date' not in account_presence_columns:
        c.execute("ALTER TABLE account_presences ADD COLUMN contract_end_date TEXT")
    if 'contract_duration_months' not in account_presence_columns:
        c.execute("ALTER TABLE account_presences ADD COLUMN contract_duration_months INTEGER")
    if 'early_terminated' not in account_presence_columns:
        c.execute("ALTER TABLE account_presences ADD COLUMN early_terminated INTEGER DEFAULT 0")

    # Configuracoes padrao da faixa de status
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('status_green_days', '7'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('status_yellow_days', '14'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('target_green_days', '5'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('target_yellow_days', '10'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('cold_green_days', '45'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('cold_yellow_days', '60'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('iata_video_path', '/videos/TocaVideo.mp4'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('tavily_api_key', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('openrouter_api_key', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('openrouter_model', 'stepfun/step-3.5-flash:free'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('openrouter_site_url', 'http://localhost'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('openrouter_app_name', 'TocaDoCoelho'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('update_github_owner', DEFAULT_GITHUB_OWNER))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('update_github_repo', DEFAULT_GITHUB_REPO))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('update_github_token', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('update_snooze_until', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_base_snapshot', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_base_updated_at', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_sai_api_key', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_sai_template_id', '69ac3c87024adc2d2bdc19f5'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_sai_base_url', 'https://sai-library.saiapplications.com'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_action_detector_template_id', '69b1c662485ca1e93db65015'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_sai_simple_template_id', '69bc155d7462bf7c702e9295'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_sai_geral_claude_api_key', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('itoca_sai_geral_claude_template_id', '6a45658f1615d7b89d76c4ac'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('waha_api_url', 'http://localhost:3001'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('waha_api_key', ''))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('waha_session_name', 'default'))
    # Auto-configurar a partir de variáveis de ambiente injetadas pelo launcher (build bundled)
    _waha_url_env = os.environ.get('WAHA_API_URL', '').strip()
    _waha_key_env = os.environ.get('WAHA_API_KEY', '').strip()
    _waha_session_env = os.environ.get('WAHA_SESSION_NAME', '').strip()
    if _waha_url_env:
        c.execute("INSERT INTO app_settings (key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value WHERE value='' OR value='http://localhost:3001'", ('waha_api_url', _waha_url_env))
    if _waha_key_env:
        c.execute("INSERT INTO app_settings (key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value WHERE value=''", ('waha_api_key', _waha_key_env))
    if _waha_session_env:
        c.execute("INSERT INTO app_settings (key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value WHERE value='' OR value='default'", ('waha_session_name', _waha_session_env))
    # Cura URLs já salvas apontando para a porta do próprio app (causa loop Flask→Flask, 404/405).
    # A validação no PUT impede salvar novas, mas valores antigos no banco precisam ser corrigidos aqui.
    _app_port = str(os.environ.get('PORT', '3000'))
    _healed_url = _waha_url_env or 'http://localhost:3001'
    c.execute(
        "UPDATE app_settings SET value=? WHERE key='waha_api_url' AND (value LIKE ? OR value LIKE ?)",
        (_healed_url, f'%localhost:{_app_port}%', f'%127.0.0.1:{_app_port}%')
    )
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('outlook_sync_timeout_seconds', '120'))
    c.execute('INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)', ('outlook_sync_days', '30'))
    # Histórico de conversas do iToca (30 dias)
    c.execute('''
        CREATE TABLE IF NOT EXISTS itoca_chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT NOT NULL,
            role TEXT NOT NULL CHECK(role IN ('user','assistant')),
            content TEXT NOT NULL,
            confidence_percent INTEGER,
            needs_refinement INTEGER DEFAULT 0,
            refinement_hint TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # ── Campanha Ofensiva (planejamento + plano de ação comercial) ──────────
    # Uma campanha = uma ofensiva comercial sobre um conjunto de contas, cruzando
    # o desafio do usuário com o portfólio (Soluções STF) e o mapeamento de contatos.
    c.execute('''CREATE TABLE IF NOT EXISTS campaigns (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        objective_text TEXT,
        llm_summary TEXT,
        areas_json TEXT,
        decision_levels_json TEXT,
        offer_ids_json TEXT,
        has_solution INTEGER DEFAULT 1,
        portfolio_note TEXT,
        start_date TEXT,
        weekly_capacity INTEGER DEFAULT 5,
        llm_source TEXT,
        status TEXT DEFAULT 'Ativo',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # Contas selecionadas para a campanha, já classificadas por bloco (A/B/C/D).
    c.execute('''CREATE TABLE IF NOT EXISTS campaign_accounts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        campaign_id INTEGER NOT NULL,
        account_id INTEGER NOT NULL,
        account_name TEXT NOT NULL,
        block_type TEXT NOT NULL DEFAULT 'C',
        readiness_score INTEGER DEFAULT 0,
        offer_id INTEGER,
        offer_title TEXT,
        rationale TEXT,
        sort_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(campaign_id) REFERENCES campaigns(id) ON DELETE CASCADE,
        FOREIGN KEY(account_id) REFERENCES accounts(id) ON DELETE CASCADE,
        UNIQUE(campaign_id, account_id)
    )''')

    # Ações individuais do plano (apresentação, mapeamento, prospecção, portfólio).
    c.execute('''CREATE TABLE IF NOT EXISTS campaign_actions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        campaign_account_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        action_type TEXT DEFAULT 'apresentacao',
        contact_id INTEGER,
        contact_name TEXT,
        scheduled_week INTEGER DEFAULT 1,
        due_date TEXT,
        status TEXT DEFAULT 'pending',
        sort_order INTEGER DEFAULT 0,
        completed_at TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(campaign_account_id) REFERENCES campaign_accounts(id) ON DELETE CASCADE,
        FOREIGN KEY(contact_id) REFERENCES clients(id) ON DELETE SET NULL
    )''')

    # Logs/updates manuais das ações — append-only (nunca deletados em update do plano).
    c.execute('''CREATE TABLE IF NOT EXISTS campaign_action_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        action_id INTEGER NOT NULL,
        log_text TEXT NOT NULL,
        log_type TEXT DEFAULT 'user',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(action_id) REFERENCES campaign_actions(id) ON DELETE CASCADE
    )''')

    conn.commit()

    # Migração: adicionar colunas se não existirem
    try:
        c.execute("PRAGMA table_info(user_profile)")
        profile_cols = [col[1] for col in c.fetchall()]
        if 'email' not in profile_cols:
            c.execute('ALTER TABLE user_profile ADD COLUMN email TEXT')
        if 'phone' not in profile_cols:
            c.execute('ALTER TABLE user_profile ADD COLUMN phone TEXT')
        if 'boss_name' not in profile_cols:
            c.execute('ALTER TABLE user_profile ADD COLUMN boss_name TEXT')
        if 'boss_email' not in profile_cols:
            c.execute('ALTER TABLE user_profile ADD COLUMN boss_email TEXT')

        c.execute("PRAGMA table_info(activities)")
        columns = [col[1] for col in c.fetchall()]
        if 'contact_type' not in columns:
            c.execute('ALTER TABLE activities ADD COLUMN contact_type TEXT DEFAULT "Outro"')
        if 'information' not in columns:
            c.execute('ALTER TABLE activities ADD COLUMN information TEXT')

        # Garantir schema mínimo do WikiToca para bases antigas/parciais
        c.execute("PRAGMA table_info(wiki_entries)")
        wiki_entry_columns = [col[1] for col in c.fetchall()]
        if 'category' not in wiki_entry_columns:
            c.execute('ALTER TABLE wiki_entries ADD COLUMN category TEXT')
        if 'tags' not in wiki_entry_columns:
            c.execute('ALTER TABLE wiki_entries ADD COLUMN tags TEXT')
        if 'created_at' not in wiki_entry_columns:
            c.execute('ALTER TABLE wiki_entries ADD COLUMN created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP')
        if 'updated_at' not in wiki_entry_columns:
            c.execute('ALTER TABLE wiki_entries ADD COLUMN updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP')

        c.execute("PRAGMA table_info(wiki_documents)")
        wiki_doc_columns = [col[1] for col in c.fetchall()]
        if 'file_ext' not in wiki_doc_columns:
            c.execute('ALTER TABLE wiki_documents ADD COLUMN file_ext TEXT')
        if 'file_size' not in wiki_doc_columns:
            c.execute('ALTER TABLE wiki_documents ADD COLUMN file_size INTEGER')
        if 'created_at' not in wiki_doc_columns:
            c.execute('ALTER TABLE wiki_documents ADD COLUMN created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP')
        if 'updated_at' not in wiki_doc_columns:
            c.execute('ALTER TABLE wiki_documents ADD COLUMN updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP')

        # Campanha Ofensiva — colunas adicionais (insight de mercado e ângulo da ação)
        c.execute("PRAGMA table_info(campaign_accounts)")
        camp_acc_cols = [col[1] for col in c.fetchall()]
        if 'market_insight' not in camp_acc_cols:
            c.execute('ALTER TABLE campaign_accounts ADD COLUMN market_insight TEXT')
        if 'insight_sources' not in camp_acc_cols:
            c.execute('ALTER TABLE campaign_accounts ADD COLUMN insight_sources TEXT')
        c.execute("PRAGMA table_info(campaign_actions)")
        camp_act_cols = [col[1] for col in c.fetchall()]
        if 'angle_area' not in camp_act_cols:
            c.execute('ALTER TABLE campaign_actions ADD COLUMN angle_area TEXT')
        conn.commit()
    except Exception as e:
        logger.debug(f'[init_db] exceção ignorada: {e}')
    
    conn.close()
    logger.info('[Database] Banco de dados inicializado')


def run_automatic_db_backup(interval_days=3):
    try:
        conn = sqlite3.connect(str(DB_PATH))
        c = conn.cursor()
        c.execute('SELECT value FROM app_settings WHERE key = ?', ('db_last_backup_at',))
        row = c.fetchone()

        now = datetime.now()
        should_backup = True

        if row and row[0]:
            try:
                last_backup = datetime.fromisoformat(row[0])
                should_backup = (now - last_backup) >= timedelta(days=interval_days)
            except Exception:
                should_backup = True

        if should_backup and DB_PATH.exists():
            backup_name = f"toca-do-coelho-backup-{now.strftime('%Y%m%d-%H%M%S')}.db"
            backup_path = BACKUP_DIR / backup_name
            shutil.copy2(str(DB_PATH), str(backup_path))
            c.execute(
                'INSERT INTO app_settings (key, value) VALUES (?, ?) '
                'ON CONFLICT(key) DO UPDATE SET value = excluded.value, updated_at = CURRENT_TIMESTAMP',
                ('db_last_backup_at', now.isoformat())
            )
            conn.commit()
            logger.info(f'[Backup] Backup automático criado em: {backup_path}')
        else:
            logger.info('[Backup] Backup automático não necessário neste momento.')

        conn.close()
    except Exception as e:
        logger.exception(f'[Backup] Falha ao executar backup automático: {e}')

# ---------------------------------------------------------------------------
# Sistema de migração versionada do schema.
# A migração 1 é a "baseline" legada: o init_db() idempotente (CREATE TABLE IF
# NOT EXISTS + ALTER TABLE condicionais via PRAGMA table_info), que cobre tanto
# bancos novos quanto bancos antigos dos usuários (o instalador atualiza os
# binários mas mantém o .db em %AppData%). Novas mudanças de schema devem ser
# adicionadas como entradas numeradas em SCHEMA_MIGRATIONS — nunca mais como
# ALTER TABLE ad-hoc no init_db.
# ---------------------------------------------------------------------------
SCHEMA_MIGRATIONS = [
    (1, 'baseline_legacy_init', None),  # None => roda init_db()
    (2, 'indices_consultas_frequentes', [
        'CREATE INDEX IF NOT EXISTS idx_activities_client_date ON activities(client_id, activity_date)',
        'CREATE INDEX IF NOT EXISTS idx_clients_company ON clients(company)',
        'CREATE INDEX IF NOT EXISTS idx_clients_last_activity ON clients(last_activity_date)',
        'CREATE INDEX IF NOT EXISTS idx_commitments_due_date ON commitments(due_date)',
        'CREATE INDEX IF NOT EXISTS idx_account_activities_account ON account_activities(account_id)',
    ]),
    (3, 'background_tasks_persistentes', [
        '''CREATE TABLE IF NOT EXISTS background_tasks (
            task_id TEXT PRIMARY KEY,
            kind TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'processing',
            step TEXT,
            progress INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''',
    ]),
    (4, 'account_planning_runs', [
        '''CREATE TABLE IF NOT EXISTS account_planning_runs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company TEXT NOT NULL,
            segment TEXT,
            result_json TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''',
        'CREATE INDEX IF NOT EXISTS idx_account_planning_created ON account_planning_runs(created_at)',
    ]),
    (5, 'radar_do_dia_score_snooze', [
        'ALTER TABLE daily_suggestions ADD COLUMN score REAL DEFAULT 0',
        'ALTER TABLE daily_suggestions ADD COLUMN snoozed_until TEXT',
    ]),
    (6, 'inbound_messages', [
        '''CREATE TABLE IF NOT EXISTS inbound_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER NOT NULL,
            channel TEXT NOT NULL DEFAULT 'whatsapp',
            received_at TEXT NOT NULL,
            preview TEXT,
            responded_at TEXT,
            source_msg_id TEXT UNIQUE,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE
        )''',
        'CREATE INDEX IF NOT EXISTS idx_inbound_pending ON inbound_messages(responded_at, received_at)',
        'CREATE INDEX IF NOT EXISTS idx_inbound_client ON inbound_messages(client_id, channel)',
    ]),
    (7, 'whatsapp_sends', [
        '''CREATE TABLE IF NOT EXISTS whatsapp_sends (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER,
            phone TEXT,
            message TEXT,
            status TEXT NOT NULL DEFAULT 'sent',
            error TEXT,
            sent_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''',
        "CREATE INDEX IF NOT EXISTS idx_wa_sends_date ON whatsapp_sends(sent_at)",
    ]),
    (8, 'clients_birthday', [
        'ALTER TABLE clients ADD COLUMN birthday TEXT',
    ]),
    (9, 'job_change_events', [
        '''CREATE TABLE IF NOT EXISTS job_change_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER NOT NULL,
            empresa_antiga TEXT,
            empresa_nova TEXT,
            cargo_novo TEXT,
            potential_new_account INTEGER DEFAULT 0,
            status TEXT NOT NULL DEFAULT 'pendente',
            detected_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE
        )''',
        'CREATE INDEX IF NOT EXISTS idx_job_change_status ON job_change_events(status)',
    ]),
    (10, 'meeting_briefings', [
        '''CREATE TABLE IF NOT EXISTS meeting_briefings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            commitment_id INTEGER NOT NULL UNIQUE,
            content_md TEXT NOT NULL,
            generated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(commitment_id) REFERENCES commitments(id) ON DELETE CASCADE
        )''',
    ]),
    (11, 'scheduled_sends', [
        '''CREATE TABLE IF NOT EXISTS scheduled_sends (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            channel TEXT NOT NULL DEFAULT 'whatsapp',
            client_id INTEGER,
            phone TEXT,
            email_to TEXT,
            subject TEXT,
            message TEXT NOT NULL,
            scheduled_for TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            error TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            sent_at TIMESTAMP
        )''',
        "CREATE INDEX IF NOT EXISTS idx_scheduled_sends_due ON scheduled_sends(status, scheduled_for)",
    ]),
    (12, 'scheduled_sends_activity_link', [
        'ALTER TABLE scheduled_sends ADD COLUMN activity_id INTEGER',
    ]),
    (13, 'inbound_purge_system_notifications', [
        # Avisos de sistema do WhatsApp (ex.: criptografia ponta-a-ponta) eram
        # registrados como pendência de resposta por engano — só geram esse
        # preview quando não há corpo de texto nem mídia (ver _waha_is_content_message).
        "DELETE FROM inbound_messages WHERE preview = '(mensagem sem texto)' AND responded_at IS NULL",
    ]),
]


def _run_schema_migrations():
    conn = sqlite3.connect(str(DB_PATH), timeout=5.0)
    try:
        conn.execute('PRAGMA journal_mode=WAL')
        conn.execute('PRAGMA busy_timeout=5000')
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS schema_version (
            version INTEGER PRIMARY KEY,
            name TEXT,
            applied_at TEXT
        )''')
        conn.commit()
        c.execute('SELECT MAX(version) FROM schema_version')
        row = c.fetchone()
        current = row[0] if row and row[0] else 0
        for version, name, statements in SCHEMA_MIGRATIONS:
            if version <= current:
                continue
            if statements is None:
                # Baseline legada: init_db() usa a própria conexão.
                init_db()
            else:
                for stmt in statements:
                    c.execute(stmt)
            c.execute(
                'INSERT INTO schema_version (version, name, applied_at) VALUES (?, ?, ?)',
                (version, name, datetime.now().isoformat(timespec='seconds'))
            )
            conn.commit()
            logger.info(f'[Database] Migração {version} ({name}) aplicada')
    finally:
        conn.close()


def _mark_interrupted_background_tasks():
    """Ao subir, marca tasks 'processing' de execuções anteriores como interrompidas."""
    try:
        conn = sqlite3.connect(str(DB_PATH), timeout=5.0)
        conn.execute(
            "UPDATE background_tasks SET status = 'interrupted', updated_at = CURRENT_TIMESTAMP "
            "WHERE status = 'processing'"
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.debug(f'[Tasks] Falha ao marcar tasks interrompidas: {e}')


_run_schema_migrations()
_mark_interrupted_background_tasks()
run_automatic_db_backup(interval_days=3)

# Funcoes auxiliares
def get_db():
    conn = sqlite3.connect(str(DB_PATH), timeout=5.0)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute('PRAGMA journal_mode=WAL')
        conn.execute('PRAGMA busy_timeout=5000')
        conn.execute('PRAGMA foreign_keys=ON')
    except Exception as e:
        logger.debug(f'[Database] Falha ao aplicar PRAGMAs de conexão: {e}')
    # Garante que o arquivo de banco só seja acessível pelo usuário atual do SO
    try:
        import stat as _stat
        os.chmod(str(DB_PATH), _stat.S_IRUSR | _stat.S_IWUSR)
    except Exception as e:
        logger.debug(f'[get_db] exceção ignorada: {e}')
    return conn

def dict_from_row(row):
    if row is None:
        return None
    return dict(row)


def _load_app_settings_map(keys):
    conn = get_db()
    c = conn.cursor()
    placeholders = ','.join(['?'] * len(keys))
    c.execute(f'SELECT key, value FROM app_settings WHERE key IN ({placeholders})', tuple(keys))
    mapping = {row['key']: row['value'] for row in c.fetchall()}
    conn.close()
    return mapping


def _normalize_version(version):
    normalized = re.sub(r'^[vV]', '', str(version or '').strip())
    return normalized


def _version_key(version):
    normalized = _normalize_version(version)
    if not normalized:
        return tuple()
    parts = re.split(r'[.+\-]', normalized)
    parsed = []
    for part in parts:
        if part.isdigit():
            parsed.append((0, int(part)))
        else:
            parsed.append((1, part.lower()))
    return tuple(parsed)

def _resolve_setting(secret_key, env_key):
    try:
        db_value = (_load_app_settings_map([secret_key]).get(secret_key) or '').strip()
    except Exception:
        db_value = ''
    if db_value:
        return db_value
    return (os.environ.get(env_key, '') or '').strip()


def _sai_execute_question_template(base_url, template_id, api_key, question, log_tag):
    """Executa um template SAI que aceita {"inputs": {"question": ...}} com retry em HTTP 429
    (rate limit transitório). Retorna o texto da resposta ou None."""
    max_attempts = 3
    backoff_seconds = (2, 4)
    for attempt in range(max_attempts):
        try:
            resp = requests.post(
                f'{base_url}/api/templates/{template_id}/execute',
                json={'inputs': {'question': question}},
                headers={'X-Api-Key': api_key},
                timeout=45
            )
            if resp.status_code == 200:
                logger.info(f'[SAI][{log_tag}] OK ({len(resp.text)} chars)')
                return resp.text
            if resp.status_code == 429 and attempt < max_attempts - 1:
                wait = backoff_seconds[min(attempt, len(backoff_seconds) - 1)]
                logger.warning(f'[SAI][{log_tag}] HTTP 429 (rate limit), tentativa {attempt + 1}/{max_attempts}, aguardando {wait}s...')
                time.sleep(wait)
                continue
            logger.warning(f'[SAI][{log_tag}] HTTP {resp.status_code}: {resp.text[:200]}')
            return None
        except Exception as e:
            logger.warning(f'[SAI][{log_tag}] falha: {e}')
            return None
    return None


def _sai_simple_prompt(question: str) -> str | None:
    """Envia uma pergunta ao template SAI de prompt simples e retorna o texto da resposta.

    Este é o helper padrão para chamar o LLM via SAI no TocaDoCoelho.
    Use esta função sempre que precisar de uma resposta de LLM para uma pergunta livre.
    A chave e a URL base são lidas automaticamente das configurações do app (itoca_sai_api_key,
    itoca_sai_base_url). O template usado é o 'simple prompt' (itoca_sai_simple_template_id),
    que aceita apenas o campo 'question' como entrada.

    Como o template 'simple prompt' é compartilhado por muitas features do app e pode saturar
    sua cota de uso (HTTP 429) sob carga concorrente, esta função tenta primeiro o template SAI
    'Geral Claude' (itoca_sai_geral_claude_template_id) — uma integração SAI separada com chave
    e cota próprias — e só cai para o 'simple prompt' (com retry em 429) se o Geral Claude falhar.

    Retorna o texto bruto da resposta do LLM, ou None se nenhum dos dois estiver configurado
    ou disponível.

    Exemplo de uso:
        raw = _sai_simple_prompt("Qual o faturamento anual da Petrobras? Responda em JSON.")
        # raw pode ser '{"faturamento": 500000000000}' ou None
    """
    base_url = (_load_app_settings_map(['itoca_sai_base_url']).get('itoca_sai_base_url') or '').strip() or 'https://sai-library.saiapplications.com'

    geral_claude_key = _resolve_setting('itoca_sai_geral_claude_api_key', 'ITOCA_SAI_GERAL_CLAUDE_API_KEY') or 'fuBoUPGL+UmrErevVE6VWQ'
    if geral_claude_key:
        geral_claude_template_id = (_load_app_settings_map(['itoca_sai_geral_claude_template_id']).get('itoca_sai_geral_claude_template_id') or '').strip() or '6a45658f1615d7b89d76c4ac'
        result = _sai_execute_question_template(base_url, geral_claude_template_id, geral_claude_key, question, 'geral_claude')
        if result:
            return result

    api_key = _resolve_setting('itoca_sai_api_key', 'ITOCA_SAI_API_KEY')
    if api_key:
        template_id = (_load_app_settings_map(['itoca_sai_simple_template_id']).get('itoca_sai_simple_template_id') or '').strip() or '69bc155d7462bf7c702e9295'
        result = _sai_execute_question_template(base_url, template_id, api_key, question, 'simple_prompt')
        if result:
            return result

    return None


def api_error(status, code, message, details=None, hint=None):
    payload = {
        'error': message,
        'error_code': code
    }
    if details:
        payload['details'] = str(details)
    if hint:
        payload['hint'] = hint
    return jsonify(payload), status


def _extract_bing_image_urls(raw_html, log_prefix='[AutoPic]'):
    """Extrai as URLs originais das imagens do HTML do Bing Images.
    O Bing codifica os dados de imagem com HTML entities nos atributos data-*,
    por isso o padrão usa &quot; ao invés de aspas diretas.
    """
    # Padrão principal: Bing usa HTML entities nos atributos data-*
    matches = re.findall(r'&quot;murl&quot;:&quot;(https?://[^&]+)&quot;', raw_html)
    logger.debug(f'{log_prefix} _extract_bing_image_urls: padrão HTML entities encontrou {len(matches)} resultado(s).')

    if not matches:
        # Fallback: tentar sem HTML entities (caso o Bing mude o formato)
        matches = re.findall(r'"murl":"(https?://[^"]+)"', raw_html)
        logger.debug(f'{log_prefix} _extract_bing_image_urls: padrão aspas diretas (fallback) encontrou {len(matches)} resultado(s).')

    if not matches:
        logger.warning(f'{log_prefix} _extract_bing_image_urls: NENHUM resultado encontrado com nenhum padrão. '
                       f'Tamanho do HTML: {len(raw_html)} chars. '
                       f'Contém "murl": {"murl" in raw_html}. '
                       f'Primeiros 300 chars: {repr(raw_html[:300])}')

    urls = []
    for item in matches:
        url = html.unescape(item).replace('\\/', '/')
        if url.startswith('http://') or url.startswith('https://'):
            urls.append(url)
    return urls


def _llm_prompt(question, log_tag='llm', temperature=0.1, web=False):
    """REGRA DE DESENVOLVIMENTO (ver CLAUDE.md): toda automação com IA usa
    SAI primeiro e OpenRouter como fallback. Única exceção: perguntas que
    exigem busca ativa na internet (web=True) — aí o OpenRouter (com plugin
    de web) vem primeiro, já que os templates SAI não têm acesso à web.
    Retorna str com a resposta, ou None se nenhum provider respondeu."""
    if web:
        raw = _openrouter_web_prompt(question)
        if raw and str(raw).strip():
            logger.info(f'[{log_tag}][OpenRouter/web] Resposta gerada com sucesso')
            return raw
        return _sai_simple_prompt(question)

    raw = _sai_simple_prompt(question)
    if raw and str(raw).strip():
        logger.info(f'[{log_tag}][SAI] Resposta gerada com sucesso')
        return raw

    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
        model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
        site_url = (or_settings.get('openrouter_site_url') or 'http://localhost').strip() or 'http://localhost'
        app_name = (or_settings.get('openrouter_app_name') or 'TocaDoCoelho').strip() or 'TocaDoCoelho'
        try:
            resp = requests.post(
                'https://openrouter.ai/api/v1/chat/completions',
                json={
                    'model': model,
                    'messages': [{'role': 'user', 'content': question}],
                    'temperature': temperature
                },
                headers={
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                timeout=90
            )
            data = resp.json()
            choices = data.get('choices') or []
            raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
            if raw and str(raw).strip():
                logger.info(f'[{log_tag}][OpenRouter] Resposta gerada com sucesso (fallback)')
                return raw
            logger.warning(f'[{log_tag}][OpenRouter] Resposta vazia.')
        except Exception as e:
            logger.warning(f'[{log_tag}][OpenRouter] Falha no fallback: {e}')
    return None


def _find_image_candidates_on_web(query, limit=3):
    """Busca as primeiras imagens do Bing Images para o query fornecido.
    Retorna lista de URLs das imagens originais (as primeiras que aparecem na tela).
    """
    log_prefix = '[AutoPic]'
    if not query:
        logger.warning(f'{log_prefix} _find_image_candidates_on_web: query vazio, abortando.')
        return []

    limit = max(1, int(limit or 3))
    user_agent = (
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
        'AppleWebKit/537.36 (KHTML, like Gecko) '
        'Chrome/124.0.0.0 Safari/537.36'
    )
    search_url = f"https://www.bing.com/images/search?q={urllib.parse.quote(query)}&form=HDRSC2&first=1"
    logger.info(f'{log_prefix} Buscando imagens para query="{query}" limit={limit} url={search_url}')

    # Contexto SSL sem verificação para evitar erros de certificado em ambientes restritos
    ssl_ctx = ssl.create_default_context()
    ssl_ctx.check_hostname = False
    ssl_ctx.verify_mode = ssl.CERT_NONE

    # Solicitar apenas gzip (não brotli) para evitar dependência do módulo brotli
    req = urllib.request.Request(
        search_url,
        headers={
            'User-Agent': user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7',
            'Accept-Encoding': 'gzip, deflate',  # sem 'br' para evitar brotli
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
    )

    import gzip
    try:
        with urllib.request.urlopen(req, timeout=15, context=ssl_ctx) as response:
            raw_data = response.read()
            http_status = response.status
            encoding = (response.headers.get('Content-Encoding') or '').lower().strip()
            content_type = (response.headers.get('Content-Type') or '').lower()
            logger.info(f'{log_prefix} Resposta do Bing: HTTP {http_status}, '
                        f'Content-Encoding={encoding!r}, Content-Type={content_type!r}, '
                        f'tamanho bruto={len(raw_data)} bytes')

            if encoding == 'gzip':
                try:
                    content = gzip.decompress(raw_data).decode('utf-8', errors='ignore')
                    logger.debug(f'{log_prefix} Descomprimido gzip: {len(content)} chars')
                except Exception as gz_err:
                    logger.error(f'{log_prefix} Falha ao descomprimir gzip: {gz_err}. Tentando decodificar direto.')
                    content = raw_data.decode('utf-8', errors='ignore')
            elif encoding == 'br':
                # brotli pode não estar instalado — tentar e logar claramente
                try:
                    import brotli
                    content = brotli.decompress(raw_data).decode('utf-8', errors='ignore')
                    logger.debug(f'{log_prefix} Descomprimido brotli: {len(content)} chars')
                except ImportError:
                    logger.error(f'{log_prefix} ERRO CRÍTICO: O servidor retornou encoding brotli (br) mas o '
                                 f'módulo "brotli" não está instalado. '
                                 f'Execute: pip install brotli  (ou adicione ao requirements.txt). '
                                 f'O AutoPic não conseguirá encontrar imagens até isso ser resolvido.')
                    return []
                except Exception as br_err:
                    logger.error(f'{log_prefix} Falha ao descomprimir brotli: {br_err}. Tentando decodificar direto.')
                    content = raw_data.decode('utf-8', errors='ignore')
            elif encoding in ('deflate', ''):
                content = raw_data.decode('utf-8', errors='ignore')
                logger.debug(f'{log_prefix} Conteúdo sem compressão: {len(content)} chars')
            else:
                logger.warning(f'{log_prefix} Encoding desconhecido "{encoding}", tentando decodificar direto.')
                content = raw_data.decode('utf-8', errors='ignore')

    except urllib.error.HTTPError as http_err:
        logger.error(f'{log_prefix} HTTPError ao acessar Bing: {http_err.code} {http_err.reason}')
        raise
    except urllib.error.URLError as url_err:
        logger.error(f'{log_prefix} URLError ao acessar Bing: {url_err.reason}')
        raise
    except Exception as e:
        logger.error(f'{log_prefix} Erro inesperado ao acessar Bing: {type(e).__name__}: {e}')
        raise

    urls = _extract_bing_image_urls(content, log_prefix=log_prefix)
    logger.info(f'{log_prefix} Total de URLs extraídas antes de deduplicar: {len(urls)}')

    # Remover duplicatas mantendo a ordem (as primeiras são as que aparecem na tela)
    seen = set()
    unique_urls = []
    for u in urls:
        if u not in seen:
            seen.add(u)
            unique_urls.append(u)

    result = unique_urls[:limit]
    logger.info(f'{log_prefix} Retornando {len(result)} candidato(s) para query="{query}": {result}')
    return result


def _download_remote_image_to_uploads(image_url, prefix='autofind'):
    log_prefix = '[AutoPic]'
    logger.info(f'{log_prefix} Download de imagem: url={image_url[:120]!r} prefix={prefix!r}')

    parsed = urllib.parse.urlparse(image_url)
    if parsed.scheme not in {'http', 'https'}:
        logger.error(f'{log_prefix} URL inválida (scheme={parsed.scheme!r}): {image_url[:120]}')
        raise ValueError('URL de imagem inválida.')

    ssl_ctx = ssl.create_default_context()
    ssl_ctx.check_hostname = False
    ssl_ctx.verify_mode = ssl.CERT_NONE

    req = urllib.request.Request(image_url, headers={
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36',
        'Referer': 'https://www.bing.com/',
        'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
    })
    try:
        with urllib.request.urlopen(req, timeout=15, context=ssl_ctx) as response:
            http_status = response.status
            content_type = (response.headers.get('Content-Type') or '').lower()
            data = response.read(6 * 1024 * 1024 + 1)
            logger.info(f'{log_prefix} Download concluído: HTTP {http_status}, '
                        f'Content-Type={content_type!r}, tamanho={len(data)} bytes')
    except urllib.error.HTTPError as http_err:
        logger.error(f'{log_prefix} HTTPError ao baixar imagem {image_url[:120]!r}: '
                     f'{http_err.code} {http_err.reason}')
        raise
    except urllib.error.URLError as url_err:
        logger.error(f'{log_prefix} URLError ao baixar imagem {image_url[:120]!r}: {url_err.reason}')
        raise
    except Exception as e:
        logger.error(f'{log_prefix} Erro inesperado ao baixar imagem {image_url[:120]!r}: '
                     f'{type(e).__name__}: {e}')
        raise

    if len(data) > 6 * 1024 * 1024:
        logger.warning(f'{log_prefix} Imagem muito grande: {len(data)} bytes (máx 6MB)')
        raise ValueError('Imagem muito grande (máximo 6MB).')
    if not content_type.startswith('image/'):
        logger.error(f'{log_prefix} Content-Type inválido: {content_type!r} para url={image_url[:120]!r}')
        raise ValueError('A URL selecionada não retornou uma imagem válida.')

    ext = '.jpg'
    if 'png' in content_type:
        ext = '.png'
    elif 'webp' in content_type:
        ext = '.webp'
    elif 'gif' in content_type:
        ext = '.gif'

    filename = secure_filename(f"{prefix}-{int(time.time()*1000)}{ext}")
    path = UPLOAD_DIR / filename
    with open(path, 'wb') as f:
        f.write(data)
    logger.info(f'{log_prefix} Imagem salva em: {path}')
    return f'/uploads/{filename}'


def _linkedin_extract_og_image(linkedin_url):
    """Tenta extrair og:image da página pública do LinkedIn."""
    if not linkedin_url or 'linkedin.com' not in linkedin_url:
        return None
    try:
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        req = urllib.request.Request(linkedin_url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'pt-BR,pt;q=0.9,en;q=0.8',
        })
        with urllib.request.urlopen(req, context=ctx, timeout=15) as resp:
            raw_html = resp.read().decode('utf-8', errors='ignore')
        meta_match = re.search(
            r'<meta[^>]+property=["\']og:image["\'][^>]+content=["\']([^"\']+)["\']',
            raw_html,
            flags=re.IGNORECASE
        )
        if not meta_match:
            meta_match = re.search(
                r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+property=["\']og:image["\']',
                raw_html,
                flags=re.IGNORECASE
            )
        if not meta_match:
            return None
        candidate = html.unescape(meta_match.group(1)).strip()
        if not candidate.startswith(('http://', 'https://')):
            return None
        return candidate
    except Exception as e:
        logger.debug(f'[LinkedIn] Falha ao extrair og:image: {e}')
        return None


def parse_currency_to_cents(value):
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    text = re.sub(r'[^\d,.-]', '', text)
    if ',' in text:
        text = text.replace('.', '').replace(',', '.')
    try:
        return int(round(float(text) * 100))
    except Exception:
        digits = re.sub(r'\D', '', str(value))
        return int(digits) if digits else None


def format_currency_br(cents):
    if cents is None:
        return 'Não informado'
    try:
        value = (int(cents) or 0) / 100.0
    except Exception:
        return 'Não informado'
    formatted = f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    return f'R$ {formatted}'


def _relation_report_hex_to_color(value, fallback='#047857'):
    if not REPORTLAB_AVAILABLE:
        return None
    text = (value or fallback or '').strip()
    if not text:
        text = fallback
    if not text.startswith('#'):
        text = f'#{text}'
    if not re.fullmatch(r'#[0-9a-fA-F]{6}', text):
        text = fallback
    return colors.HexColor(text)


def _relation_report_pick_system_colors():
    settings = _load_app_settings_map([
        'theme_primary_color', 'theme_secondary_color', 'theme_accent_color',
        'brand_primary_color', 'brand_secondary_color', 'brand_accent_color'
    ])
    primary = settings.get('theme_primary_color') or settings.get('brand_primary_color') or '#047857'
    secondary = settings.get('theme_secondary_color') or settings.get('brand_secondary_color') or '#065f46'
    accent = settings.get('theme_accent_color') or settings.get('brand_accent_color') or '#34d399'
    return {
        'primary': _relation_report_hex_to_color(primary, '#047857'),
        'secondary': _relation_report_hex_to_color(secondary, '#065f46'),
        'accent': _relation_report_hex_to_color(accent, '#34d399'),
        'primary_hex': primary if str(primary).startswith('#') else f'#{primary}',
        'secondary_hex': secondary if str(secondary).startswith('#') else f'#{secondary}',
        'accent_hex': accent if str(accent).startswith('#') else f'#{accent}'
    }


def _relation_report_resolve_local_image(candidate_url):
    if not candidate_url:
        return None
    try:
        parsed = urlparse(candidate_url)
        path = parsed.path or candidate_url
    except Exception:
        path = candidate_url
    path = str(path).strip()
    if not path:
        return None
    if path.startswith('/uploads/accounts/'):
        local = ACCOUNT_UPLOAD_DIR / Path(path).name
        return local if local.exists() else None
    if path.startswith('/uploads/'):
        local = UPLOAD_DIR / path.replace('/uploads/', '', 1)
        return local if local.exists() else None
    local_path = Path(path)
    if local_path.exists():
        return local_path
    project_public = Path(__file__).resolve().parent / 'public'
    for rel in ['logo-coelho.png', 'favicon.png', 'coelho-sugestoes.png']:
        probe = project_public / rel
        if probe.exists() and Path(path).name == rel:
            return probe
    return None


def _relation_report_system_logo_path():
    project_public = Path(__file__).resolve().parent / 'public'
    for rel in ['logo-coelho.png', 'favicon.png']:
        probe = project_public / rel
        if probe.exists():
            return probe
    return None


def _relation_report_safe_image(image_path, max_width=220, max_height=120):
    if not image_path or not PIL_AVAILABLE:
        return None
    try:
        img = PILImage.open(str(image_path)).convert('RGBA')
        img.thumbnail((max_width, max_height))
        background = PILImage.new('RGBA', img.size, (255, 255, 255, 255))
        background.alpha_composite(img)
        return ImageReader(background.convert('RGB'))
    except Exception:
        return None


def _save_avatar_photo(file_storage, filepath, size=320):
    """Salva uma foto de avatar já redimensionada e nitidificada.

    Fotos de celular chegam em resolução muito maior que o círculo de
    avatar exibido na tela (ex.: 44-56px). O downscale "cru" feito pelo
    navegador nessa proporção (>15x) borra detalhes finos (cabelo, barba)
    porque não aplica nenhum sharpening pós-redução. Pré-processar aqui
    evita esse borrão em qualquer lugar do app onde o avatar apareça.
    """
    if not PIL_AVAILABLE:
        file_storage.save(str(filepath))
        return
    try:
        img = PILImage.open(file_storage.stream)
        img = ImageOps.exif_transpose(img)
        img = img.convert('RGB')
        img = ImageOps.fit(img, (size, size), PILImage.LANCZOS)
        img = img.filter(ImageFilter.UnsharpMask(radius=1.2, percent=60, threshold=2))
        img.save(str(filepath), format='JPEG', quality=90, optimize=True)
    except Exception:
        file_storage.stream.seek(0)
        file_storage.save(str(filepath))


def _relation_report_parse_dt(value):
    if not value:
        return None
    raw = str(value).strip()
    for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%dT%H:%M:%S', '%Y-%m-%d', '%Y-%m', '%Y-%m-%d %H:%M'):
        try:
            sample = raw[:19] if ('%H' in fmt or 'T' in fmt) else raw[:10]
            return datetime.strptime(sample, fmt)
        except Exception:
            continue
    try:
        return datetime.fromisoformat(raw)
    except Exception:
        return None


def _relation_report_format_dt(value):
    dt = _relation_report_parse_dt(value)
    if not dt:
        return 'Não informado'
    raw = str(value).strip()
    if len(raw) <= 10:
        return dt.strftime('%d/%m/%Y')
    return dt.strftime('%d/%m/%Y %H:%M')


def _relation_report_period_clause(date_column, start_date=None, end_date=None):
    clauses = []
    params = []
    if start_date:
        clauses.append(f"date({date_column}) >= date(?)")
        params.append(start_date)
    if end_date:
        clauses.append(f"date({date_column}) <= date(?)")
        params.append(end_date)
    return (' AND ' + ' AND '.join(clauses)) if clauses else '', params


RELATION_REPORT_TOPICS = ['DWP e Infraestrutura', 'Application', 'Marketing', 'Cyber', 'Cloud', 'Outros']

RELATION_REPORT_TOPIC_RULES = [
    ('DWP e Infraestrutura', [
        ' dwp', ' dws', 'digital workplace', 'digital workspace', 'workplace',
        'workspace', 'modern workplace', 'm365', 'microsoft 365', 'office 365',
        ' o365', 'intune', 'endpoint', ' vdi', 'service desk', 'servicedesk',
        'help desk', 'helpdesk', 'service-desk', 'field service', 'field services',
        'haas', 'hardware as a service', 'daas', 'device as a service',
        'alocação de máquina', 'alocacao de maquina', 'venda de máquina',
        'venda de maquina', 'locação de máquina', 'locacao de maquina',
        'parque de máquinas', 'parque de maquinas', 'parque tecnológico',
        'notebook', 'desktop', 'estação de trabalho', 'estacao de trabalho',
        'workstation', 'infraestrutura', ' infra ', 'datacenter', 'data center',
        'servidor', ' rede ', 'network', 'on-premises', 'on premise',
        'backup', 'storage',
    ]),
    ('Application', [
        'aplica', 'application', 'software', 'sistema', ' app ', ' erp', ' crm',
        ' sap', 'oracle', 'totvs', 'salesforce', 'desenvolvimento', ' dev ',
        'low-code', 'low code', 'integração', 'integracao', 'modernização',
        'modernizacao', 'sustentação', 'sustentacao',
    ]),
    ('Marketing', [
        'marketing', 'campanha', 'mídia', 'midia', ' lead', 'brand', 'marca',
        'comunicação', 'comunicacao', 'evento', 'branding', 'publicidade',
    ]),
    ('Cyber', [
        'cyber', 'segurança', 'seguranca', 'security', ' soc ', 'siem',
        'firewall', ' iam ', 'identity', 'phishing', 'ransomware',
        'vulnerabilidade', 'pentest', 'lgpd', 'compliance',
    ]),
    ('Cloud', [
        'cloud', 'nuvem', ' aws', 'azure', ' gcp', 'google cloud',
        'oracle cloud', 'multicloud', 'finops', 'kubernetes', 'container',
        'devops', 'landing zone',
    ]),
]


def _relation_report_topic_from_text(text):
    content = f" {str(text or '').lower()} "
    best_topic = 'Outros'
    best_score = 0
    for topic, keywords in RELATION_REPORT_TOPIC_RULES:
        score = sum(1 for keyword in keywords if keyword in content)
        if score > best_score:
            best_score = score
            best_topic = topic
    return best_topic


def _relation_report_collect_data(account_id, start_date=None, end_date=None):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM accounts WHERE id = ?", (account_id,))
    account = dict_from_row(c.fetchone())
    if not account:
        conn.close()
        return None

    c.execute("""SELECT c.*, CASE WHEN amc.client_id IS NOT NULL THEN 1 ELSE 0 END AS is_main_contact
                 FROM clients c
                 LEFT JOIN account_main_contacts amc ON amc.client_id = c.id AND amc.account_id = ?
                 WHERE LOWER(TRIM(c.company)) = LOWER(TRIM(?))
                 ORDER BY is_main_contact DESC, c.name COLLATE NOCASE""", (account_id, account['name']))
    contacts = [dict_from_row(r) for r in c.fetchall()]
    contact_ids = [row['id'] for row in contacts]

    c.execute("""SELECT ap.*, fc.name AS focal_client_name
                 FROM account_presences ap
                 LEFT JOIN clients fc ON fc.id = ap.focal_client_id
                 WHERE ap.account_id = ?
                 ORDER BY ap.delivery_name COLLATE NOCASE""", (account_id,))
    presences = [dict_from_row(r) for r in c.fetchall()]

    period_sql_activities, activity_params = _relation_report_period_clause('a.activity_date', start_date, end_date)
    activities = []
    if contact_ids:
        placeholders = ','.join(['?'] * len(contact_ids))
        c.execute(f"""SELECT a.*, cl.name AS client_name, cl.position AS client_position, cl.company AS client_company
                      FROM activities a
                      JOIN clients cl ON cl.id = a.client_id
                      WHERE a.client_id IN ({placeholders}) {period_sql_activities}
                      ORDER BY datetime(a.activity_date) DESC, a.id DESC""", tuple(contact_ids) + tuple(activity_params))
        activities = [dict_from_row(r) for r in c.fetchall()]

    period_sql_account_acts, account_activity_params = _relation_report_period_clause('activity_date', start_date, end_date)
    c.execute(f"""SELECT id, account_id, description, activity_date, created_at
                  FROM account_activities
                  WHERE account_id = ? {period_sql_account_acts}
                  ORDER BY datetime(activity_date) DESC, created_at DESC""", (account_id, *account_activity_params))
    account_activities = [dict_from_row(r) for r in c.fetchall()]

    environment_responses = []
    if contact_ids:
        placeholders = ','.join(['?'] * len(contact_ids))
        c.execute(f"""SELECT er.*, ec.title AS card_title, ec.description AS card_description,
                             cl.name AS client_name, cl.position AS client_position
                      FROM environment_responses er
                      JOIN environment_cards ec ON ec.id = er.card_id
                      JOIN clients cl ON cl.id = er.client_id
                      WHERE er.client_id IN ({placeholders})
                      ORDER BY ec.display_order ASC, ec.title COLLATE NOCASE, cl.name COLLATE NOCASE""", tuple(contact_ids))
        environment_responses = [dict_from_row(r) for r in c.fetchall()]

    period_sql_kanban, kanban_params = _relation_report_period_clause('kc.updated_at', start_date, end_date)
    contact_condition = ''
    params = [account_id]
    if contact_ids:
        contact_condition = ' OR kc.contact_id IN ({})'.format(','.join(['?'] * len(contact_ids)))
        params.extend(contact_ids)
    c.execute(f"""SELECT kc.*, kcol.title AS column_title, cl.name AS contact_name, ac.name AS account_name
                  FROM kanban_cards kc
                  LEFT JOIN kanban_columns kcol ON kcol.id = kc.column_id
                  LEFT JOIN clients cl ON cl.id = kc.contact_id
                  LEFT JOIN accounts ac ON ac.id = kc.account_id
                  WHERE (kc.account_id = ? {contact_condition}) {period_sql_kanban}
                  ORDER BY datetime(kc.updated_at) DESC, kc.id DESC""", tuple(params) + tuple(kanban_params))
    kanban_cards = [dict_from_row(r) for r in c.fetchall()]

    for card in kanban_cards:
        c.execute("SELECT content, created_at FROM kanban_card_activities WHERE card_id = ? ORDER BY created_at DESC", (card['id'],))
        card['activities'] = [dict_from_row(r) for r in c.fetchall()]

    latest_interaction = None
    latest_candidates = []
    for item in activities:
        latest_candidates.append({
            'date': item.get('activity_date') or item.get('created_at'),
            'person': item.get('client_name'),
            'source': 'Atividade',
            'description': item.get('description') or item.get('information')
        })
    for item in account_activities:
        latest_candidates.append({
            'date': item.get('activity_date') or item.get('created_at'),
            'person': account.get('name'),
            'source': 'Atividade da conta',
            'description': item.get('description')
        })
    latest_candidates = [x for x in latest_candidates if x.get('date')]
    latest_candidates.sort(key=lambda x: _relation_report_parse_dt(x.get('date')) or datetime.min, reverse=True)
    if latest_candidates:
        latest_interaction = latest_candidates[0]

    relationship_cards = []
    for contact in contacts:
        contact_activities = [a for a in activities if a.get('client_id') == contact.get('id')]
        contact_responses = [r for r in environment_responses if r.get('client_id') == contact.get('id')]
        contact_kanban = [k for k in kanban_cards if k.get('contact_id') == contact.get('id')]
        last_contact = None
        if contact_activities:
            act = sorted(contact_activities, key=lambda x: _relation_report_parse_dt(x.get('activity_date')) or datetime.min, reverse=True)[0]
            last_contact = {
                'date': act.get('activity_date') or act.get('created_at'),
                'summary': act.get('description') or act.get('information') or 'Interação registrada'
            }
        relationship_cards.append({
            'contact': contact,
            'activities_count': len(contact_activities),
            'mapping_count': len(contact_responses),
            'kanban_count': len(contact_kanban),
            'last_contact': last_contact,
            'topics': sorted(set(_relation_report_topic_from_text(' '.join(filter(None, [a.get('description') or '', a.get('information') or '']))) for a in contact_activities if (a.get('description') or a.get('information'))))
        })

    topic_buckets = {key: [] for key in RELATION_REPORT_TOPICS}
    topic_sources = []
    for item in activities:
        topic_sources.append({
            'text': ' '.join(filter(None, [item.get('description'), item.get('information')])),
            'date': item.get('activity_date') or item.get('created_at'),
            'person': item.get('client_name'),
            'source': 'atividade'
        })
    for item in account_activities:
        topic_sources.append({
            'text': item.get('description') or '',
            'date': item.get('activity_date') or item.get('created_at'),
            'person': account.get('name'),
            'source': 'atividade_conta'
        })
    for item in kanban_cards:
        card_activity_text = ' '.join(
            str(a.get('content') or '') for a in (item.get('activities') or [])
        )
        topic_sources.append({
            'text': ' '.join(filter(None, [item.get('title'), item.get('description'), card_activity_text])),
            'date': item.get('updated_at') or item.get('created_at'),
            'person': item.get('contact_name') or account.get('name'),
            'source': 'kanban'
        })
    for item in environment_responses:
        topic_sources.append({
            'text': ' '.join(filter(None, [item.get('card_title'), item.get('card_description'), item.get('response')])),
            'date': item.get('updated_at'),
            'person': item.get('client_name'),
            'source': 'mapeamento'
        })
    for source in topic_sources:
        if not str(source.get('text') or '').strip():
            continue
        topic = _relation_report_topic_from_text(source.get('text'))
        topic_buckets[topic].append(source)

    summary_counts = {
        'contacts': len(contacts),
        'main_contacts': len([c for c in contacts if c.get('is_main_contact')]),
        'activities': len(activities),
        'account_activities': len(account_activities),
        'mapping_items': len(environment_responses),
        'kanban_cards': len(kanban_cards),
        'presences': len(presences)
    }

    conn.close()
    return {
        'account': account,
        'contacts': contacts,
        'presences': presences,
        'activities': activities,
        'account_activities': account_activities,
        'environment_responses': environment_responses,
        'kanban_cards': kanban_cards,
        'latest_interaction': latest_interaction,
        'relationship_cards': relationship_cards,
        'topics': topic_buckets,
        'summary_counts': summary_counts,
        'start_date': start_date,
        'end_date': end_date,
        'full_period': not start_date and not end_date
    }


def _relation_report_build_llm_context(report_data):
    account = report_data['account']
    lines = []
    lines.append(f"Conta: {account.get('name')}")
    lines.append(f"Setor: {account.get('sector') or 'Não informado'}")
    lines.append(f"Receita média: {format_currency_br(account.get('average_revenue_cents'))}")
    lines.append(f"Profissionais: {account.get('professionals_count') or 'Não informado'}")
    lines.append(f"Presença global: {account.get('global_presence') or 'Não informado'}")
    lines.append('')
    lines.append('Resumo quantitativo:')
    for key, value in report_data['summary_counts'].items():
        lines.append(f"- {key}: {value}")
    lines.append('')
    lines.append('Contatos da conta:')
    for contact in report_data['contacts'][:50]:
        lines.append(f"- {contact.get('name')} | cargo: {contact.get('position') or 'Não informado'} | área: {contact.get('area_of_activity') or 'Não informado'} | email: {contact.get('email') or 'Não informado'} | principal: {'sim' if contact.get('is_main_contact') else 'não'}")
    lines.append('')
    lines.append('Powermapping / cards resumidos por contato:')
    for card in report_data['relationship_cards'][:50]:
        c = card['contact']
        lines.append(f"- {c.get('name')}: atividades={card.get('activities_count',0)}, mapeamentos={card.get('mapping_count',0)}, kanban={card.get('kanban_count',0)}, último contato={_relation_report_format_dt((card.get('last_contact') or {}).get('date'))}")
    lines.append('')
    lines.append('Presenças / entregas:')
    for presence in report_data['presences'][:50]:
        lines.append(f"- {presence.get('delivery_name')}: owner={presence.get('stf_owner') or 'Não informado'}, receita atual={format_currency_br(presence.get('current_revenue_cents'))}, validade={presence.get('validity_month') or 'Não informado'}, focal={presence.get('focal_client_name') or 'Não informado'}")
    lines.append('')
    lines.append('Últimas interações:')
    for item in report_data['activities'][:25]:
        lines.append(f"- {item.get('client_name')} em {_relation_report_format_dt(item.get('activity_date') or item.get('created_at'))}: {item.get('description') or item.get('information') or 'Sem detalhe'}")
    for item in report_data['account_activities'][:15]:
        lines.append(f"- Conta em {_relation_report_format_dt(item.get('activity_date') or item.get('created_at'))}: {item.get('description') or 'Sem detalhe'}")
    lines.append('')
    lines.append('Kanban:')
    for item in report_data['kanban_cards'][:25]:
        lines.append(f"- [{item.get('column_title') or 'Sem coluna'}] {item.get('title')} | contato={item.get('contact_name') or 'Não informado'} | urgência={item.get('urgency') or 'Não informada'} | descrição={item.get('description') or 'Sem descrição'}")
    lines.append('')
    lines.append('Mapeamento de ambiente:')
    for item in report_data['environment_responses'][:30]:
        lines.append(f"- {item.get('client_name')} | card={item.get('card_title')} | resposta={item.get('response') or 'Sem resposta'}")
    lines.append('')
    lines.append('Tópicos estratégicos:')
    for topic, items in report_data['topics'].items():
        lines.append(f"- {topic}: {len(items)} evidência(s)")
        for sample in items[:5]:
            lines.append(f"  * {_relation_report_format_dt(sample.get('date'))} | {sample.get('person') or 'Não informado'} | {str(sample.get('text') or '')[:240]}")
    return '\n'.join(lines)


def _relation_report_build_account_snapshot(report_data):
    account = report_data.get('account') or {}
    counts = report_data.get('summary_counts') or {}
    latest = report_data.get('latest_interaction') or {}
    presences = report_data.get('presences') or []
    total_stefanini_cents = sum(int(p.get('current_revenue_cents') or 0) for p in presences if p.get('current_revenue_cents') is not None)
    parts = []
    segment = account.get('sector') or account.get('segment') or 'Não informado'
    parts.append(f"Conta do segmento {segment}.")
    parts.append(
        f"{counts.get('contacts', 0)} contatos relacionados. "
        f"{counts.get('activities', 0) + counts.get('account_activities', 0)} atividades registradas no período. "
        f"{counts.get('kanban_cards', 0)} cards no Kanban. "
        f"{counts.get('mapping_items', 0)} itens de mapeamento de ambiente. "
        f"{counts.get('presences', 0)} presenças/entregas ativas ou históricas."
    )
    parts.append(
        f"Última interação identificada em {_relation_report_format_dt(latest.get('date'))} "
        f"com {latest.get('person') or 'contato não identificado'}, via {latest.get('source') or 'registro interno'}."
    )
    extra = []
    if account.get('average_revenue'):
        extra.append(f"Faturamento de mercado da conta: {account.get('average_revenue')}")
    elif account.get('average_revenue_cents') is not None:
        extra.append(f"Faturamento de mercado da conta: {format_currency_br(account.get('average_revenue_cents'))}")
    if total_stefanini_cents > 0:
        extra.append(f"Serviços Stefanini atualmente mapeados: {format_currency_br(total_stefanini_cents)}")
    if account.get('number_of_professionals'):
        extra.append(f"Profissionais: {account.get('number_of_professionals')}")
    elif account.get('professionals_count'):
        extra.append(f"Profissionais: {account.get('professionals_count')}")
    if account.get('presence'):
        extra.append(f"Presença geográfica: {account.get('presence')}")
    elif account.get('global_presence'):
        extra.append(f"Presença geográfica: {account.get('global_presence')}")
    if extra:
        parts.append('. '.join(extra) + '.')
    parts.append('Ao interpretar a conta, trate faturamento de mercado da conta e receita de Serviços Stefanini como informações distintas.')
    return '\n'.join([p for p in parts if p])


def _relation_report_build_relationship_snapshot(report_data):
    cards = report_data.get('relationship_cards') or []
    activities = report_data.get('activities') or []
    account_activities = report_data.get('account_activities') or []
    kanban_cards = report_data.get('kanban_cards') or []
    lines = []
    concrete_signals = 0
    scheduling_signals = 0
    for item in activities[:40] + account_activities[:20]:
        text = ' '.join([
            str(item.get('activity_type') or ''),
            str(item.get('description') or ''),
            str(item.get('notes') or ''),
            str(item.get('subject') or '')
        ]).lower()
        if any(token in text for token in ['agenda', 'agendar', 'tentativa', 'follow-up', 'follow up', 'cobrança']) and not any(token in text for token in ['reunião realizada', 'workshop', 'apresentação', 'assessment', 'kickoff', 'discussão', 'definição', 'próximo passo definido', 'proposta apresentada']):
            scheduling_signals += 1
        if any(token in text for token in ['reunião realizada', 'workshop', 'apresentação', 'assessment', 'kickoff', 'discussão', 'alinhamento', 'proposta', 'entrega', 'roadmap', 'diagnóstico']):
            concrete_signals += 1
    for card in cards[:8]:
        contact = card.get('contact') or {}
        lines.append(
            f"{contact.get('name') or 'Contato sem nome'}, {contact.get('position') or 'cargo não informado'}, "
            f"atua em {contact.get('area_of_activity') or 'área não informada'}; "
            f"{card.get('activities_count', 0)} atividades, {card.get('mapping_count', 0)} mapeamentos, "
            f"{card.get('kanban_count', 0)} itens no Kanban; último contato em {_relation_report_format_dt((card.get('last_contact') or {}).get('date'))}."
        )
    if not lines:
        lines.append('Não há contatos suficientes mapeados para descrever o Power Mapping da conta.')
    lines.append(
        f"Sinais de profundidade relacional: {concrete_signals} registro(s) de interação concreta e {scheduling_signals} registro(s) predominantemente voltados a pedido de agenda ou tentativa de contato. "
        f"Não trate volume de interação como sinônimo automático de relacionamento profundo."
    )
    if kanban_cards:
        kanban_lines = []
        for card in kanban_cards[:6]:
            title = str(card.get('title') or 'Card sem título').strip()
            col = str(card.get('column_name') or 'coluna não informada').strip()
            desc = str(card.get('description') or '').strip().replace('\n', ' ')
            if len(desc) > 140:
                desc = desc[:137] + '...'
            snippet = f"{title} [{col}]"
            if desc:
                snippet += f": {desc}"
            kanban_lines.append(snippet)
        lines.append('Leitura de Kanban: ' + ' | '.join(kanban_lines))
    lines.append('Considere a densidade de interação, sinais de influência, cobertura dos stakeholders e conteúdo real das interações ao redigir a análise.')
    return '\n'.join(lines)


def _relation_report_build_topic_evidence(report_data):
    sections = []
    for topic in RELATION_REPORT_TOPICS:
        items = (report_data.get('topics') or {}).get(topic) or []
        if not items:
            sections.append(f"{topic}:\nSem evidências relevantes identificadas no período.")
            continue
        lines = []
        for item in items[:5]:
            when = _relation_report_format_dt(item.get('date'))
            who = item.get('person') or 'não informado'
            source = item.get('source') or 'registro interno'
            snippet = str(item.get('text') or '').strip().replace('\n', ' ')
            if len(snippet) > 220:
                snippet = snippet[:217] + '...'
            lines.append(f"- {when} | {who} | {source}: {snippet}")
        sections.append(f"{topic}:\n" + '\n'.join(lines))
    return '\n\n'.join(sections)


def _openrouter_web_prompt(question: str) -> str | None:
    """Envia uma pergunta ao OpenRouter com o plugin de busca web ativado (necessário para
    informações atuais como notícias/momento de mercado — os templates SAI não têm acesso
    à web em tempo real, então não servem para esse tipo de pergunta).
    Retorna o texto da resposta, ou None se o OpenRouter não estiver configurado ou falhar."""
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if not or_key:
        return None
    or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
    model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
    site_url = (or_settings.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
    app_name = (or_settings.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
    try:
        payload = {
            'model': model,
            'messages': [{'role': 'user', 'content': question}],
            'plugins': [{'id': 'web', 'max_results': 3}],
            'temperature': 0.2
        }
        req = urllib.request.Request(
            'https://openrouter.ai/api/v1/chat/completions',
            data=json.dumps(payload, ensure_ascii=False).encode('utf-8'),
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {or_key}',
                'HTTP-Referer': site_url,
                'X-Title': app_name
            },
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read().decode('utf-8'))
        choices = data.get('choices') or []
        raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
        if raw and raw.strip():
            logger.info('[OpenRouter][web] Resposta gerada com sucesso')
            return raw.strip()
    except Exception as e:
        logger.warning(f'[OpenRouter][web] Falha: {e}')
    return None


def _relation_report_fetch_market_context(account_name: str) -> str | None:
    search = f"Me de um resumo do momento atual da empresa {account_name} no mercado, resumido, em 1 paragrafo"

    # OpenRouter (com busca web) é a fonte primária aqui: o momento de mercado precisa de
    # dados atuais/recentes, e os templates SAI não têm acesso à web em tempo real.
    or_result = _openrouter_web_prompt(search)
    if or_result and len(or_result) >= 30:
        return or_result

    # Fallback: template SAI dedicado (pode não trazer dados atualizados, mas é melhor
    # que deixar o bloco vazio caso o OpenRouter não esteja configurado/disponível).
    api_key = _resolve_setting('itoca_sai_api_key', 'ITOCA_SAI_API_KEY')
    if not api_key:
        return None
    base_url = (_load_app_settings_map(['itoca_sai_base_url']).get('itoca_sai_base_url') or '').strip() or 'https://sai-library.saiapplications.com'

    # Esse template SAI dedicado também pode sofrer rate limit (HTTP 429) sob uso
    # concorrente de outras chamadas SAI do app. Tenta novamente com backoff antes de desistir.
    max_attempts = 3
    backoff_seconds = (2, 4)
    for attempt in range(max_attempts):
        try:
            resp = requests.post(
                f'{base_url}/api/templates/67dc479828232c97f38a887f/execute',
                json={'inputs': {'search': search}},
                headers={'X-Api-Key': api_key},
                timeout=45,
            )
            if resp.status_code == 429 and attempt < max_attempts - 1:
                wait = backoff_seconds[min(attempt, len(backoff_seconds) - 1)]
                logger.warning(f'[RelationReport] HTTP 429 (rate limit) no contexto de mercado, tentativa {attempt + 1}/{max_attempts}, aguardando {wait}s...')
                time.sleep(wait)
                continue
            resp.raise_for_status()
            text = resp.text.strip()
            if not text or len(text) < 30:
                return None
            return text
        except Exception as e:
            logger.warning(f'[RelationReport] Falha ao buscar contexto de mercado: {e}')
            return None
    return None


def _relation_report_generate_highlights(report_data: dict) -> list[str]:
    account_name = ((report_data.get('account') or {}).get('name') or 'Conta').strip()
    activities = report_data.get('activities') or []
    account_activities = report_data.get('account_activities') or []

    lines = []
    for activity in (account_activities + activities)[:40]:
        desc = (activity.get('description') or activity.get('notes') or activity.get('information') or '').strip()
        date = (
            activity.get('date')
            or activity.get('activity_date')
            or activity.get('created_at')
            or ''
        )
        atype = (activity.get('type') or activity.get('category') or activity.get('origin') or '').strip()
        if desc:
            lines.append(f"[{str(date)[:10]}][{atype}] {desc}")

    if not lines:
        return []

    activities_text = '\n'.join(lines)
    question = (
        f"Analise as atividades abaixo registradas na conta '{account_name}'. "
        "Gere de 4 a 6 bullets executivos destacando os pontos mais relevantes: "
        "temas estratégicos discutidos, avanços concretos de relacionamento, "
        "pendências ou oportunidades identificadas, e padrão de engajamento. "
        "Seja direto, use verbos no passado ou presente, sem inventar fatos. "
        "Retorne SOMENTE os bullets, um por linha, começando cada linha com '- '. "
        "Atividades:\n" + activities_text
    )
    raw = _sai_simple_prompt(question)
    if not raw:
        return []

    bullets = [
        line.lstrip('- •*').strip()
        for line in raw.strip().splitlines()
        if line.strip().startswith(('-', '•', '*')) and len(line.strip()) > 10
    ]
    return bullets[:6]


def _relation_report_call_sai_narrative_template(
    account_name,
    report_period,
    account_snapshot,
    relationship_snapshot,
    topic_evidence,
    output_style,
):
    settings_map = _load_app_settings_map([
        'relation_report_sai_api_key',
        'relation_report_sai_template_id',
        'relation_report_sai_base_url'
    ])
    api_key = (settings_map.get('relation_report_sai_api_key') or '').strip() or (os.environ.get('RELATION_REPORT_SAI_API_KEY', '') or '').strip() or 'RuWKlxg1Sk+/3PpzUKof+w'
    template_id = (settings_map.get('relation_report_sai_template_id') or '').strip() or '69b83e37025459101ee6735d'
    base_url = (settings_map.get('relation_report_sai_base_url') or '').strip() or 'https://sai-library.saiapplications.com'

    if not api_key:
        return None

    url = f'{base_url}/api/templates/{template_id}/execute'
    headers = {
        'Content-Type': 'application/json',
        'X-Api-Key': api_key
    }
    payload = {
        'inputs': {
            'account_name': account_name,
            'report_period': report_period,
            'account_snapshot': account_snapshot,
            'relationship_snapshot': relationship_snapshot,
            'topic_evidence': topic_evidence,
            'output_style': output_style,
        }
    }
    resp = requests.post(url, json=payload, headers={'X-Api-Key': api_key}, timeout=60)
    resp.raise_for_status()
    raw = resp.text

    logger.info(f'[RelationReport][SAI] OK ({len(raw)} chars)')
    logger.debug(f'[RelationReport][SAI] raw response (primeiros 500 chars): {raw[:500]}')
    parsed_outer = None
    try:
        parsed_outer = json.loads(raw)
    except Exception:
        parsed_outer = None

    candidate_texts = []
    if isinstance(parsed_outer, dict):
        for key in ['answer', 'output', 'result', 'text', 'response', 'data']:
            value = parsed_outer.get(key)
            if isinstance(value, str):
                candidate_texts.append(value)
            elif isinstance(value, dict):
                candidate_texts.append(json.dumps(value, ensure_ascii=False))
        candidate_texts.append(raw)
    else:
        candidate_texts.append(raw)

    parsed = None
    for candidate in candidate_texts:
        try:
            obj = json.loads((candidate or '').strip())
            if isinstance(obj, dict):
                parsed = obj
                break
        except Exception:
            obj = _extract_json_object_from_text(candidate or '')
            if isinstance(obj, dict):
                parsed = obj
                break

    if not isinstance(parsed, dict):
        return None

    return {
        'executive_summary': (parsed.get('executive_summary') or '').strip(),
        'relationship_maturity': (parsed.get('relationship_maturity') or 'Em evolução').strip(),
        'next_steps': [str(x).strip() for x in (parsed.get('next_steps') or []) if str(x).strip()][:5],
        'topic_breakdown': parsed.get('topic_breakdown') or {},
        'highlights': [str(x).strip() for x in (parsed.get('highlights') or []) if str(x).strip()][:6],
        'llm_used': True
    }


def _relation_report_generate_narrative(report_data):
    account = report_data.get('account') or {}
    period = report_data.get('period') or {}
    account_name = (account.get('name') or 'Conta sem nome').strip()
    report_period = 'Todo o período' if period.get('full_period') else f"{period.get('start_date') or 'Início não informado'} a {period.get('end_date') or 'Fim não informado'}"
    account_snapshot = _relation_report_build_account_snapshot(report_data)
    relationship_snapshot = _relation_report_build_relationship_snapshot(report_data)
    topic_evidence = _relation_report_build_topic_evidence(report_data)
    output_style = 'Tom executivo, objetivo, claro, elegante e sem inventar fatos. Use sempre a expressão Power Mapping. Não trate faturamento de mercado da conta como receita da Stefanini. Só classifique o relacionamento como profundo quando houver evidências concretas de avanço, reuniões realizadas, discussões de conteúdo, entregas, proposta, assessment, workshop, roadmap ou desdobramentos objetivos. Se o histórico indicar apenas tentativa de agenda, cobrança ou follow-up sem avanço real, deixe isso explícito. Considere também o conteúdo dos cards de Kanban para explicar como eles se relacionam com a conta.'

    narrative_data = None
    market_context = None
    llm_highlights = []

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            f_narrative = executor.submit(
                _relation_report_call_sai_narrative_template,
                account_name,
                report_period,
                account_snapshot,
                relationship_snapshot,
                topic_evidence,
                output_style,
            )
            f_market = executor.submit(_relation_report_fetch_market_context, account_name)
            f_highlights = executor.submit(_relation_report_generate_highlights, report_data)

            try:
                narrative_data = f_narrative.result()
            except Exception as e:
                logger.warning(f'[RelationReport] Falha ao gerar narrativa com SAI: {e}')
            try:
                market_context = f_market.result()
            except Exception as e:
                logger.warning(f'[RelationReport] Falha ao gerar contexto de mercado: {e}')
            try:
                llm_highlights = f_highlights.result() or []
            except Exception as e:
                logger.warning(f'[RelationReport] Falha ao gerar destaques via LLM: {e}')
    except Exception as e:
        logger.warning(f'[RelationReport] Falha no executor paralelo da narrativa: {e}')

    if isinstance(narrative_data, dict):
        if len(llm_highlights) >= 2:
            narrative_data['highlights'] = llm_highlights
        narrative_data['market_context'] = market_context
        return narrative_data

    latest = report_data.get('latest_interaction') or {}
    counts = report_data.get('summary_counts') or {}
    executive_summary = (
        f"A conta {report_data['account'].get('name')} possui {counts.get('contacts', 0)} contato(s) relacionado(s), "
        f"{counts.get('presences', 0)} presença(s)/entrega(s) mapeada(s), {counts.get('activities', 0)} atividade(s) de relacionamento "
        f"e {counts.get('kanban_cards', 0)} card(s) ativos ou históricos no Kanban dentro do recorte consultado. "
        f"O conjunto de registros indica um relacionamento {'mais estruturado' if counts.get('activities', 0) >= 5 or counts.get('mapping_items', 0) >= 5 else 'em consolidação'}, "
        f"com evidências distribuídas entre interações registradas, mapeamentos de ambiente e acompanhamento operacional.\n\n"
        f"A interação mais recente identificada ocorreu em {_relation_report_format_dt(latest.get('date'))}, "
        f"associada a {latest.get('person') or report_data['account'].get('name')}, via {latest.get('source') or 'registro interno'}. "
        f"Isso sugere {'continuidade recente do relacionamento' if latest.get('date') else 'baixa clareza temporal sobre a última interação'}, "
        f"e reforça a importância de manter cadência de contato, atualização do Kanban e aprofundamento do powermapping por contato-chave."
    )
    topic_breakdown = {}
    for topic, items in report_data['topics'].items():
        if items:
            sample = items[0]
            topic_breakdown[topic] = f"Há {len(items)} evidência(s) relacionadas a {topic.lower()} no período, com destaque para registros de {sample.get('source')} e menções envolvendo {sample.get('person') or 'contatos da conta'}."
        else:
            topic_breakdown[topic] = f"Não foram encontradas evidências relevantes sobre {topic.lower()} no período analisado."
    fallback_narrative = {
        'executive_summary': executive_summary,
        'relationship_maturity': 'Em evolução' if counts.get('activities', 0) < 8 else 'Estruturado',
        'next_steps': [
            'Validar os contatos principais e atualizar os responsáveis por influência e decisão.',
            'Revisar os cards do Kanban com foco em prioridades e próximos passos claros.',
            'Usar o resumo temático para orientar novas conversas consultivas com a conta.'
        ],
        'topic_breakdown': topic_breakdown,
        'highlights': [
            f"{counts.get('contacts', 0)} contato(s) associado(s) à conta.",
            f"{counts.get('mapping_items', 0)} item(ns) de mapeamento de ambiente registrados.",
            f"Última interação em {_relation_report_format_dt(latest.get('date'))}."
        ],
        'llm_used': False,
        'market_context': market_context,
    }
    if len(llm_highlights) >= 2:
        fallback_narrative['highlights'] = llm_highlights
    return fallback_narrative


def _relation_report_draw_paragraph(c, text, x, y, width, style):
    para = Paragraph(text or '', style)
    _, h = para.wrap(width, 10000)
    para.drawOn(c, x, y - h)
    return y - h


def _relation_report_draw_header(c, report_data, colors_map, page_width, page_height):
    c.setFillColor(colors_map['primary'])
    c.roundRect(18 * mm, page_height - 42 * mm, page_width - 36 * mm, 24 * mm, 6 * mm, fill=1, stroke=0)
    account = report_data['account']
    system_logo = _relation_report_safe_image(_relation_report_system_logo_path(), 180, 90)
    account_logo = _relation_report_safe_image(_relation_report_resolve_local_image(account.get('logo_url')), 180, 90)
    if system_logo:
        c.drawImage(system_logo, 22 * mm, page_height - 38 * mm, width=18 * mm, height=18 * mm, mask='auto', preserveAspectRatio=True)
    if account_logo:
        c.drawImage(account_logo, page_width - 40 * mm, page_height - 38 * mm, width=18 * mm, height=18 * mm, mask='auto', preserveAspectRatio=True)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 18)
    c.drawString(44 * mm, page_height - 27 * mm, 'Relationship Report')
    c.setFont('Helvetica', 10)
    c.drawString(44 * mm, page_height - 33 * mm, 'Toca do Coelho')
    c.setFont('Helvetica-Bold', 14)
    title = account.get('name') or 'Cliente'
    max_width = 90 * mm
    while stringWidth(title, 'Helvetica-Bold', 14) > max_width and len(title) > 20:
        title = title[:-4] + '...'
    c.drawRightString(page_width - 44 * mm, page_height - 27 * mm, title)
    c.setFont('Helvetica', 9)
    subtitle = 'Período completo' if report_data.get('full_period') else f"{report_data.get('start_date') or '...'} a {report_data.get('end_date') or '...'}"
    c.drawRightString(page_width - 44 * mm, page_height - 33 * mm, subtitle)


def _relation_report_build_browser_html(report_data, profile=None, embed_images=False):
    profile = profile or {}

    def _local_image_from_url(url):
        if not url:
            return None
        raw_url = str(url).strip()
        if not raw_url:
            return None
        try:
            parsed = urlparse(raw_url)
            path = (parsed.path or raw_url).strip()
        except Exception:
            parsed = None
            path = raw_url
        if raw_url.startswith('file://') and parsed and parsed.path:
            file_path = Path(parsed.path)
            return file_path if file_path.exists() else None
        if not path:
            return None
        normalized = path.replace('\\', '/').strip()
        if normalized.startswith('/uploads/'):
            local = UPLOAD_DIR / normalized.replace('/uploads/', '', 1)
            return local if local.exists() else None
        if normalized.startswith('uploads/'):
            local = UPLOAD_DIR / normalized.replace('uploads/', '', 1)
            return local if local.exists() else None
        if normalized.startswith('/public/'):
            local = Path(BASE_DIR) / normalized.lstrip('/')
            return local if local.exists() else None
        if normalized.startswith('/'):
            local = Path(BASE_DIR) / 'public' / normalized.lstrip('/')
            return local if local.exists() else None
        local = Path(normalized)
        if local.exists():
            return local
        local_public = Path(BASE_DIR) / 'public' / normalized
        if local_public.exists():
            return local_public
        resolved = _relation_report_resolve_local_image(raw_url) or _relation_report_resolve_local_image(path)
        return resolved if resolved and Path(resolved).exists() else None

    def _inline_image_url(url):
        url = (url or '').strip()
        if not url or url.startswith('data:') or not embed_images:
            return url
        try:
            local_path = _local_image_from_url(url)
            if local_path:
                with open(local_path, 'rb') as fh:
                    raw = fh.read()
                mime, _ = mimetypes.guess_type(str(local_path))
                mime = mime or 'application/octet-stream'
                return f"data:{mime};base64,{base64.b64encode(raw).decode('ascii')}"
            if url.startswith('/static/'):
                static_path = os.path.join(BASE_DIR, url.lstrip('/'))
                if os.path.exists(static_path):
                    with open(static_path, 'rb') as fh:
                        raw = fh.read()
                    mime, _ = mimetypes.guess_type(static_path)
                    mime = mime or 'application/octet-stream'
                    return f"data:{mime};base64,{base64.b64encode(raw).decode('ascii')}"
            if url.startswith('http://') or url.startswith('https://'):
                resp = requests.get(url, timeout=10)
                resp.raise_for_status()
                mime = resp.headers.get('Content-Type', '').split(';')[0].strip() or mimetypes.guess_type(url)[0] or 'application/octet-stream'
                return f"data:{mime};base64,{base64.b64encode(resp.content).decode('ascii')}"
        except Exception:
            return url
        return url

    narrative = report_data.get('narrative') or _relation_report_generate_narrative(report_data)
    report_data['narrative'] = narrative
    account = report_data.get('account') or {}
    latest = report_data.get('latest_interaction') or {}
    counts = report_data.get('summary_counts') or {}
    profile = profile or {}

    def esc(value):
        return html.escape(str(value or ''))

    def fmt_money(value):
        try:
            if value in (None, ''):
                return 'Não informado'
            cents = int(value)
            return f"R$ {cents/100:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        except Exception:
            return 'Não informado'

    def topic_card(topic, bg):
        txt = ((narrative.get('topic_breakdown') or {}).get(topic) or f'Sem evidências relevantes para {topic.lower()}.').strip()
        return f"<div class='rr-topic-card' style='background:{bg};'><div class='rr-topic-title'>{esc(topic)}</div><div class='rr-topic-text'>{esc(txt)}</div></div>"

    profile_name = profile.get('nickname') or profile.get('full_name') or 'Usuário'
    profile_photo = profile.get('photo_url') or ''
    profile_role = profile.get('role') or profile.get('job_title') or profile.get('position') or 'Responsável pelo relacionamento'
    account_logo = account.get('logo_url') or ''
    account_name = account.get('name') or 'Conta'
    period_label = 'Todo o período' if report_data.get('full_period') else f"{report_data.get('start_date') or '...'} a {report_data.get('end_date') or '...'}"
    latest_text = 'Sem interações registradas'
    if latest:
        latest_text = f"{_relation_report_format_dt(latest.get('date'))} · {latest.get('person') or 'Contato não identificado'}"

    def company_rank(position):
        p = str(position or '').strip().lower()
        if p == 'ceo':
            return 1
        if p == 'cio':
            return 2
        if p.startswith('c') and len(p) <= 5:
            return 3
        if 'diretor' in p or 'superintendente' in p:
            return 4
        if 'gerente' in p:
            return 5
        if 'coordenador' in p:
            return 6
        return 7

    sorted_contacts = sorted(
        report_data.get('contacts') or [],
        key=lambda c: (company_rank(c.get('position')), str(c.get('name') or '').lower())
    )

    contacts_html = []
    for contact in sorted_contacts:
        photo = _inline_image_url(contact.get('photo_url') or '')
        initials = (str(contact.get('name') or '?').strip()[:1] or '?').upper()
        badges = []
        if contact.get('is_main_contact'):
            badges.append("<span class='rr-badge rr-badge-primary'>Contato-chave</span>")
        if contact.get('is_target'):
            badges.append("<span class='rr-badge rr-badge-soft'>Target</span>")
        if contact.get('is_cold_contact'):
            badges.append("<span class='rr-badge rr-badge-cold'>Cold</span>")
        meta = []
        if contact.get('position'):
            meta.append(esc(contact.get('position')))
        if contact.get('email'):
            meta.append(esc(contact.get('email')))
        if contact.get('phone'):
            meta.append(esc(contact.get('phone')))
        linkedin_raw = str(contact.get('linkedin') or '').strip()
        linkedin_html = ''
        if linkedin_raw:
            linkedin_safe = esc(linkedin_raw)
            linkedin_html = f"<div class='rr-contact-linkedin'><a href='{linkedin_safe}' target='_blank' rel='noopener noreferrer'>LinkedIn</a></div>"
        activities = [a for a in (report_data.get('activities') or []) if a.get('client_id') == contact.get('id')]
        last_contact = activities[0] if activities else None
        last_contact_line = f"Última interação: {esc(_relation_report_format_dt(last_contact.get('activity_date')))}" if last_contact else 'Última interação: não registrada'
        photo_html = f"<img src='{esc(photo)}' class='rr-contact-photo'/>" if photo else f"<div class='rr-contact-photo rr-contact-photo-fallback'>{esc(initials)}</div>"
        contacts_html.append(f"""
        <div class='rr-contact-card'>
            <div class='rr-contact-head'>
                {photo_html}
                <div>
                    <div class='rr-contact-name'>{esc(contact.get('name'))}</div>
                    <div class='rr-contact-meta'>{' · '.join(meta) if meta else 'Sem detalhes complementares'}</div>
                    {linkedin_html}
                </div>
            </div>
            <div class='rr-badge-row'>{''.join(badges)}</div>
            <div class='rr-contact-kpi'>{esc(last_contact_line)}</div>
        </div>
        """)

    highlights_html = ''.join([f"<li>{esc(item)}</li>" for item in (narrative.get('highlights') or [])]) or '<li>Sem destaques adicionais.</li>'
    market_context_text = (narrative.get('market_context') or '').strip()
    market_context_html = ''
    if market_context_text:
        market_context_html = f"""
        <div class='rr-market-context'>
          <h3 style='font-size:13px; color:#ffffff; text-transform:uppercase; letter-spacing:.05em; margin:18px 0 6px;'>
            Contexto de Mercado
          </h3>
          <p style='font-size:14px; line-height:1.7; color:#ffffff;'>
            {esc(market_context_text)}
          </p>
        </div>
        """
    next_steps_html = ''.join([f"<li>{esc(item)}</li>" for item in (narrative.get('next_steps') or [])]) or '<li>Sem próximos passos sugeridos.</li>'
    account_logo = _inline_image_url(account_logo)
    account_logo_html = f"<img src='{esc(account_logo)}' alt='Logo da conta'>" if account_logo else "📊"
    context_badges_html = ''.join([
        f"<span class='rr-context-pill'><strong>Período</strong> {esc(period_label)}</span>",
        f"<span class='rr-context-pill'><strong>Última interação</strong> {esc(latest_text)}</span>",
        f"<span class='rr-context-pill'><strong>Maturidade</strong> {esc(narrative.get('relationship_maturity') or 'Não classificada')}</span>",
        f"<span class='rr-context-pill'><strong>Presença</strong> {esc(account.get('global_presence') or 'Não informada')}</span>"
    ])
    profile_photo = _inline_image_url(profile_photo)
    profile_photo_html = f"<img src='{esc(profile_photo)}' class='rr-user-photo' alt='Foto do usuário'>" if profile_photo else f"<div class='rr-user-photo rr-user-fallback'>{esc(str(profile_name)[:1].upper())}</div>"

    return f"""<!DOCTYPE html>
<html lang='pt-BR'>
<head>
<meta charset='UTF-8'>
<meta name='viewport' content='width=device-width, initial-scale=1.0'>
<title>Relationship Report - {esc(account_name)}</title>
<style>
:root {{ --green:#059669; --green-dark:#065f46; --mint:#d1fae5; --bg:#f6fffb; --text:#1f2937; --muted:#6b7280; --card:#ffffff; --line:#d1fae5; }}
* {{ box-sizing:border-box; }}
body {{ margin:0; font-family:Inter,Segoe UI,Arial,sans-serif; background:linear-gradient(180deg,#ecfdf5 0%, #ffffff 38%, #f8fafc 100%); color:var(--text); }}
.rr-shell {{ max-width:1360px; margin:0 auto; padding:32px 28px 60px; }}
.rr-hero {{ background:linear-gradient(135deg,#064e3b 0%, #047857 50%, #10b981 100%); color:#fff; border-radius:28px; padding:28px; box-shadow:0 24px 80px rgba(5,150,105,.22); position:relative; overflow:hidden; }}
.rr-hero:after {{ content:''; position:absolute; inset:auto -80px -80px auto; width:240px; height:240px; background:rgba(255,255,255,.08); border-radius:50%; }}
.rr-top {{ display:grid; grid-template-columns:minmax(0,1fr) auto; gap:20px; align-items:flex-start; }}
.rr-brand {{ display:flex; gap:16px; align-items:center; min-width:0; }}
.rr-brand-mark {{ min-width:72px; width:auto; height:72px; border-radius:22px; background:rgba(255,255,255,.12); display:flex; align-items:center; justify-content:center; font-size:28px; border:1px solid rgba(255,255,255,.18); backdrop-filter:blur(8px); overflow:hidden; padding:8px 12px; max-width:240px; }}
.rr-brand-mark img {{ max-width:216px; max-height:56px; width:auto; height:auto; object-fit:contain; }}
.rr-title {{ font-size:34px; font-weight:800; line-height:1.05; margin:0 0 8px; }}
.rr-subtitle {{ margin:0; color:rgba(255,255,255,.88); font-size:15px; max-width:760px; line-height:1.6; }}
.rr-user {{ display:flex; gap:14px; align-items:center; justify-self:end; background:rgba(255,255,255,.12); border:1px solid rgba(255,255,255,.18); padding:12px 14px; border-radius:20px; backdrop-filter:blur(8px); min-width:260px; }}
.rr-user-photo {{ width:88px; height:88px; border-radius:50%; object-fit:cover; border:2px solid rgba(255,255,255,.35); background:rgba(255,255,255,.16); }}
.rr-user-fallback {{ display:flex; align-items:center; justify-content:center; color:#fff; font-weight:800; }}
.rr-user-name {{ font-size:16px; font-weight:700; margin:0; }}
.rr-user-role {{ margin:4px 0 0; font-size:12px; color:rgba(255,255,255,.82); }}
.rr-hero-grid {{ display:grid; grid-template-columns:1fr; gap:18px; margin-top:22px; }}
.rr-panel {{ background:rgba(255,255,255,.12); border:1px solid rgba(255,255,255,.16); border-radius:24px; padding:18px; backdrop-filter:blur(8px); }}
.rr-panel-summary {{ grid-column:1 / -1; }}

.rr-panel h3 {{ margin:0 0 10px; font-size:14px; text-transform:uppercase; letter-spacing:.08em; color:rgba(255,255,255,.98); }}
.rr-kpis {{ display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:14px; margin-top:24px; }}
.rr-kpi {{ background:var(--card); border:1px solid var(--line); border-radius:22px; padding:18px; box-shadow:0 10px 30px rgba(16,185,129,.08); }}
.rr-kpi-label {{ font-size:12px; color:var(--muted); text-transform:uppercase; letter-spacing:.06em; margin-bottom:8px; }}
.rr-kpi-value {{ font-size:28px; font-weight:800; color:var(--green-dark); }}
.rr-grid {{ display:grid; grid-template-columns:1.2fr .8fr; gap:22px; margin-top:24px; }}
.rr-section {{ background:rgba(255,255,255,.86); border:1px solid rgba(209,250,229,.9); border-radius:24px; padding:24px; box-shadow:0 18px 55px rgba(15,118,110,.08); }}
.rr-section h2 {{ margin:0 0 14px; font-size:22px; color:var(--green-dark); }}
.rr-lead {{ color:rgba(255,255,255,.96); font-size:15px; line-height:1.75; white-space:pre-line; }}
.rr-context-pills {{ display:flex; flex-wrap:wrap; gap:10px; margin-bottom:14px; }}
.rr-context-pill {{ display:inline-flex; align-items:center; gap:8px; padding:9px 12px; border-radius:999px; background:rgba(255,255,255,.14); border:1px solid rgba(255,255,255,.18); color:#f8fafc; font-size:12px; line-height:1.35; }}
.rr-context-pill strong {{ color:#ffffff; font-size:11px; text-transform:uppercase; letter-spacing:.04em; }}
.rr-info-list {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:10px; }}
.rr-info-item {{ background:#fff; border:1px solid #ecfdf5; border-radius:14px; padding:10px 12px; }}
.rr-info-item-label {{ font-size:10px; color:var(--muted); text-transform:uppercase; letter-spacing:.06em; margin-bottom:6px; }}
.rr-info-item-value {{ font-size:13px; font-weight:700; color:#111827; }}
.rr-contact-grid {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:16px; align-items:start; }}
.rr-contact-card {{ background:#fff; border:1px solid #e5e7eb; border-radius:20px; padding:16px; box-shadow:0 8px 24px rgba(15,23,42,.05); }}
.rr-contact-head {{ display:flex; gap:12px; align-items:center; margin-bottom:12px; }}
.rr-contact-photo {{ width:58px; height:58px; border-radius:18px; object-fit:cover; background:#ecfdf5; }}
.rr-contact-photo-fallback {{ display:flex; align-items:center; justify-content:center; font-weight:800; color:var(--green-dark); font-size:22px; }}
.rr-contact-name {{ font-size:17px; font-weight:800; color:#111827; }}
.rr-contact-meta {{ font-size:13px; color:#6b7280; margin-top:3px; line-height:1.4; }}
.rr-contact-linkedin {{ margin-top:6px; }}
.rr-contact-linkedin a {{ color:#0a66c2; text-decoration:none; font-size:12px; font-weight:600; }}
.rr-contact-linkedin a:hover {{ text-decoration:underline; }}
.rr-badge-row {{ display:flex; flex-wrap:wrap; gap:8px; margin-bottom:10px; }}
.rr-badge {{ font-size:11px; font-weight:700; border-radius:999px; padding:6px 10px; display:inline-flex; align-items:center; }}
.rr-badge-primary {{ background:#dcfce7; color:#166534; }}
.rr-badge-soft {{ background:#ecfeff; color:#155e75; }}
.rr-badge-cold {{ background:#eff6ff; color:#1d4ed8; }}
.rr-contact-kpi {{ font-size:13px; color:#4b5563; line-height:1.5; }}
.rr-topic-grid {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:14px; }}
.rr-topic-card {{ border-radius:18px; padding:16px; border:1px solid rgba(255,255,255,.35); min-height:132px; }}
.rr-topic-title {{ font-size:16px; font-weight:800; margin-bottom:10px; color:#0f172a; }}
.rr-topic-text {{ font-size:14px; line-height:1.65; color:#334155; }}
.rr-list {{ padding-left:20px; margin:0; color:#374151; line-height:1.8; }}
.rr-toolbar {{ position:sticky; top:0; z-index:20; backdrop-filter:blur(12px); background:rgba(255,255,255,.74); border-bottom:1px solid rgba(209,250,229,.9); padding:12px 28px; display:flex; justify-content:space-between; gap:12px; flex-wrap:wrap; }}
.rr-toolbar-title {{ font-weight:700; color:var(--green-dark); }}
.rr-toolbar-actions {{ display:flex; gap:10px; flex-wrap:wrap; }}
.rr-btn {{ border:none; border-radius:14px; padding:10px 16px; font-weight:700; cursor:pointer; }}
.rr-btn-primary {{ background:#10b981; color:#fff; }}
.rr-btn-secondary {{ background:#e5e7eb; color:#111827; }}
/* ── FireShot Modal ── */
#rr-fireshot-modal {{
  display: none;
  position: fixed;
  inset: 0;
  background: rgba(0,0,0,0.45);
  z-index: 9999;
  align-items: center;
  justify-content: center;
}}
#rr-fireshot-modal.active {{
  display: flex;
}}
.rr-fs-box {{
  background: #fff;
  border-radius: 18px;
  padding: 32px 28px 24px;
  max-width: 480px;
  width: 92%;
  box-shadow: 0 24px 64px rgba(4,120,87,0.18);
  border-top: 4px solid #10b981;
}}
.rr-fs-header {{
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  margin-bottom: 18px;
}}
.rr-fs-title {{
  font-size: 18px;
  font-weight: 700;
  color: #047857;
  display: flex;
  align-items: center;
  gap: 8px;
}}
.rr-fs-close {{
  background: none;
  border: none;
  font-size: 22px;
  color: #6b7280;
  cursor: pointer;
  line-height: 1;
  padding: 0;
}}
.rr-fs-close:hover {{ color: #1f2937; }}
.rr-fs-steps {{
  list-style: none;
  padding: 0;
  margin: 0 0 20px;
  display: flex;
  flex-direction: column;
  gap: 12px;
}}
.rr-fs-steps li {{
  display: flex;
  gap: 12px;
  align-items: flex-start;
  font-size: 13.5px;
  color: #1f2937;
  line-height: 1.5;
}}
.rr-fs-step-num {{
  background: #10b981;
  color: #fff;
  font-weight: 700;
  font-size: 12px;
  border-radius: 50%;
  min-width: 22px;
  height: 22px;
  display: flex;
  align-items: center;
  justify-content: center;
  margin-top: 1px;
}}
.rr-fs-note {{
  background: #ecfdf5;
  border: 1px solid #d1fae5;
  border-radius: 10px;
  padding: 10px 14px;
  font-size: 12.5px;
  color: #047857;
  margin-bottom: 20px;
}}
.rr-fs-actions {{
  display: flex;
  gap: 10px;
  justify-content: flex-end;
}}
.rr-fs-btn-install {{
  background: #10b981;
  color: #fff;
  border: none;
  border-radius: 12px;
  padding: 10px 18px;
  font-weight: 700;
  font-size: 13px;
  cursor: pointer;
  text-decoration: none;
  display: inline-flex;
  align-items: center;
  gap: 6px;
}}
.rr-fs-btn-install:hover {{ background: #047857; }}
.rr-fs-btn-ok {{
  background: #e5e7eb;
  color: #1f2937;
  border: none;
  border-radius: 12px;
  padding: 10px 18px;
  font-weight: 700;
  font-size: 13px;
  cursor: pointer;
}}
.rr-fs-btn-ok:hover {{ background: #d1d5db; }}
@media (max-width: 1024px) {{ .rr-hero-grid,.rr-grid,.rr-kpis,.rr-contact-grid,.rr-topic-grid,.rr-info-list {{ grid-template-columns:1fr; }} }}
@page {{ size: landscape; margin: 12mm 10mm; }}
@media print {{ .rr-toolbar {{ display:none !important; }} html, body {{ background:#fff !important; width:100%; height:auto; margin:0 !important; padding:0 !important; -webkit-print-color-adjust:exact; print-color-adjust:exact; }} .rr-shell {{ max-width:none !important; width:100% !important; padding:0 !important; margin:0 !important; }} .rr-hero {{ min-height:auto !important; padding:16px !important; margin-top:0 !important; border-radius:18px !important; overflow:visible !important; }} .rr-hero:after {{ display:none !important; }} .rr-top {{ display:grid !important; grid-template-columns:minmax(0,1fr) 220px !important; gap:14px !important; align-items:start !important; }} .rr-brand {{ gap:12px !important; align-items:flex-start !important; min-width:0 !important; }} .rr-user {{ justify-self:end !important; align-self:start !important; width:220px !important; }} .rr-brand-copy p {{ max-width:none !important; }} .rr-section,.rr-kpi,.rr-panel,.rr-contact-card,.rr-topic-card,.rr-info-item {{ box-shadow:none !important; break-inside:avoid; page-break-inside:avoid; }} .rr-hero {{ break-inside:avoid !important; page-break-inside:avoid !important; }} .rr-grid,.rr-kpis,.rr-contact-grid,.rr-topic-grid,.rr-info-list {{ display:grid !important; }} .rr-kpis {{ grid-template-columns:repeat(4,minmax(0,1fr)) !important; gap:10px !important; margin-top:12px !important; }} .rr-grid {{ grid-template-columns:1.1fr .9fr !important; gap:14px !important; margin-top:16px !important; }} .rr-hero-grid {{ display:block !important; margin-top:10px !important; }} .rr-panel-summary {{ display:block !important; width:100% !important; background:rgba(255,255,255,.14) !important; border:1px solid rgba(255,255,255,.18) !important; padding:14px !important; border-radius:18px !important; margin-top:8px !important; break-inside:avoid !important; page-break-inside:avoid !important; }} .rr-contact-grid {{ grid-template-columns:repeat(2,minmax(0,1fr)) !important; gap:12px !important; }} .rr-topic-grid {{ grid-template-columns:repeat(2,minmax(0,1fr)) !important; gap:12px !important; }} .rr-info-list {{ grid-template-columns:repeat(2,minmax(0,1fr)) !important; gap:10px !important; }} .rr-kpi,.rr-section,.rr-contact-card,.rr-topic-card,.rr-info-item,.rr-panel {{ margin-bottom:10px !important; }} .rr-section {{ padding:16px !important; border-radius:18px !important; }} .rr-contact-card {{ min-height:0 !important; padding:14px !important; }} .rr-topic-card {{ min-height:0 !important; padding:14px !important; }} .rr-kpi {{ padding:14px !important; }} .rr-user-photo {{ width:68px !important; height:68px !important; }} .rr-brand-mark {{ max-width:170px !important; height:52px !important; border-radius:16px !important; }} .rr-brand-mark img {{ max-width:150px !important; max-height:38px !important; }} .rr-title {{ font-size:26px !important; line-height:1.08 !important; margin-bottom:6px !important; }} .rr-subtitle {{ font-size:12px !important; line-height:1.4 !important; margin-top:4px !important; max-width:none !important; }} .rr-panel-summary h3 {{ margin:0 0 10px !important; font-size:16px !important; color:#ffffff !important; }} .rr-lead {{ font-size:12px !important; line-height:1.5 !important; color:#ffffff !important; display:block !important; visibility:visible !important; opacity:1 !important; }} .rr-context-pills {{ gap:8px !important; margin-bottom:10px !important; }} .rr-context-pill {{ padding:6px 9px !important; font-size:10px !important; }} .rr-contact-meta, .rr-muted, .rr-user-role, .rr-info-item-value {{ font-size:10px !important; }} .rr-page-break-before {{ break-before:page; page-break-before:always; }} }}
</style>
</head>
<body>
<div class='rr-toolbar'>
  <div class='rr-toolbar-title'>Relationship Report · {esc(account_name)}</div>
  <div class='rr-toolbar-actions'>
    <button class='rr-btn rr-btn-secondary' onclick='window.close()'>✕ Fechar</button>
    <button class='rr-btn rr-btn-primary' onclick='rrShowFireshotModal()'>
      <span style='font-size:15px;'>📷</span> Exportar JPG
    </button>
  </div>
</div>
<div class='rr-shell'>
  <section class='rr-hero'>
    <div class='rr-top'>
      <div class='rr-brand'>
        <div class='rr-brand-mark'>{account_logo_html}</div>
        <div>
          <p style='margin:0 0 6px; font-size:12px; text-transform:uppercase; letter-spacing:.14em; color:rgba(255,255,255,.72);'>Toca do Coelho · Executive Relationship Report</p>
          <h1 class='rr-title'>{esc(account_name)}</h1>
          <p class='rr-subtitle'>Visão executiva do relacionamento da conta, combinando Power Mapping, histórico de interação, presença operacional, leitura temática e próximos passos recomendados.</p>
        </div>
      </div>
      <div class='rr-user'>{profile_photo_html}
        <div>
          <p class='rr-user-name'>{esc(profile_name)}</p>
          <p class='rr-user-role'>{esc(profile_role)}</p>
          <p class='rr-user-role'>Relatório gerado em {esc(datetime.now().strftime('%d/%m/%Y %H:%M'))}</p>
        </div>
      </div>
    </div>
    <div class='rr-hero-grid'>
      <div class='rr-panel rr-panel-summary'>
        <h3>Resumo executivo</h3>
        <div class='rr-context-pills'>{context_badges_html}</div>
        <div class='rr-lead'>{esc((narrative.get('executive_summary') or 'Sem resumo gerado.').replace('powermapping', 'Power Mapping').replace('Powermapping', 'Power Mapping'))}</div>
        {market_context_html}
      </div>
    </div>
  </section>
  <section class='rr-kpis'>
    <div class='rr-kpi'><div class='rr-kpi-label'>Contatos</div><div class='rr-kpi-value'>{counts.get('contacts', 0)}</div></div>
    <div class='rr-kpi'><div class='rr-kpi-label'>Atividades</div><div class='rr-kpi-value'>{counts.get('activities', 0) + counts.get('account_activities', 0)}</div></div>
    <div class='rr-kpi'><div class='rr-kpi-label'>Kanban</div><div class='rr-kpi-value'>{counts.get('kanban_cards', 0)}</div></div>
    <div class='rr-kpi'><div class='rr-kpi-label'>Mapeamentos</div><div class='rr-kpi-value'>{counts.get('mapping_items', 0)}</div></div>
  </section>
  <section class='rr-grid'>
    <div class='rr-section'>
      <h2>Power Mapping e contatos-chave</h2>
      <div class='rr-contact-grid'>{''.join(contacts_html) or '<div class="rr-contact-card">Nenhum contato encontrado para a conta.</div>'}</div>
    </div>
    <div class='rr-section'>
      <h2>Contexto da conta</h2>
      <div class='rr-info-list'>
        <div class='rr-info-item'><div class='rr-info-item-label'>Setor</div><div class='rr-info-item-value'>{esc(account.get('sector') or 'Não informado')}</div></div>
        <div class='rr-info-item'><div class='rr-info-item-label'>Receita média</div><div class='rr-info-item-value'>{esc(fmt_money(account.get('average_revenue_cents')))}</div></div>
        <div class='rr-info-item'><div class='rr-info-item-label'>Profissionais</div><div class='rr-info-item-value'>{esc(account.get('professionals_count') or 'Não informado')}</div></div>
        <div class='rr-info-item'><div class='rr-info-item-label'>Conta-alvo</div><div class='rr-info-item-value'>{'Sim' if account.get('is_target') else 'Não'}</div></div>
      </div>
      <div style='margin-top:18px;'>
        <h2 style='font-size:18px;'>Destaques</h2>
        <ul class='rr-list'>{highlights_html}</ul>
      </div>
    </div>
  </section>
  <section class='rr-grid'>
    <div class='rr-section'>
      <h2>Leitura temática</h2>
      <div class='rr-topic-grid'>
        {topic_card('DWP e Infraestrutura', '#ecfeff')}
        {topic_card('Application', '#f0fdf4')}
        {topic_card('Marketing', '#fff7ed')}
        {topic_card('Cyber', '#f8fafc')}
        {topic_card('Cloud', '#eff6ff')}
        {topic_card('Outros', '#f9fafb')}
      </div>
    </div>
    <div class='rr-section'>
      <h2>Próximos passos sugeridos</h2>
      <ul class='rr-list'>{next_steps_html}</ul>
    </div>
  </section>
</div>
<div id='rr-fireshot-modal' onclick="if(event.target===this)rrCloseFireshotModal()">
  <div class='rr-fs-box'>
    <div class='rr-fs-header'>
      <div class='rr-fs-title'>
        <span>📷</span> Exportar como JPG
      </div>
      <button class='rr-fs-close' onclick='rrCloseFireshotModal()'>&#215;</button>
    </div>
    <ol class='rr-fs-steps'>
      <li>
        <span class='rr-fs-step-num'>1</span>
        <span>Instale a extensão <strong>FireShot</strong> no Google Chrome (gratuita).</span>
      </li>
      <li>
        <span class='rr-fs-step-num'>2</span>
        <span>Feche este popup e, com o relatório aberto, clique com o botão <strong>direito</strong> em qualquer área da página.</span>
      </li>
      <li>
        <span class='rr-fs-step-num'>3</span>
        <span>No menu de contexto, selecione <strong>"Capturar página inteira"</strong> → <strong>"Salvar como imagem"</strong>.</span>
      </li>
      <li>
        <span class='rr-fs-step-num'>4</span>
        <span>Escolha o formato <strong>JPG</strong> e salve o arquivo.</span>
      </li>
    </ol>
    <div class='rr-fs-note'>
      💡 Se preferir, use o ícone do FireShot na barra de extensões do Chrome e selecione <em>"Capturar página inteira"</em>.
    </div>
    <div class='rr-fs-actions'>
      <a class='rr-fs-btn-install'
         href='https://chrome.google.com/webstore/detail/fireshot/mcbpblocgmgfnpjjppndjkmgjaogfceg'
         target='_blank' rel='noopener'>
        ⬇ Instalar FireShot
      </a>
      <button class='rr-fs-btn-ok' onclick='rrCloseFireshotModal()'>Entendido</button>
    </div>
  </div>
</div>
<script>
function rrShowFireshotModal() {{
  document.getElementById('rr-fireshot-modal').classList.add('active');
}}
function rrCloseFireshotModal() {{
  document.getElementById('rr-fireshot-modal').classList.remove('active');
}}
// Fechar com ESC
document.addEventListener('keydown', function(e) {{
  if (e.key === 'Escape') rrCloseFireshotModal();
}});
</script>
</body>
</html>"""


def _relation_report_render_pdf(report_data):
    global REPORTLAB_AVAILABLE, REPORTLAB_IMPORT_ERROR, colors, TA_LEFT, TA_CENTER, A4, ParagraphStyle, getSampleStyleSheet, mm, ImageReader, stringWidth, Paragraph
    if not REPORTLAB_AVAILABLE:
        try:
            from reportlab.lib import colors as _colors
            from reportlab.lib.enums import TA_LEFT as _TA_LEFT, TA_CENTER as _TA_CENTER
            from reportlab.lib.pagesizes import A4 as _A4
            from reportlab.lib.styles import ParagraphStyle as _ParagraphStyle, getSampleStyleSheet as _getSampleStyleSheet
            from reportlab.lib.units import mm as _mm
            from reportlab.lib.utils import ImageReader as _ImageReader
            from reportlab.pdfbase.pdfmetrics import stringWidth as _stringWidth
            from reportlab.platypus import Paragraph as _Paragraph
            colors = _colors
            TA_LEFT = _TA_LEFT
            TA_CENTER = _TA_CENTER
            A4 = _A4
            ParagraphStyle = _ParagraphStyle
            getSampleStyleSheet = _getSampleStyleSheet
            mm = _mm
            ImageReader = _ImageReader
            stringWidth = _stringWidth
            Paragraph = _Paragraph
            REPORTLAB_AVAILABLE = True
            REPORTLAB_IMPORT_ERROR = None
        except Exception as e:
            REPORTLAB_IMPORT_ERROR = e
            raise RuntimeError(f'ReportLab não está disponível para gerar o PDF. Detalhe: {e}')

    from reportlab.pdfgen import canvas
    colors_map = _relation_report_pick_system_colors()
    buffer = BytesIO()
    page_width, page_height = A4
    c = canvas.Canvas(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle('RelationBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=9.5, leading=13, textColor=colors.HexColor('#1f2937'))
    small_style = ParagraphStyle('RelationSmall', parent=body_style, fontSize=8.5, leading=11.5)

    def new_page():
        c.showPage()
        _relation_report_draw_header(c, report_data, colors_map, page_width, page_height)
        return page_height - 52 * mm

    def ensure_space(y, needed):
        return new_page() if y - needed < 18 * mm else y

    _relation_report_draw_header(c, report_data, colors_map, page_width, page_height)
    y = page_height - 52 * mm

    account = report_data['account']
    latest = report_data.get('latest_interaction') or {}
    counts = report_data.get('summary_counts') or {}
    narrative = report_data.get('narrative') or _relation_report_generate_narrative(report_data)

    cards = [
        ('Contatos', str(counts.get('contacts', 0))),
        ('Atividades', str(counts.get('activities', 0) + counts.get('account_activities', 0))),
        ('Kanban', str(counts.get('kanban_cards', 0))),
        ('Mapeamentos', str(counts.get('mapping_items', 0))),
    ]
    x_positions = [18 * mm, 67 * mm, 116 * mm, 165 * mm]
    for idx, (label, value) in enumerate(cards):
        c.setFillColor(colors.white)
        c.setStrokeColor(colors_map['accent'])
        c.roundRect(x_positions[idx], y - 16 * mm, 42 * mm, 14 * mm, 4 * mm, fill=1, stroke=1)
        c.setFillColor(colors_map['secondary'])
        c.setFont('Helvetica-Bold', 14)
        c.drawCentredString(x_positions[idx] + 21 * mm, y - 8 * mm, value)
        c.setFont('Helvetica', 8.5)
        c.drawCentredString(x_positions[idx] + 21 * mm, y - 12 * mm, label)
    y -= 22 * mm

    c.setFillColor(colors_map['secondary'])
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, y, 'Visão geral da conta')
    y -= 5 * mm
    c.setStrokeColor(colors.HexColor('#d1d5db'))
    c.line(18 * mm, y, page_width - 18 * mm, y)
    y -= 6 * mm

    general_lines = [
        f"Setor: {account.get('sector') or 'Não informado'}",
        f"Receita média: {format_currency_br(account.get('average_revenue_cents'))}",
        f"Profissionais: {account.get('professionals_count') or 'Não informado'}",
        f"Presença global: {account.get('global_presence') or 'Não informado'}",
        f"Última interação: {_relation_report_format_dt(latest.get('date'))}",
        f"Com quem: {latest.get('person') or 'Não identificado'}"
    ]
    for idx, line in enumerate(general_lines):
        col = 18 * mm if idx < 3 else 110 * mm
        row_y = y - (idx % 3) * 6 * mm
        c.setFillColor(colors.HexColor('#111827'))
        c.setFont('Helvetica', 9.2)
        c.drawString(col, row_y, line)
    y -= 22 * mm

    y = ensure_space(y, 40 * mm)
    c.setFillColor(colors_map['secondary'])
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, y, 'Resumo executivo do relacionamento')
    y -= 4 * mm
    c.setStrokeColor(colors_map['accent'])
    c.line(18 * mm, y, page_width - 18 * mm, y)
    y -= 5 * mm
    y = _relation_report_draw_paragraph(c, narrative.get('executive_summary') or 'Sem resumo disponível.', 18 * mm, y, page_width - 36 * mm, body_style)
    y -= 6 * mm

    if narrative.get('highlights'):
        y = ensure_space(y, 24 * mm)
        c.setFillColor(colors.HexColor('#ecfdf5'))
        c.roundRect(18 * mm, y - 18 * mm, page_width - 36 * mm, 16 * mm, 4 * mm, fill=1, stroke=0)
        c.setFillColor(colors_map['secondary'])
        c.setFont('Helvetica-Bold', 9.5)
        c.drawString(22 * mm, y - 6 * mm, 'Destaques')
        c.setFillColor(colors.HexColor('#065f46'))
        c.setFont('Helvetica', 8.5)
        for i, item in enumerate(narrative['highlights'][:3]):
            c.drawString(22 * mm, y - (10 + i * 4) * mm, f'• {item[:120]}')
        y -= 22 * mm

    y = ensure_space(y, 52 * mm)
    c.setFillColor(colors_map['secondary'])
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, y, 'Powermapping da conta')
    y -= 4 * mm
    c.line(18 * mm, y, page_width - 18 * mm, y)
    y -= 6 * mm
    for rel in report_data['relationship_cards'][:18]:
        y = ensure_space(y, 22 * mm)
        c.setFillColor(colors.white)
        c.setStrokeColor(colors.HexColor('#d1d5db'))
        c.roundRect(18 * mm, y - 17 * mm, page_width - 36 * mm, 15 * mm, 4 * mm, fill=1, stroke=1)
        contact = rel['contact']
        c.setFillColor(colors.HexColor('#111827'))
        c.setFont('Helvetica-Bold', 10)
        c.drawString(22 * mm, y - 6 * mm, f"{contact.get('name') or 'Contato'} — {contact.get('position') or 'Cargo não informado'}")
        c.setFont('Helvetica', 8.3)
        c.setFillColor(colors.HexColor('#4b5563'))
        linkedin_value = str(contact.get('linkedin') or '').strip() or 'Não informado'
        c.drawString(22 * mm, y - 10 * mm, f"Área: {contact.get('area_of_activity') or 'Não informada'} | Email: {contact.get('email') or 'Não informado'}")
        c.drawString(22 * mm, y - 12.5 * mm, f"LinkedIn: {linkedin_value[:95]}")
        c.drawString(22 * mm, y - 15 * mm, f"Cards: atividades {rel.get('activities_count',0)} | mapeamento {rel.get('mapping_count',0)} | kanban {rel.get('kanban_count',0)} | último contato {_relation_report_format_dt((rel.get('last_contact') or {}).get('date'))}")
        y -= 20 * mm

    y = ensure_space(y, 45 * mm)
    c.setFillColor(colors_map['secondary'])
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, y, 'Entregas, presenças e Kanban')
    y -= 4 * mm
    c.line(18 * mm, y, page_width - 18 * mm, y)
    y -= 5 * mm
    for presence in report_data['presences'][:8]:
        y = ensure_space(y, 9 * mm)
        txt = f"• {presence.get('delivery_name')} | owner {presence.get('stf_owner') or 'N/I'} | receita {format_currency_br(presence.get('current_revenue_cents'))} | validade {presence.get('validity_month') or 'N/I'} | focal {presence.get('focal_client_name') or 'N/I'}"
        y = _relation_report_draw_paragraph(c, txt, 18 * mm, y, page_width - 36 * mm, small_style) - 2 * mm
    if not report_data['presences']:
        y = _relation_report_draw_paragraph(c, 'Nenhuma presença/entrega mapeada para esta conta.', 18 * mm, y, page_width - 36 * mm, small_style) - 2 * mm
    y -= 2 * mm
    y = _relation_report_draw_paragraph(c, f"Kanban mapeado: {len(report_data['kanban_cards'])} card(s).", 18 * mm, y, page_width - 36 * mm, body_style) - 3 * mm
    for card in report_data['kanban_cards'][:10]:
        y = ensure_space(y, 8 * mm)
        txt = f"• [{card.get('column_title') or 'Sem coluna'}] {card.get('title')} — contato: {card.get('contact_name') or 'N/I'} — urgência: {card.get('urgency') or 'N/I'}"
        y = _relation_report_draw_paragraph(c, txt, 22 * mm, y, page_width - 40 * mm, small_style) - 1 * mm

    y = ensure_space(y, 55 * mm)
    c.setFillColor(colors_map['secondary'])
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, y, 'Quebra resumida por tópico')
    y -= 4 * mm
    c.line(18 * mm, y, page_width - 18 * mm, y)
    y -= 6 * mm
    for topic in RELATION_REPORT_TOPICS:
        y = ensure_space(y, 15 * mm)
        c.setFillColor(colors.HexColor('#ecfeff') if topic in ('DWP e Infraestrutura', 'Cloud') else colors.HexColor('#f9fafb'))
        c.roundRect(18 * mm, y - 12 * mm, page_width - 36 * mm, 10 * mm, 3 * mm, fill=1, stroke=0)
        c.setFillColor(colors_map['secondary'])
        c.setFont('Helvetica-Bold', 9.5)
        c.drawString(21 * mm, y - 6 * mm, topic)
        c.setFillColor(colors.HexColor('#374151'))
        topic_text = (narrative.get('topic_breakdown') or {}).get(topic) or f"Sem evidências relevantes para {topic.lower()}."
        y = _relation_report_draw_paragraph(c, topic_text, 42 * mm, y - 2 * mm, page_width - 60 * mm, small_style) - 4 * mm

    y = ensure_space(y, 30 * mm)
    c.setFillColor(colors_map['secondary'])
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, y, 'Próximos passos sugeridos')
    y -= 4 * mm
    c.line(18 * mm, y, page_width - 18 * mm, y)
    y -= 6 * mm
    for item in narrative.get('next_steps') or []:
        y = ensure_space(y, 8 * mm)
        y = _relation_report_draw_paragraph(c, f"• {item}", 18 * mm, y, page_width - 36 * mm, body_style) - 2 * mm

    c.setFont('Helvetica', 8)
    c.setFillColor(colors.HexColor('#6b7280'))
    c.drawRightString(page_width - 18 * mm, 12 * mm, f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    c.save()
    buffer.seek(0)
    return buffer


@app.route('/api/report/relation', methods=['GET'])
def export_relation_report():
    try:
        account_id_raw = (request.args.get('account_id') or '').strip()
        if not account_id_raw.isdigit():
            return jsonify({'error': 'account_id é obrigatório'}), 400
        account_id = int(account_id_raw)
        full_period = (request.args.get('full_period', 'false') or 'false').lower() == 'true'
        start_date = (request.args.get('start_date') or '').strip() or None
        end_date = (request.args.get('end_date') or '').strip() or None
        if not full_period and (not start_date or not end_date):
            return jsonify({'error': 'Informe start_date e end_date ou marque full_period=true'}), 400
        if full_period:
            start_date = None
            end_date = None
        report_data = _relation_report_collect_data(account_id, start_date=start_date, end_date=end_date)
        if not report_data:
            return jsonify({'error': 'Conta não encontrada'}), 404
        pdf_buffer = _relation_report_render_pdf(report_data)
        safe_name = re.sub(r'[^a-zA-Z0-9_-]+', '-', (report_data['account'].get('name') or 'cliente').strip()).strip('-').lower() or 'cliente'
        file_name = f'relation-report-{safe_name}-{datetime.now().strftime("%Y%m%d-%H%M%S")}.pdf'
        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=file_name
        )
    except Exception as e:
        logger.exception(f'[RelationReport] Falha ao gerar relatório: {e}')
        return jsonify({'error': str(e)}), 500


@app.route('/report/relation/view', methods=['GET'])
def view_relation_report_browser():
    try:
        account_id_raw = (request.args.get('account_id') or '').strip()
        if not account_id_raw.isdigit():
            return 'account_id é obrigatório', 400
        account_id = int(account_id_raw)
        full_period = (request.args.get('full_period', 'false') or 'false').lower() == 'true'
        start_date = (request.args.get('start_date') or '').strip() or None
        end_date = (request.args.get('end_date') or '').strip() or None
        if full_period:
            start_date = None
            end_date = None
        report_data = _relation_report_collect_data(account_id, start_date=start_date, end_date=end_date)
        if not report_data:
            return 'Conta não encontrada', 404
        report_data['narrative'] = _relation_report_generate_narrative(report_data)
        profile_response = get_profile_config()
        profile = profile_response.get_json(silent=True) if hasattr(profile_response, 'get_json') else {}
        html_doc = _relation_report_build_browser_html(report_data, profile=profile or {}, embed_images=True)
        return Response(html_doc, mimetype='text/html; charset=utf-8')
    except Exception as e:
        logger.exception(f'[RelationReport] Falha ao montar visualização HTML: {e}')
        return f'<h1>Erro ao gerar relatório</h1><pre>{html.escape(str(e))}</pre>', 500


@app.route('/report/relation/export-html', methods=['GET'])
def export_relation_report_html():
    try:
        account_id_raw = (request.args.get('account_id') or '').strip()
        if not account_id_raw.isdigit():
            return 'account_id é obrigatório', 400
        account_id = int(account_id_raw)
        full_period = (request.args.get('full_period', 'false') or 'false').lower() == 'true'
        start_date = (request.args.get('start_date') or '').strip() or None
        end_date = (request.args.get('end_date') or '').strip() or None
        if full_period:
            start_date = None
            end_date = None
        report_data = _relation_report_collect_data(account_id, start_date=start_date, end_date=end_date)
        if not report_data:
            return 'Conta não encontrada', 404
        report_data['narrative'] = _relation_report_generate_narrative(report_data)
        profile_response = get_profile_config()
        profile = profile_response.get_json(silent=True) if hasattr(profile_response, 'get_json') else {}
        html_doc = _relation_report_build_browser_html(report_data, profile=profile or {}, embed_images=True)
        account_name = (report_data.get('account') or {}).get('name') or f'account-{account_id}'
        safe_name = re.sub(r'[^a-zA-Z0-9_-]+', '-', account_name.strip()).strip('-').lower() or f'account-{account_id}'
        response = Response(html_doc, mimetype='text/html; charset=utf-8')
        response.headers['Content-Disposition'] = f'attachment; filename=relation-report-{safe_name}.html'
        return response
    except Exception as e:
        logger.exception(f'[RelationReport] Falha ao exportar HTML: {e}')
        return f'<h1>Erro ao exportar relatório</h1><pre>{html.escape(str(e))}</pre>', 500


@app.route('/api/report/relation/preview', methods=['GET'])
def preview_relation_report_data():
    try:
        account_id_raw = (request.args.get('account_id') or '').strip()
        if not account_id_raw.isdigit():
            return jsonify({'error': 'account_id é obrigatório'}), 400
        account_id = int(account_id_raw)
        full_period = (request.args.get('full_period', 'false') or 'false').lower() == 'true'
        start_date = (request.args.get('start_date') or '').strip() or None
        end_date = (request.args.get('end_date') or '').strip() or None
        if full_period:
            start_date = None
            end_date = None
        report_data = _relation_report_collect_data(account_id, start_date=start_date, end_date=end_date)
        if not report_data:
            return jsonify({'error': 'Conta não encontrada'}), 404
        narrative = _relation_report_generate_narrative(report_data)
        return jsonify({
            'account': report_data['account'],
            'summary_counts': report_data['summary_counts'],
            'latest_interaction': report_data['latest_interaction'],
            'relationship_cards': report_data['relationship_cards'][:12],
            'topics': {k: len(v) for k, v in report_data['topics'].items()},
            'highlights': narrative.get('highlights') or [],
            'narrative': narrative,
            'period': {
                'full_period': report_data['full_period'],
                'start_date': report_data['start_date'],
                'end_date': report_data['end_date']
            }
        })
    except Exception as e:
        logger.exception(f'[RelationReport] Falha ao gerar preview: {e}')
        return jsonify({'error': str(e)}), 500


ITOCA_EXCLUDED_TABLES = {
    'sqlite_sequence',
    'app_settings',
    'itoca_chat_history',
    'automapping_runs',
    'status_rules',
    'job_groupings',
    'job_grouping_positions',
}


# Mapa de sinônimos para expansão semântica nas buscas do RAG
_ITOCA_SYNONYMS = {
    'cnpj': ['cnpj', 'cadastro nacional', 'pessoa juridica', 'registro empresa', 'inscricao federal'],
    'cpf': ['cpf', 'cadastro pessoa fisica', 'registro pessoa'],
    'email': ['email', 'e-mail', 'correio eletronico', 'contato'],
    'telefone': ['telefone', 'celular', 'fone', 'contato', 'whatsapp'],
    'agenda': ['agenda', 'compromisso', 'reuniao', 'evento', 'encontro', 'meeting'],
    'reuniao': ['reuniao', 'meeting', 'encontro', 'compromisso', 'agenda'],
    'evento': ['evento', 'compromisso', 'reuniao', 'agenda', 'encontro'],
    'proxima': ['proxima', 'proximo', 'futuro', 'upcoming', 'semana'],
    'semana': ['semana', 'week', 'proximos dias', 'agenda'],
    'contato': ['contato', 'cliente', 'pessoa', 'lead'],
    'empresa': ['empresa', 'companhia', 'organizacao', 'account', 'conta'],
    'wiki': ['wiki', 'documento', 'conhecimento', 'wikitoca'],
    'documento': ['documento', 'arquivo', 'pdf', 'doc', 'wiki'],
    'atividade': ['atividade', 'activity', 'interacao', 'historico'],
    'contas': ['contas', 'conta', 'empresa', 'account', 'accounts', 'clientes', 'relacionamento'],
    'conta': ['conta', 'contas', 'empresa', 'account', 'accounts'],
    'clientes': ['clientes', 'cliente', 'contatos', 'contato', 'leads', 'prospects', 'contas'],
    'cliente': ['cliente', 'clientes', 'contato', 'lead', 'prospect'],
    'relacionamento': ['relacionamento', 'relacao', 'historico', 'atividade', 'interacao', 'contas', 'clientes'],
    'pipeline': ['pipeline', 'kanban', 'funil', 'oportunidade', 'negociacao'],
    'resumo': ['resumo', 'status', 'situacao', 'overview', 'panorama', 'visao geral'],
    'status': ['status', 'situacao', 'resumo', 'estado', 'andamento'],
}

def _itoca_tokenize(question):
    base_tokens = [
        token for token in re.findall(r'[a-zA-ZÀ-ÿ0-9_-]{3,}', (question or '').lower())
        if token not in {'para', 'com', 'sem', 'sobre', 'como', 'qual', 'quais', 'onde', 'quando', 'que', 'uma', 'uns', 'das', 'dos', 'nos', 'nas', 'por', 'foi', 'sao', 'tem', 'ter'}
    ]
    # Expande com sinônimos semânticos
    expanded = list(base_tokens)
    for token in base_tokens:
        if token in _ITOCA_SYNONYMS:
            for syn in _ITOCA_SYNONYMS[token]:
                if syn not in expanded:
                    expanded.append(syn)
    return expanded[:20]  # limita para não explodir a query SQL


def _itoca_text_columns(cursor, table_name):
    cursor.execute(f'PRAGMA table_info("{table_name}")')
    columns = []
    for row in cursor.fetchall():
        column_name = row['name']
        column_type = (row['type'] or '').upper()
        if any(t in column_type for t in ['CHAR', 'TEXT', 'CLOB']) or column_type == '':
            columns.append(column_name)
    return columns


def _itoca_build_snippet(row_dict):
    # Campos de baixa relevância semântica que não devem aparecer no snippet genérico
    _SKIP_KEYS = {'id', 'created_at', 'updated_at', 'logo_url', 'photo_url', 'file_url',
                  'file_name', 'file_size', 'display_order', 'is_system', 'is_locked',
                  'completed', 'completed_at',
                  'account_id', 'client_id', 'card_id', 'column_id', 'presence_id',
                  'activity_id', 'grouping_id', 'focal_client_id', 'contact_id', 'target_id'}
    # Campos de data que devem ser formatados
    _DATE_KEYS = {'activity_date', 'due_date', 'validity_month'}
    parts = []
    # Traduz flags booleanas para rótulos semânticos legíveis pelo LLM
    if row_dict.get('is_target'):
        parts.append('classificacao: conta-alvo (target)')
    if row_dict.get('is_cold_contact'):
        parts.append('classificacao: cold contact')
    for key, value in row_dict.items():
        if key in _SKIP_KEYS:
            continue
        if key in ('is_target', 'is_cold_contact'):
            continue  # já tratados acima
            continue
        if value is None:
            continue
        text = str(value).strip()
        if not text:
            continue
        # Formata datas para o padrão brasileiro
        if key in _DATE_KEYS:
            try:
                dt = datetime.strptime(text[:10], '%Y-%m-%d')
                text = dt.strftime('%d/%m/%Y')
            except Exception as e:
                logger.debug(f'[_itoca_build_snippet] exceção ignorada: {e}')
        # Trunca campos muito longos
        if len(text) > 500:
            text = text[:500] + '...'
        parts.append(f'{key}: {text}')
        if len(parts) >= 8:
            break
    return ' | '.join(parts)


def _itoca_search_context(question, limit=18):
    tokens = _itoca_tokenize(question)
    if not tokens:
        return []
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")
    tables = [row['name'] for row in cursor.fetchall() if row['name'] not in ITOCA_EXCLUDED_TABLES]
    results = []
    seen = set()
    like_values = [f'%{token}%' for token in tokens[:6]]
    _generic_limit_reached = False
    for table in tables:
        if _generic_limit_reached:
            break
        text_columns = _itoca_text_columns(cursor, table)
        if not text_columns:
            continue
        search_clauses = []
        params = []
        for col in text_columns[:8]:
            for like_value in like_values:
                search_clauses.append(f'LOWER(COALESCE("{col}", "")) LIKE ?')
                params.append(like_value)
        if not search_clauses:
            continue
        sql = f'SELECT * FROM "{table}" WHERE ' + ' OR '.join(search_clauses) + ' LIMIT 6'
        try:
            cursor.execute(sql, params)
            rows = cursor.fetchall()
        except Exception:
            continue
        for row in rows:
            row_dict = dict_from_row(row)
            # Enriquece o row_dict resolvendo FKs para nomes legíveis antes de gerar o snippet
            row_dict = _itoca_enrich_snippet_with_joins(cursor, table, row_dict)
            snippet = _itoca_build_snippet(row_dict)
            if not snippet:
                continue
            fingerprint = f"{table}:{snippet[:160]}"
            if fingerprint in seen:
                continue
            seen.add(fingerprint)
            results.append({
                'table': table,
                'id': row_dict.get('id'),
                'snippet': snippet,
                'search_text': snippet.lower()
            })
            if len(results) >= limit:
                _generic_limit_reached = True
                break  # Quebra o for row; a flag quebra o for table na próxima iteração.
                       # As buscas especializadas de wiki, agenda, accounts continuam abaixo.

    # Busca especializada de accounts/clients quando a pergunta menciona contas, clientes, relacionamento ou resumo
    account_client_keywords = {'contas', 'conta', 'clientes', 'cliente', 'relacionamento', 'relacao', 'resumo', 'status', 'pipeline', 'accounts', 'account'}
    is_account_query = any(t in account_client_keywords for t in tokens[:10])
    if is_account_query and len(results) < limit:
        try:
            # Busca accounts com última atividade, is_target e contagem de serviços
            cursor.execute('''
                SELECT ac.id, ac.name, ac.sector, ac.description, ac.is_target,
                       MAX(act.activity_date) as last_activity,
                       COUNT(DISTINCT act.id) as total_activities,
                       COUNT(DISTINCT ap.id) as total_services
                FROM accounts ac
                LEFT JOIN activities act ON act.account_id = ac.id
                LEFT JOIN account_presences ap ON ap.account_id = ac.id
                GROUP BY ac.id
                ORDER BY ac.is_target DESC, last_activity DESC NULLS LAST
                LIMIT 30
            ''')
            for row in cursor.fetchall():
                rd = dict_from_row(row)
                parts = []
                if rd.get('is_target'):
                    parts.append('classificacao: conta-alvo (target)')
                if rd.get('name'):
                    parts.append(f'empresa: {rd["name"]}')
                if rd.get('sector'):
                    parts.append(f'setor: {rd["sector"]}')
                if rd.get('description'):
                    parts.append(f'descricao: {rd["description"][:200]}')
                if rd.get('last_activity'):
                    try:
                        dt = datetime.strptime(rd['last_activity'][:10], '%Y-%m-%d')
                        parts.append(f'ultimo_contato: {dt.strftime("%d/%m/%Y")}')
                    except Exception:
                        parts.append(f'ultimo_contato: {rd["last_activity"]}')
                if rd.get('total_activities'):
                    parts.append(f'total_interacoes: {rd["total_activities"]}')
                if rd.get('total_services'):
                    parts.append(f'servicos_stefanini_cadastrados: {rd["total_services"]}')
                snippet = ' | '.join(parts)
                if not snippet:
                    continue
                fp = f'accounts:{snippet[:160]}'
                if fp in seen:
                    continue
                seen.add(fp)
                results.append({'table': 'accounts', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
        except Exception as e:
            logger.warning(f'[iToca] Erro ao buscar accounts para query de contas: {e}')

        try:
            # Busca clients com última atividade, is_target e is_cold_contact
            cursor.execute('''
                SELECT cl.id, cl.name, cl.company, cl.position, cl.email,
                       cl.last_activity_date, cl.is_target, cl.is_cold_contact,
                       COUNT(act.id) as total_activities
                FROM clients cl
                LEFT JOIN activities act ON act.client_id = cl.id
                GROUP BY cl.id
                ORDER BY cl.is_target DESC, cl.last_activity_date DESC NULLS LAST
                LIMIT 25
            ''')
            for row in cursor.fetchall():
                rd = dict_from_row(row)
                parts = []
                if rd.get('is_target'):
                    parts.append('classificacao: contato-alvo (target)')
                if rd.get('is_cold_contact'):
                    parts.append('classificacao: cold contact')
                if rd.get('name'):
                    parts.append(f'contato: {rd["name"]}')
                if rd.get('company'):
                    parts.append(f'empresa: {rd["company"]}')
                if rd.get('position'):
                    parts.append(f'cargo: {rd["position"]}')
                if rd.get('email'):
                    parts.append(f'email: {rd["email"]}')
                if rd.get('last_activity_date'):
                    try:
                        dt = datetime.strptime(rd['last_activity_date'][:10], '%Y-%m-%d')
                        parts.append(f'ultimo_contato: {dt.strftime("%d/%m/%Y")}')
                    except Exception:
                        parts.append(f'ultimo_contato: {rd["last_activity_date"]}')
                if rd.get('total_activities'):
                    parts.append(f'total_interacoes: {rd["total_activities"]}')
                snippet = ' | '.join(parts)
                if not snippet:
                    continue
                fp = f'clients:{snippet[:160]}'
                if fp in seen:
                    continue
                seen.add(fp)
                results.append({'table': 'clients', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
        except Exception as e:
            logger.warning(f'[iToca] Erro ao buscar clients para query de contas: {e}')

    # Busca dedicada na agenda (commitments + account_renewal_events)
    try:
        agenda_tokens = tokens[:8]
        # Sempre inclui eventos futuros quando a pergunta menciona agenda/evento/semana/próximo
        agenda_keywords = {'agenda', 'compromisso', 'reuniao', 'evento', 'encontro', 'meeting', 'semana', 'proxima', 'proximo', 'futuro', 'upcoming', 'proximos'}
        is_agenda_query = any(t in agenda_keywords for t in agenda_tokens)
        if is_agenda_query:
            # Retorna todos os eventos futuros (próximos 90 dias)
            future_limit = (datetime.now() + timedelta(days=90)).strftime('%Y-%m-%d')
            today_str = datetime.now().strftime('%Y-%m-%d')
            cursor.execute('''
                SELECT cm.id, cm.title, cm.notes, cm.due_date, cm.due_time,
                       cl.name as client_name, cl.company as client_company
                FROM commitments cm
                LEFT JOIN clients cl ON cm.client_id = cl.id
                WHERE DATE(cm.due_date) >= ?
                ORDER BY cm.due_date ASC
                LIMIT 20
            ''', (today_str,))
            for row in cursor.fetchall():
                rd = dict_from_row(row)
                parts = []
                if rd.get('title'):
                    parts.append(f'titulo: {rd["title"]}')
                if rd.get('due_date'):
                    try:
                        dt = datetime.strptime(rd['due_date'][:10], '%Y-%m-%d')
                        parts.append(f'data: {dt.strftime("%d/%m/%Y")}')
                    except Exception:
                        parts.append(f'data: {rd["due_date"]}')
                if rd.get('due_time'):
                    parts.append(f'hora: {rd["due_time"]}')
                if rd.get('client_name'):
                    parts.append(f'contato: {rd["client_name"]}')
                if rd.get('client_company'):
                    parts.append(f'empresa: {rd["client_company"]}')
                if rd.get('notes') and rd['notes'] != rd.get('title'):
                    parts.append(f'notas: {rd["notes"][:200]}')
                snippet = ' | '.join(parts)
                if not snippet:
                    continue
                fp = f'commitments:{snippet[:160]}'
                if fp in seen:
                    continue
                seen.add(fp)
                results.append({'table': 'commitments', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
        else:
            # Busca por tokens no título/notas dos compromissos
            if agenda_tokens:
                like_clauses_a = ' OR '.join(['LOWER(cm.title) LIKE ? OR LOWER(cm.notes) LIKE ?' for _ in agenda_tokens[:4]])
                params_a = []
                for t in agenda_tokens[:4]:
                    params_a += [f'%{t}%', f'%{t}%']
                cursor.execute(f'''
                    SELECT cm.id, cm.title, cm.notes, cm.due_date, cm.due_time,
                           cl.name as client_name, cl.company as client_company
                    FROM commitments cm
                    LEFT JOIN clients cl ON cm.client_id = cl.id
                    WHERE {like_clauses_a}
                    ORDER BY cm.due_date ASC LIMIT 6
                ''', params_a)
                for row in cursor.fetchall():
                    rd = dict_from_row(row)
                    parts = []
                    if rd.get('title'):
                        parts.append(f'titulo: {rd["title"]}')
                    if rd.get('due_date'):
                        try:
                            dt = datetime.strptime(rd['due_date'][:10], '%Y-%m-%d')
                            parts.append(f'data: {dt.strftime("%d/%m/%Y")}')
                        except Exception:
                            parts.append(f'data: {rd["due_date"]}')
                    if rd.get('due_time'):
                        parts.append(f'hora: {rd["due_time"]}')
                    if rd.get('client_name'):
                        parts.append(f'contato: {rd["client_name"]}')
                    snippet = ' | '.join(parts)
                    if not snippet:
                        continue
                    fp = f'commitments:{snippet[:160]}'
                    if fp in seen:
                        continue
                    seen.add(fp)
                    results.append({'table': 'commitments', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao buscar agenda: {e}')

    # Busca adicional em wiki_entries (conteúdo completo)
    try:
        like_clauses = ' OR '.join(['LOWER(title) LIKE ? OR LOWER(content) LIKE ? OR LOWER(tags) LIKE ?' for _ in tokens[:6]])
        params_wiki = []
        for t in tokens[:6]:
            params_wiki += [f'%{t}%', f'%{t}%', f'%{t}%']
        cursor.execute(f'SELECT id, title, category, content, tags FROM wiki_entries WHERE {like_clauses} LIMIT 6', params_wiki)
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'):
                parts.append(f'titulo: {rd["title"]}')
            if rd.get('category'):
                parts.append(f'categoria: {rd["category"]}')
            if rd.get('content'):
                parts.append(f'conteudo: {rd["content"][:1500]}')
            snippet = ' | '.join(parts)
            if not snippet:
                continue
            fp = f'wiki_entries:{snippet[:160]}'
            if fp in seen:
                continue
            seen.add(fp)
            results.append({'table': 'wiki_entries', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao buscar wiki_entries: {e}')

    # Busca adicional em wiki_documents (metadados + busca no texto extraído do arquivo)
    try:
        # Busca por título/nome do arquivo
        like_clauses_d = ' OR '.join(['LOWER(title) LIKE ? OR LOWER(original_name) LIKE ?' for _ in tokens[:6]])
        params_docs = []
        for t in tokens[:6]:
            params_docs += [f'%{t}%', f'%{t}%']
        cursor.execute(f'SELECT id, title, original_name, file_name, file_ext FROM wiki_documents WHERE {like_clauses_d} LIMIT 4', params_docs)
        doc_rows_by_title = {row['id']: dict_from_row(row) for row in cursor.fetchall()}

        # Busca também em TODOS os documentos pelo conteúdo extraído (para termos como CNPJ que não estão no título)
        cursor.execute('SELECT id, title, original_name, file_name, file_ext FROM wiki_documents LIMIT 20')
        all_doc_rows = {row['id']: dict_from_row(row) for row in cursor.fetchall()}
        # Mescla: prioriza os encontrados por título, depois verifica os demais pelo conteúdo
        candidate_docs = {**all_doc_rows, **doc_rows_by_title}  # doc_rows_by_title sobrescreve

        for doc_id, rd in candidate_docs.items():
            file_path = WIKI_UPLOAD_DIR / (rd.get('file_name') or '')
            doc_text = ''
            if file_path.exists():
                doc_text = _itoca_extract_text_from_file(str(file_path))
            # Verifica se algum token aparece no conteúdo extraído
            doc_text_lower = doc_text.lower()
            found_in_content = any(t in doc_text_lower for t in tokens[:10])
            found_by_title = doc_id in doc_rows_by_title
            if not found_in_content and not found_by_title:
                continue
            doc_title = rd.get('title') or rd.get('original_name', '')
            doc_ext = rd.get('file_ext', '')
            if doc_text.strip():
                snippet = f'[WikiToca Doc] titulo: {doc_title} | tipo: {doc_ext} | conteudo: {doc_text[:3000]}'
            else:
                snippet = (f'[WikiToca Doc] titulo: {doc_title} | tipo: {doc_ext} | '
                           f'CONFIRMADO: o documento "{rd.get("original_name", "")}" EXISTE no WikiToca. '
                           f'Conteúdo interno não disponível para leitura automática (biblioteca de extração não instalada). '
                           f'Ao responder, informe que o documento existe mas não é possível detalhar o conteúdo.')
            fp = f'wiki_documents:{snippet[:160]}'
            if fp in seen:
                continue
            seen.add(fp)
            results.append({'table': 'wiki_documents', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao buscar wiki_documents: {e}')

    conn.close()
    return results


def _itoca_enrich_snippet_with_joins(cursor, table, row_dict):
    """Enriquece um row_dict de busca ao vivo resolvendo FKs para nomes legíveis.
    Isso garante que qualquer tabela retornada pela busca genérica (SELECT *)
    nunca exponha IDs brutos ao LLM quando o nome correspondente está disponível.
    Aplica-se a qualquer tabela que contenha client_id, account_id, contact_id,
    column_id, card_id, grouping_id, etc.
    """
    enriched = dict(row_dict)  # cópia para não mutar o original

    # Resolve client_id / contact_id / focal_client_id -> nome do contato e empresa
    for fk_field in ('client_id', 'contact_id', 'focal_client_id'):
        fk_val = enriched.get(fk_field)
        if fk_val and not enriched.get('client_name') and not enriched.get('contact_name'):
            try:
                cursor.execute('SELECT name, company, position FROM clients WHERE id = ? LIMIT 1', (fk_val,))
                r = cursor.fetchone()
                if r:
                    rd_cl = dict_from_row(r)
                    if rd_cl.get('name'):
                        enriched['nome_contato'] = rd_cl['name']
                    if rd_cl.get('company'):
                        enriched['empresa'] = rd_cl['company']
                    if rd_cl.get('position'):
                        enriched['cargo'] = rd_cl['position']
            except Exception as e:
                logger.debug(f'[_itoca_enrich_snippet_with_joins] exceção ignorada: {e}')
        # Remove o campo de FK bruto para não aparecer no snippet
        enriched.pop(fk_field, None)

    # Resolve account_id -> nome da conta
    acct_id = enriched.get('account_id')
    if acct_id and not enriched.get('account_name') and not enriched.get('nome_conta'):
        try:
            cursor.execute('SELECT name, sector FROM accounts WHERE id = ? LIMIT 1', (acct_id,))
            r = cursor.fetchone()
            if r:
                rd_ac = dict_from_row(r)
                if rd_ac.get('name'):
                    enriched['nome_conta'] = rd_ac['name']
                if rd_ac.get('sector'):
                    enriched['setor'] = rd_ac['sector']
        except Exception as e:
            logger.debug(f'[_itoca_enrich_snippet_with_joins] exceção ignorada: {e}')
    enriched.pop('account_id', None)

    # Resolve column_id -> título da coluna do Kanban
    col_id = enriched.get('column_id')
    if col_id and not enriched.get('column_title') and not enriched.get('coluna'):
        try:
            cursor.execute('SELECT title FROM kanban_columns WHERE id = ? LIMIT 1', (col_id,))
            r = cursor.fetchone()
            if r:
                rd_col = dict_from_row(r)
                if rd_col.get('title'):
                    enriched['coluna_kanban'] = rd_col['title']
        except Exception as e:
            logger.debug(f'[_itoca_enrich_snippet_with_joins] exceção ignorada: {e}')
    enriched.pop('column_id', None)

    # Resolve card_id -> título do card
    card_id = enriched.get('card_id')
    if card_id and not enriched.get('card_title'):
        try:
            cursor.execute('SELECT title FROM environment_cards WHERE id = ? LIMIT 1', (card_id,))
            r = cursor.fetchone()
            if r:
                rd_card = dict_from_row(r)
                if rd_card.get('title'):
                    enriched['mapeamento'] = rd_card['title']
        except Exception as e:
            logger.debug(f'[_itoca_enrich_snippet_with_joins] exceção ignorada: {e}')
    enriched.pop('card_id', None)

    # Resolve grouping_id -> nome do agrupamento
    grp_id = enriched.get('grouping_id')
    if grp_id and not enriched.get('agrupamento'):
        try:
            cursor.execute('SELECT name FROM job_groupings WHERE id = ? LIMIT 1', (grp_id,))
            r = cursor.fetchone()
            if r:
                rd_grp = dict_from_row(r)
                if rd_grp.get('name'):
                    enriched['agrupamento'] = rd_grp['name']
        except Exception as e:
            logger.debug(f'[_itoca_enrich_snippet_with_joins] exceção ignorada: {e}')
    enriched.pop('grouping_id', None)

    # Remove outros campos de FK que não foram resolvidos
    for leftover in ('activity_id', 'presence_id', 'target_id'):
        enriched.pop(leftover, None)

    return enriched


def _itoca_find_tesseract_cmd():
    """Localiza o binário do tesseract no sistema.
    Ordem de busca:
      1. Diretório do próprio executável (bundled com o Toca do Coelho via NSIS)
      2. PATH do sistema
      3. Caminhos padrão do Windows
    """
    import subprocess

    # 1. Bundled junto ao executável do Toca do Coelho
    #    Quando empacotado com PyInstaller, sys.executable aponta para TocaDoCoelho.exe
    #    O installer.nsi instala o Tesseract em $INSTDIR\tesseract\
    bundled_candidates = []
    exe_dir = Path(sys.executable).parent if getattr(sys, 'frozen', False) else Path(__file__).resolve().parent
    bundled_candidates.append(exe_dir / 'tesseract' / 'tesseract.exe')
    # PyInstaller _MEIPASS
    meipass = getattr(sys, '_MEIPASS', None)
    if meipass:
        bundled_candidates.append(Path(meipass) / 'tesseract' / 'tesseract.exe')
    # Diretório pai (caso o app.py esteja em subpasta)
    bundled_candidates.append(Path(__file__).resolve().parent / 'tesseract' / 'tesseract.exe')
    for candidate in bundled_candidates:
        if candidate.exists():
            # Configura TESSDATA_PREFIX para o tessdata bundled
            tessdata_dir = candidate.parent / 'tessdata'
            if tessdata_dir.exists():
                os.environ['TESSDATA_PREFIX'] = str(tessdata_dir)
            return str(candidate)

    # 2. PATH do sistema
    try:
        result = subprocess.run(['tesseract', '--version'], capture_output=True, timeout=5)
        if result.returncode == 0:
            return 'tesseract'
    except Exception as e:
        logger.debug(f'[_itoca_find_tesseract_cmd] exceção ignorada: {e}')

    # 3. Caminhos padrão do Windows
    if sys.platform == 'win32':
        windows_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Users\{}\AppData\Local\Tesseract-OCR\tesseract.exe'.format(os.environ.get('USERNAME', '')),
            r'C:\TocaDoCoelho\tesseract\tesseract.exe',
        ]
        for p in windows_paths:
            if Path(p).exists():
                return p

    return None


def _itoca_extract_text_from_file(file_path_str):
    """Extrai texto de PDF, DOCX ou XLSX para indexação no RAG.
    Estratégia para PDFs:
      1. pdfplumber (texto digital)
      2. pdftotext via subprocess (poppler, Windows/Linux)
      3. OCR com pdf2image + pytesseract (PDFs escaneados/imagens)
    """
    path = Path(file_path_str)
    if not path.exists():
        logger.warning(f'[iToca] Arquivo não encontrado: {file_path_str}')
        return ''
    ext = path.suffix.lower()
    text_parts = []
    try:
        if ext == '.pdf':
            extracted = False

            # --- Tentativa 1: pdfplumber (PDFs com texto digital) ---
            if PDFPLUMBER_AVAILABLE:
                try:
                    with pdfplumber.open(str(path)) as pdf:
                        for page in pdf.pages[:30]:
                            page_text = page.extract_text() or ''
                            if page_text.strip():
                                text_parts.append(page_text.strip())
                    extracted = bool(text_parts)
                except Exception as e1:
                    logger.warning(f'[iToca] pdfplumber falhou em {path.name}: {e1}')

            # --- Tentativa 2: pdftotext (poppler-utils) ---
            if not extracted:
                try:
                    import subprocess
                    # Tenta localizar pdftotext no Windows e Linux
                    pdftotext_cmd = 'pdftotext'
                    if sys.platform == 'win32':
                        win_poppler_paths = [
                            r'C:\poppler\bin\pdftotext.exe',
                            r'C:\Program Files\poppler\bin\pdftotext.exe',
                            r'C:\Program Files (x86)\poppler\bin\pdftotext.exe',
                        ]
                        for wp in win_poppler_paths:
                            if Path(wp).exists():
                                pdftotext_cmd = wp
                                break
                    result = subprocess.run(
                        [pdftotext_cmd, '-enc', 'UTF-8', str(path), '-'],
                        capture_output=True, timeout=20
                    )
                    if result.returncode == 0 and result.stdout:
                        decoded = result.stdout.decode('utf-8', errors='replace')
                        if decoded.strip():
                            text_parts.append(decoded)
                            extracted = True
                except Exception as e2:
                    logger.debug(f'[iToca] pdftotext não disponível para {path.name}: {e2}')

            # --- Tentativa 3: OCR com pypdfium2/pdf2image + pytesseract (PDFs escaneados) ---
            if not extracted and PYTESSERACT_AVAILABLE and (PYPDFIUM2_AVAILABLE or PDF2IMAGE_AVAILABLE):
                tess_cmd = _itoca_find_tesseract_cmd()
                if tess_cmd:
                    pytesseract.pytesseract.tesseract_cmd = tess_cmd
                    images = []
                    try:
                        if PYPDFIUM2_AVAILABLE:
                            doc = pdfium.PdfDocument(str(path))
                            for i in range(min(len(doc), 15)):
                                page = doc[i]
                                bitmap = page.render(scale=200/72)
                                images.append(bitmap.to_pil())
                        elif PDF2IMAGE_AVAILABLE:
                            images = convert_from_path(str(path), dpi=200, first_page=1, last_page=15)
                    except Exception as e_render:
                        logger.warning(f'[iToca] OCR render falhou em {path.name}: {e_render}')
                    if images:
                        try:
                            ocr_parts = []
                            for img in images:
                                ocr_text = pytesseract.image_to_string(img, lang='por+eng')
                                if ocr_text.strip():
                                    ocr_parts.append(ocr_text.strip())
                            if ocr_parts:
                                text_parts.extend(ocr_parts)
                                extracted = True
                                logger.info(f'[iToca] OCR extraiu texto de {path.name} ({len(images)} páginas)')
                        except Exception as e3:
                            logger.warning(f'[iToca] OCR falhou em {path.name}: {e3}')
                else:
                    logger.info(f'[iToca] Tesseract não encontrado. Instale em https://github.com/UB-Mannheim/tesseract/wiki para OCR de PDFs escaneados.')

            if not extracted:
                logger.warning(f'[iToca] Não foi possível extrair texto de {path.name} '
                               f'(pdfplumber={PDFPLUMBER_AVAILABLE}, ocr={PDF2IMAGE_AVAILABLE and PYTESSERACT_AVAILABLE})')

        elif ext in ('.docx', '.doc'):
            if PYTHON_DOCX_AVAILABLE:
                try:
                    doc = python_docx.Document(str(path))
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text.strip())
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                text_parts.append(row_text)
                except Exception as e4:
                    logger.warning(f'[iToca] python-docx falhou em {path.name}: {e4}')

        elif ext in ('.xlsx', '.xls'):
            if OPENPYXL_AVAILABLE:
                try:
                    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            row_text = ' | '.join(str(c) for c in row if c is not None and str(c).strip())
                            if row_text:
                                text_parts.append(row_text)
                except Exception as e5:
                    logger.warning(f'[iToca] openpyxl falhou em {path.name}: {e5}')

        elif ext == '.txt':
            try:
                text_parts.append(path.read_text(encoding='utf-8', errors='replace'))
            except Exception as e:
                logger.debug(f'[_itoca_extract_text_from_file] exceção ignorada: {e}')

    except Exception as e:
        logger.warning(f'[iToca] Erro geral ao extrair texto de {file_path_str}: {e}')

    result_text = '\n'.join(text_parts)
    if result_text:
        logger.info(f'[iToca] Extraiu {len(result_text)} chars de {path.name}')
    return result_text


def _itoca_build_agenda_items(conn):
    """Gera itens da agenda (commitments + account_renewal_events) para o snapshot RAG com contexto rico."""
    cursor = conn.cursor()
    items = []
    today = datetime.now().strftime('%Y-%m-%d')
    # Compromissos da agenda (próximos 90 dias e últimos 30 dias)
    try:
        cursor.execute('''
            SELECT cm.id, cm.title, cm.notes, cm.due_date, cm.due_time, cm.source_type,
                   cl.name as client_name, cl.company as client_company
            FROM commitments cm
            LEFT JOIN clients cl ON cm.client_id = cl.id
            ORDER BY cm.due_date ASC
        ''')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'):
                parts.append(f'titulo: {rd["title"]}')
            if rd.get('due_date'):
                try:
                    dt = datetime.strptime(rd['due_date'][:10], '%Y-%m-%d')
                    parts.append(f'data: {dt.strftime("%d/%m/%Y")}')
                except Exception:
                    parts.append(f'data: {rd["due_date"]}')
            if rd.get('due_time'):
                parts.append(f'hora: {rd["due_time"]}')
            if rd.get('client_name'):
                parts.append(f'contato: {rd["client_name"]}')
            if rd.get('client_company'):
                parts.append(f'empresa: {rd["client_company"]}')
            if rd.get('notes') and rd['notes'] != rd.get('title'):
                parts.append(f'notas: {rd["notes"][:300]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({
                    'table': 'commitments',
                    'id': rd.get('id'),
                    'snippet': snippet,
                    'search_text': snippet.lower()
                })
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar commitments: {e}')

    # Eventos de renovação de contas
    try:
        cursor.execute('''
            SELECT ev.id, ev.title, ev.due_date, ev.due_time,
                   ac.name as account_name
            FROM account_renewal_events ev
            LEFT JOIN accounts ac ON ev.account_id = ac.id
            ORDER BY ev.due_date ASC
        ''')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'):
                parts.append(f'titulo: {rd["title"]}')
            if rd.get('due_date'):
                try:
                    dt = datetime.strptime(rd['due_date'][:10], '%Y-%m-%d')
                    parts.append(f'data: {dt.strftime("%d/%m/%Y")}')
                except Exception:
                    parts.append(f'data: {rd["due_date"]}')
            if rd.get('due_time'):
                parts.append(f'hora: {rd["due_time"]}')
            if rd.get('account_name'):
                parts.append(f'conta: {rd["account_name"]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({
                    'table': 'account_renewal_events',
                    'id': rd.get('id'),
                    'snippet': snippet,
                    'search_text': snippet.lower()
                })
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar account_renewal_events: {e}')

    return items


def _itoca_build_wiki_items(conn, max_chars_per_doc=8000):
    """Gera itens do WikiToca (entradas de conhecimento + documentos) para o snapshot RAG."""
    cursor = conn.cursor()
    items = []

    # 1. Entradas de conhecimento (wiki_entries)
    try:
        cursor.execute('SELECT id, title, category, content, tags FROM wiki_entries ORDER BY updated_at DESC')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'):
                parts.append(f'titulo: {rd["title"]}')
            if rd.get('category'):
                parts.append(f'categoria: {rd["category"]}')
            if rd.get('tags'):
                parts.append(f'tags: {rd["tags"]}')
            if rd.get('content'):
                parts.append(f'conteudo: {rd["content"][:2000]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({
                    'table': 'wiki_entries',
                    'id': rd.get('id'),
                    'snippet': snippet,
                    'search_text': snippet.lower()
                })
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar wiki_entries: {e}')

    # 2. Documentos (wiki_documents) — extrai texto do arquivo
    try:
        cursor.execute('SELECT id, title, original_name, file_name, file_ext FROM wiki_documents ORDER BY updated_at DESC')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            file_path = WIKI_UPLOAD_DIR / (rd.get('file_name') or '')
            doc_text = ''
            if file_path.exists():
                doc_text = _itoca_extract_text_from_file(str(file_path))

            doc_title = rd.get('title') or rd.get('original_name', '')
            doc_ext = rd.get('file_ext', '')

            if not doc_text.strip():
                # Sem conteúdo extraível: indexa metadados com nota clara ao LLM
                snippet = (f'[WikiToca Doc] titulo: {doc_title} | tipo: {doc_ext} | '
                           f'CONFIRMADO: o documento "{rd.get("original_name", "")}" EXISTE no WikiToca. '
                           f'Conteúdo interno não disponível para leitura automática. '
                           f'Ao responder sobre este tema, informe que o documento existe mas não é possível detalhar o conteúdo.')
                # search_text inclui o título completo + nome original para melhor matching por tokens
                _search_base = (doc_title + ' ' + rd.get('original_name', '') + ' ' + doc_ext).lower()
                items.append({
                    'table': 'wiki_documents',
                    'id': rd.get('id'),
                    'snippet': snippet,
                    'search_text': _search_base
                })
            else:
                # Divide em chunks de max_chars_per_doc para não sobrecarregar o contexto
                max_total = max_chars_per_doc * 5
                chunks = [doc_text[i:i+max_chars_per_doc] for i in range(0, min(len(doc_text), max_total), max_chars_per_doc)]
                for idx, chunk in enumerate(chunks):
                    chunk_label = f'parte {idx+1}' if len(chunks) > 1 else 'conteúdo'
                    snippet = f'[WikiToca Doc] titulo: {doc_title} | {chunk_label} | {chunk.strip()[:3000]}'
                    items.append({
                        'table': 'wiki_documents',
                        'id': rd.get('id'),
                        'snippet': snippet,
                        'search_text': snippet.lower()
                    })
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar wiki_documents: {e}')

    return items


def _itoca_build_activities_items(conn):
    """Gera itens de atividades com JOIN em clients e commitments para contexto rico.
    Limita a 5 atividades mais recentes por empresa para evitar sobrecarga de tokens.
    Inclui o título do compromisso vinculado quando existir (via notes ou title match).
    """
    cursor = conn.cursor()
    items = []
    try:
        # Busca as 500 mais recentes; depois limita por empresa no Python
        cursor.execute('''
            SELECT ac.id, ac.contact_type, ac.information, ac.description, ac.activity_date,
                   cl.name as client_name, cl.company as client_company, cl.position as client_position,
                   cl.id as client_id
            FROM activities ac
            LEFT JOIN clients cl ON ac.client_id = cl.id
            ORDER BY ac.activity_date DESC
            LIMIT 500
        ''')
        rows = cursor.fetchall()

        # Conta por empresa para limitar a 5 mais recentes
        company_count = {}
        for row in rows:
            rd = dict_from_row(row)
            company = (rd.get('client_company') or '').strip().lower()
            if not company:
                company = '__sem_empresa__'
            if company_count.get(company, 0) >= 5:
                continue
            company_count[company] = company_count.get(company, 0) + 1

            parts = []
            if rd.get('client_name'):
                parts.append(f'contato: {rd["client_name"]}')
            if rd.get('client_company'):
                parts.append(f'empresa: {rd["client_company"]}')
            if rd.get('client_position'):
                parts.append(f'cargo: {rd["client_position"]}')
            if rd.get('contact_type'):
                parts.append(f'tipo: {rd["contact_type"]}')
            if rd.get('activity_date'):
                try:
                    dt = datetime.strptime(str(rd['activity_date'])[:10], '%Y-%m-%d')
                    parts.append(f'data: {dt.strftime("%d/%m/%Y")}')
                except Exception:
                    parts.append(f'data: {rd["activity_date"]}')
            if rd.get('information'):
                parts.append(f'informacao: {str(rd["information"])[:300]}')
            if rd.get('description') and rd['description'] != rd.get('information'):
                parts.append(f'descricao: {str(rd["description"])[:300]}')

            # Tenta encontrar compromisso vinculado pela data + contato
            # (commitments com mesmo client_id e due_date próxima da activity_date)
            if rd.get('client_id') and rd.get('activity_date'):
                try:
                    act_date = str(rd['activity_date'])[:10]
                    cursor.execute('''
                        SELECT title FROM commitments
                        WHERE client_id = ?
                          AND ABS(julianday(due_date) - julianday(?)) <= 7
                        ORDER BY ABS(julianday(due_date) - julianday(?)) ASC
                        LIMIT 1
                    ''', (rd['client_id'], act_date, act_date))
                    cm_row = cursor.fetchone()
                    if cm_row:
                        cm_title = (cm_row[0] if not isinstance(cm_row, sqlite3.Row) else cm_row['title'] if 'title' in cm_row.keys() else cm_row[0]) or ''
                        if cm_title:
                            parts.append(f'compromisso_vinculado: {cm_title[:120]}')
                except Exception as e:
                    logger.debug(f'[_itoca_build_activities_items] exceção ignorada: {e}')

            snippet = ' | '.join(parts)
            if snippet:
                items.append({'table': 'activities', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar activities: {e}')
    return items


def _itoca_build_presences_items(conn):
    """Gera itens de presenças em contas com JOIN em accounts e clients."""
    cursor = conn.cursor()
    items = []
    try:
        cursor.execute('''
            SELECT ap.id, ap.delivery_name, ap.stf_owner, ap.current_revenue_cents,
                   ap.validity_month,
                   ac.name as account_name, ac.sector as account_sector,
                   cl.name as focal_contact_name
            FROM account_presences ap
            LEFT JOIN accounts ac ON ap.account_id = ac.id
            LEFT JOIN clients cl ON ap.focal_client_id = cl.id
            ORDER BY ap.updated_at DESC
        ''')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('account_name'):
                parts.append(f'conta: {rd["account_name"]}')
            if rd.get('account_sector'):
                parts.append(f'setor: {rd["account_sector"]}')
            if rd.get('delivery_name'):
                parts.append(f'entrega: {rd["delivery_name"]}')
            if rd.get('stf_owner'):
                parts.append(f'responsavel_stf: {rd["stf_owner"]}')
            if rd.get('focal_contact_name'):
                parts.append(f'contato_focal: {rd["focal_contact_name"]}')
            if rd.get('current_revenue_cents'):
                try:
                    receita = int(rd['current_revenue_cents']) / 100
                    parts.append(f'receita_atual: R$ {receita:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'))
                except Exception as e:
                    logger.debug(f'[_itoca_build_presences_items] exceção ignorada: {e}')
            if rd.get('validity_month'):
                parts.append(f'validade: {rd["validity_month"]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({'table': 'account_presences', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar account_presences: {e}')
    return items


def _itoca_build_kanban_items(conn):
    """Gera itens do Kanban com JOIN em accounts, clients e kanban_columns."""
    cursor = conn.cursor()
    items = []
    try:
        cursor.execute('''
            SELECT kc.id, kc.title, kc.description, kc.tag, kc.urgency, kc.activity,
                   kc.updated_at,
                   col.title as column_title,
                   ac.name as account_name,
                   cl.name as contact_name
            FROM kanban_cards kc
            LEFT JOIN kanban_columns col ON kc.column_id = col.id
            LEFT JOIN accounts ac ON kc.account_id = ac.id
            LEFT JOIN clients cl ON kc.contact_id = cl.id
            ORDER BY kc.updated_at DESC
        ''')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'):
                parts.append(f'titulo: {rd["title"]}')
            if rd.get('column_title'):
                parts.append(f'coluna: {rd["column_title"]}')
            if rd.get('account_name'):
                parts.append(f'conta: {rd["account_name"]}')
            if rd.get('contact_name'):
                parts.append(f'contato: {rd["contact_name"]}')
            if rd.get('urgency'):
                parts.append(f'urgencia: {rd["urgency"]}')
            if rd.get('tag'):
                parts.append(f'tag: {rd["tag"]}')
            if rd.get('description'):
                parts.append(f'descricao: {str(rd["description"])[:300]}')
            if rd.get('activity'):
                parts.append(f'atividade: {str(rd["activity"])[:200]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({'table': 'kanban_cards', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar kanban_cards: {e}')
    return items


def _itoca_build_environment_items(conn):
    """Gera itens de mapeamento de ambiente com JOIN em environment_cards, clients e accounts."""
    cursor = conn.cursor()
    items = []
    try:
        cursor.execute('''
            SELECT er.id, er.response,
                   ec.title as card_title,
                   cl.name as client_name, cl.company as client_company
            FROM environment_responses er
            LEFT JOIN environment_cards ec ON er.card_id = ec.id
            LEFT JOIN clients cl ON er.client_id = cl.id
            WHERE er.response IS NOT NULL AND TRIM(er.response) != ""
            ORDER BY er.updated_at DESC
        ''')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('client_company'):
                parts.append(f'empresa: {rd["client_company"]}')
            if rd.get('client_name'):
                parts.append(f'contato: {rd["client_name"]}')
            if rd.get('card_title'):
                parts.append(f'mapeamento: {rd["card_title"]}')
            if rd.get('response'):
                parts.append(f'resposta: {str(rd["response"])[:400]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({'table': 'environment_responses', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar environment_responses: {e}')
    return items


def _itoca_build_user_profile_item(conn):
    """Gera item fixo com o perfil do usuário para sempre estar no contexto."""
    cursor = conn.cursor()
    try:
        cursor.execute('SELECT full_name, nickname, position, email, phone, boss_name, boss_email FROM user_profile WHERE id = 1')
        row = cursor.fetchone()
        if not row:
            return None
        rd = dict_from_row(row)
        parts = []
        if rd.get('full_name'):
            parts.append(f'nome: {rd["full_name"]}')
        if rd.get('nickname'):
            parts.append(f'apelido: {rd["nickname"]}')
        if rd.get('position'):
            parts.append(f'cargo: {rd["position"]}')
        if rd.get('email'):
            parts.append(f'email: {rd["email"]}')
        if rd.get('phone'):
            parts.append(f'telefone: {rd["phone"]}')
        if rd.get('boss_name'):
            parts.append(f'gestor: {rd["boss_name"]}')
        if rd.get('boss_email'):
            parts.append(f'email_gestor: {rd["boss_email"]}')
        snippet = ' | '.join(parts)
        if snippet:
            return {'table': 'user_profile', 'id': 1, 'snippet': snippet, 'search_text': snippet.lower()}
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar user_profile: {e}')
    return None


def _itoca_build_message_templates_items(conn):
    """Gera itens dos templates de mensagem para consulta pelo iToca."""
    cursor = conn.cursor()
    items = []
    try:
        cursor.execute('SELECT id, title, description FROM message_templates ORDER BY title')
        for row in cursor.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'):
                parts.append(f'titulo: {rd["title"]}')
            if rd.get('description'):
                parts.append(f'conteudo: {str(rd["description"])[:600]}')
            snippet = ' | '.join(parts)
            if snippet:
                items.append({'table': 'message_templates', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar message_templates: {e}')
    return items


def _itoca_build_base_snapshot(max_tables=120, max_rows_per_table=250, max_items=6000, progress_cb=None):
    """Constrói o snapshot RAG. progress_cb(percent, message) é chamado durante o processo."""
    def _progress(pct, msg):
        if progress_cb:
            progress_cb(pct, msg)

    # Tabelas que têm funções especializadas e não devem ser varridas genericamente
    _SPECIALIZED_TABLES = {
        'activities', 'commitments', 'account_renewal_events', 'account_presences',
        'kanban_cards', 'environment_responses', 'wiki_entries', 'wiki_documents',
        'user_profile', 'message_templates'
    }

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")
    tables = [
        row['name'] for row in cursor.fetchall()
        if row['name'] not in ITOCA_EXCLUDED_TABLES and row['name'] not in _SPECIALIZED_TABLES
    ][:max_tables]

    _progress(5, f'Lendo {len(tables)} tabelas genéricas do banco de dados...')
    items = []
    total_tables = max(len(tables), 1)
    for t_idx, table in enumerate(tables):
        text_columns = _itoca_text_columns(cursor, table)
        if not text_columns:
            continue
        try:
            cursor.execute(f'SELECT * FROM "{table}" LIMIT {int(max_rows_per_table)}')
            rows = cursor.fetchall()
        except Exception:
            continue

        for row in rows:
            row_dict = dict_from_row(row)
            snippet = _itoca_build_snippet(row_dict)
            if not snippet:
                continue
            items.append({
                'table': table,
                'id': row_dict.get('id'),
                'snippet': snippet,
                'search_text': snippet.lower()
            })
            if len(items) >= max_items:
                conn.close()
                _progress(100, f'Concluído: {len(items)} registros indexados.')
                return items

        pct_tables = 5 + int((t_idx + 1) / total_tables * 20)  # 5% a 25%
        _progress(pct_tables, f'Tabela {table} indexada ({t_idx+1}/{total_tables})...')

    # Perfil do usuário (item fixo)
    _progress(26, 'Indexando perfil do usuário...')
    profile_item = _itoca_build_user_profile_item(conn)
    if profile_item:
        items.append(profile_item)

    # Atividades com JOIN enriquecido
    _progress(30, 'Indexando atividades com contatos...')
    activity_items = _itoca_build_activities_items(conn)
    items.extend(activity_items)
    _progress(38, f'{len(activity_items)} atividades indexadas.')

    # Agenda
    _progress(40, 'Indexando agenda e compromissos...')
    agenda_items = _itoca_build_agenda_items(conn)
    items.extend(agenda_items)
    _progress(48, f'{len(agenda_items)} eventos de agenda indexados.')

    # Presenças em contas com JOIN
    _progress(50, 'Indexando presenças em contas...')
    presences_items = _itoca_build_presences_items(conn)
    items.extend(presences_items)
    _progress(55, f'{len(presences_items)} presenças indexadas.')

    # Kanban com JOIN
    _progress(57, 'Indexando cards do Kanban...')
    kanban_items = _itoca_build_kanban_items(conn)
    items.extend(kanban_items)
    _progress(62, f'{len(kanban_items)} cards do Kanban indexados.')

    # Mapeamento de ambiente com JOIN
    _progress(64, 'Indexando mapeamento de ambiente...')
    env_items = _itoca_build_environment_items(conn)
    items.extend(env_items)
    _progress(68, f'{len(env_items)} respostas de mapeamento indexadas.')

    # Templates de mensagem
    _progress(69, 'Indexando templates de mensagem...')
    template_items = _itoca_build_message_templates_items(conn)
    items.extend(template_items)
    _progress(70, f'{len(template_items)} templates indexados.')

    # WikiToca — mais demorado por causa da extração de PDF
    _progress(72, 'Indexando WikiToca (entradas de conhecimento)...')
    conn2 = get_db()
    cursor2 = conn2.cursor()
    wiki_entry_items = []
    try:
        cursor2.execute('SELECT id, title, category, content, tags FROM wiki_entries ORDER BY updated_at DESC')
        for row in cursor2.fetchall():
            rd = dict_from_row(row)
            parts = []
            if rd.get('title'): parts.append(f'titulo: {rd["title"]}')
            if rd.get('category'): parts.append(f'categoria: {rd["category"]}')
            if rd.get('tags'): parts.append(f'tags: {rd["tags"]}')
            if rd.get('content'): parts.append(f'conteudo: {rd["content"][:2000]}')
            snippet = ' | '.join(parts)
            if snippet:
                wiki_entry_items.append({'table': 'wiki_entries', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar wiki_entries: {e}')
    items.extend(wiki_entry_items)
    _progress(78, f'{len(wiki_entry_items)} entradas WikiToca indexadas. Processando documentos...')

    # Documentos WikiToca com extração de texto (pode ser lento por OCR)
    wiki_doc_items = []
    try:
        cursor2.execute('SELECT id, title, original_name, file_name, file_ext FROM wiki_documents ORDER BY updated_at DESC')
        all_docs = cursor2.fetchall()
        total_docs = max(len(all_docs), 1)
        for d_idx, row in enumerate(all_docs):
            rd = dict_from_row(row)
            doc_title = rd.get('title') or rd.get('original_name', '')
            doc_ext = rd.get('file_ext', '')
            file_path = WIKI_UPLOAD_DIR / (rd.get('file_name') or '')
            doc_text = ''
            pct_docs = 78 + int((d_idx + 1) / total_docs * 18)  # 78% a 96%
            _progress(pct_docs, f'Processando documento {d_idx+1}/{total_docs}: {doc_title[:40]}...')
            if file_path.exists():
                doc_text = _itoca_extract_text_from_file(str(file_path))
            if not doc_text.strip():
                snippet = (f'[WikiToca Doc] titulo: {doc_title} | tipo: {doc_ext} | '
                           f'CONFIRMADO: o documento "{rd.get("original_name", "")}" EXISTE no WikiToca. '
                           f'Conteúdo interno não disponível para leitura automática. '
                           f'Ao responder sobre este tema, informe que o documento existe mas não é possível detalhar o conteúdo.')
                _search_base = (doc_title + ' ' + rd.get('original_name', '') + ' ' + doc_ext).lower()
                wiki_doc_items.append({'table': 'wiki_documents', 'id': rd.get('id'), 'snippet': snippet, 'search_text': _search_base})
            else:
                max_total = 8000 * 5
                chunks = [doc_text[i:i+8000] for i in range(0, min(len(doc_text), max_total), 8000)]
                for idx, chunk in enumerate(chunks):
                    chunk_label = f'parte {idx+1}' if len(chunks) > 1 else 'conteúdo'
                    snippet = f'[WikiToca Doc] titulo: {doc_title} | {chunk_label} | {chunk.strip()[:3000]}'
                    wiki_doc_items.append({'table': 'wiki_documents', 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
    except Exception as e:
        logger.warning(f'[iToca] Erro ao indexar wiki_documents: {e}')
    conn2.close()
    items.extend(wiki_doc_items)

    conn.close()
    _progress(100, f'Concluído! {len(items)} registros indexados no total.')
    return items


def _itoca_update_cached_base(progress_cb=None, incremental=False):
    """Atualiza o snapshot RAG.
    Se incremental=True, re-indexa apenas registros modificados após o último update.
    Se incremental=False (padrão), faz a indexação completa.
    """
    def _progress(pct, msg):
        if progress_cb:
            progress_cb(pct, msg)

    if incremental:
        # Busca o timestamp do último update
        settings_map = _load_app_settings_map(['itoca_base_updated_at', 'itoca_base_snapshot'])
        last_updated = (settings_map.get('itoca_base_updated_at') or '').strip()
        raw_snapshot = (settings_map.get('itoca_base_snapshot') or '').strip()

        if not last_updated or not raw_snapshot:
            # Nunca foi indexado: faz completo
            _progress(0, 'Primeira indexação: executando indexação completa...')
            return _itoca_update_cached_base(progress_cb=progress_cb, incremental=False)

        try:
            existing_items = json.loads(raw_snapshot)
        except Exception:
            existing_items = []

        _progress(5, f'Indexação incremental desde {last_updated}...')
        conn = get_db()
        cursor = conn.cursor()
        new_items = []
        seen_keys = set()

        # Tabelas com coluna updated_at ou created_at para filtro incremental
        _INCREMENTAL_TABLES = [
            'clients', 'accounts', 'account_presences', 'account_main_contacts',
            'kanban_cards', 'kanban_card_activities', 'daily_suggestions',
            'environment_cards', 'account_sectors'
        ]
        _INCREMENTAL_SPECIALIZED = [
            ('activities', '_itoca_build_activities_items'),
            ('commitments', '_itoca_build_agenda_items'),
            ('account_renewal_events', '_itoca_build_agenda_items'),
            ('account_presences', '_itoca_build_presences_items'),
            ('kanban_cards', '_itoca_build_kanban_items'),
            ('environment_responses', '_itoca_build_environment_items'),
            ('wiki_entries', None),
            ('wiki_documents', None),
            ('message_templates', '_itoca_build_message_templates_items'),
        ]

        # Verifica se houve alterações em tabelas genéricas
        generic_changed = False
        for table in _INCREMENTAL_TABLES:
            try:
                cursor.execute(
                    f'SELECT COUNT(*) as cnt FROM "{table}" WHERE updated_at > ? OR created_at > ?',
                    (last_updated, last_updated)
                )
                row = cursor.fetchone()
                if row and (dict_from_row(row).get('cnt') or 0) > 0:
                    generic_changed = True
                    break
            except Exception as e:
                logger.debug(f'[_progress] exceção ignorada: {e}')

        # Verifica se houve alterações nas tabelas especializadas
        specialized_changed_tables = set()
        for table, _ in _INCREMENTAL_SPECIALIZED:
            try:
                col = 'updated_at'
                cursor.execute(f"PRAGMA table_info('{table}')")
                cols = [r['name'] for r in cursor.fetchall()]
                if 'updated_at' not in cols:
                    col = 'created_at'
                if col not in cols:
                    specialized_changed_tables.add(table)
                    continue
                cursor.execute(
                    f'SELECT COUNT(*) as cnt FROM "{table}" WHERE {col} > ?',
                    (last_updated,)
                )
                row = cursor.fetchone()
                if row and (dict_from_row(row).get('cnt') or 0) > 0:
                    specialized_changed_tables.add(table)
            except Exception:
                specialized_changed_tables.add(table)

        # Se nada mudou, retorna o snapshot existente sem reprocessar
        if not generic_changed and not specialized_changed_tables:
            _progress(100, f'Nenhuma alteração detectada. Base já está atualizada ({len(existing_items)} registros).')
            conn.close()
            updated_at = datetime.now().isoformat(timespec='seconds')
            c = get_db()
            cc = c.cursor()
            cc.execute('UPDATE app_settings SET value = ?, updated_at = CURRENT_TIMESTAMP WHERE key = ?', (updated_at, 'itoca_base_updated_at'))
            c.commit()
            c.close()
            return {'items': existing_items, 'updated_at': updated_at, 'incremental': True, 'changed': False}

        # Remove do snapshot existente os registros das tabelas que mudaram
        tables_to_refresh = set()
        if generic_changed:
            tables_to_refresh.update(_INCREMENTAL_TABLES)
        tables_to_refresh.update(specialized_changed_tables)

        # Sempre re-indexa user_profile (item fixo, muito leve)
        tables_to_refresh.add('user_profile')

        _progress(10, f'Atualizando {len(tables_to_refresh)} tabelas modificadas...')

        # Mantém registros de tabelas não alteradas
        kept_items = [item for item in existing_items if item.get('table') not in tables_to_refresh]
        seen_keys = {f"{item.get('table')}:{item.get('id')}" for item in kept_items}

        # Re-indexa tabelas genéricas alteradas
        if generic_changed:
            for table in _INCREMENTAL_TABLES:
                if table not in tables_to_refresh:
                    continue
                text_columns = _itoca_text_columns(cursor, table)
                if not text_columns:
                    continue
                try:
                    cursor.execute(f'SELECT * FROM "{table}" LIMIT 250')
                    for row in cursor.fetchall():
                        rd = dict_from_row(row)
                        snippet = _itoca_build_snippet(rd)
                        if snippet:
                            kept_items.append({'table': table, 'id': rd.get('id'), 'snippet': snippet, 'search_text': snippet.lower()})
                except Exception as e:
                    logger.debug(f'[_progress] exceção ignorada: {e}')

        # Re-indexa tabelas especializadas alteradas
        if 'user_profile' in tables_to_refresh:
            p = _itoca_build_user_profile_item(conn)
            if p:
                kept_items.append(p)
        if 'activities' in tables_to_refresh:
            kept_items.extend(_itoca_build_activities_items(conn))
        if 'commitments' in tables_to_refresh or 'account_renewal_events' in tables_to_refresh:
            kept_items.extend(_itoca_build_agenda_items(conn))
        if 'account_presences' in tables_to_refresh:
            kept_items.extend(_itoca_build_presences_items(conn))
        if 'kanban_cards' in tables_to_refresh:
            kept_items.extend(_itoca_build_kanban_items(conn))
        if 'environment_responses' in tables_to_refresh:
            kept_items.extend(_itoca_build_environment_items(conn))
        if 'message_templates' in tables_to_refresh:
            kept_items.extend(_itoca_build_message_templates_items(conn))

        # Re-indexa wiki se necessário
        if 'wiki_entries' in tables_to_refresh or 'wiki_documents' in tables_to_refresh:
            _progress(60, 'Atualizando WikiToca...')
            kept_items.extend(_itoca_build_wiki_items(conn))

        conn.close()
        snapshot_items = kept_items
        _progress(95, f'Incremental concluído: {len(snapshot_items)} registros na base.')

    else:
        # Indexação completa
        snapshot_items = _itoca_build_base_snapshot(progress_cb=progress_cb)

    updated_at = datetime.now().isoformat(timespec='seconds')
    conn = get_db()
    c = conn.cursor()
    c.execute('UPDATE app_settings SET value = ?, updated_at = CURRENT_TIMESTAMP WHERE key = ?', (json.dumps(snapshot_items, ensure_ascii=False), 'itoca_base_snapshot'))
    c.execute('UPDATE app_settings SET value = ?, updated_at = CURRENT_TIMESTAMP WHERE key = ?', (updated_at, 'itoca_base_updated_at'))
    conn.commit()
    conn.close()
    _progress(100, f'Base atualizada: {len(snapshot_items)} registros indexados.')
    return {'items': snapshot_items, 'updated_at': updated_at, 'incremental': incremental}


def _itoca_get_cached_base():
    settings_map = _load_app_settings_map(['itoca_base_snapshot', 'itoca_base_updated_at'])
    raw_snapshot = (settings_map.get('itoca_base_snapshot') or '').strip()
    updated_at = (settings_map.get('itoca_base_updated_at') or '').strip()
    if not raw_snapshot:
        return [], updated_at
    try:
        snapshot = json.loads(raw_snapshot)
    except Exception:
        snapshot = []
    return snapshot if isinstance(snapshot, list) else [], updated_at


def _itoca_search_in_cached_snapshot(question, snapshot_items, limit=18):
    tokens = _itoca_tokenize(question)
    if not tokens or not snapshot_items:
        return []

    scored = []
    for item in snapshot_items:
        haystack = str(item.get('search_text') or item.get('snippet') or '').lower()
        if not haystack:
            continue
        score = sum(1 for token in tokens if token in haystack)
        if score <= 0:
            continue
        scored.append((score, item))

    scored.sort(key=lambda x: (-x[0], str(x[1].get('table') or ''), str(x[1].get('id') or '')))
    return [item for _, item in scored[:max(1, int(limit or 18))]]


# Rótulos legíveis por tabela para o contexto do LLM
_ITOCA_TABLE_LABELS = {
    'clients': 'Contato/Cliente',
    'accounts': 'Conta/Empresa',
    'activities': 'Atividade/Interação com Contato',
    'commitments': 'Compromisso/Evento de Agenda',
    'account_renewal_events': 'Evento de Renovação de Contrato',
    'account_presences': 'Presença/Entrega em Conta',
    'account_main_contacts': 'Contato Principal de Conta',
    'wiki_entries': 'Conhecimento do WikiToca',
    'wiki_documents': 'Documento do WikiToca',
    'environment_cards': 'Card de Mapeamento de Ambiente',
    'environment_responses': 'Resposta de Mapeamento de Ambiente',
    'kanban_columns': 'Coluna do Kanban',
    'kanban_cards': 'Card do Kanban',
    'daily_suggestions': 'Sugestão Diária',
    'user_profile': 'Perfil do Usuário',
    'message_templates': 'Template de Mensagem',
    'account_sectors': 'Setor de Conta',
    'job_groupings': 'Agrupamento de Cargos',
}


def infer_kanban_tag(description):
    text = (description or '').strip().lower()
    if not text:
        return ''

    tag_rules = [
        ('Comercial', ['proposta', 'orcamento', 'negoci', 'cliente', 'venda', 'pipeline', 'lead']),
        ('Técnico', ['api', 'erro', 'sistema', 'infra', 'deploy', 'integracao', 'bug', 'arquitetura']),
        ('Financeiro', ['fatur', 'pagamento', 'custo', 'contrato', 'budget', 'receita']),
        ('Relacionamento', ['reuniao', 'follow-up', 'contato', 'alinhamento', 'feedback', 'stakeholder']),
        ('Produto', ['produto', 'feature', 'roadmap', 'release', 'sprint']),
        ('Prioridade Alta', ['urgente', 'critico', 'hoje', 'bloqueio'])
    ]

    found_tags = []
    for tag, keywords in tag_rules:
        if any(keyword in text for keyword in keywords):
            found_tags.append(tag)

    if found_tags:
        return ' | '.join(found_tags)

    tokens = re.findall(r"[a-zA-ZÀ-ÿ]{3,}", text)
    stopwords = {
        'para', 'com', 'sobre', 'entre', 'pelos', 'pelas', 'isso', 'essa', 'esse', 'dele', 'dela',
        'card', 'tarefa', 'atividade', 'fazer', 'sera', 'será', 'esta', 'está', 'mais', 'muito', 'pouco',
        'uma', 'uns', 'das', 'dos', 'que', 'por', 'sem'
    }
    keywords = []
    for token in tokens:
        tk = token.strip().lower()
        if tk in stopwords:
            continue
        label = tk.capitalize()
        if label not in keywords:
            keywords.append(label)
        if len(keywords) >= 4:
            break
    if keywords:
        return ' | '.join(keywords)

    raw_tokens = [t.strip().capitalize() for t in re.findall(r"[a-zA-ZÀ-ÿ]{2,}", text) if t.strip()]
    raw_tokens = [t for t in raw_tokens if t.lower() not in stopwords][:3]
    return ' | '.join(raw_tokens) if raw_tokens else ''

def _itoca_compose_context_text(context_rows, history_rows=None):
    """Formata as linhas de contexto em texto estruturado e semântico para envio ao LLM.
    Inclui: data/hora atual, instruções de formatação, histórico de sessão e registros agrupados por categoria.
    """
    from datetime import datetime as _dt
    import locale as _locale

    # Data e hora atual em português
    now = _dt.now()
    dias_semana = ['segunda-feira', 'terça-feira', 'quarta-feira', 'quinta-feira', 'sexta-feira', 'sábado', 'domingo']
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    dia_semana = dias_semana[now.weekday()]
    data_str = f'{dia_semana}, {now.day} de {meses[now.month - 1]} de {now.year}'
    hora_str = now.strftime('%H:%M')

    lines = []
    lines.append('=== CONTEXTO INTERNO DO SISTEMA TOCA DO COELHO ===')
    lines.append(f'DATA E HORA ATUAL: {data_str}, {hora_str}')
    lines.append('')
    lines.append('INSTRUÇÕES DE FORMATAÇÃO E PROFUNDIDADE DA RESPOSTA:')
    lines.append('- Use linguagem natural, profissional e direta. Evite termos técnicos do banco de dados.')
    lines.append('- Não repita os rótulos internos (ex: "titulo:", "conta:", "nome_contato:", "classificacao:") na resposta final; traduza para linguagem natural.')
    lines.append('- O rótulo "classificacao: conta-alvo (target)" significa que a conta/contato é prioritário (marcado como Target no sistema).')
    lines.append('- O rótulo "classificacao: cold contact" significa que o contato está em prospecção fria.')
    lines.append('- O rótulo "servicos_stefanini_cadastrados" indica quantos serviços/entregas Stefanini estão mapeados naquela conta.')
    lines.append('- O rótulo "PAINEL_GERAL" contém estatísticas globais do sistema — use para contextualizar totais e proporções.')
    lines.append('- Para listas com 3 ou mais itens, use tópicos com "•" (bullet) ou números.')
    lines.append('- TABELAS: use tabela Markdown APENAS para comparações diretas de atributos similares (ex: lista de contatos com cargo e e-mail). Para análise de relacionamento, use seções por conta com bullets — NÃO use tabela plana que mistura contas com contatos em linhas.')
    lines.append('- Datas devem ser apresentadas no formato DD/MM/AAAA. Horários no formato HH:MM.')
    lines.append('- Valores monetários devem usar o formato R$ X.XXX,XX.')
    lines.append('- Ao citar um compromisso, sempre informe: título, data, horário e contato/empresa envolvidos.')
    lines.append('- Se a pergunta envolver "próximos eventos" ou "esta semana", use a DATA ATUAL acima como referência.')
    lines.append('- ANÁLISE DE RELACIONAMENTO: quando perguntado sobre relacionamento com contas target ou múltiplas contas, organize a resposta COM UMA SEÇÃO POR CONTA, mostrando para cada uma:')
    lines.append('  * Nome da conta e se é target ou não')
    lines.append('  * Último contato realizado (data e tipo)')
    lines.append('  * Total de interações registradas')
    lines.append('  * Se possui serviços/entregas Stefanini cadastrados (quantos)')
    lines.append('  * Contatos/pessoas mapeadas (se disponível)')
    lines.append('  * Avaliação resumida: "evoluindo bem" se há atividades recentes, "só tentativas" se há poucas atividades, "sem histórico" se não há interações')
    lines.append('- NOMES: sempre use o nome da pessoa e da empresa, nunca IDs numéricos.')
    lines.append('- COMPLETUDE: se houver dados de múltiplas contas target, apresente TODAS, não apenas as primeiras.')
    lines.append('')

    # Histórico da sessão (se fornecido)
    if history_rows:
        lines.append('=== HISTÓRICO DA CONVERSA ATUAL ===')
        for h in history_rows[-6:]:  # últimas 3 trocas (user + assistant)
            role_label = 'Usuário' if h.get('role') == 'user' else 'iToca'
            lines.append(f'[{role_label}]: {str(h.get("content", ""))[:400]}')
        lines.append('')

    if not context_rows:
        lines.append('(Nenhum dado encontrado na base interna para esta pergunta.)')
        return '\n'.join(lines)

    lines.append('=== REGISTROS ENCONTRADOS NA BASE INTERNA ===')
    lines.append('(Use SOMENTE estas informações. CNPJ = Cadastro Nacional de Pessoa Jurídica. CPF = Cadastro de Pessoa Física.)')
    lines.append('')

    # Agrupa registros por categoria para facilitar a leitura do LLM
    from collections import defaultdict as _dd
    grouped = _dd(list)
    for item in context_rows[:25]:
        table = item.get('table', 'outros')
        grouped[table].append(item)

    # Ordem de apresentação das categorias (mais relevantes primeiro)
    _CATEGORY_ORDER = [
        'user_profile', 'commitments', 'account_renewal_events',
        'activities', 'clients', 'accounts', 'account_presences',
        'account_main_contacts', 'kanban_cards', 'environment_responses',
        'wiki_entries', 'wiki_documents', 'message_templates', 'daily_suggestions'
    ]
    ordered_tables = _CATEGORY_ORDER + [t for t in grouped if t not in _CATEGORY_ORDER]

    reg_num = 1
    for table in ordered_tables:
        if table not in grouped:
            continue
        label = _ITOCA_TABLE_LABELS.get(table, table)
        lines.append(f'--- [{label}] ---')
        for item in grouped[table]:
            row_id = f" #{item['id']}" if item.get('id') is not None else ''
            lines.append(f'  Registro {reg_num}{row_id}: {item["snippet"]}')
            reg_num += 1
        lines.append('')

    return '\n'.join(lines)


def _itoca_call_sai_llm(question, context_rows, history_rows=None):
    """Chama a API SAI LLM com a pergunta e o contexto. Retorna dict com answer, confidence_percent, needs_refinement, refinement_hint."""
    settings_map = _load_app_settings_map(['itoca_sai_api_key', 'itoca_sai_template_id', 'itoca_sai_base_url'])
    api_key = (settings_map.get('itoca_sai_api_key') or '').strip() or (os.environ.get('ITOCA_SAI_API_KEY', '') or '').strip()
    template_id = (settings_map.get('itoca_sai_template_id') or '').strip() or '69ac3c87024adc2d2bdc19f5'
    base_url = (settings_map.get('itoca_sai_base_url') or '').strip() or 'https://sai-library.saiapplications.com'

    context_text = _itoca_compose_context_text(context_rows, history_rows=history_rows)

    if not api_key:
        # Fallback sem LLM: retorna resposta estruturada básica
        if not context_rows:
            return {
                'answer': 'Não encontrei dados suficientes na base interna para responder com segurança. Tente perguntar com mais detalhes (nome, empresa, cargo, data ou módulo).',
                'confidence_percent': 0,
                'needs_refinement': True,
                'refinement_hint': 'Forneça mais detalhes como nome, empresa, cargo ou data.',
                'llm_used': False
            }
        answer_lines = [f'Encontrei {len(context_rows)} registro(s) relevante(s):']
        for item in context_rows[:8]:
            row_id = f"#{item['id']}" if item.get('id') is not None else ''
            answer_lines.append(f"• [{item['table']}{row_id}] {item['snippet']}")
        answer_lines.append('')
        answer_lines.append('Configure a chave da API SAI em Configurações > Integrações para respostas em linguagem natural.')
        return {
            'answer': '\n'.join(answer_lines),
            'confidence_percent': 50,
            'needs_refinement': False,
            'refinement_hint': '',
            'llm_used': False
        }

    url = f'{base_url}/api/templates/{template_id}/execute'
    headers = {
        'Content-Type': 'application/json',
        'X-Api-Key': api_key
    }
    payload = {
        'inputs': {
            'question': question,
            'context_sources': context_text if context_text else 'Nenhum dado encontrado na base interna para esta pergunta.'
        }
    }
    try:
        resp = requests.post(url, json=payload, headers={'X-Api-Key': api_key}, timeout=60)
        if not resp.ok:
            logger.error(f'[iToca][SAI] HTTPError {resp.status_code}: {resp.text[:400]}')
            raise RuntimeError(f'Erro na API SAI (HTTP {resp.status_code}): {resp.text[:200]}')
        logger.info(f'[iToca][SAI] OK ({len(resp.text)} chars)')
        raw = resp.text
    except RuntimeError:
        raise
    except Exception as e:
        logger.error(f'[iToca][SAI] Erro de conexão: {e}')
        raise RuntimeError(f'Falha ao conectar com a API SAI: {str(e)}')

    # ─── Parsing robusto da resposta da API SAI ───────────────────────────────────────────────
    # A API SAI pode retornar o JSON do LLM de várias formas:
    #   1. JSON direto: {"answer": "...", "confidence_percent": 80, ...}
    #   2. Wrapper com "output": {"output": "{\"answer\":\"...\"}"}
    #   3. Wrapper com "result": {"result": "{\"answer\":\"...\"}"}
    #   4. Wrapper com "text": {"text": "{\"answer\":\"...\"}"}
    #   5. O campo answer já contém um JSON serializado como string
    # ───────────────────────────────────────────────────────────────────────
    logger.debug(f'[iToca][SAI] raw response (primeiros 500 chars): {raw[:500]}')

    def _try_parse_llm_json(text):
        """Tenta extrair o dict com 'answer' de uma string que pode conter JSON em vários níveis."""
        if not text:
            return None
        # Tentativa 1: parse direto
        try:
            obj = json.loads(text.strip())
            if isinstance(obj, dict):
                return obj
        except Exception as e:
            logger.debug(f'[_try_parse_llm_json] exceção ignorada: {e}')
        # Tentativa 2: extrai o primeiro objeto JSON da string
        obj = _extract_json_object_from_text(text)
        if obj and isinstance(obj, dict):
            return obj
        return None

    # Passo 1: parse do envelope externo retornado pela API
    outer = _try_parse_llm_json(raw)

    # Passo 2: se o envelope externo não tem 'answer', procura em campos comuns de wrapper
    inner = None
    if outer and isinstance(outer, dict):
        if 'answer' in outer:
            # Caso ideal: o JSON do LLM já está no nível raiz
            inner = outer
        else:
            # Tenta extrair de campos de wrapper conhecidos
            for wrapper_key in ('output', 'result', 'text', 'content', 'response', 'data', 'message'):
                candidate = outer.get(wrapper_key)
                if candidate:
                    if isinstance(candidate, dict) and 'answer' in candidate:
                        inner = candidate
                        break
                    if isinstance(candidate, str):
                        parsed_inner = _try_parse_llm_json(candidate)
                        if parsed_inner and isinstance(parsed_inner, dict) and 'answer' in parsed_inner:
                            inner = parsed_inner
                            break
            # Se ainda não achou, tenta buscar qualquer JSON com 'answer' na string bruta
            if not inner:
                inner = _try_parse_llm_json(raw)
    else:
        # Não conseguiu parsear o envelope; tenta direto na string bruta
        inner = _try_parse_llm_json(raw)

    # Passo 3: verifica se o campo 'answer' em si é um JSON serializado (double-encoding)
    if inner and isinstance(inner, dict):
        raw_answer = inner.get('answer') or ''
        if isinstance(raw_answer, str) and raw_answer.strip().startswith('{'):
            double_parsed = _try_parse_llm_json(raw_answer)
            if double_parsed and isinstance(double_parsed, dict) and 'answer' in double_parsed:
                logger.debug('[iToca][SAI] Detectado double-encoding no campo answer — desencapsulando.')
                inner = double_parsed

    if not inner or not isinstance(inner, dict):
        # Fallback final: tenta extrair o campo 'answer' mesmo de um JSON truncado
        fallback_text = None
        if '"answer"' in raw:
            # Primeiro tenta parse completo
            extracted = _try_parse_llm_json(raw)
            if extracted and isinstance(extracted, dict) and 'answer' in extracted:
                fallback_text = (extracted.get('answer') or '').strip()
            else:
                # JSON truncado: extrai o texto do 'answer' até onde chegou
                partial = _extract_answer_from_partial_json(raw)
                if partial:
                    logger.warning('[iToca][SAI] JSON truncado detectado — usando extração parcial do campo answer.')
                    fallback_text = partial
        if not fallback_text:
            fallback_text = raw.strip()
        # Converte \n literais em quebras de linha
        fallback_text = fallback_text.replace('\\n', '\n')
        return {
            'answer': fallback_text or 'Não foi possível interpretar a resposta do assistente.',
            'confidence_percent': 50,
            'needs_refinement': False,
            'refinement_hint': '',
            'llm_used': True
        }

    # Passo 4: limpa o campo answer de \n literais (sequencia de barra-n) que o LLM pode gerar
    answer_text = (inner.get('answer') or '').strip()
    # Converte \n literais (como string de 2 chars) em quebras de linha reais
    answer_text = answer_text.replace('\\n', '\n')
    # Remove aspas externas se o LLM retornou a string entre aspas
    if answer_text.startswith('"') and answer_text.endswith('"') and len(answer_text) > 2:
        try:
            answer_text = json.loads(answer_text)
        except Exception as e:
            logger.debug(f'[_try_parse_llm_json] exceção ignorada: {e}')

    if not answer_text:
        answer_text = 'Sem resposta disponível.'

    return {
        'answer': answer_text,
        'confidence_percent': max(0, min(100, int(inner.get('confidence_percent') or 0))),
        'needs_refinement': bool(inner.get('needs_refinement')),
        'refinement_hint': (inner.get('refinement_hint') or '').strip(),
        'llm_used': True
    }


# ─────────────────────────────────────────────────────────────────────────────
# iToca — Detector de Intenção de Ação (segundo LLM SAI)
# ─────────────────────────────────────────────────────────────────────────────

# Tipos de ação reconhecidos e seus rótulos legíveis
_ITOCA_ACTION_LABELS = {
    'kanban_card':          'Criar card no Kanban',
    'activity':             'Registrar atividade com contato',
    'new_contact':          'Adicionar novo contato',
    'environment_mapping':  'Registrar mapeamento de ambiente',
    'wiki_entry':           'Salvar conhecimento no WikiToca',
    'commitment':           'Agendar compromisso',
}


def _itoca_detect_action_intent(question: str, answer: str) -> dict:
    """Chama o segundo LLM SAI (template 69b1c662485ca1e93db65015) para detectar
    se a mensagem do usuário expressa intenção de criar um registro no sistema.

    Retorna um dict com:
        action_type  (str | None)  — tipo da ação detectada ou None
        confidence   (float)       — 0.0 a 1.0
        label        (str)         — rótulo legível da ação
        fields       (dict)        — campos extraídos pelo LLM
        raw          (str)         — resposta bruta da API (para debug)
    """
    try:
        settings_map = _load_app_settings_map([
            'itoca_sai_api_key',
            'itoca_sai_base_url',
            'itoca_action_detector_template_id'
        ])
        api_key = (settings_map.get('itoca_sai_api_key') or '').strip() \
                  or (os.environ.get('ITOCA_SAI_API_KEY', '') or '').strip()
        base_url = (settings_map.get('itoca_sai_base_url') or '').strip() \
                   or 'https://sai-library.saiapplications.com'
        template_id = (settings_map.get('itoca_action_detector_template_id') or '').strip() \
                      or '69b1c662485ca1e93db65015'

        if not api_key:
            logger.debug('[iToca][ActionDetector] API key não configurada — detector desativado.')
            return {'action_type': None, 'confidence': 0.0, 'label': '', 'fields': {}, 'raw': ''}

        url = f'{base_url}/api/templates/{template_id}/execute'
        headers = {
            'Content-Type': 'application/json',
            'X-Api-Key': api_key
        }
        payload = {
            'inputs': {
                'question': question[:1000],   # limita para não inflar o payload
                'answer': answer[:1000]
            }
        }

        resp = requests.post(url, json=payload, headers={'X-Api-Key': api_key}, timeout=30)
        raw = resp.text
        logger.info(f'[iToca][ActionDetector] OK ({len(raw)} chars)')

        parsed = _extract_json_object_from_text(raw)
        if not parsed or not isinstance(parsed, dict):
            logger.warning(f'[iToca][ActionDetector] Resposta não-JSON: {raw[:200]}')
            return {'action_type': None, 'confidence': 0.0, 'label': '', 'fields': {}, 'raw': raw}

        action_type = (parsed.get('action_type') or '').strip().lower()
        if action_type in ('none', '', 'null'):
            action_type = None

        # Valida que o tipo retornado é um dos reconhecidos
        if action_type and action_type not in _ITOCA_ACTION_LABELS:
            logger.warning(f'[iToca][ActionDetector] Tipo desconhecido: {action_type!r}')
            action_type = None

        confidence = float(parsed.get('confidence') or 0.0)
        fields = parsed.get('fields') or {}
        if not isinstance(fields, dict):
            fields = {}

        label = _ITOCA_ACTION_LABELS.get(action_type, '') if action_type else ''

        logger.info(f'[iToca][ActionDetector] action={action_type!r} confidence={confidence:.2f} fields={list(fields.keys())}')
        return {
            'action_type': action_type,
            'confidence': confidence,
            'label': label,
            'fields': fields,
            'raw': raw
        }

    except Exception as e:
        logger.warning(f'[iToca][ActionDetector] Erro ao detectar intenção: {e}')
        return {'action_type': None, 'confidence': 0.0, 'label': '', 'fields': {}, 'raw': ''}


def ensure_account_for_company(cursor, company_name):
    name = (company_name or '').strip()
    if not name:
        return None
    cursor.execute('SELECT id FROM accounts WHERE LOWER(TRIM(name)) = LOWER(TRIM(?))', (name,))
    row = cursor.fetchone()
    if row:
        return row['id'] if isinstance(row, sqlite3.Row) else row[0]
    cursor.execute('INSERT INTO accounts (name, updated_at) VALUES (?, CURRENT_TIMESTAMP)', (name,))
    return cursor.lastrowid


def sync_accounts_from_clients():
    try:
        conn = get_db()
        c = conn.cursor()
        c.execute('SELECT DISTINCT company FROM clients WHERE company IS NOT NULL AND TRIM(company) != ""')
        companies = [row['company'] for row in c.fetchall()]
        for company in companies:
            ensure_account_for_company(c, company)
        conn.commit()
        conn.close()
    except Exception as e:
        logger.warning(f'[WARN] sync_accounts_from_clients: {e}')


sync_accounts_from_clients()


def _normalize_automapping_key(company, country, industry):
    def clean(v):
        return re.sub(r'\s+', ' ', (v or '').strip().lower())
    return f"{clean(company)}|{clean(country)}|{clean(industry)}"


def _extract_tavily_evidence(results, max_items=5):
    evidences = []
    seen_urls = set()
    for item in (results or []):
        url = (item.get('url') or '').strip()
        if url and url in seen_urls:
            continue
        if url:
            seen_urls.add(url)

        snippet = (item.get('content') or item.get('snippet') or '').strip()
        title = (item.get('title') or '').strip()
        score = item.get('score')
        if title and snippet:
            snippet = f"{title}: {snippet}"
        evidences.append({
            'url': url,
            'snippet': snippet[:500],
            'title': title,
            'score': score
        })
        if len(evidences) >= max_items:
            break
    return evidences


def _detect_keywords_in_evidence(evidence, keywords):
    haystack = ' '.join([(e.get('snippet') or '').lower() for e in evidence])
    return [kw for kw in keywords if kw in haystack]


def _build_automapping_search_plan(company, country, industry):
    company_hint = f'"{company}"'
    base_context = f'{industry} {country}'
    return {
        'cloud': {
            'type': 'value',
            'keywords': ['aws', 'azure', 'microsoft azure', 'google cloud', 'gcp', 'oracle cloud', 'nuvem'],
            'query': f'{company_hint} {base_context} infraestrutura cloud provedor de nuvem aws azure gcp'
        },
        'erp': {
            'type': 'value',
            'keywords': ['sap', 'oracle erp', 'totvs', 'protheus', 'microsoft dynamics', 's/4hana'],
            'query': f'{company_hint} {base_context} sistema erp sap totvs oracle dynamics'
        },
        'crm': {
            'type': 'value',
            'keywords': ['salesforce', 'hubspot', 'dynamics 365', 'zoho crm', 'crm'],
            'query': f'{company_hint} {base_context} crm salesforce hubspot dynamics 365'
        },
        'observability': {
            'type': 'tools',
            'keywords': ['datadog', 'splunk', 'new relic', 'dynatrace', 'grafana', 'observability', 'observabilidade'],
            'query': f'{company_hint} {base_context} observabilidade datadog splunk dynatrace grafana'
        },
        'security': {
            'type': 'tools',
            'keywords': ['crowdstrike', 'palo alto', 'sentinelone', 'fortinet', 'okta', 'siem', 'edr', 'xdr'],
            'query': f'{company_hint} {base_context} cibersegurança siem edr xdr crowdstrike fortinet palo alto'
        },
        'data_analytics': {
            'type': 'tools',
            'keywords': ['snowflake', 'databricks', 'power bi', 'tableau', 'bigquery', 'data lake', 'analytics'],
            'query': f'{company_hint} {base_context} analytics bi power bi tableau databricks snowflake'
        },
        'ai': {
            'type': 'tools',
            'keywords': ['openai', 'copilot', 'gemini', 'claude', 'machine learning', 'ia generativa', 'inteligência artificial'],
            'query': f'{company_hint} {base_context} inteligência artificial ia generativa openai copilot gemini'
        }
    }


def _calculate_section_confidence(evidence, matched_keywords):
    evidence_count = len(evidence or [])
    keyword_hits = len(matched_keywords or [])
    if evidence_count == 0:
        return 'unknown'
    if keyword_hits >= 2:
        return 'high'
    if keyword_hits == 1:
        return 'medium'
    if evidence_count >= 3:
        return 'low'
    return 'unknown'


def _calculate_evidence_quality(evidence):
    if not evidence:
        return 0
    quality = 0
    for ev in evidence:
        url = (ev.get('url') or '').lower()
        snippet = (ev.get('snippet') or '')
        if url.startswith('https://'):
            quality += 2
        if any(domain in url for domain in ['.gov', '.edu', 'microsoft.com', 'oracle.com', 'sap.com', 'aws.amazon.com', 'salesforce.com']):
            quality += 2
        if len(snippet) >= 120:
            quality += 1
    return quality


def _company_mentioned_in_evidence(evidence, company):
    """Verifica se o nome da empresa aparece em pelo menos um snippet da evidência."""
    if not company:
        return False
    company_lower = company.lower().strip()
    for ev in (evidence or []):
        snippet = (ev.get('snippet') or '').lower()
        if company_lower in snippet:
            return True
    return False


def _build_section_result(section_key, section_cfg, evidence, error_message=None, company=''):
    matched_keywords = _detect_keywords_in_evidence(evidence, section_cfg['keywords'])
    confidence = _calculate_section_confidence(evidence, matched_keywords)
    # Só marca como 'identified' se a empresa for mencionada nos snippets + keyword encontrado
    company_in_evidence = _company_mentioned_in_evidence(evidence, company)
    if matched_keywords and company_in_evidence:
        status = 'identified'
    elif evidence:
        status = 'investigate'
    else:
        status = 'unknown'

    if error_message:
        status = 'error'
        confidence = 'unknown'

    evidence_quality = _calculate_evidence_quality(evidence)

    base = {
        'status': status,
        'confidence': confidence,
        'evidence': evidence,
        'matched_keywords': matched_keywords,
        'query_used': section_cfg['query'],
        'evidence_quality': evidence_quality,
        'error': error_message
    }

    if section_cfg['type'] == 'value':
        value = ', '.join(matched_keywords) if matched_keywords else 'Não identificado'
        base['value'] = value
        if section_key in {'cloud', 'erp', 'crm'}:
            base['partners'] = []
    else:
        base['tools'] = matched_keywords

    return base


def _build_automapping_sections(company, country, industry, search_results_by_section, section_errors=None):
    plan = _build_automapping_search_plan(company, country, industry)
    sections = {}
    section_errors = section_errors or {}
    for section_key, section_cfg in plan.items():
        raw_results = search_results_by_section.get(section_key, [])
        evidence = _extract_tavily_evidence(raw_results, max_items=5)
        sections[section_key] = _build_section_result(section_key, section_cfg, evidence, section_errors.get(section_key), company)
    return sections


def _build_automapping_payload(company, country, industry, search_results_by_section, section_errors=None, execution_meta=None):
    sections = _build_automapping_sections(company, country, industry, search_results_by_section, section_errors)
    identified_sections = [k for k, v in sections.items() if v.get('status') == 'identified']
    error_sections = [k for k, v in sections.items() if v.get('status') == 'error']
    return {
        'schema_version': '2.0',
        'company': company,
        'country': country,
        'industry': industry,
        'sections': sections,
        'execution_summary': {
            'identified_sections_count': len(identified_sections),
            'error_sections_count': len(error_sections),
            'identified_sections': identified_sections,
            'error_sections': error_sections,
            'mode': 'partial' if error_sections else 'complete'
        },
        'strategic_reading': {
            'competitive_space': [],
            'lock_in_risk': [],
            'entry_points': []
        },
        'execution_meta': execution_meta or {},
        'generated_at': datetime.utcnow().isoformat() + 'Z'
    }


def _run_tavily_request(api_key, query, max_results=6):
    payload = {
        'api_key': api_key,
        'query': query,
        'search_depth': 'advanced',
        'max_results': max_results,
        'include_answer': False,
        'include_raw_content': False
    }

    req = urllib.request.Request(
        'https://api.tavily.com/search',
        data=json.dumps(payload).encode('utf-8'),
        headers={'Content-Type': 'application/json'},
        method='POST'
    )

    with urllib.request.urlopen(req, timeout=25) as resp:
        data = json.loads(resp.read().decode('utf-8'))
    return data.get('results') or []


def _run_tavily_search(company, country, industry):
    api_key = _resolve_setting('tavily_api_key', 'TAVILY_API_KEY')
    if not api_key:
        raise RuntimeError('A chave da Tavily não está configurada. Configure em Configurações > Integrações ou defina TAVILY_API_KEY no ambiente.')

    plan = _build_automapping_search_plan(company, country, industry)
    all_results = {}
    section_errors = {}
    attempts_by_section = {}

    for section_key, section_cfg in plan.items():
        max_attempts = 2
        attempts = 0
        last_error = None
        while attempts < max_attempts:
            attempts += 1
            try:
                all_results[section_key] = _run_tavily_request(api_key, section_cfg['query'])
                last_error = None
                break
            except Exception as e:
                last_error = str(e)
                if attempts < max_attempts:
                    time.sleep(0.4 * attempts)
        attempts_by_section[section_key] = attempts
        if last_error is not None:
            section_errors[section_key] = f'Falha ao consultar provider: {last_error[:180]}'
            all_results[section_key] = []

    execution_meta = {
        'provider': 'tavily',
        'attempts_by_section': attempts_by_section,
        'searched_sections_count': len(plan)
    }
    return all_results, section_errors, execution_meta


def _default_llm_summary(sections):
    summary = {}
    for section_key in (sections or {}).keys():
        summary[section_key] = {
            'final_answer': 'Inconclusivo com as evidências atuais.',
            'confidence_percent': 20,
            'certainty': 'uncertain',
            'reasoning': 'Não houve evidência específica suficiente para consolidar uma resposta final.'
        }
    return summary


def _extract_answer_from_partial_json(text):
    """Extrai o valor do campo 'answer' de um JSON potencialmente truncado.
    Útil quando a API retorna um JSON incompleto (cortado no meio).
    Estratégia: encontra '"answer"' na string, pula o ':', lê o valor da string
    até encontrar um '"' de fechamento não escapado.
    """
    if not text:
        return None
    # Procura a chave 'answer' no texto
    key_pos = text.find('"answer"')
    if key_pos == -1:
        return None
    # Avança até o ':'
    colon_pos = text.find(':', key_pos + 8)
    if colon_pos == -1:
        return None
    # Avança até a abertura da string de valor
    value_start = text.find('"', colon_pos + 1)
    if value_start == -1:
        return None
    # Lê a string até o fechamento (respeitando escapes)
    chars = []
    i = value_start + 1
    while i < len(text):
        ch = text[i]
        if ch == '\\' and i + 1 < len(text):
            # Caractere escapado
            next_ch = text[i + 1]
            if next_ch == 'n':
                chars.append('\n')
            elif next_ch == 't':
                chars.append('\t')
            elif next_ch == 'r':
                chars.append('\r')
            elif next_ch == '"':
                chars.append('"')
            elif next_ch == '\\':
                chars.append('\\')
            else:
                chars.append(next_ch)
            i += 2
            continue
        if ch == '"':
            # Fim da string — retorna o que foi lido até aqui
            return ''.join(chars).strip()
        chars.append(ch)
        i += 1
    # JSON truncado: retorna o que foi lido até o fim do texto
    result = ''.join(chars).strip()
    return result if result else None


def _extract_json_object_from_text(text):
    """Extrai o primeiro objeto JSON válido de uma string usando balanceamento de chaves.
    Mais robusto que rfind('}') para JSONs aninhados ou com múltiplos objetos.
    Também tenta detectar e limpar prefixos de texto antes do JSON.
    """
    if not text:
        return None
    # Tenta parse direto primeiro (caso mais comum e mais rápido)
    stripped = text.strip()
    if stripped.startswith('{'):
        try:
            return json.loads(stripped)
        except Exception as e:
            logger.debug(f'[_extract_json_object_from_text] exceção ignorada: {e}')

    # Busca o primeiro '{' e tenta balancear as chaves
    start = text.find('{')
    if start == -1:
        return None

    depth = 0
    in_string = False
    escape_next = False
    for i, ch in enumerate(text[start:], start=start):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                chunk = text[start:i+1]
                try:
                    return json.loads(chunk)
                except Exception:
                    # Se falhou, tenta o próximo '{'
                    next_start = text.find('{', start + 1)
                    if next_start == -1:
                        return None
                    start = next_start
                    depth = 0
                    in_string = False
                    escape_next = False
    return None


def _validate_openrouter_api_key(api_key):
    key = (api_key or '').strip()
    if not key:
        return False, 'OPENROUTER_API_KEY vazia.'
    # Formato mais comum das chaves OpenRouter: sk-or-v1-...
    if key.startswith('sk-or-'):
        return True, ''
    return False, 'OPENROUTER_API_KEY com formato inválido (esperado prefixo sk-or-).'


def _run_openrouter_synthesis(result_payload):
    settings_map = _load_app_settings_map([
        'openrouter_model',
        'openrouter_site_url',
        'openrouter_app_name'
    ])

    api_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if not api_key:
        raise RuntimeError('A chave da OpenRouter não está configurada. Configure em Configurações > Integrações ou defina OPENROUTER_API_KEY no ambiente.')

    key_ok, key_msg = _validate_openrouter_api_key(api_key)
    if not key_ok:
        raise RuntimeError(f'Chave OpenRouter inválida: {key_msg}')

    model = (settings_map.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
    site_url = (settings_map.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
    app_name = (settings_map.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
    sections = result_payload.get('sections') or {}

    company = result_payload.get('company') or ''
    compact_sections = {}
    for section_key, section_data in sections.items():
        evidence = section_data.get('evidence') or []
        company_lower = company.lower()
        compact_sections[section_key] = {
            'matched_keywords': section_data.get('matched_keywords') or [],
            'evidence': [
                {
                    'url': ev.get('url'),
                    'snippet': (ev.get('snippet') or '')[:500],
                    'mentions_company': company_lower in (ev.get('snippet') or '').lower()
                }
                for ev in evidence[:4]
            ]
        }

    user_payload = {
        'company': company,
        'country': result_payload.get('country'),
        'industry': result_payload.get('industry'),
        'sections': compact_sections
    }

    system_prompt = (
        'Você é um analista de inteligência de mercado. Responda APENAS em JSON válido. '
        'REGRA ABSOLUTA: só gere um final_answer específico se pelo menos um trecho '
        '(campo "snippet") mencionar EXPLICITAMENTE o nome da empresa junto com a solução confirmada. '
        'Indícios proibidos: "parece utilizar", "pode estar usando", "provavelmente usa", listas de opções. '
        'Se nenhum trecho com mentions_company=true confirmar a solução, use '
        'final_answer:"Não foi possível identificar com as evidências disponíveis" e confidence_percent:0. '
        'confidence_percent deve refletir a qualidade da evidência: 0 se genérica, >70 só se confirmada. '
        'Formato obrigatório: '
        '{"sections":{"cloud":{"final_answer":"...","confidence_percent":0,"certainty":"confirmed|uncertain","reasoning":"..."}, ...}}'
    )

    request_payload = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': json.dumps(user_payload, ensure_ascii=False)}
        ],
        'temperature': 0.1
    }

    req = urllib.request.Request(
        'https://openrouter.ai/api/v1/chat/completions',
        data=json.dumps(request_payload).encode('utf-8'),
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}',
            'HTTP-Referer': site_url,
            'X-Title': app_name
        },
        method='POST'
    )

    try:
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = json.loads(resp.read().decode('utf-8'))
    except urllib.error.HTTPError as e:
        detail = e.read().decode('utf-8', errors='ignore') if hasattr(e, 'read') else str(e)
        diagnostics = {
            'status': getattr(e, 'code', None),
            'model': model,
            'has_api_key': bool(api_key),
            'api_key_prefix': api_key[:7] if api_key else '',
            'sent_headers': ['Content-Type', 'Authorization', 'HTTP-Referer', 'X-Title']
        }
        logger.exception(f'[ERROR][OpenRouter] HTTPError diagnostics={diagnostics} body={detail[:500]}')
        hint = ' Verifique se OPENROUTER_API_KEY é a chave da OpenRouter (prefixo sk-or-) e se foi exportada no mesmo terminal do app.' if diagnostics.get('status') == 401 else ''
        raise RuntimeError(f'OpenRouter HTTP {diagnostics["status"]} - body: {detail[:400]} | diagnostics: {json.dumps(diagnostics, ensure_ascii=False)}{hint}')
    except Exception as e:
        diagnostics = {
            'model': model,
            'has_api_key': bool(api_key),
            'api_key_prefix': api_key[:7] if api_key else ''
        }
        logger.exception(f'[ERROR][OpenRouter] Exception diagnostics={diagnostics} error={e}')
        raise RuntimeError(f'Falha inesperada OpenRouter: {str(e)} | diagnostics: {json.dumps(diagnostics, ensure_ascii=False)}')

    choices = data.get('choices') or []
    message_content = ''
    if choices:
        content = ((choices[0] or {}).get('message') or {}).get('content')
        if isinstance(content, list):
            message_content = ''.join([part.get('text', '') for part in content if isinstance(part, dict)])
        else:
            message_content = str(content or '')

    parsed = _extract_json_object_from_text(message_content)
    llm_sections = ((parsed or {}).get('sections') or {}) if isinstance(parsed, dict) else {}

    final_sections = _default_llm_summary(sections)
    for section_key in final_sections.keys():
        candidate = llm_sections.get(section_key)
        if not isinstance(candidate, dict):
            continue
        answer = (candidate.get('final_answer') or '').strip() or final_sections[section_key]['final_answer']
        reasoning = (candidate.get('reasoning') or '').strip() or final_sections[section_key]['reasoning']
        certainty = candidate.get('certainty') if candidate.get('certainty') in {'confirmed', 'uncertain'} else 'uncertain'
        try:
            confidence = int(candidate.get('confidence_percent'))
        except Exception:
            confidence = final_sections[section_key]['confidence_percent']
        confidence = max(0, min(100, confidence))

        final_sections[section_key] = {
            'final_answer': answer,
            'confidence_percent': confidence,
            'certainty': certainty,
            'reasoning': reasoning
        }

    meta = {
        'provider': 'openrouter',
        'model': model,
        'raw_response_received': bool(message_content)
    }
    return final_sections, meta

def normalize_phone(phone):
    if not phone:
        return None
    digits = re.sub(r'\D', '', str(phone))
    if not digits:
        return None
    if len(digits) > 11:
        digits = digits[-11:]
    if len(digits) < 10:
        return digits
    if len(digits) == 10:
        digits = digits[:2] + '9' + digits[2:]
    return f"{digits[:2]} {digits[2:7]}.{digits[7:11]}"


def extract_time_from_text(text):
    if not text:
        return None
    lower = text.lower()
    patterns = [
        r'\b([01]?\d|2[0-3])\s*[:h]\s*([0-5]\d)\b',
        r'\b([01]?\d|2[0-3])\s*h\s*([0-5]\d)\b',
    ]
    for pattern in patterns:
        m = re.search(pattern, lower)
        if m:
            return f"{int(m.group(1)):02d}:{int(m.group(2)):02d}"
    m = re.search(r'\b([01]?\d|2[0-3])\s*h\b', lower)
    if m:
        return f"{int(m.group(1)):02d}:00"
    return None


def extract_future_commitment_dates(text):
    if not text:
        return []

    now = datetime.now()
    matches = []
    month_map = {
        'janeiro': 1, 'jan': 1,
        'fevereiro': 2, 'fev': 2,
        'marco': 3, 'março': 3, 'mar': 3,
        'abril': 4, 'abr': 4,
        'maio': 5, 'mai': 5,
        'junho': 6, 'jun': 6,
        'julho': 7, 'jul': 7,
        'agosto': 8, 'ago': 8,
        'setembro': 9, 'set': 9,
        'outubro': 10, 'out': 10,
        'novembro': 11, 'nov': 11,
        'dezembro': 12, 'dez': 12,
    }

    def add_date(day, month, year=None):
        y = year or now.year
        if y < 100:
            y += 2000
        try:
            dt = datetime(y, month, day)
            if dt.date() < now.date() and year is None:
                dt = datetime(y + 1, month, day)
            if dt.date() >= now.date():
                matches.append(dt.date().isoformat())
        except Exception as e:
            logger.debug(f'[add_date] exceção ignorada: {e}')

    for m in re.finditer(r'\b(\d{1,2})[\/\.\-](\d{1,2})(?:[\/\.\-](\d{2,4}))?\b', text):
        add_date(int(m.group(1)), int(m.group(2)), int(m.group(3)) if m.group(3) else None)

    lowered = text.lower()
    for m in re.finditer(r'\b(\d{1,2})\s+de\s+([a-zç]+)(?:\s+de\s+(\d{2,4}))?\b', lowered):
        day = int(m.group(1))
        month = month_map.get((m.group(2) or '').strip())
        if month:
            add_date(day, month, int(m.group(3)) if m.group(3) else None)

    for m in re.finditer(r'\bem\s+(\d{1,3})\s+dias?\b', lowered):
        days = int(m.group(1))
        if days >= 0:
            matches.append((now.date() + timedelta(days=days)).isoformat())

    if 'depois de amanhã' in lowered or 'depois de amanha' in lowered:
        matches.append((now.date() + timedelta(days=2)).isoformat())
    elif 'amanh' in lowered:
        matches.append((now.date() + timedelta(days=1)).isoformat())

    seen = set()
    ordered = []
    for d in matches:
        if d not in seen:
            seen.add(d)
            ordered.append(d)
    return ordered


def create_commitments_from_activity(cursor, client_id, activity_id, text):
    dates = extract_future_commitment_dates(text)
    if not dates:
        return []

    safe_title = (text or '').strip().replace('\n', ' ')
    if len(safe_title) > 120:
        safe_title = safe_title[:117] + '...'

    parsed_time = extract_time_from_text(text)
    created = []
    for due_date in dates:
        cursor.execute(
            '''INSERT INTO commitments (client_id, activity_id, title, notes, due_date, due_time, source_type)
               VALUES (?, ?, ?, ?, ?, ?, ?)''',
            (client_id, activity_id, safe_title or 'Retorno com cliente', text, due_date, parsed_time, 'activity')
        )
        created.append({
            'id': cursor.lastrowid,
            'client_id': client_id,
            'due_date': due_date,
            'due_time': parsed_time,
            'title': safe_title or 'Retorno com cliente',
            'notes': text
        })
    return created


def enrich_commitments_with_client_data(cursor, commitments, client_id):
    if not commitments:
        return commitments
    cursor.execute('SELECT name, company, position, email, photo_url FROM clients WHERE id = ?', (client_id,))
    client = dict_from_row(cursor.fetchone()) or {}
    for item in commitments:
        item['client_name'] = client.get('name')
        item['client_company'] = client.get('company')
        item['client_position'] = client.get('position')
        item['client_email'] = client.get('email')
        item['client_photo'] = client.get('photo_url')
    return commitments


def _col_index(cell_ref):
    letters = ''.join(ch for ch in cell_ref if ch.isalpha()).upper()
    idx = 0
    for ch in letters:
        idx = (idx * 26) + (ord(ch) - 64)
    return max(idx - 1, 0)


def parse_xlsx_without_openpyxl(file_storage):
    """Le arquivo XLSX usando apenas bibliotecas nativas (fallback)."""
    import io

    file_bytes = file_storage.read()
    file_storage.seek(0)

    ns = {'x': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}

    with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
        # shared strings (quando as celulas usam indice de string)
        shared_strings = []
        if 'xl/sharedStrings.xml' in zf.namelist():
            root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
            for si in root.findall('x:si', ns):
                parts = [t.text or '' for t in si.findall('.//x:t', ns)]
                shared_strings.append(''.join(parts))

        sheet_name = 'xl/worksheets/sheet1.xml'
        if sheet_name not in zf.namelist():
            # fallback para o primeiro worksheet encontrado
            worksheet_files = [n for n in zf.namelist() if n.startswith('xl/worksheets/sheet') and n.endswith('.xml')]
            if not worksheet_files:
                raise ValueError('Planilha inválida: worksheet não encontrada')
            sheet_name = sorted(worksheet_files)[0]

        sheet_root = ET.fromstring(zf.read(sheet_name))
        rows = []

        for row in sheet_root.findall('.//x:sheetData/x:row', ns):
            row_data = []
            for cell in row.findall('x:c', ns):
                cell_ref = cell.get('r', '')
                col_idx = _col_index(cell_ref) if cell_ref else len(row_data)

                while len(row_data) < col_idx:
                    row_data.append('')

                cell_type = cell.get('t')
                value = ''

                if cell_type == 's':
                    v = cell.find('x:v', ns)
                    if v is not None and (v.text or '').isdigit():
                        s_idx = int(v.text)
                        if 0 <= s_idx < len(shared_strings):
                            value = shared_strings[s_idx]
                elif cell_type == 'inlineStr':
                    t = cell.find('.//x:t', ns)
                    value = t.text if t is not None and t.text is not None else ''
                else:
                    v = cell.find('x:v', ns)
                    value = v.text if v is not None and v.text is not None else ''

                row_data.append(str(value).strip())

            rows.append(row_data)

    return rows


@app.route('/api/transcribe-audio', methods=['POST', 'OPTIONS'])
@app.route('/api/transcribe-audio/', methods=['POST', 'OPTIONS'])
def transcribe_audio():
    if request.method == 'OPTIONS':
        return ('', 204)

    if TRANSCRIPTION_DEBUG:
        logger.debug('[Transcription][DEBUG] Requisição recebida em /api/transcribe-audio')
        logger.debug(f"[Transcription][DEBUG] Content-Type: {request.content_type}")
        logger.debug(f"[Transcription][DEBUG] Content-Length: {request.content_length}")

    if not WHISPER_AVAILABLE:
        details = f' ({WHISPER_IMPORT_ERROR})' if WHISPER_IMPORT_ERROR else ''
        logger.error(f'[Transcription][ERROR] Biblioteca faster-whisper indisponível neste ambiente{details}.')
        return jsonify({'error': 'Biblioteca faster-whisper não está disponível neste ambiente.'}), 503

    audio_file = request.files.get('audio')
    if not audio_file:
        return jsonify({'error': 'Arquivo de áudio não enviado.'}), 400

    ffmpeg_path = configure_ffmpeg_for_whisper()
    if TRANSCRIPTION_DEBUG and not ffmpeg_path:
        logger.debug('[Transcription][DEBUG] FFmpeg externo não encontrado; continuando com decoder da stack local.')

    suffix = Path(audio_file.filename or 'audio.webm').suffix or '.webm'
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            audio_file.save(tmp)
            temp_path = tmp.name

        if TRANSCRIPTION_DEBUG:
            logger.debug(f"[Transcription][DEBUG] Arquivo salvo temporariamente: {temp_path}")
            logger.debug(f"[Transcription][DEBUG] Nome original: {audio_file.filename}")
            logger.debug(f"[Transcription][DEBUG] FFmpeg em uso: {ffmpeg_path or 'decoder interno'}")

        model = get_whisper_model()
        if model is None:
            details = f' ({WHISPER_IMPORT_ERROR})' if WHISPER_IMPORT_ERROR else ''
            return jsonify({'error': f'Backend de transcrição indisponível{details}.'}), 503
        segments, _ = model.transcribe(temp_path, language='pt')
        text = ''.join(segment.text for segment in segments).strip()

        if TRANSCRIPTION_DEBUG:
            logger.debug(f"[Transcription][DEBUG] Texto transcrito (primeiros 200 chars): {text[:200]}")

        return jsonify({'text': text})
    except Exception as e:
        logger.error(f'[Transcription][ERROR] POST /api/transcribe-audio: {e}')
        if TRANSCRIPTION_DEBUG:
            traceback.print_exc()
        return jsonify({'error': f'Falha ao transcrever áudio com faster-whisper: {e}'}), 500
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
            if TRANSCRIPTION_DEBUG:
                logger.debug(f"[Transcription][DEBUG] Arquivo temporário removido: {temp_path}")

# API - Clientes (rotas alternativas para compatibilidade)


_VALID_THEMES = {'verde-classico', 'baby-pink', 'black-cat', 'white-pearl', 'blue-space'}


# =====================================================
# Logs de Depuração
# Endpoints para o painel "Logs de Depuração" em Configurações.
# - GET  /api/config/logs        -> retorna as últimas N linhas de LOG_FILE
# - POST /api/config/logs/client -> recebe entradas de log do cliente
#                                   (tocaDebug) e grava no mesmo LOG_FILE
# =====================================================

_CLIENT_LOGGER = logging.getLogger('toca-do-coelho.client')


# ---------------------------------------------------------------------------
# Atualização automática baseada em release: baixa o instalador anexado à
# release e o executa, encerrando o app para que os arquivos sejam liberados.
# ---------------------------------------------------------------------------
UPDATE_DOWNLOAD_DIR = DATA_DIR / 'updates'


def _download_update_async(task_id, url, name):
    try:
        _bg_task_set(task_id, {'status': 'processing', 'step': 'Iniciando download...', 'progress': 3})
        UPDATE_DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

        safe_name = os.path.basename(name or '') or 'TocaDoCoelho-Setup.exe'
        if not safe_name.lower().endswith('.exe'):
            safe_name += '.exe'
        dest = UPDATE_DOWNLOAD_DIR / safe_name

        req = urllib.request.Request(url, headers={'User-Agent': 'TocaDoCoelho-Updater'})
        with urllib.request.urlopen(req, timeout=30) as resp:
            total = int(resp.headers.get('Content-Length') or 0)
            downloaded = 0
            with open(dest, 'wb') as fh:
                while True:
                    chunk = resp.read(262144)
                    if not chunk:
                        break
                    fh.write(chunk)
                    downloaded += len(chunk)
                    if total > 0:
                        pct = 3 + int(downloaded / total * 94)
                        _bg_task_set(task_id, {
                            'progress': min(97, pct),
                            'step': f'Baixando... {downloaded // (1024 * 1024)} MB de {total // (1024 * 1024)} MB'
                        })

        _bg_task_set(task_id, {
            'status': 'done', 'progress': 100, 'step': 'Download concluído.',
            'installer_path': str(dest)
        })
        _bg_task_cleanup(task_id, delay=600)
    except Exception as e:
        logger.exception(f'[Update] Falha no download da atualização: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': f'Falha ao baixar a atualização: {e}'})
        _bg_task_cleanup(task_id, delay=600)


def _outlook_fetch_via_powershell(days=60):
    """Desabilitado por recomendação de segurança SAST — use Microsoft Graph API."""
    raise RuntimeError(
        'Sincronização via PowerShell/COM foi desabilitada. Use "Sync Microsoft 365" (Graph API).'
    )


def _DISABLED_outlook_fetch_via_powershell_legacy(days=60):
    """Mantido apenas como referência histórica — NÃO EXECUTAR."""
    import subprocess, tempfile, json, os

    ps_script = """\
param([int]$Days = 60)
$ErrorActionPreference = 'Continue'
$cutoff = (Get-Date).AddDays(-$Days)

function Get-SmtpAddress($r) {
    $a = $null; try { $a = $r.Address } catch {}
    if ($a -and $a -match '@' -and $a -notmatch '^/o=') { return $a.ToLower() }
    try { return ($r.PropertyAccessor.GetProperty('http://schemas.microsoft.com/mapi/proptag/0x39FE001E')).ToLower() } catch {}
    return ''
}

function Get-SenderSmtp($m) {
    $a = $null; try { $a = $m.SenderEmailAddress } catch {}
    if ($a -and $a -match '@' -and $a -notmatch '^/o=') { return $a.ToLower() }
    try { return ($m.PropertyAccessor.GetProperty('http://schemas.microsoft.com/mapi/proptag/0x5D01001E')).ToLower() } catch {}
    return ''
}

function Get-AllFolders($folder) {
    $list = [System.Collections.Generic.List[object]]::new()
    $list.Add($folder)
    try { foreach ($s in $folder.Folders) { (Get-AllFolders $s) | ForEach-Object { $list.Add($_) } } } catch {}
    return ,$list
}

function Make-Item($item, $dt, $dir) {
    $rcpts = [System.Collections.Generic.List[hashtable]]::new()
    try { foreach ($r in $item.Recipients) { try { $rcpts.Add(@{ name=($r.Name+''); email=(Get-SmtpAddress $r) }) } catch {} } } catch {}
    $bp = ''
    try { $bp = ($item.Body -replace '[\\r\\n\\t]+',' ').Trim(); if ($bp.Length -gt 1500) { $bp = $bp.Substring(0,1500) } } catch {}
    return [PSCustomObject]@{
        subject      = ($item.Subject+'')
        date         = $dt.ToString('yyyy-MM-ddTHH:mm:ss')
        direction    = $dir
        sender       = @{ name=($item.SenderName+''); email=(Get-SenderSmtp $item) }
        recipients   = @($rcpts)
        body_preview = $bp
    }
}

$ol = $null
try {
    $ol = [System.Runtime.InteropServices.Marshal]::GetActiveObject('Outlook.Application')
} catch {
    try { $ol = New-Object -ComObject Outlook.Application } catch {
        [Console]::Error.WriteLine("Nao foi possivel conectar ao Outlook: $_")
        exit 1
    }
}
$ns = $ol.GetNamespace('MAPI')
$results = [System.Collections.Generic.List[object]]::new()

try {
    $inbox = $ns.GetDefaultFolder(6)
    foreach ($folder in (Get-AllFolders $inbox)) {
        try {
            foreach ($item in $folder.Items) {
                try {
                    if ($item.Class -ne 43) { continue }
                    $dt = $null; try { $dt = $item.ReceivedTime } catch { continue }
                    if ($dt -lt $cutoff) { continue }
                    $results.Add((Make-Item $item $dt 'received'))
                } catch {}
            }
        } catch {}
    }
} catch { [Console]::Error.WriteLine("Erro inbox: $_") }

try {
    $sent = $ns.GetDefaultFolder(5)
    foreach ($item in $sent.Items) {
        try {
            if ($item.Class -ne 43) { continue }
            $dt = $null; try { $dt = $item.SentOn } catch { continue }
            if ($dt -lt $cutoff) { continue }
            $results.Add((Make-Item $item $dt 'sent'))
        } catch {}
    }
} catch { [Console]::Error.WriteLine("Erro sent: $_") }

if ($results.Count -eq 0) { '[]'; exit 0 }
$results | ConvertTo-Json -Depth 5 -Compress
"""

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(
            mode='w', suffix='.ps1', delete=False, encoding='utf-8', errors='replace'
        ) as tmp:
            tmp.write(ps_script)
            tmp_path = tmp.name

        timeout_s = int(_resolve_setting('outlook_sync_timeout_seconds', 'OUTLOOK_SYNC_TIMEOUT') or 120)
        try:
            proc = subprocess.run(
                ['powershell', '-ExecutionPolicy', 'Bypass', '-NoProfile',
                 '-File', tmp_path, '-Days', str(int(days))],
                capture_output=True, text=True, encoding='utf-8', errors='replace',
                timeout=timeout_s
            )
        except subprocess.TimeoutExpired as timeout_err:
            raise RuntimeError(
                f'A leitura via Outlook COM excedeu {timeout_s}s. '
                'Tente reduzir o período (ex.: 7 dias) ou configure o OAuth Graph para acesso sem limite de tempo.'
            ) from timeout_err
        stderr = (proc.stderr or '').strip()
        stdout = (proc.stdout or '').strip()
        if proc.returncode != 0 and not stdout:
            raise RuntimeError(stderr or f'PowerShell retornou código {proc.returncode}')
        if not stdout or stdout == '[]':
            return []
        data = json.loads(stdout)
        if isinstance(data, dict):
            data = [data]
        for item in data:
            rcpts = item.get('recipients') or []
            if isinstance(rcpts, dict):
                rcpts = [rcpts]
            item['recipients'] = rcpts
            if not isinstance(item.get('sender'), dict):
                item['sender'] = {'name': '', 'email': ''}
        return data
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception as e:
                logger.debug(f'[_DISABLED_outlook_fetch_via_powershell_legacy] exceção ignorada: {e}')


def _outlook_extract_smtp_from_recipient(recipient):
    """Extrai endereço SMTP de um recipient do Outlook (trata Exchange/EX)."""
    try:
        addr = (recipient.Address or '').strip().lower()
        if addr and '@' in addr and not addr.startswith('/o='):
            return addr
        return recipient.PropertyAccessor.GetProperty(
            'http://schemas.microsoft.com/mapi/proptag/0x39FE001E'
        ).strip().lower()
    except Exception:
        return ''


def _outlook_extract_smtp_from_sender(mail_item):
    """Extrai endereço SMTP do remetente (trata Exchange/EX)."""
    try:
        addr = (mail_item.SenderEmailAddress or '').strip().lower()
        if addr and '@' in addr and not addr.startswith('/o='):
            return addr
        return mail_item.PropertyAccessor.GetProperty(
            'http://schemas.microsoft.com/mapi/proptag/0x5D01001E'
        ).strip().lower()
    except Exception:
        return ''


def _outlook_get_all_subfolders(folder):
    result = [folder]
    try:
        for sub in folder.Folders:
            result.extend(_outlook_get_all_subfolders(sub))
    except Exception as e:
        logger.debug(f'[_outlook_get_all_subfolders] exceção ignorada: {e}')
    return result


def _outlook_extract_email(item, direction, cutoff_dt):
    """Extrai metadados de um MailItem. Retorna dict ou None."""
    try:
        if item.Class != 43:  # 43 = olMail
            return None
        raw_dt = item.ReceivedTime if direction == 'received' else item.SentOn
        dt = datetime(raw_dt.year, raw_dt.month, raw_dt.day,
                      raw_dt.hour, raw_dt.minute, raw_dt.second)
        if dt < cutoff_dt:
            return None
        recipients = []
        try:
            for r in item.Recipients:
                try:
                    email = _outlook_extract_smtp_from_recipient(r)
                    recipients.append({'name': (r.Name or '').strip(), 'email': email})
                except Exception as e:
                    logger.debug(f'[_outlook_extract_email] exceção ignorada: {e}')
        except Exception as e:
            logger.debug(f'[_outlook_extract_email] exceção ignorada: {e}')
        body_preview = ''
        try:
            body_preview = (item.Body or '')[:1500].strip()
        except Exception as e:
            logger.debug(f'[_outlook_extract_email] exceção ignorada: {e}')
        return {
            'subject': (item.Subject or '').strip(),
            'date': dt.strftime('%Y-%m-%dT%H:%M:%S'),
            'direction': direction,
            'sender': {
                'name': (item.SenderName or '').strip(),
                'email': _outlook_extract_smtp_from_sender(item)
            },
            'recipients': recipients,
            'body_preview': body_preview
        }
    except Exception:
        return None


def _outlook_process_folder(folder, direction, cutoff_dt, emails):
    count = 0
    try:
        items = folder.Items
        date_str = cutoff_dt.strftime('%m/%d/%Y %H:%M %p')
        field = 'ReceivedTime' if direction == 'received' else 'SentOn'
        try:
            items = items.Restrict(f"[{field}] >= '{date_str}'")
        except Exception as e:
            logger.debug(f'[_outlook_process_folder] exceção ignorada: {e}')
        for item in items:
            try:
                data = _outlook_extract_email(item, direction, cutoff_dt)
                if data:
                    emails.append(data)
                    count += 1
            except Exception as e:
                logger.debug(f'[_outlook_process_folder] exceção ignorada: {e}')
    except Exception as e:
        logger.debug(f'[_outlook_process_folder] exceção ignorada: {e}')
    return count


def _outlook_import_emails(emails_data, conn):
    """
    Importa lista de emails já extraídos para o banco.
    Retorna (imported, skipped_duplicates, skipped_no_match).
    """
    c = conn.cursor()
    c.execute('SELECT id, email FROM clients WHERE email IS NOT NULL AND TRIM(email) != ""')
    email_to_client = {}
    for row in c.fetchall():
        normalized = (row['email'] or '').strip().lower()
        if normalized:
            email_to_client[normalized] = row['id']

    imported = 0
    skipped_duplicates = 0
    skipped_no_match = 0

    for email_data in emails_data:
        subject = (email_data.get('subject') or '').strip()
        email_date = (email_data.get('date') or '').strip()
        direction = email_data.get('direction', 'received')
        sender = email_data.get('sender') or {}
        recipients = email_data.get('recipients') or []
        body_preview = (email_data.get('body_preview') or '').strip()

        if not email_date:
            continue

        if direction == 'received':
            candidates = [sender] if sender.get('email') else []
            counterpart_label = 'De'
        else:
            candidates = [r for r in recipients if r.get('email')]
            counterpart_label = 'Para'

        matched_any = False
        for candidate in candidates:
            candidate_email = (candidate.get('email') or '').strip().lower()
            if not candidate_email or candidate_email not in email_to_client:
                continue

            matched_any = True
            client_id = email_to_client[candidate_email]

            date_minute = email_date[:16]
            subject_key = subject[:100]
            c.execute(
                '''SELECT id FROM activities
                   WHERE client_id = ?
                     AND contact_type = 'Email'
                     AND strftime('%Y-%m-%dT%H:%M', activity_date) = ?
                     AND information LIKE ?
                   LIMIT 1''',
                (client_id, date_minute, f'{subject_key}%')
            )
            if c.fetchone():
                skipped_duplicates += 1
                continue

            # Gerar resumo via LLM somente para emails novos
            summary = None
            if body_preview:
                raw_summary = _sai_simple_prompt(
                    f'Resuma em 1 a 2 frases o conteúdo do email abaixo em português, '
                    f'de forma objetiva, sem mencionar remetente ou destinatário:\n'
                    f'Assunto: {subject}\n'
                    f'Texto: {body_preview[:1000]}\n\n'
                    'Retorne SOMENTE o resumo, sem introdução, aspas ou formatação extra.'
                )
                if raw_summary:
                    summary = raw_summary.strip()

            candidate_name = (candidate.get('name') or candidate_email).strip()
            info_parts = [subject, f'{counterpart_label}: {candidate_name} <{candidate_email}>']
            if summary:
                info_parts.append(f'Resumo: {summary}')
            elif body_preview:
                info_parts.append(body_preview[:300])
            information = '\n'.join(info_parts)

            c.execute(
                '''INSERT INTO activities (client_id, contact_type, information, activity_date)
                   VALUES (?, 'Email', ?, ?)''',
                (client_id, information, email_date)
            )
            _addin_activity_id = c.lastrowid
            c.execute(
                '''UPDATE clients SET last_activity_date = ?
                   WHERE id = ? AND (last_activity_date IS NULL OR last_activity_date < ?)''',
                (email_date, client_id, email_date)
            )
            # Follow-up combinado detectado via LLM (Bloco 7 — ingest do add-in)
            try:
                _detect_followup_from_text(c, client_id, _addin_activity_id, candidate_name,
                                           f'{subject}\n{body_preview}')
            except Exception as e:
                logger.debug(f'[FollowupDetect] falha no ingest do add-in: {e}')
            imported += 1

        if not matched_any:
            skipped_no_match += 1

    conn.commit()
    return imported, skipped_duplicates, skipped_no_match


def _outlook_call_llm(prompt):
    """Chama SAI para análise de email. Restrito a SAI por compliance — conteúdo de
    e-mail não deve ser enviado a provedores externos não homologados (LGPD)."""
    raw = _sai_simple_prompt(prompt)
    return raw.strip() if raw else None


def _outlook_sync_stream_com():
    def _err():
        yield f"data: {json.dumps({'phase': 'error', 'message': 'Sincronização via PowerShell/COM foi desabilitada. Use o botão Sync Microsoft 365 (Graph API).'})}\n\n"
    return Response(stream_with_context(_err()), mimetype='text/event-stream')


try:
    import graph_credentials as _gc
    _GRAPH_DEFAULT_TENANT = getattr(_gc, 'GRAPH_TENANT_ID', '')
    _GRAPH_DEFAULT_CLIENT_ID = getattr(_gc, 'GRAPH_CLIENT_ID', '')
except ImportError:
    _GRAPH_DEFAULT_TENANT = ''
    _GRAPH_DEFAULT_CLIENT_ID = ''


def _graph_redirect_uri():
    return f"{request.scheme}://{request.host}/api/outlook/oauth/callback"


def _graph_make_settings(redirect_uri=''):
    return {
        'tenant': (_resolve_setting('outlook_graph_tenant_id', 'OUTLOOK_GRAPH_TENANT_ID') or _GRAPH_DEFAULT_TENANT).strip(),
        'client_id': (_resolve_setting('outlook_graph_client_id', 'OUTLOOK_GRAPH_CLIENT_ID') or _GRAPH_DEFAULT_CLIENT_ID).strip(),
        'redirect_uri': redirect_uri or (_resolve_setting('outlook_graph_redirect_uri', 'OUTLOOK_GRAPH_REDIRECT_URI') or '').strip(),
        'scope': (_resolve_setting('outlook_graph_scope', 'OUTLOOK_GRAPH_SCOPE') or 'offline_access Mail.Read Mail.Send User.Read').strip(),
    }


def _outlook_sync_stream_graph():
    default_days = int(_resolve_setting('outlook_sync_days', 'OUTLOOK_SYNC_DAYS') or 30)
    days = max(1, min(int(request.args.get('days', default_days)), 365))
    page_size = max(1, min(int(request.args.get('page_size', 50)), 200))
    max_pages = max(1, min(int(request.args.get('max_pages', 10)), 50))
    user_id = max(1, int(request.args.get('user_id', 1)))
    graph_settings = _graph_make_settings(redirect_uri=_graph_redirect_uri())
    return _build_outlook_stream_response(days=days, source='graph', page_size=page_size, max_pages=max_pages, user_id=user_id, graph_settings=graph_settings)


_OUTLOOK_SUBJECT_PREFIX_RE = re.compile(r'^\s*(re|res|fw|fwd|enc|encaminhada|encaminhado)\s*:\s*', re.IGNORECASE)


def _outlook_normalize_subject(subject):
    """Remove prefixos de resposta/encaminhamento e normaliza para agrupar threads."""
    s = (subject or '').strip()
    prev = None
    while s and s != prev:
        prev = s
        s = _OUTLOOK_SUBJECT_PREFIX_RE.sub('', s).strip()
    return ' '.join(s.lower().split())


def _outlook_thread_key(email_data):
    """Chave de agrupamento de thread: conversation_id se houver, senão assunto normalizado."""
    conv = (email_data.get('conversation_id') or '').strip()
    if conv:
        return f'conv:{conv}'
    return f'subj:{_outlook_normalize_subject(email_data.get("subject"))}'


def _graph_get_me_email(access_token):
    """Retorna o email da conta conectada (para excluir o domínio próprio do matching)."""
    try:
        req = urllib.request.Request(
            'https://graph.microsoft.com/v1.0/me?$select=mail,userPrincipalName', method='GET')
        req.add_header('Authorization', f'Bearer {access_token}')
        req.add_header('Accept', 'application/json')
        with urllib.request.urlopen(req, timeout=10) as resp:
            me = json.loads(resp.read())
        return ((me.get('mail') or me.get('userPrincipalName') or '')).strip().lower()
    except Exception:
        return ''


def _build_outlook_stream_response(days=60, source='com', page_size=50, max_pages=10, user_id=1, graph_settings=None):
    def generate():
        own_domain = ''
        def evt(d):
            return f"data: {json.dumps(d, ensure_ascii=False)}\n\n"
        try:
            connecting_message = 'Conectando via Microsoft Graph...' if source == 'graph' else 'Lendo emails do Outlook via COM...'
            yield evt({'phase': 'connecting', 'message': connecting_message})

            try:
                if source == 'graph':
                    end_date = datetime.utcnow()
                    start_date = end_date - timedelta(days=days)
                    conn = get_db()
                    access_token = outlook_graph_get_valid_access_token(conn=conn, user_id=user_id, settings=graph_settings)
                    own_email = _graph_get_me_email(access_token)
                    own_domain = own_email.split('@', 1)[-1] if '@' in own_email else ''
                    emails = outlook_graph_fetch_messages(
                        access_token=access_token,
                        start_date=start_date,
                        end_date=end_date,
                        page_size=page_size,
                        max_pages=max_pages
                    )
                    conn.close()
                else:
                    yield evt({
                        'phase': 'reading',
                        'message': 'Consulta COM iniciada. Isso pode levar alguns minutos em caixas postais grandes.',
                        'count': 0
                    })
                    result_holder = {'emails': None, 'error': None}

                    def _run_com_fetch():
                        try:
                            result_holder['emails'] = _outlook_fetch_via_powershell(days)
                        except Exception as run_err:
                            result_holder['error'] = run_err

                    worker = threading.Thread(target=_run_com_fetch, daemon=True)
                    started_at = time.time()
                    worker.start()
                    while worker.is_alive():
                        elapsed = int(time.time() - started_at)
                        yield evt({
                            'phase': 'reading',
                            'message': f'Consulta COM em andamento... {elapsed}s decorridos.',
                            'count': 0
                        })
                        worker.join(timeout=5)

                    if result_holder.get('error') is not None:
                        raise result_holder['error']
                    emails = result_holder.get('emails') or []
            except OutlookOAuthError as e:
                logger.error(f'[Outlook][OAuth] Falha de autenticação no sync-stream ({source}): {e}')
                yield evt({'phase': 'error', 'error_type': 'oauth_authentication', 'message': str(e)})
                return
            except OutlookSyncError as e:
                logger.error(f'[Outlook][Sync] Falha na leitura Graph ({source}): {e}')
                yield evt({'phase': 'error', 'error_type': 'sync_failure', 'message': str(e)})
                return
            except Exception as e:
                logger.exception(f'[Outlook][Sync] Falha no conector {source}: {e}')
                yield evt({'phase': 'error', 'error_type': 'sync_failure', 'message': str(e)})
                return

            total_read = len(emails)
            yield evt({'phase': 'reading', 'message': f'{total_read} email(s) lidos.', 'count': total_read})

            if not emails:
                yield evt({'phase': 'done', 'total_read': 0, 'activities': [],
                           'message': f'Nenhum email encontrado nos últimos {days} dias.'})
                return

            yield evt({'phase': 'matching', 'message': f'Identificando contatos em {total_read} email(s)...', 'count': total_read})

            conn = get_db()
            c = conn.cursor()
            c.execute('SELECT id, name, email, photo_url FROM clients WHERE email IS NOT NULL AND TRIM(email) != ""')
            clients_map = {}
            domain_map = {}
            for row in c.fetchall():
                norm = (row['email'] or '').strip().lower()
                if not norm:
                    continue
                entry = {'id': row['id'], 'name': row['name'], 'photo_url': row['photo_url'] or ''}
                clients_map[norm] = entry
                domain = norm.split('@', 1)[-1] if '@' in norm else ''
                if domain and '.' in domain:
                    domain_map.setdefault(domain, []).append(entry)

            all_clients_for_select = []
            c.execute('SELECT id, name, company FROM clients ORDER BY name COLLATE NOCASE')
            for row in c.fetchall():
                all_clients_for_select.append({'id': row['id'], 'name': row['name'], 'company': row['company'] or ''})

            # IDs e conversas já processadas em sincronizações anteriores (dedup persistente)
            processed_ids = set()
            processed_conv_ids = set()
            try:
                c.execute('SELECT message_id, conversation_id FROM outlook_processed_emails WHERE user_id = ?', (user_id,))
                for _pr in c.fetchall():
                    if _pr['message_id']:
                        processed_ids.add(_pr['message_id'])
                    if _pr['conversation_id']:
                        processed_conv_ids.add(_pr['conversation_id'])
            except Exception:
                processed_ids = set()
                processed_conv_ids = set()

            def _match_client(cand_email):
                cand_email = (cand_email or '').strip().lower()
                if not cand_email:
                    return None
                client = clients_map.get(cand_email)
                if client:
                    return client
                domain = cand_email.split('@', 1)[-1] if '@' in cand_email else ''
                # Não usa fallback por domínio para o domínio corporativo do próprio usuário,
                # evitando atribuir emails internos a um único cliente que compartilha o domínio.
                if domain and domain != own_domain:
                    domain_clients = domain_map.get(domain, [])
                    if len(domain_clients) == 1:
                        return domain_clients[0]
                return None

            matched_groups = {}    # (client_id, thread_key) -> grupo
            unmatched_groups = {}  # (counterpart_email, thread_key) -> grupo

            for email_data in emails:
                subject = (email_data.get('subject') or '').strip()
                email_date = (email_data.get('date') or '').strip()
                direction = email_data.get('direction', 'received')
                sender = email_data.get('sender') or {}
                recipients = email_data.get('recipients') or []
                body_preview = (email_data.get('body_preview') or '').strip()
                message_id = (email_data.get('message_id') or '').strip()
                conversation_id = (email_data.get('conversation_id') or '').strip()
                if not email_date:
                    continue
                # Pula mensagens já importadas em sincronizações anteriores
                if message_id and message_id in processed_ids:
                    continue

                # Candidatos = remetente (recebido) ou destinatários TO (enviado).
                # CC é ignorado de propósito (item de quem está só em cópia não vira atividade).
                if direction == 'received':
                    candidates = [sender] if sender.get('email') else []
                    label = 'De'
                else:
                    candidates = [r for r in recipients if r.get('email')]
                    label = 'Para'

                tkey = _outlook_thread_key(email_data)

                # Em emails enviados todos os destinatários TO são candidatos;
                # cria um grupo por cliente correspondente (multi-recipiente).
                matched_for_email = []
                for candidate in candidates:
                    c_match = _match_client(candidate.get('email'))
                    if c_match:
                        matched_for_email.append((c_match, candidate))

                if matched_for_email:
                    for m_client, chosen in matched_for_email:
                        msg_entry = {
                            'direction': direction,
                            'name': (chosen.get('name') or '').strip(),
                            'email': (chosen.get('email') or '').strip().lower(),
                            'date': email_date,
                            'subject': subject,
                            'body_preview': body_preview,
                            'message_id': message_id,
                            'conversation_id': conversation_id,
                        }
                        gkey = (m_client['id'], tkey)
                        grp = matched_groups.get(gkey)
                        if not grp:
                            grp = {
                                'client_id': m_client['id'],
                                'client_name': m_client['name'],
                                'client_photo_url': m_client['photo_url'],
                                'counterpart_label': label,
                                'conversation_id': conversation_id,
                                'messages': [],
                                'message_ids': [],
                            }
                            matched_groups[gkey] = grp
                        grp['messages'].append(msg_entry)
                        if message_id:
                            grp['message_ids'].append(message_id)
                elif candidates:
                    ref = candidates[0]
                    cemail = (ref.get('email') or '').strip().lower()
                    cdomain = cemail.split('@', 1)[-1] if '@' in cemail else ''
                    # Ignora ruído interno (mesmo domínio do usuário) na lista "sem cliente"
                    if cdomain and cdomain == own_domain:
                        continue
                    msg_entry = {
                        'direction': direction,
                        'name': (ref.get('name') or '').strip(),
                        'email': cemail,
                        'date': email_date,
                        'subject': subject,
                        'body_preview': body_preview,
                        'message_id': message_id,
                        'conversation_id': conversation_id,
                    }
                    gkey = (cemail, tkey)
                    grp = unmatched_groups.get(gkey)
                    if not grp:
                        grp = {
                            'counterpart_label': label,
                            'counterpart_name': (ref.get('name') or cemail).strip(),
                            'counterpart_email': cemail,
                            'conversation_id': conversation_id,
                            'messages': [],
                            'message_ids': [],
                        }
                        unmatched_groups[gkey] = grp
                    grp['messages'].append(msg_entry)
                    if message_id:
                        grp['message_ids'].append(message_id)

            # Dedup de grupo: remove threads cujo conteúdo todo já foi importado antes.
            # grp['message_ids'] só contém IDs *novos* (os já processados foram filtrados
            # individualmente acima). Se a lista está vazia, confirma via conversation_id.
            def _is_group_new(grp):
                if grp.get('message_ids'):
                    return True
                conv_id = (grp.get('conversation_id') or '').strip()
                return not (conv_id and conv_id in processed_conv_ids)

            matched_groups = {k: v for k, v in matched_groups.items() if v['messages'] and _is_group_new(v)}
            unmatched_groups = {k: v for k, v in unmatched_groups.items() if v['messages'] and _is_group_new(v)}

            def _finalize(grp):
                msgs = sorted(grp['messages'], key=lambda m: m['date'])
                latest = msgs[-1]
                grp['messages'] = msgs[:12]
                grp['subject'] = latest['subject']
                grp['date'] = latest['date']
                grp['direction'] = latest['direction']
                grp['counterpart_name'] = latest['name'] or grp.get('counterpart_name') or latest['email']
                grp['counterpart_email'] = latest['email'] or grp.get('counterpart_email') or ''
                grp['body_preview'] = latest['body_preview']
                grp['email_count'] = len(msgs)
                return grp

            activities = [_finalize(g) for g in matched_groups.values() if g['messages']]
            unmatched = [_finalize(g) for g in unmatched_groups.values() if g['messages']]
            activities.sort(key=lambda g: g['date'], reverse=True)
            unmatched.sort(key=lambda g: g['date'], reverse=True)

            conn.close()

            msg = f'{len(activities)} atividade(s) agrupada(s) por assunto em {total_read} email(s) lidos.'
            if unmatched:
                msg += f' {len(unmatched)} thread(s) sem cliente correspondente.'
            yield evt({
                'phase': 'done',
                'total_read': total_read,
                'activities': activities,
                'unmatched': unmatched,
                'all_clients': all_clients_for_select,
                'message': msg
            })
        except Exception as e:
            logger.exception(f'[ERROR] SSE /api/outlook/sync-stream ({source}): {e}')
            yield evt({'phase': 'error', 'error_type': 'sync_failure', 'message': f'Erro inesperado: {str(e)}'})
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception as e:
                logger.debug(f'[_finalize] exceção ignorada: {e}')

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'}
    )


_outlook_confirm_tasks = {}
_outlook_confirm_tasks_lock = threading.Lock()


def _outlook_confirm_task_set(task_id, update):
    with _outlook_confirm_tasks_lock:
        if task_id not in _outlook_confirm_tasks:
            _outlook_confirm_tasks[task_id] = {}
        _outlook_confirm_tasks[task_id].update(update)
        snapshot = dict(_outlook_confirm_tasks[task_id])
    _bg_task_persist(task_id, 'outlook_confirm', snapshot)


def _outlook_confirm_task_get(task_id):
    with _outlook_confirm_tasks_lock:
        return dict(_outlook_confirm_tasks.get(task_id) or {})


def _outlook_confirm_task_cleanup(task_id, delay=300):
    def _do():
        time.sleep(delay)
        with _outlook_confirm_tasks_lock:
            _outlook_confirm_tasks.pop(task_id, None)
    threading.Thread(target=_do, daemon=True).start()


def _detect_followup_from_text(c, client_id, activity_id, client_name, text):
    """Detecta follow-up combinado via LLM no mesmo formato JSON do sync WhatsApp
    ({"followup": {"data", "titulo"}}) e cria o compromisso na agenda (Bloco 7).
    Retorna 1 se criou compromisso, 0 caso contrário."""
    if not text:
        return 0
    today_str = datetime.now().strftime('%Y-%m-%d')
    raw = _llm_prompt(
        f"A data de hoje é {today_str}. Analise a comunicação abaixo entre o usuário e '{client_name}'. "
        'Retorne SOMENTE um JSON válido no formato '
        '{"followup": {"data": "YYYY-MM-DD", "titulo": "descrição curta do retorno/compromisso combinado"}}. '
        "Preencha 'followup' apenas se houver data combinada de retorno, reunião ou ligação "
        '(resolva datas relativas, ex. "amanhã", a partir de hoje). '
        'Se NÃO houver compromisso combinado, retorne {"followup": null}.\n\n'
        f'{text[:3000]}',
        log_tag='FollowupDetect'
    )
    if not raw:
        return 0
    try:
        parsed = _extract_json_object_from_text(raw) or json.loads(raw)
    except Exception:
        return 0
    fu = (parsed or {}).get('followup')
    if not isinstance(fu, dict):
        return 0
    due_date = (fu.get('data') or fu.get('date') or '').strip()
    if not re.match(r'^\d{4}-\d{2}-\d{2}$', due_date):
        return 0
    title = ((fu.get('titulo') or fu.get('title') or '').strip() or f'Retorno combinado com {client_name}')[:120]
    # Não duplica compromisso já criado para a mesma atividade/data (ex.: via regex)
    c.execute('SELECT id FROM commitments WHERE activity_id = ? AND due_date = ?', (activity_id, due_date))
    if c.fetchone():
        return 0
    c.execute(
        '''INSERT INTO commitments (client_id, activity_id, title, notes, due_date, source_type)
           VALUES (?, ?, ?, ?, ?, ?)''',
        (client_id, activity_id, title, (text or '')[:500], due_date, 'outlook')
    )
    return 1


def _outlook_send_mail(to, subject, body_html, attachments=None):
    """Envia e-mail pelo Outlook do próprio usuário via Microsoft Graph
    (Bloco 13). Se `to` for vazio, envia para o e-mail do próprio usuário.
    attachments: [{'name', 'content_bytes' (base64), 'content_type'}]."""
    graph_settings = _graph_make_settings(redirect_uri=_graph_redirect_uri())
    conn = get_db()
    try:
        access_token = outlook_graph_get_valid_access_token(conn=conn, user_id=1, settings=graph_settings)
    finally:
        conn.close()
    recipient = (to or '').strip() or _graph_get_me_email(access_token)
    if not recipient:
        raise OutlookSyncError('Não foi possível determinar o destinatário do e-mail.')
    outlook_graph_send_mail(access_token, recipient, subject, body_html, attachments)
    logger.info(f'[Outlook] E-mail "{subject}" enviado para {recipient}')
    return recipient


def _outlook_confirm_async(task_id, activities_to_import):
    try:
        total = len(activities_to_import)
        imported = 0
        skipped = 0
        commitments_created = 0
        status_suggestions = []
        kanban_suggestions = []
        imported_items = []
        skipped_items = []

        conn = get_db()
        c = conn.cursor()

        for i, act in enumerate(activities_to_import):
            _outlook_confirm_task_set(task_id, {
                'step': f'Processando thread {i + 1}/{total}...',
                'progress': 10 + int((i / total) * 75)
            })

            client_id = act.get('client_id')
            subject = (act.get('subject') or '').strip()
            email_date = (act.get('date') or '').strip()
            label = act.get('counterpart_label', 'De')
            cname = (act.get('counterpart_name') or '').strip()
            cemail = (act.get('counterpart_email') or '').strip()
            body_preview = (act.get('body_preview') or '').strip()
            messages = act.get('messages') or []
            message_ids = [m for m in (act.get('message_ids') or []) if m]
            conv_id = (act.get('conversation_id') or '').strip()

            if not client_id or not email_date:
                continue

            # Dedup persistente: se todas as mensagens da thread já foram processadas, pula
            if message_ids:
                placeholders = ','.join('?' * len(message_ids))
                c.execute(
                    f'SELECT COUNT(*) AS n FROM outlook_processed_emails WHERE message_id IN ({placeholders})',
                    tuple(message_ids)
                )
                row_n = c.fetchone()
                if row_n and row_n['n'] >= len(message_ids):
                    skipped += 1
                    skipped_items.append({'client_name': act.get('client_name') or cname, 'subject': subject})
                    continue

            # Dedup contra atividade já existente (mesma thread/cliente no mesmo minuto)
            date_minute = email_date[:16]
            c.execute(
                '''SELECT id FROM activities WHERE client_id = ? AND contact_type = 'Email'
                   AND strftime('%Y-%m-%dT%H:%M', activity_date) = ?
                   AND information LIKE ? LIMIT 1''',
                (client_id, date_minute, f'{subject[:100]}%')
            )
            if c.fetchone():
                skipped += 1
                skipped_items.append({'client_name': act.get('client_name') or cname, 'subject': subject})
                for mid in message_ids:
                    try:
                        c.execute('INSERT OR IGNORE INTO outlook_processed_emails '
                                  '(user_id, message_id, conversation_id, client_id) VALUES (1, ?, ?, ?)',
                                  (mid, conv_id, client_id))
                    except Exception as e:
                        logger.debug(f'[_outlook_confirm_async] exceção ignorada: {e}')
                continue

            # Monta o texto consolidado da thread para o resumo do LLM
            if not messages:
                messages = [{'direction': act.get('direction', 'received'), 'name': cname,
                             'email': cemail, 'date': email_date, 'subject': subject,
                             'body_preview': body_preview}]
            thread_lines = []
            participants = []
            for m in messages:
                nm = (m.get('name') or m.get('email') or '').strip()
                if nm and nm not in participants:
                    participants.append(nm)
                d = 'enviado' if m.get('direction') == 'sent' else 'recebido'
                thread_lines.append(
                    f"- ({d}) {nm}: {(m.get('subject') or '').strip()} — "
                    f"{(m.get('body_preview') or '').strip()[:400]}"
                )
            thread_text = '\n'.join(thread_lines)

            summary = _outlook_call_llm(
                'Você é um assistente de CRM. Resuma a thread de emails abaixo em 1 a 2 frases objetivas, '
                'em português, dizendo quem fez o quê e qual o assunto ou decisão principal. '
                'Se vários participantes fizeram a mesma ação (ex.: aceitaram um convite), agrupe-os numa só frase. '
                'Não comece com "Resumo:" e não use aspas.\n\n'
                f'Contraparte principal: {cname}\n'
                f'Thread:\n{thread_text}'
            )
            summary = summary.strip() if summary else None

            participants_str = ', '.join(participants[:6]) or cname
            first_line = subject if len(messages) <= 1 else f'{subject}  ({len(messages)} emails na thread)'
            info_parts = [first_line, f'{label}: {participants_str}']
            if summary:
                info_parts.append(f'Resumo: {summary}')
            elif body_preview:
                info_parts.append(body_preview[:300])
            information = '\n'.join(info_parts)

            analysis_text = (summary or body_preview or subject or '')[:600]

            c.execute(
                '''INSERT INTO activities (client_id, contact_type, information, activity_date)
                   VALUES (?, 'Email', ?, ?)''',
                (client_id, information, email_date)
            )
            activity_id = c.lastrowid
            c.execute(
                '''UPDATE clients SET last_activity_date = ?
                   WHERE id = ? AND (last_activity_date IS NULL OR last_activity_date < ?)''',
                (email_date, client_id, email_date)
            )
            imported += 1
            imported_items.append({
                'client_name': act.get('client_name') or cname,
                'subject': subject,
                'date': email_date[:10] if email_date else '',
            })

            # Registra mensagens processadas para não reimportar nas próximas sincronizações
            for mid in message_ids:
                try:
                    c.execute('INSERT OR IGNORE INTO outlook_processed_emails '
                              '(user_id, message_id, conversation_id, client_id, activity_id) VALUES (1, ?, ?, ?, ?)',
                              (mid, conv_id, client_id, activity_id))
                except Exception as e:
                    logger.debug(f'[_outlook_confirm_async] exceção ignorada: {e}')

            new_commitments = create_commitments_from_activity(c, client_id, activity_id, information)
            commitments_created += len(new_commitments)
            # Follow-up combinado detectado via LLM (mesmo formato do sync WhatsApp)
            try:
                commitments_created += _detect_followup_from_text(c, client_id, activity_id, cname, thread_text)
            except Exception as e:
                logger.debug(f'[FollowupDetect] falha na thread de {cname}: {e}')

            if analysis_text:
                stage_raw = _outlook_call_llm(
                    f'Analise esta interação e classifique o momento do relacionamento comercial.\n'
                    f'Assunto: {subject}\nTexto: {analysis_text}\n\n'
                    'Retorne SOMENTE um destes valores (sem mais texto):\n'
                    'lead_quente | negociando | pos_venda | sem_mudanca\n'
                    'lead_quente = cliente demonstrou interesse ativo\n'
                    'negociando = discussão de proposta, contrato ou condições comerciais\n'
                    'pos_venda = suporte, onboarding, entrega pós-compra\n'
                    'sem_mudanca = email de rotina sem impacto comercial'
                )
                if stage_raw:
                    stage = stage_raw.strip().split()[0].lower().replace('-', '_')
                    if stage in {'lead_quente', 'negociando', 'pos_venda'}:
                        c.execute('SELECT name, relationship_stage FROM clients WHERE id = ?', (client_id,))
                        row = c.fetchone()
                        if row:
                            current = (row['relationship_stage'] or '').strip()
                            if current != stage:
                                stage_labels = {
                                    'lead_quente': 'Lead Quente',
                                    'negociando': 'Negociando',
                                    'pos_venda': 'Pós-venda'
                                }
                                status_suggestions.append({
                                    'client_id': client_id,
                                    'client_name': row['name'],
                                    'current_stage': current,
                                    'suggested_stage': stage,
                                    'suggested_label': stage_labels[stage],
                                    'reason': f'Detectado via email: {subject[:60]}'
                                })

                c.execute(
                    '''SELECT kc.id, kc.title, kcol.title as col_title, kc.column_id
                       FROM kanban_cards kc
                       JOIN kanban_columns kcol ON kcol.id = kc.column_id
                       WHERE kc.contact_id = ?
                       ORDER BY kc.updated_at DESC LIMIT 5''',
                    (client_id,)
                )
                cards = [dict_from_row(r) for r in c.fetchall()]
                if cards:
                    c.execute('SELECT id, title FROM kanban_columns ORDER BY display_order')
                    columns = [dict_from_row(r) for r in c.fetchall()]
                    cards_text = '\n'.join([f'- ID {card["id"]}: "{card["title"]}" (coluna: {card["col_title"]})' for card in cards])
                    cols_text = ' | '.join([col['title'] for col in columns])
                    kanban_raw = _outlook_call_llm(
                        f'Interação:\nAssunto: {subject}\nTexto: {analysis_text}\n\n'
                        f'Cards do Kanban deste cliente:\n{cards_text}\n\n'
                        f'Colunas disponíveis: {cols_text}\n\n'
                        'Algum card precisa ser movido com base neste email? '
                        'Retorne SOMENTE JSON válido: '
                        '{"card_id": <id inteiro ou null>, "new_column": "<título exato da coluna ou null>", "reason": "<motivo curto>"} '
                        'Use null se não houver mudança necessária.'
                    )
                    if kanban_raw:
                        try:
                            raw_str = kanban_raw.strip()
                            if raw_str.startswith('{'):
                                kanban_data = json.loads(raw_str)
                            else:
                                import re as _re
                                m = _re.search(r'\{[^}]+\}', raw_str)
                                kanban_data = json.loads(m.group()) if m else {}
                            if isinstance(kanban_data, dict) and kanban_data.get('card_id') and kanban_data.get('new_column'):
                                card_id = int(kanban_data['card_id'])
                                new_col_title = kanban_data['new_column'].strip()
                                matching_card = next((card for card in cards if card['id'] == card_id), None)
                                matching_col = next((col for col in columns if col['title'].lower() == new_col_title.lower()), None)
                                if matching_card and matching_col and matching_card['column_id'] != matching_col['id']:
                                    kanban_suggestions.append({
                                        'card_id': card_id,
                                        'card_title': matching_card['title'],
                                        'current_column': matching_card['col_title'],
                                        'suggested_column': matching_col['title'],
                                        'suggested_column_id': matching_col['id'],
                                        'reason': (kanban_data.get('reason') or '')[:120],
                                        'client_id': client_id
                                    })
                        except Exception as e:
                            logger.debug(f'[_outlook_confirm_async] exceção ignorada: {e}')

        conn.commit()
        conn.close()

        msg = f'{imported} atividade(s) registrada(s).'
        if commitments_created:
            msg += f' {commitments_created} compromisso(s) criado(s) na agenda.'
        if skipped:
            msg += f' {skipped} duplicata(s) ignorada(s).'

        _outlook_confirm_task_set(task_id, {
            'status': 'done',
            'step': 'Concluído!',
            'progress': 100,
            'result': {
                'imported': imported,
                'imported_items': imported_items,
                'skipped': skipped,
                'skipped_items': skipped_items,
                'commitments_created': commitments_created,
                'status_suggestions': status_suggestions,
                'kanban_suggestions': kanban_suggestions,
                'message': msg
            }
        })
        _outlook_confirm_task_cleanup(task_id)
    except Exception as e:
        logger.exception(f'[ERROR] _outlook_confirm_async task={task_id}: {e}')
        _outlook_confirm_task_set(task_id, {'status': 'error', 'step': str(e), 'progress': 0})
        _outlook_confirm_task_cleanup(task_id)


# ── Outlook Add-in (task pane EWS) ──────────────────────────────────────────

_addin_pending_data = None
_addin_pending_lock = threading.Lock()
_addin_pending_expires = 0.0


def _addin_set_pending(data):
    global _addin_pending_data, _addin_pending_expires
    with _addin_pending_lock:
        _addin_pending_data = data
        _addin_pending_expires = time.time() + 600  # 10 min TTL


def _addin_get_pending():
    global _addin_pending_data, _addin_pending_expires
    with _addin_pending_lock:
        if _addin_pending_data is None or time.time() > _addin_pending_expires:
            _addin_pending_data = None
            return None
        return dict(_addin_pending_data)


def _addin_clear_pending():
    global _addin_pending_data
    with _addin_pending_lock:
        _addin_pending_data = None


def _outlook_match_emails(emails_data, conn):
    """Faz matching de uma lista de emails contra clientes cadastrados.
    Retorna (activities, unmatched, all_clients_for_select).
    """
    c = conn.cursor()
    c.execute('SELECT id, name, email, photo_url FROM clients WHERE email IS NOT NULL AND TRIM(email) != ""')
    clients_map = {}
    domain_map = {}
    for row in c.fetchall():
        norm = (row['email'] or '').strip().lower()
        if not norm:
            continue
        entry = {'id': row['id'], 'name': row['name'], 'photo_url': row['photo_url'] or ''}
        clients_map[norm] = entry
        domain = norm.split('@', 1)[-1] if '@' in norm else ''
        if domain and '.' in domain:
            domain_map.setdefault(domain, []).append(entry)

    c.execute('SELECT id, name, company FROM clients ORDER BY name COLLATE NOCASE')
    all_clients = [{'id': r['id'], 'name': r['name'], 'company': r['company'] or ''} for r in c.fetchall()]

    activities = []
    unmatched = []
    seen_keys = set()

    for email_data in emails_data:
        subject = (email_data.get('subject') or '').strip()
        email_date = (email_data.get('date') or '').strip()
        direction = email_data.get('direction', 'received')
        sender = email_data.get('sender') or {}
        recipients = email_data.get('recipients') or []
        body_preview = (email_data.get('body_preview') or '').strip()
        if not email_date:
            continue

        if direction == 'received':
            candidates = [sender] if sender.get('email') else []
            label = 'De'
        else:
            candidates = [r for r in recipients if r.get('email')]
            label = 'Para'

        matched = False
        for candidate in candidates:
            cand_email = (candidate.get('email') or '').strip().lower()
            if not cand_email:
                continue
            client = clients_map.get(cand_email)
            if not client:
                domain = cand_email.split('@', 1)[-1] if '@' in cand_email else ''
                domain_clients = domain_map.get(domain, [])
                if len(domain_clients) == 1:
                    client = domain_clients[0]
            if not client:
                continue

            dedup_key = (client['id'], email_date[:16], subject[:100])
            if dedup_key in seen_keys:
                matched = True
                continue
            c.execute(
                '''SELECT id FROM activities WHERE client_id = ? AND contact_type = 'Email'
                   AND strftime('%Y-%m-%dT%H:%M', activity_date) = ?
                   AND information LIKE ? LIMIT 1''',
                (client['id'], email_date[:16], f'{subject[:100]}%')
            )
            if c.fetchone():
                matched = True
                continue
            seen_keys.add(dedup_key)
            matched = True
            activities.append({
                'client_id': client['id'],
                'client_name': client['name'],
                'client_photo_url': client['photo_url'],
                'subject': subject,
                'date': email_date,
                'direction': direction,
                'counterpart_label': label,
                'counterpart_name': (candidate.get('name') or cand_email).strip(),
                'counterpart_email': cand_email,
                'body_preview': body_preview
            })
            break

        if not matched and candidates:
            first = candidates[0]
            unmatched.append({
                'subject': subject,
                'date': email_date,
                'direction': direction,
                'counterpart_label': label,
                'counterpart_name': (first.get('name') or first.get('email') or '').strip(),
                'counterpart_email': (first.get('email') or '').strip().lower(),
                'body_preview': body_preview[:180]
            })

    # Caixa de respostas pendentes (Bloco 6): e-mails de clientes sem resposta sua
    try:
        _inbound_feed_from_outlook(c, emails_data, clients_map)
        conn.commit()
    except Exception as e:
        logger.debug(f'[Inbound] Falha ao alimentar pendências via Outlook: {e}')

    return activities, unmatched, all_clients


def _itoca_base_update_async(task_id, incremental):
    try:
        def progress_cb(pct, msg):
            _bg_task_set(task_id, {'step': msg, 'progress': pct})

        _bg_task_set(task_id, {'step': 'Iniciando indexação...', 'progress': 5})
        result = _itoca_update_cached_base(progress_cb=progress_cb, incremental=incremental)

        changed = result.get('changed', True)
        msg = 'Base iToca atualizada com sucesso.' if changed else 'Base já estava atualizada. Nenhuma alteração detectada.'
        _bg_task_set(task_id, {
            'step': 'Concluído!', 'progress': 100, 'status': 'done',
            'result': {
                'message': msg,
                'items_count': len(result.get('items', [])),
                'updated_at': result.get('updated_at', ''),
                'incremental': result.get('incremental', False),
                'changed': changed
            }
        })
    except Exception as e:
        logger.exception(f'[iToca][BaseUpdate][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _bg_task_cleanup(task_id)


def _itoca_ask_async(task_id, question, session_id, snapshot_items, updated_at, history_rows):
    """Processa a pergunta do iToca em background. Atualiza _bg_tasks com progresso em tempo real."""
    try:
        _q_lower = question.lower()
        _analytical_kws = {'target', 'todas', 'todos', 'resumo', 'resumir', 'panorama', 'geral',
                           'visao geral', 'visão geral', 'analise', 'análise', 'situação', 'situacao',
                           'relacionamento', 'evolucao', 'evolução', 'overview', 'mapa', 'balanço', 'balanco'}
        is_analytical = any(kw in _q_lower for kw in _analytical_kws)
        snapshot_limit = 30 if is_analytical else 15
        live_limit = 15 if is_analytical else 8
        context_max = 45 if is_analytical else 25

        # 1. Busca na base de conhecimento Wiki (rápido — já indexado)
        _bg_task_set(task_id, {'step': '📚 Consultando base de conhecimento...', 'progress': 20})
        snapshot_rows = _itoca_search_in_cached_snapshot(question, snapshot_items, limit=snapshot_limit)

        # 2. Busca ao vivo no banco de dados (contatos, atividades, agenda...)
        _bg_task_set(task_id, {'step': '📊 Buscando no banco de dados...', 'progress': 35})
        live_rows = _itoca_search_context(question, limit=live_limit)

        # 3. Para queries analíticas/target: força inclusão de todas as contas target
        target_rows = []
        if is_analytical and 'target' in _q_lower:
            _bg_task_set(task_id, {'step': '🎯 Carregando contas target...', 'progress': 45})
            try:
                _conn_t = get_db()
                _cur_t = _conn_t.cursor()
                _cur_t.execute('''
                    SELECT ac.id, ac.name, ac.sector, ac.is_target,
                           MAX(act.activity_date) as last_activity,
                           COUNT(DISTINCT act.id) as total_activities,
                           COUNT(DISTINCT ap.id) as total_services,
                           COUNT(DISTINCT amc.id) as total_contacts
                    FROM accounts ac
                    LEFT JOIN activities act ON act.account_id = ac.id
                    LEFT JOIN account_presences ap ON ap.account_id = ac.id
                    LEFT JOIN account_main_contacts amc ON amc.account_id = ac.id
                    WHERE ac.is_target = 1
                    GROUP BY ac.id
                    ORDER BY last_activity DESC NULLS LAST
                ''')
                for _row in _cur_t.fetchall():
                    _rd = dict_from_row(_row)
                    _parts = ['classificacao: conta-alvo (target)']
                    if _rd.get('name'): _parts.append(f'empresa: {_rd["name"]}')
                    if _rd.get('sector'): _parts.append(f'setor: {_rd["sector"]}')
                    if _rd.get('last_activity'):
                        try:
                            _dt = datetime.strptime(_rd['last_activity'][:10], '%Y-%m-%d')
                            _parts.append(f'ultimo_contato: {_dt.strftime("%d/%m/%Y")}')
                        except Exception:
                            _parts.append(f'ultimo_contato: {_rd["last_activity"]}')
                    else:
                        _parts.append('ultimo_contato: sem registros')
                    if _rd.get('total_activities'): _parts.append(f'total_interacoes: {_rd["total_activities"]}')
                    if _rd.get('total_services'): _parts.append(f'servicos_stefanini_cadastrados: {_rd["total_services"]}')
                    if _rd.get('total_contacts'): _parts.append(f'total_contatos_mapeados: {_rd["total_contacts"]}')
                    _snip = ' | '.join(_parts)
                    if _snip:
                        target_rows.append({'table': 'accounts', 'id': _rd.get('id'), 'snippet': _snip, 'search_text': _snip.lower()})
                _conn_t.close()
            except Exception as _te:
                logger.warning(f'[iToca] Erro ao buscar contas target: {_te}')

        # 4. Para queries analíticas gerais: painel de stats do banco
        stats_rows = []
        if is_analytical:
            try:
                _conn_s = get_db()
                _cur_s = _conn_s.cursor()
                _cur_s.execute('SELECT COUNT(*) as n FROM accounts')
                _total_ac = (_cur_s.fetchone() or [0])[0]
                _cur_s.execute('SELECT COUNT(*) as n FROM accounts WHERE is_target = 1')
                _total_target = (_cur_s.fetchone() or [0])[0]
                _cur_s.execute('SELECT COUNT(*) as n FROM clients')
                _total_cl = (_cur_s.fetchone() or [0])[0]
                _cur_s.execute('SELECT COUNT(*) as n FROM activities')
                _total_act = (_cur_s.fetchone() or [0])[0]
                _cur_s.execute('SELECT COUNT(*) as n FROM account_presences')
                _total_srv = (_cur_s.fetchone() or [0])[0]
                _cur_s.execute("SELECT COUNT(*) as n FROM activities WHERE activity_date >= date('now', '-30 days')")
                _act_30d = (_cur_s.fetchone() or [0])[0]
                _conn_s.close()
                _stats_snip = (
                    f'total_contas: {_total_ac} | contas_target: {_total_target} | '
                    f'total_contatos: {_total_cl} | total_interacoes: {_total_act} | '
                    f'interacoes_ultimos_30dias: {_act_30d} | servicos_cadastrados: {_total_srv}'
                )
                stats_rows.append({'table': 'user_profile', 'id': None, 'snippet': f'PAINEL_GERAL: {_stats_snip}', 'search_text': _stats_snip.lower()})
            except Exception as _se:
                logger.warning(f'[iToca] Erro ao calcular stats analíticos: {_se}')

        # Mescla: prioriza target > stats > live (banco) > snapshot, sem duplicatas
        seen_keys = set()
        context_rows = []
        for item in (target_rows + stats_rows + live_rows):
            key = f"{item.get('table')}:{item.get('id')}:{str(item.get('snippet',''))[:60]}"
            if key not in seen_keys:
                seen_keys.add(key)
                context_rows.append(item)
        for item in snapshot_rows:
            key = f"{item.get('table')}:{item.get('id')}:{str(item.get('snippet',''))[:60]}"
            if key not in seen_keys:
                seen_keys.add(key)
                context_rows.append(item)
        context_rows = context_rows[:context_max]

        # 5. Chamada ao LLM (parte mais demorada)
        _bg_task_set(task_id, {'step': '🤖 Integrando com IA...', 'progress': 60})
        llm_result = _itoca_call_sai_llm(question, context_rows, history_rows=history_rows)

        _bg_task_set(task_id, {'step': '💬 Processando resposta...', 'progress': 85})
        answer = llm_result.get('answer', '')
        confidence = llm_result.get('confidence_percent', 0)
        needs_ref = llm_result.get('needs_refinement', False)
        ref_hint = llm_result.get('refinement_hint', '')

        # 6. Salva resposta da IA no histórico (usuário já foi salvo antes de iniciar a task)
        if session_id:
            try:
                conn_h = get_db()
                c_h = conn_h.cursor()
                c_h.execute(
                    'INSERT INTO itoca_chat_history (session_id, role, content, confidence_percent, needs_refinement, refinement_hint) VALUES (?, ?, ?, ?, ?, ?)',
                    (session_id, 'assistant', answer, confidence, 1 if needs_ref else 0, ref_hint)
                )
                conn_h.commit()
                conn_h.close()
            except Exception as he:
                logger.warning(f'[iToca] Erro ao salvar resposta no histórico: {he}')

        # 7. Detecção de intenção de ação (não bloqueia a resposta)
        suggested_action = None
        if llm_result.get('llm_used') and not needs_ref:
            try:
                import queue as _queue
                result_queue = _queue.Queue()
                def _run_detector():
                    try:
                        result_queue.put(_itoca_detect_action_intent(question, answer))
                    except Exception as _e:
                        result_queue.put(None)
                t = threading.Thread(target=_run_detector, daemon=True)
                t.start()
                t.join(timeout=12)
                if not result_queue.empty():
                    det = result_queue.get_nowait()
                    if det and det.get('action_type') and float(det.get('confidence', 0)) >= 0.75:
                        action_type = det['action_type']
                        fields = det.get('fields') or {}
                        _REQUIRED_FIELDS = {
                            'kanban_card': ['title'], 'activity': ['contact_name', 'description'],
                            'new_contact': ['name', 'company'], 'environment_mapping': ['company', 'information'],
                            'wiki_entry': ['title', 'content'], 'commitment': ['title', 'due_date'],
                        }
                        required = _REQUIRED_FIELDS.get(action_type, [])
                        has_minimum = all(fields.get(f) and str(fields[f]).strip() for f in required)
                        desc_field = (fields.get('description') or fields.get('information') or '').strip().lower()
                        question_lower = question.strip().lower()
                        is_echo = desc_field and (desc_field == question_lower or (len(desc_field) > 20 and question_lower.startswith(desc_field[:40])))
                        if has_minimum and not is_echo:
                            suggested_action = {'action_type': action_type, 'label': det['label'], 'confidence': det['confidence'], 'fields': fields}
                        else:
                            logger.debug(f'[iToca][ActionDetector] Ação {action_type!r} descartada: has_minimum={has_minimum}, is_echo={is_echo}')
            except Exception as det_err:
                logger.warning(f'[iToca] Erro no detector de intenção: {det_err}')

        _bg_task_set(task_id, {
            'status': 'done', 'step': '✅ Concluído!', 'progress': 100,
            'result': {
                'answer': answer, 'confidence_percent': confidence,
                'needs_refinement': needs_ref, 'refinement_hint': ref_hint,
                'llm_used': llm_result.get('llm_used', False),
                'items_found': len(context_rows), 'sources': context_rows,
                'base_updated_at': updated_at, 'base_ready': True,
                'suggested_action': suggested_action,
            }
        })
        _bg_task_cleanup(task_id, delay=300)

    except Exception as e:
        logger.exception(f'[iToca][Ask][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
        _bg_task_cleanup(task_id, delay=300)


def _autopic_get_role_via_llm(name, company):
    """Usa o SAI (ou OpenRouter como fallback) para descobrir o cargo da pessoa na empresa.
    Retorna string com o cargo, ou None se não conseguir.
    """
    llm_prompt = (
        f"Qual é o cargo ou função profissional de '{name}' na empresa '{company}'? "
        "Retorne SOMENTE o cargo em português, sem texto adicional, sem aspas, sem ponto final. "
        "Exemplos de resposta: CEO, Diretor Comercial, Gerente de Marketing, Engenheiro de Software. "
        "Se não souber com certeza, retorne null."
    )
    raw = _sai_simple_prompt(llm_prompt)
    if raw:
        role = raw.strip().split('\n')[0].strip(' "\'.')
        if role.lower() not in ('null', 'none', '', 'não sei', 'desconhecido'):
            logger.info(f'[AutoPic][LLM] Cargo encontrado via SAI: {role!r}')
            return role[:80]

    # Fallback: OpenRouter
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
        model = (or_settings.get('openrouter_model') or 'stepfun/step-3.5-flash:free').strip() or 'stepfun/step-3.5-flash:free'
        site_url = (or_settings.get('openrouter_site_url') or 'http://localhost').strip() or 'http://localhost'
        app_name = (or_settings.get('openrouter_app_name') or 'TocaDoCoelho').strip() or 'TocaDoCoelho'
        try:
            or_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'Você é um especialista em mercado corporativo. Responda SOMENTE com o cargo, sem texto adicional.'},
                    {'role': 'user', 'content': llm_prompt}
                ],
                'temperature': 0.1
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(or_payload, ensure_ascii=False).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=20) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            choices = data.get('choices') or []
            role = ((choices[0].get('message') or {}).get('content', '') if choices else '').strip().split('\n')[0].strip(' "\'.')
            if role and role.lower() not in ('null', 'none', '', 'não sei', 'desconhecido'):
                logger.info(f'[AutoPic][LLM] Cargo encontrado via OpenRouter: {role!r}')
                return role[:80]
        except Exception as e:
            logger.warning(f'[AutoPic][LLM] Falha ao buscar cargo via OpenRouter: {e}')

    return None


def _save_photo_from_base64(data_url, person_name='linkedin'):
    """Recebe um data URL base64 (data:image/...;base64,...), salva localmente e retorna '/uploads/<filename>'.
    Retorna None se o data_url não for válido."""
    try:
        if not data_url or not data_url.startswith('data:image/'):
            return None
        header, encoded = data_url.split(',', 1)
        mime_type = header.split(';')[0].replace('data:', '')
        ext_map = {'image/jpeg': '.jpg', 'image/png': '.png', 'image/webp': '.webp', 'image/gif': '.gif'}
        ext = ext_map.get(mime_type, '.jpg')
        img_data = base64.b64decode(encoded)
        if len(img_data) > 6 * 1024 * 1024:
            logger.warning(f'[_save_photo_from_base64] imagem muito grande ({len(img_data)} bytes), ignorando')
            return None
        safe_prefix = secure_filename(person_name) or 'linkedin'
        filename = secure_filename(f"{safe_prefix}-{int(time.time()*1000)}{ext}")
        path = UPLOAD_DIR / filename
        with open(path, 'wb') as f:
            f.write(img_data)
        return f'/uploads/{filename}'
    except Exception as e:
        logger.warning(f'[_save_photo_from_base64] falha: {e}')
        return None


# API - Atividades (rotas alternativas para compatibilidade)


@app.route('/api/export/group-xlsx', methods=['GET'])
def export_group_xlsx():
    try:
        group_type = (request.args.get('group_type') or '').strip().lower()
        group_value = (request.args.get('group_value') or '').strip()

        if group_type not in ('company', 'position') or not group_value:
            return jsonify({'error': 'Parâmetros inválidos'}), 400

        if not OPENPYXL_AVAILABLE:
            return jsonify({'error': 'Exportação XLSX requer openpyxl instalado'}), 500

        conn = get_db()
        c = conn.cursor()

        if group_type == 'company':
            c.execute('''
                SELECT name, position, email, phone
                FROM clients
                WHERE LOWER(company) = LOWER(?)
                ORDER BY
                    CASE
                        WHEN LOWER(TRIM(position)) = 'ceo' THEN 1
                        WHEN LOWER(TRIM(position)) LIKE 'c%' AND LENGTH(TRIM(position)) <= 4 THEN 2
                        WHEN LOWER(position) LIKE '%diretor%' OR LOWER(position) LIKE '%superintendente%' THEN 3
                        WHEN LOWER(position) LIKE '%gerente%' THEN 4
                        WHEN LOWER(position) LIKE '%coordenador%' THEN 5
                        ELSE 6
                    END,
                    LOWER(name)
            ''', (group_value,))
        else:
            c.execute('''
                SELECT name, position, email, phone
                FROM clients
                WHERE LOWER(position) = LOWER(?)
                ORDER BY LOWER(name)
            ''', (group_value,))

        rows = c.fetchall()
        conn.close()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Grupo'
        ws.append(['Nome', 'Cargo', 'Email', 'Telefone'])

        for row in rows:
            ws.append([
                row['name'] or '',
                row['position'] or '',
                row['email'] or '',
                row['phone'] or ''
            ])

        from io import BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        from flask import send_file
        safe_group = secure_filename(group_value) or 'grupo'
        file_name = f'grupo-{group_type}-{safe_group}.xlsx'
        return send_file(
            output,
            as_attachment=True,
            download_name=file_name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        logger.exception(f'[ERROR] GET /api/export/group-xlsx: {e}')
        return jsonify({'error': str(e)}), 500


# Rotas de exportacao


# Template Excel para importacao de clientes/contatos

# Importar clientes via CSV/Excel (suporta CSV e XLSX) COM VALIDACOES

# Rotas de Kanban


# Rotas de mapeamento de ambiente
@app.route('/api/environment/cards', methods=['GET'])
def get_environment_cards():
    try:
        conn = get_db()
        c = conn.cursor()
        c.execute('SELECT * FROM environment_cards ORDER BY display_order, id')
        cards = [dict_from_row(row) for row in c.fetchall()]
        conn.close()
        return jsonify(cards)
    except Exception as e:
        logger.exception(f'[ERROR] GET /api/environment/cards: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/environment/cards', methods=['POST'])
def create_environment_card():
    try:
        title = request.json.get('title', '').strip()
        description = request.json.get('description', '').strip()
        
        if not title:
            return jsonify({'error': 'Título é obrigatório'}), 400
        
        conn = get_db()
        c = conn.cursor()
        
        # Obter a próxima posição de exibição
        c.execute('SELECT MAX(display_order) FROM environment_cards')
        max_order = c.fetchone()[0] or 0
        
        c.execute('INSERT INTO environment_cards (title, description, display_order) VALUES (?, ?, ?)',
                  (title, description, max_order + 1))
        conn.commit()
        card_id = c.lastrowid
        conn.close()
        
        return jsonify({'message': 'Card criado com sucesso', 'id': card_id}), 201
    except Exception as e:
        logger.exception(f'[ERROR] POST /api/environment/cards: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/environment/cards/<int:card_id>', methods=['PUT'])
def update_environment_card(card_id):
    try:
        title = request.json.get('title', '').strip()
        description = request.json.get('description', '').strip()
        
        if not title:
            return jsonify({'error': 'Título é obrigatório'}), 400
        
        conn = get_db()
        c = conn.cursor()
        c.execute('UPDATE environment_cards SET title = ?, description = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?',
                  (title, description, card_id))
        conn.commit()
        conn.close()
        
        return jsonify({'message': 'Card atualizado com sucesso'})
    except Exception as e:
        logger.exception(f'[ERROR] PUT /api/environment/cards/{card_id}: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/environment/cards/<int:card_id>', methods=['DELETE'])
def delete_environment_card(card_id):
    try:
        conn = get_db()
        c = conn.cursor()
        
        # Deletar todas as respostas associadas ao card
        c.execute('DELETE FROM environment_responses WHERE card_id = ?', (card_id,))
        
        # Deletar o card
        c.execute('DELETE FROM environment_cards WHERE id = ?', (card_id,))
        
        conn.commit()
        conn.close()
        
        return jsonify({'message': 'Card deletado com sucesso'})
    except Exception as e:
        logger.exception(f'[ERROR] DELETE /api/environment/cards/{card_id}: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/environment/responses', methods=['GET'])
def get_environment_responses():
    try:
        client_id = request.args.get('client_id')
        
        conn = get_db()
        c = conn.cursor()
        
        if client_id:
            # Buscar respostas de um cliente específico
            c.execute('''SELECT er.*, ec.title as card_title 
                        FROM environment_responses er
                        JOIN environment_cards ec ON er.card_id = ec.id
                        WHERE er.client_id = ?
                        ORDER BY ec.display_order, ec.id''', (client_id,))
        else:
            # Buscar todas as respostas
            c.execute('''SELECT er.*, ec.title as card_title, cl.name as client_name, cl.company as client_company
                        FROM environment_responses er
                        JOIN environment_cards ec ON er.card_id = ec.id
                        JOIN clients cl ON er.client_id = cl.id
                        ORDER BY ec.display_order, ec.id, cl.company, cl.name''')
        
        responses = [dict_from_row(row) for row in c.fetchall()]
        conn.close()
        return jsonify(responses)
    except Exception as e:
        logger.exception(f'[ERROR] GET /api/environment/responses: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/environment/responses', methods=['POST'])
def save_environment_response():
    try:
        card_id = request.json.get('card_id')
        client_id = request.json.get('client_id')
        response_text = request.json.get('response', '').strip()
        
        if not card_id or not client_id:
            return jsonify({'error': 'Card e cliente são obrigatórios'}), 400
        
        # Limitar a 1000 caracteres
        if len(response_text) > 1000:
            return jsonify({'error': 'Resposta deve ter no máximo 1000 caracteres'}), 400
        
        conn = get_db()
        c = conn.cursor()
        
        # Inserir ou atualizar resposta
        c.execute('''INSERT INTO environment_responses (card_id, client_id, response, updated_at)
                     VALUES (?, ?, ?, CURRENT_TIMESTAMP)
                     ON CONFLICT(card_id, client_id) DO UPDATE SET
                        response = excluded.response,
                        updated_at = CURRENT_TIMESTAMP''',
                  (card_id, client_id, response_text))
        
        conn.commit()
        conn.close()
        
        return jsonify({'message': 'Resposta salva com sucesso'})
    except Exception as e:
        logger.exception(f'[ERROR] POST /api/environment/responses: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/environment/card/<int:card_id>/all-responses', methods=['GET'])
def get_card_all_responses(card_id):
    try:
        conn = get_db()
        c = conn.cursor()
        
        # Buscar card
        c.execute('SELECT * FROM environment_cards WHERE id = ?', (card_id,))
        card = dict_from_row(c.fetchone())
        
        if not card:
            return jsonify({'error': 'Card não encontrado'}), 404
        
        # Buscar todos os clientes com suas respostas para este card
        # Agrupar por empresa (pegar apenas a primeira ocorrência de cada empresa)
        c.execute('''SELECT cl.id, cl.name, cl.company, 
                            COALESCE(er.response, '') as response
                     FROM clients cl
                     LEFT JOIN environment_responses er ON er.client_id = cl.id AND er.card_id = ?
                     ORDER BY cl.company, cl.name''', (card_id,))
        
        all_clients = [dict_from_row(row) for row in c.fetchall()]
        
        # Filtrar para pegar apenas uma empresa por vez (primeira ocorrência)
        seen_companies = set()
        clients_responses = []
        for client in all_clients:
            if client['company'] and client['company'] not in seen_companies:
                seen_companies.add(client['company'])
                clients_responses.append(client)
        conn.close()
        
        return jsonify({
            'card': card,
            'responses': clients_responses
        })
    except Exception as e:
        logger.exception(f'[ERROR] GET /api/environment/card/{card_id}/all-responses: {e}')
        return jsonify({'error': str(e)}), 500

def _environment_extract_suggestions(company, industry, cards, card_results):
    """Usa LLM para extrair respostas para cada card a partir dos resultados Tavily."""
    cards_context = []
    for card in cards:
        card_id = card['id']
        data = card_results.get(card_id, {})
        snippets = []
        sources = []
        for r in (data.get('results') or [])[:5]:
            text = (r.get('snippet') or r.get('content') or '').strip()
            if text:
                snippets.append(text[:400])
            if r.get('url'):
                sources.append(r['url'])
        cards_context.append({
            'card_id': card_id,
            'title': card['title'],
            'description': card.get('description') or '',
            'snippets': snippets,
            'sources': sources,
        })

    context_text = ''
    for cc in cards_context:
        context_text += f'\n### Card {cc["card_id"]}: {cc["title"]}'
        if cc['description']:
            context_text += f'\nContexto: {cc["description"]}'
        if cc['snippets']:
            context_text += '\nTrechos encontrados na web:'
            for s in cc['snippets']:
                context_text += f'\n- {s}'
        else:
            context_text += '\nNenhum resultado encontrado na web para este card.'

    prompt_system = (
        'Você é um analista de inteligência de mercado extremamente rigoroso. '
        'Responda APENAS em JSON válido, sem texto extra. '
        'REGRA ABSOLUTA: só retorne uma resposta se os trechos confirmarem EXPLICITAMENTE qual produto/informação '
        'a empresa NOMEADA usa ou possui. Exemplos de respostas PROIBIDAS: '
        '"pode usar SAP ou Oracle", "utiliza sistemas como...", "entre outros", listas de possibilidades. '
        'Se o trecho apenas cita a empresa em contexto genérico, ou lista opções sem confirmar qual ela usa, '
        'retorne null. Prefira null a qualquer resposta vaga ou especulativa. '
        'Quando a resposta for encontrada, seja direto: cite o nome específico confirmado pelos trechos. '
        'Máximo 500 caracteres por resposta. '
        'Formato obrigatório: {"answers": {"<card_id>": "<resposta específica confirmada ou null>", ...}}'
    )
    prompt_user = (
        f'Empresa: {company}\nSetor: {industry}\n\n'
        f'Para cada card abaixo, extraia uma resposta baseada nos trechos encontrados na web:\n'
        f'{context_text}'
    )

    answer_map = {}
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        try:
            or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
            model = (or_settings.get('openrouter_model') or 'stepfun/step-3.5-flash:free').strip()
            site_url = (or_settings.get('openrouter_site_url') or 'http://localhost').strip()
            app_name = (or_settings.get('openrouter_app_name') or 'TocaDoCoelho').strip()
            req_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': prompt_system},
                    {'role': 'user', 'content': prompt_user},
                ],
                'temperature': 0.1,
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(req_payload).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name,
                },
                method='POST',
            )
            with urllib.request.urlopen(req, timeout=45) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            choices = data.get('choices') or []
            if choices:
                content = ((choices[0] or {}).get('message') or {}).get('content') or ''
                if isinstance(content, list):
                    content = ''.join([p.get('text', '') for p in content if isinstance(p, dict)])
                parsed = _extract_json_object_from_text(content.strip())
                if isinstance(parsed, dict):
                    answer_map = parsed.get('answers') or {}
        except Exception as e:
            logger.warning(f'[EnvAutoFill] OpenRouter error: {e}')

    if not answer_map:
        raw = _sai_simple_prompt(
            prompt_system + '\n\n' + prompt_user +
            '\nRetorne APENAS JSON no formato: {"answers": {"<card_id>": "<resposta ou null>", ...}}'
        )
        if raw:
            parsed = _extract_json_object_from_text(raw)
            if isinstance(parsed, dict):
                answer_map = parsed.get('answers') or {}

    suggestions = []
    for cc in cards_context:
        raw_answer = answer_map.get(str(cc['card_id'])) or answer_map.get(cc['card_id'])
        if raw_answer and str(raw_answer).strip().lower() not in ('null', 'none', 'não encontrado', ''):
            answer = str(raw_answer).strip()[:1000]
            found = True
        else:
            answer = None
            found = False
        suggestions.append({
            'card_id': cc['card_id'],
            'card_title': cc['title'],
            'suggestion': answer,
            'sources': cc['sources'][:3],
            'found': found,
        })
    return suggestions


def _environment_autofill_process_async(task_id, account_id):
    try:
        _bg_task_set(task_id, {'step': 'Buscando dados da conta...', 'progress': 5})
        conn = get_db()
        c = conn.cursor()

        c.execute('SELECT name, sector FROM accounts WHERE id = ?', (account_id,))
        row = c.fetchone()
        if not row:
            conn.close()
            _bg_task_set(task_id, {'status': 'error', 'error': 'Conta não encontrada'})
            return
        account = dict_from_row(row)
        company = account['name']
        industry = (account.get('sector') or '').strip()

        c.execute('SELECT id, title, description FROM environment_cards ORDER BY display_order, id')
        cards = [dict_from_row(r) for r in c.fetchall()]
        conn.close()

        if not cards:
            _bg_task_set(task_id, {'status': 'error', 'error': 'Nenhum card de mapeamento cadastrado. Crie cards antes de usar o Auto-preencher.'})
            return

        api_key = _resolve_setting('tavily_api_key', 'TAVILY_API_KEY')
        if not api_key:
            _bg_task_set(task_id, {'status': 'error', 'error': 'Chave Tavily não configurada. Configure em Configurações > Integrações.'})
            return

        card_results = {}
        total = len(cards)
        for i, card in enumerate(cards):
            progress = 15 + int((i / total) * 50)
            _bg_task_set(task_id, {
                'step': f'Pesquisando: {card["title"]}... ({i + 1}/{total})',
                'progress': progress,
            })
            # Busca direcionada a evidências concretas de uso, não artigos genéricos
            query = f'"{company}" {card["title"]} implementação adotou contratou utiliza'
            if card.get('description'):
                query += f' {card["description"][:60]}'
            try:
                results = _run_tavily_request(api_key, query, max_results=5)
                card_results[card['id']] = {'card': card, 'results': results}
            except Exception as e:
                logger.warning(f'[EnvAutoFill] Tavily error for card {card["id"]}: {e}')
                card_results[card['id']] = {'card': card, 'results': [], 'error': str(e)}

        _bg_task_set(task_id, {'step': 'Gerando sugestões com IA...', 'progress': 70})
        suggestions = _environment_extract_suggestions(company, industry or 'tecnologia', cards, card_results)

        found_count = sum(1 for s in suggestions if s['found'])
        _bg_task_set(task_id, {
            'step': 'Concluído!',
            'progress': 100,
            'status': 'done',
            'result': {
                'company': company,
                'total_cards': total,
                'found_count': found_count,
                'suggestions': suggestions,
            },
        })
    except Exception as e:
        logger.exception(f'[EnvAutoFill][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _bg_task_cleanup(task_id)


@app.route('/api/environment/auto-fill', methods=['POST'])
def environment_auto_fill():
    try:
        data = request.get_json() or {}
        account_id = data.get('account_id')
        if not account_id:
            return jsonify({'error': 'account_id é obrigatório'}), 400
        task_id = uuid.uuid4().hex
        _bg_task_set(task_id, {'status': 'processing', 'step': 'Iniciando...', 'progress': 5})
        threading.Thread(
            target=_environment_autofill_process_async,
            args=(task_id, int(account_id)),
            daemon=True,
        ).start()
        return jsonify({'task_id': task_id}), 202
    except Exception as e:
        logger.exception(f'[ERROR] POST /api/environment/auto-fill: {e}')
        return jsonify({'error': str(e)}), 500


# Rotas de backup e restore do banco de dados
def _normalize_merge_text(value):
    return str(value or '').strip().lower()


def _normalize_merge_phone(value):
    digits = re.sub(r'\D+', '', str(value or ''))
    if len(digits) >= 11:
        return digits[-11:]
    if len(digits) == 10:
        return f"{digits[:2]}9{digits[2:]}"
    return digits


def _is_empty_merge_value(value):
    return value is None or str(value).strip() == ''


def _merge_clients_from_db(temp_db_path):
    incoming_conn = sqlite3.connect(str(temp_db_path))
    incoming_conn.row_factory = sqlite3.Row
    current_conn = sqlite3.connect(str(DB_PATH))
    current_conn.row_factory = sqlite3.Row

    try:
        incoming_cur = incoming_conn.cursor()
        current_cur = current_conn.cursor()

        incoming_cur.execute('PRAGMA table_info(clients)')
        incoming_client_columns = [row['name'] for row in incoming_cur.fetchall()]
        if not incoming_client_columns:
            return {
                'processed': 0,
                'imported': 0,
                'updated': 0,
                'unchanged': 0,
                'conflicts': 0,
                'skipped': 0,
                'conflict_rows': []
            }

        current_cur.execute('PRAGMA table_info(clients)')
        current_client_columns = [row['name'] for row in current_cur.fetchall()]

        updatable_columns = [col for col in incoming_client_columns if col in current_client_columns and col != 'id']
        if not updatable_columns:
            return {
                'processed': 0,
                'imported': 0,
                'updated': 0,
                'unchanged': 0,
                'conflicts': 0,
                'skipped': 0,
                'conflict_rows': []
            }

        def load_current_clients():
            current_cur.execute('SELECT * FROM clients')
            rows = [dict(row) for row in current_cur.fetchall()]
            index_name = {}
            index_email = {}
            index_phone = {}
            for row in rows:
                rid = row.get('id')
                name_key = _normalize_merge_text(row.get('name'))
                email_key = _normalize_merge_text(row.get('email'))
                phone_key = _normalize_merge_phone(row.get('phone'))
                if name_key:
                    index_name.setdefault(name_key, rid)
                if email_key:
                    index_email.setdefault(email_key, rid)
                if phone_key:
                    index_phone.setdefault(phone_key, rid)
            return rows, index_name, index_email, index_phone

        rows, index_name, index_email, index_phone = load_current_clients()
        current_by_id = {row['id']: row for row in rows}

        incoming_cur.execute('SELECT * FROM clients ORDER BY id ASC')
        incoming_rows = [dict(row) for row in incoming_cur.fetchall()]

        summary = {
            'processed': len(incoming_rows),
            'imported': 0,
            'updated': 0,
            'unchanged': 0,
            'conflicts': 0,
            'skipped': 0,
            'conflict_rows': []
        }

        for incoming in incoming_rows:
            name_key = _normalize_merge_text(incoming.get('name'))
            email_key = _normalize_merge_text(incoming.get('email'))
            phone_key = _normalize_merge_phone(incoming.get('phone'))

            match_id = None
            match_basis = None
            if name_key and name_key in index_name:
                match_id = index_name[name_key]
                match_basis = 'name'
            elif email_key and email_key in index_email:
                match_id = index_email[email_key]
                match_basis = 'email'
            elif phone_key and phone_key in index_phone:
                match_id = index_phone[phone_key]
                match_basis = 'phone'

            if not match_id:
                insert_cols = [col for col in updatable_columns if col in incoming]
                insert_values = [incoming.get(col) for col in insert_cols]
                placeholders = ','.join(['?'] * len(insert_cols))
                current_cur.execute(
                    f"INSERT INTO clients ({','.join(insert_cols)}) VALUES ({placeholders})",
                    tuple(insert_values)
                )
                new_id = current_cur.lastrowid
                summary['imported'] += 1
                if name_key:
                    index_name.setdefault(name_key, new_id)
                if email_key:
                    index_email.setdefault(email_key, new_id)
                if phone_key:
                    index_phone.setdefault(phone_key, new_id)
                continue

            current_row = current_by_id.get(match_id) or {}
            updates = {}
            conflict_fields = []
            for col in updatable_columns:
                incoming_val = incoming.get(col)
                current_val = current_row.get(col)
                if _is_empty_merge_value(incoming_val):
                    continue
                if _is_empty_merge_value(current_val):
                    updates[col] = incoming_val
                    continue
                if str(incoming_val).strip() == str(current_val).strip():
                    continue
                conflict_fields.append(col)

            if conflict_fields:
                summary['conflicts'] += 1
                summary['skipped'] += 1
                summary['conflict_rows'].append({
                    'incoming_name': incoming.get('name') or '',
                    'match_id': match_id,
                    'match_basis': match_basis or 'none',
                    'fields': conflict_fields[:8]
                })
                continue

            if updates:
                set_sql = ', '.join([f"{col} = ?" for col in updates.keys()])
                params = list(updates.values()) + [match_id]
                current_cur.execute(f'UPDATE clients SET {set_sql} WHERE id = ?', tuple(params))
                summary['updated'] += 1
                for col, value in updates.items():
                    current_row[col] = value
            else:
                summary['unchanged'] += 1

        current_conn.commit()
        return summary
    finally:
        incoming_conn.close()
        current_conn.close()


@app.route('/api/backup/database', methods=['GET'])
def backup_database():
    try:
        from flask import send_file
        import tempfile
        import shutil
        include_uploads = str(request.args.get('include_uploads', '')).strip().lower() in {'1', 'true', 'yes', 'sim', 'on'}
        
        # Criar cópia temporária do banco
        temp_dir = tempfile.mkdtemp()
        temp_db = Path(temp_dir) / 'toca-do-coelho-backup.db'
        shutil.copy2(str(DB_PATH), str(temp_db))

        if include_uploads:
            temp_zip = Path(temp_dir) / 'toca-do-coelho-backup.zip'
            with zipfile.ZipFile(str(temp_zip), mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
                zf.write(str(temp_db), arcname='database/toca-do-coelho.db')
                if UPLOAD_DIR.exists():
                    for file_path in UPLOAD_DIR.rglob('*'):
                        if file_path.is_file():
                            relative = file_path.relative_to(UPLOAD_DIR)
                            zf.write(str(file_path), arcname=str(Path('uploads') / relative))

            return send_file(
                str(temp_zip),
                as_attachment=True,
                download_name=f'toca-do-coelho-backup-{datetime.now().strftime("%Y%m%d-%H%M%S")}.zip',
                mimetype='application/zip'
            )
        
        return send_file(
            str(temp_db),
            as_attachment=True,
            download_name=f'toca-do-coelho-backup-{datetime.now().strftime("%Y%m%d-%H%M%S")}.db',
            mimetype='application/x-sqlite3'
        )
    except Exception as e:
        logger.exception(f'[ERROR] GET /api/backup/database: {e}')
        return jsonify({'error': str(e)}), 500

@app.route('/api/restore/database', methods=['POST'])
def restore_database():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Arquivo inválido'}), 400
        
        mode = str(request.form.get('mode', 'replace_all') or 'replace_all').strip().lower()
        if mode not in {'replace_all', 'merge'}:
            return jsonify({'error': 'Modo de importação inválido. Use replace_all ou merge.'}), 400

        uploaded_name = (file.filename or '').strip().lower()
        is_zip_backup = uploaded_name.endswith('.zip')
        is_db_backup = uploaded_name.endswith('.db')
        if not is_zip_backup and not is_db_backup:
            return jsonify({'error': 'Formato inválido. Envie um arquivo .db ou .zip de backup.'}), 400
        
        # Criar backup do banco atual antes de restaurar
        backup_dir = DATA_DIR / 'backups'
        backup_dir.mkdir(exist_ok=True)
        backup_path = backup_dir / f'pre-restore-{datetime.now().strftime("%Y%m%d-%H%M%S")}.db'
        backup_uploads_path = backup_dir / f'pre-restore-uploads-{datetime.now().strftime("%Y%m%d-%H%M%S")}.zip'
        
        import shutil
        shutil.copy2(str(DB_PATH), str(backup_path))
        logger.info(f'[Database] Backup de segurança criado em {backup_path}')

        # Backup de segurança dos uploads atuais (garante rollback dos arquivos visuais)
        with zipfile.ZipFile(str(backup_uploads_path), mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
            if UPLOAD_DIR.exists():
                for file_path in UPLOAD_DIR.rglob('*'):
                    if file_path.is_file():
                        rel = file_path.relative_to(UPLOAD_DIR)
                        zf.write(str(file_path), arcname=str(Path('uploads') / rel))
        logger.info(f'[Database] Backup de segurança dos uploads criado em {backup_uploads_path}')
        
        # Salvar arquivo temporário
        temp_path = DATA_DIR / 'temp_restore.db'
        temp_zip_path = DATA_DIR / 'temp_restore.zip'
        if temp_path.exists():
            temp_path.unlink()
        if temp_zip_path.exists():
            temp_zip_path.unlink()
        
        if is_zip_backup:
            file.save(str(temp_zip_path))
            with zipfile.ZipFile(str(temp_zip_path), mode='r') as zf:
                db_member = None
                for name in zf.namelist():
                    normalized = name.replace('\\', '/')
                    if normalized.endswith('/'):
                        continue
                    if normalized in {'database/toca-do-coelho.db', 'toca-do-coelho.db'} or normalized.lower().endswith('.db'):
                        db_member = name
                        break
                if not db_member:
                    temp_zip_path.unlink(missing_ok=True)
                    return jsonify({'error': 'Backup .zip inválido: arquivo de banco (.db) não encontrado.'}), 400
                with zf.open(db_member, 'r') as db_src, open(temp_path, 'wb') as db_out:
                    shutil.copyfileobj(db_src, db_out)
        else:
            file.save(str(temp_path))
        
        # Validar se é um banco SQLite válido
        try:
            test_conn = sqlite3.connect(str(temp_path))
            test_conn.execute('SELECT name FROM sqlite_master WHERE type="table"')
            test_conn.close()
        except Exception as e:
            temp_path.unlink()
            return jsonify({'error': 'Arquivo não é um banco de dados SQLite válido'}), 400
        
        if mode == 'merge':
            merge_summary = _merge_clients_from_db(temp_path)
            temp_path.unlink(missing_ok=True)
            temp_zip_path.unlink(missing_ok=True)
            return jsonify({
                'message': 'Fusão concluída! Nenhum dado existente foi apagado.',
                'mode': 'merge',
                'backup_location': str(backup_path),
                'backup_uploads_location': str(backup_uploads_path),
                'merge_summary': merge_summary
            }), 200

        # Substituir banco atual
        shutil.move(str(temp_path), str(DB_PATH))
        logger.info(f'[Database] Banco de dados restaurado com sucesso')

        restored_uploads = 0
        if is_zip_backup and temp_zip_path.exists():
            extracted_files = []
            with zipfile.ZipFile(str(temp_zip_path), mode='r') as zf:
                for member in zf.infolist():
                    normalized = member.filename.replace('\\', '/')
                    if member.is_dir() or not normalized.startswith('uploads/'):
                        continue
                    relative = Path(normalized.replace('uploads/', '', 1)).as_posix().strip('/')
                    if not relative:
                        continue
                    if relative.startswith('../') or '/..' in f'/{relative}':
                        continue
                    destination = (UPLOAD_DIR / relative).resolve()
                    upload_root = UPLOAD_DIR.resolve()
                    if upload_root not in destination.parents and destination != upload_root:
                        continue
                    extracted_files.append((member, destination))

                # Para manter o estado visual EXATAMENTE como no backup, limpa uploads antes de restaurar
                if UPLOAD_DIR.exists():
                    shutil.rmtree(str(UPLOAD_DIR))
                UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

                for member, destination in extracted_files:
                    destination.parent.mkdir(parents=True, exist_ok=True)
                    with zf.open(member, 'r') as src, open(destination, 'wb') as dst:
                        shutil.copyfileobj(src, dst)
                    restored_uploads += 1
            temp_zip_path.unlink(missing_ok=True)
        
        return jsonify({
            'message': 'Banco de dados restaurado com sucesso! Recarregue a página.',
            'mode': 'replace_all',
            'backup_location': str(backup_path),
            'backup_uploads_location': str(backup_uploads_path),
            'restored_upload_files': restored_uploads
        }), 200
        
    except Exception as e:
        logger.exception(f'[ERROR] POST /api/restore/database: {e}')
        return jsonify({'error': str(e)}), 500

# Rotas de sugestões diárias


def _create_or_update_presence_event(c, account_name, account_id, presence):
    validity = (presence.get('validity_month') or '').strip()
    if not validity:
        return
    due_date = f"{validity}-01"
    title = f"Renovação {account_name}: {presence.get('delivery_name') or ''}".strip()
    c.execute('''INSERT INTO account_renewal_events (account_id, presence_id, title, due_date, due_time)
                 VALUES (?, ?, ?, ?, '09:00')
                 ON CONFLICT(presence_id) DO UPDATE SET
                    title = excluded.title,
                    due_date = excluded.due_date,
                    due_time = '09:00' ''',
              (account_id, presence['id'], title, due_date))


def _account_autofill_parse_llm_raw(raw):
    """Extrai dict de dados corporativos de uma resposta bruta de LLM (SAI ou OpenRouter)."""
    def _try_parse_json(value):
        if not value:
            return None
        if isinstance(value, dict):
            return value
        if not isinstance(value, str):
            return None
        try:
            parsed = json.loads(value.strip())
            if isinstance(parsed, dict):
                return parsed
        except Exception as e:
            logger.debug(f'[_try_parse_json] exceção ignorada: {e}')
        m = re.search(r'\{[^{}]*\}', value, re.DOTALL)
        if m:
            try:
                parsed = json.loads(m.group(0))
                if isinstance(parsed, dict):
                    return parsed
            except Exception as e:
                logger.debug(f'[_try_parse_json] exceção ignorada: {e}')
        return None

    parsed = _try_parse_json(raw) or {}
    if 'answer' not in parsed and 'average_revenue_brl' not in parsed:
        for key in ('output', 'result', 'text', 'content', 'response', 'data', 'message'):
            candidate = parsed.get(key)
            nested = _try_parse_json(candidate)
            if nested:
                parsed = nested
                break

    answer = parsed.get('answer') if isinstance(parsed, dict) else ''
    answer_obj = _try_parse_json(answer) if isinstance(answer, str) else None
    if answer_obj:
        parsed = answer_obj

    average_revenue_brl = parsed.get('average_revenue_brl') if isinstance(parsed, dict) else None
    professionals_count = parsed.get('professionals_count') if isinstance(parsed, dict) else None
    global_presence_raw = (parsed.get('global_presence') or '') if isinstance(parsed, dict) else ''

    global_presence = ''
    if global_presence_raw:
        gp_lower = str(global_presence_raw).strip().lower()
        if 'global' in gp_lower:
            global_presence = 'Global'
        elif 'latam' in gp_lower or 'latina' in gp_lower:
            global_presence = 'Latam'
        elif 'brasil' in gp_lower or 'brazil' in gp_lower:
            global_presence = 'Brasil'

    try:
        pc = int(professionals_count) if professionals_count is not None else None
    except Exception:
        pc = None

    try:
        rev = float(average_revenue_brl) if average_revenue_brl is not None else None
    except Exception:
        rev = None

    return {'average_revenue_brl': rev, 'professionals_count': pc, 'global_presence': global_presence}


def _account_autofill_via_sai(account_name):
    """Busca dados corporativos da empresa via SAI LLM com fallback para OpenRouter."""
    llm_prompt = (
        f"Quais são os dados corporativos da empresa '{account_name}'? "
        "Retorne SOMENTE um JSON válido, sem texto adicional, no formato: "
        '{"average_revenue_brl": 3500000000, "professionals_count": 50000, "global_presence": "Global"} '
        "Onde average_revenue_brl é o faturamento médio anual em reais (número inteiro, sem símbolos), "
        "professionals_count é o número de funcionários/profissionais (número inteiro), "
        "global_presence é 'Brasil' se opera principalmente no Brasil, 'Latam' se opera em países da América Latina além do Brasil, "
        "'Global' se opera em múltiplos continentes. "
        "Use null para campos dos quais não tem certeza."
    )

    # --- Tentativa 1: SAI (simple prompt template) ---
    raw = _sai_simple_prompt(llm_prompt)
    if raw is not None:
        result = _account_autofill_parse_llm_raw(raw)
        result['source'] = 'SAI'
        logger.info(f'[AccountAutoFill][SAI] resultado: {result}')
        return result

    # --- Tentativa 2: OpenRouter ---
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
        model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
        site_url = (or_settings.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
        app_name = (or_settings.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
        try:
            or_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'Você é um analista corporativo. Responda SEMPRE e SOMENTE com um JSON válido, sem texto adicional.'},
                    {'role': 'user', 'content': llm_prompt}
                ],
                'temperature': 0.1
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(or_payload, ensure_ascii=False).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=45) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            choices = data.get('choices') or []
            raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
            result = _account_autofill_parse_llm_raw(raw)
            result['source'] = 'OpenRouter'
            logger.info(f'[AccountAutoFill][OpenRouter] resultado: {result}')
            return result
        except Exception as e:
            logger.warning(f'[AccountAutoFill][OpenRouter] falha: {e}')

    logger.info(f'[AccountAutoFill] Nenhum LLM configurado para empresa: {account_name!r}')
    return {'average_revenue_brl': None, 'professionals_count': None, 'global_presence': '', 'source': 'sem_llm'}


def _portfolio_extract_pdf_text(file_storage):
    if not file_storage:
        return ''
    if not PDFPLUMBER_AVAILABLE:
        raise RuntimeError('Leitura de PDF indisponível no servidor (pdfplumber não instalado).')

    file_bytes = file_storage.read()
    if not file_bytes:
        return ''

    text_parts = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in (pdf.pages or []):
            try:
                text_parts.append(page.extract_text() or '')
            except Exception:
                continue
    return '\n'.join(part.strip() for part in text_parts if part and part.strip()).strip()


def _portfolio_read_image_data(file_storage):
    if not file_storage:
        return None, None
    file_bytes = file_storage.read()
    if not file_bytes:
        return None, None
    mime = (file_storage.mimetype or 'image/jpeg').split(';')[0].strip()
    return base64.b64encode(file_bytes).decode('utf-8'), mime


def _portfolio_parse_llm_raw(raw):
    def _try_parse_json(value):
        if not value:
            return None
        if isinstance(value, dict):
            return value
        if isinstance(value, list):
            return {'items': value}
        if not isinstance(value, str):
            return None
        stripped = value.strip()
        if stripped.startswith('```'):
            fenced_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', stripped, flags=re.IGNORECASE)
            if fenced_match:
                stripped = fenced_match.group(1).strip()
        try:
            parsed = json.loads(stripped)
            if isinstance(parsed, dict):
                return parsed
            if isinstance(parsed, list):
                return {'items': parsed}
        except Exception as e:
            logger.debug(f'[_try_parse_json] exceção ignorada: {e}')
        parsed = _extract_json_object_from_text(stripped)
        if isinstance(parsed, dict):
            return parsed
        return None

    parsed = _try_parse_json(raw)
    if parsed and 'items' not in parsed:
        for key in ('output', 'result', 'text', 'content', 'response', 'data', 'message', 'answer'):
            candidate = parsed.get(key)
            nested = _try_parse_json(candidate)
            if nested:
                parsed = nested
                break
    if not parsed:
        return None

    title = (parsed.get('title') or '').strip() if isinstance(parsed, dict) else ''
    summary = (parsed.get('summary') or '').strip() if isinstance(parsed, dict) else ''
    raw_items = parsed.get('items') if isinstance(parsed, dict) else []
    if not isinstance(raw_items, list):
        raw_items = []

    items = []
    for item in raw_items:
        if isinstance(item, dict):
            pain = (item.get('pain') or '').strip()
            solution = (item.get('solution') or '').strip()
            if pain or solution:
                items.append({'pain': pain, 'solution': solution})

    if not title and not summary and not items:
        return None
    if not title:
        title = 'Oferta gerada por IA'
    if not summary:
        summary = 'Resumo não informado.'
    if not items:
        items = [{'pain': 'Ponto principal não identificado', 'solution': 'Solução principal não identificada'}]
    return {'title': title, 'summary': summary, 'items': items}


def _portfolio_generate_offer_from_llm(raw_input, image_data=None, image_mime=None):
    text_prompt = (
        "Você é um especialista em posicionamento comercial B2B. "
        "Analise o material abaixo e extraia um portfólio comercial em português (Brasil). "
        "Retorne EXCLUSIVAMENTE um objeto JSON válido no formato exato abaixo, sem blocos de código markdown (```json), sem introdução e sem conclusão: "
        '{"title":"Título objetivo da oferta","summary":"Resumo executivo curto das ofertas/cases",'
        '"items":[{"pain":"Dor do cliente","solution":"Solução ofertada"}]}. '
        "Regras: gere entre 3 e 12 itens úteis; não invente informações fora do texto; "
        "use frases curtas e claras; não inclua markdown."
    )
    if raw_input:
        text_prompt += f"\n\nMATERIAL:\n{raw_input[:30000]}"

    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
        model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
        site_url = (or_settings.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
        app_name = (or_settings.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
        try:
            if image_data and image_mime:
                user_content = [
                    {'type': 'text', 'text': text_prompt},
                    {'type': 'image_url', 'image_url': {'url': f'data:{image_mime};base64,{image_data}'}}
                ]
            else:
                user_content = text_prompt
            or_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'Você é um analista comercial. Responda SEMPRE e SOMENTE com JSON válido.'},
                    {'role': 'user', 'content': user_content}
                ],
                'temperature': 0.2
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(or_payload, ensure_ascii=False).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            choices = data.get('choices') or []
            raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
            source = 'OpenRouter'
            parsed = _portfolio_parse_llm_raw(raw)
            if parsed:
                return parsed, source
            logger.warning('[Portfolio][OpenRouter] Resposta recebida mas inválida para parsing de oferta. Tentando SAI.')
        except Exception as e:
            logger.warning(f'[Portfolio][OpenRouter] Falha ao gerar oferta: {e}. Tentando SAI.')

    if not image_data:
        raw = _sai_simple_prompt(text_prompt)
        source = 'SAI'
        if raw and str(raw).strip():
            parsed = _portfolio_parse_llm_raw(raw)
            if parsed:
                return parsed, source
            logger.warning('[Portfolio][SAI] Resposta recebida mas inválida para parsing de oferta.')

    if not or_key and not _resolve_setting('itoca_sai_api_key', 'ITOCA_SAI_API_KEY'):
        return None, 'sem_llm'
    return None, 'llm_invalid_response'


# ---------------------------------------------------------------------------
# Unified Background Task Store (used by autofill, linkedin, automapping, iToca)
# ---------------------------------------------------------------------------
_bg_tasks: dict = {}
_bg_tasks_lock = threading.Lock()
# Tasks registradas aqui têm estado mínimo espelhado na tabela background_tasks
# (sobrevive a restart — ao subir, 'processing' antigas viram 'interrupted').
_bg_persistent_kinds: dict = {}


def _bg_task_register_persistent(task_id, kind):
    _bg_persistent_kinds[task_id] = kind


def _bg_task_persist(task_id, kind, task):
    try:
        conn = get_db()
        conn.execute(
            'INSERT INTO background_tasks (task_id, kind, status, step, progress, updated_at) '
            'VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP) '
            'ON CONFLICT(task_id) DO UPDATE SET status = excluded.status, step = excluded.step, '
            'progress = excluded.progress, updated_at = CURRENT_TIMESTAMP',
            (task_id, kind, task.get('status') or 'processing',
             task.get('step'), int(task.get('progress') or 0))
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.debug(f'[Tasks] Falha ao persistir estado da task {task_id}: {e}')


def _bg_task_set(task_id, updates):
    with _bg_tasks_lock:
        task = _bg_tasks.get(task_id, {})
        task.update(updates)
        _bg_tasks[task_id] = task
        snapshot = dict(task)
    kind = _bg_persistent_kinds.get(task_id)
    if kind:
        _bg_task_persist(task_id, kind, snapshot)


def _bg_task_get(task_id):
    with _bg_tasks_lock:
        return dict(_bg_tasks.get(task_id) or {})


def _bg_task_cleanup(task_id, delay=300):
    def _do():
        import time as _time
        _time.sleep(delay)
        with _bg_tasks_lock:
            _bg_tasks.pop(task_id, None)
    threading.Thread(target=_do, daemon=True).start()


# ---------------------------------------------------------------------------

_portfolio_tasks = {}
_portfolio_tasks_lock = threading.Lock()


def _portfolio_task_set(task_id, update):
    with _portfolio_tasks_lock:
        if task_id not in _portfolio_tasks:
            _portfolio_tasks[task_id] = {}
        _portfolio_tasks[task_id].update(update)


def _portfolio_task_get(task_id):
    with _portfolio_tasks_lock:
        return dict(_portfolio_tasks.get(task_id) or {})


def _portfolio_task_cleanup(task_id, delay=300):
    def _do():
        time.sleep(delay)
        with _portfolio_tasks_lock:
            _portfolio_tasks.pop(task_id, None)
    threading.Thread(target=_do, daemon=True).start()


def _portfolio_process_offer_async(task_id, input_text, file_bytes, file_mime, filename):
    try:
        raw_input = input_text
        image_data, image_mime = None, None

        if file_bytes and filename:
            fname = filename.lower()
            mime = (file_mime or '').lower()
            if 'pdf' in mime or fname.endswith('.pdf'):
                _portfolio_task_set(task_id, {'step': 'Extraindo texto do PDF...', 'progress': 30})
                if not PDFPLUMBER_AVAILABLE:
                    raise RuntimeError('Leitura de PDF indisponível no servidor (pdfplumber não instalado).')
                text_parts = []
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for page in (pdf.pages or []):
                        try:
                            text_parts.append(page.extract_text() or '')
                        except Exception:
                            continue
                pdf_text = '\n'.join(p.strip() for p in text_parts if p and p.strip()).strip()
                raw_input = (input_text + '\n\n' + pdf_text).strip() if input_text and pdf_text else (input_text or pdf_text)
            elif mime.startswith('image/') or fname.endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                _portfolio_task_set(task_id, {'step': 'Preparando imagem...', 'progress': 30})
                image_mime = mime.split(';')[0].strip() or 'image/jpeg'
                image_data = base64.b64encode(file_bytes).decode('utf-8')
            else:
                raise ValueError('Formato de arquivo não suportado. Envie um PDF ou imagem (JPG, PNG, WEBP).')

        if not raw_input and not image_data:
            raise ValueError('Informe um texto ou envie um PDF/imagem para análise.')

        _portfolio_task_set(task_id, {'step': 'Analisando com IA...', 'progress': 50})
        parsed, source = _portfolio_generate_offer_from_llm(raw_input, image_data=image_data, image_mime=image_mime)

        if not parsed:
            if source == 'sem_llm':
                raise RuntimeError('Nenhum serviço de IA configurado (OpenRouter ou SAI).')
            raise RuntimeError('A IA retornou uma resposta inválida. Tente novamente com mais contexto no material.')

        _portfolio_task_set(task_id, {'step': 'Salvando oferta...', 'progress': 88})
        conn = get_db()
        c = conn.cursor()
        c.execute(
            'INSERT INTO portfolio_offers (title, summary, raw_input, updated_at) VALUES (?, ?, ?, CURRENT_TIMESTAMP)',
            (parsed['title'], parsed['summary'], raw_input or '')
        )
        offer_id = c.lastrowid
        for idx, item in enumerate(parsed.get('items') or []):
            c.execute(
                '''INSERT INTO portfolio_offer_items (offer_id, pain, solution, sort_order, updated_at)
                   VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)''',
                (offer_id, (item.get('pain') or '').strip(), (item.get('solution') or '').strip(), idx)
            )
        conn.commit()
        offer = _portfolio_fetch_offer(c, offer_id)
        conn.close()
        if offer is None:
            raise RuntimeError('Falha ao carregar oferta criada.')
        offer['llm_source'] = source
        _portfolio_task_set(task_id, {'step': 'Concluído!', 'progress': 100, 'status': 'done', 'result': offer})
    except Exception as e:
        logger.exception(f'[Portfolio][Task:{task_id}] Erro: {e}')
        _portfolio_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _portfolio_task_cleanup(task_id)


def _portfolio_fetch_offer(c, offer_id):
    c.execute('SELECT * FROM portfolio_offers WHERE id = ?', (offer_id,))
    offer = dict_from_row(c.fetchone())
    if not offer:
        return None
    c.execute('SELECT * FROM portfolio_offer_items WHERE offer_id = ? ORDER BY sort_order ASC, id ASC', (offer_id,))
    offer['items'] = [dict_from_row(row) for row in c.fetchall()]
    return offer


def _parse_vtt_text(file_bytes):
    try:
        text = file_bytes.decode('utf-8', errors='replace')
    except Exception:
        return ''
    lines = text.split('\n')
    dialogue_lines = []
    timestamp_pat = re.compile(r'^\d{2}:\d{2}[\d:.,]* *--> *\d{2}:\d{2}')
    for line in lines:
        line = line.strip()
        if not line or line == 'WEBVTT' or line.startswith('NOTE') or line.isdigit():
            continue
        if timestamp_pat.match(line):
            continue
        dialogue_lines.append(line)
    return '\n'.join(dialogue_lines)


def _parse_srt_text(file_bytes):
    try:
        text = file_bytes.decode('utf-8', errors='replace')
    except Exception:
        return ''
    lines = text.split('\n')
    dialogue_lines = []
    timestamp_pat = re.compile(r'^\d{2}:\d{2}:\d{2},\d{3} *--> *\d{2}:\d{2}')
    for line in lines:
        line = line.strip()
        if not line or line.isdigit() or timestamp_pat.match(line):
            continue
        dialogue_lines.append(line)
    return '\n'.join(dialogue_lines)


def _iata_extract_file_text(file_storage):
    if not file_storage:
        return ''
    filename = (file_storage.filename or '').lower()
    file_bytes = file_storage.read()
    if not file_bytes:
        return ''

    if filename.endswith('.pdf'):
        if not PDFPLUMBER_AVAILABLE:
            raise RuntimeError('Leitura de PDF indisponível (pdfplumber não instalado).')
        text_parts = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in (pdf.pages or []):
                try:
                    text_parts.append(page.extract_text() or '')
                except Exception:
                    continue
        return '\n'.join(p.strip() for p in text_parts if p and p.strip()).strip()

    if filename.endswith('.docx'):
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError('Leitura de DOCX indisponível (python-docx não instalado).')
        doc = python_docx.Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs).strip()

    if filename.endswith('.vtt'):
        return _parse_vtt_text(file_bytes)

    if filename.endswith('.srt'):
        return _parse_srt_text(file_bytes)

    # Plain text / fallback
    try:
        if CHARDET_AVAILABLE:
            detected = chardet.detect(file_bytes)
            encoding = (detected.get('encoding') or 'utf-8')
        else:
            encoding = 'utf-8'
        return file_bytes.decode(encoding, errors='replace').strip()
    except Exception:
        return file_bytes.decode('utf-8', errors='replace').strip()


def _iata_parse_llm_ata(raw):
    if not raw:
        return None
    stripped = str(raw).strip()
    if stripped.startswith('```'):
        m = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', stripped, flags=re.IGNORECASE)
        if m:
            stripped = m.group(1).strip()
    try:
        parsed = json.loads(stripped)
    except Exception:
        parsed = _extract_json_object_from_text(stripped)
    if not isinstance(parsed, dict):
        return None
    if 'title' not in parsed:
        for key in ('output', 'result', 'text', 'content', 'response', 'data'):
            candidate = parsed.get(key)
            if isinstance(candidate, str):
                try:
                    nested = json.loads(candidate)
                    if isinstance(nested, dict) and 'title' in nested:
                        parsed = nested
                        break
                except Exception as e:
                    logger.debug(f'[_iata_parse_llm_ata] exceção ignorada: {e}')
    title = (parsed.get('title') or '').strip()
    if not title:
        return None

    def _clean_null(val):
        v = str(val or '').strip()
        return None if not v or v.lower() in ('null', 'none', 'n/a', '-') else v

    def _parse_participants(raw_list):
        result = []
        for p in (raw_list or []):
            if isinstance(p, dict):
                name = str(p.get('name') or p.get('nome') or '').strip()
                role = str(p.get('role') or p.get('cargo') or p.get('empresa') or '').strip()
                if name:
                    result.append({'name': name, 'role': role or ''})
            elif p and str(p).strip():
                result.append({'name': str(p).strip(), 'role': ''})
        return result

    return {
        'title': title,
        'meeting_date': _clean_null(parsed.get('meeting_date')),
        'meeting_time': _clean_null(parsed.get('meeting_time')),
        'location': _clean_null(parsed.get('location')),
        'topic': (parsed.get('topic') or '').strip() or title,
        'participants': _parse_participants(parsed.get('participants')),
        'objective': (parsed.get('objective') or '').strip(),
        'summary': (parsed.get('summary') or '').strip(),
        'agenda': [str(a).strip() for a in (parsed.get('agenda') or []) if str(a).strip()],
        'key_points': [str(k).strip() for k in (parsed.get('key_points') or []) if str(k).strip()],
        'decisions': [str(d).strip() for d in (parsed.get('decisions') or []) if str(d).strip()],
        'next_steps': [
            {
                'action': str(s.get('action') or '').strip(),
                'responsible': str(s.get('responsible') or 'A definir').strip(),
                'deadline': _clean_null(s.get('deadline')),
                'status': str(s.get('status') or 'pendente').strip()
            }
            for s in (parsed.get('next_steps') or []) if isinstance(s, dict)
        ],
        'observations': _clean_null(parsed.get('observations'))
    }


def _iata_parse_llm_insights(raw):
    if not raw:
        return None
    stripped = str(raw).strip()
    if stripped.startswith('```'):
        m = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', stripped, flags=re.IGNORECASE)
        if m:
            stripped = m.group(1).strip()
    try:
        parsed = json.loads(stripped)
    except Exception:
        parsed = _extract_json_object_from_text(stripped)
    if not isinstance(parsed, dict):
        return None
    insights_raw = parsed.get('insights') or []
    if not isinstance(insights_raw, list):
        return None
    insights = []
    for item in insights_raw:
        if not isinstance(item, dict):
            continue
        pain = (item.get('pain') or '').strip()
        if not pain:
            continue
        insights.append({
            'pain': pain,
            'matched_offer': item.get('matched_offer') or None,
            'matched_solution': item.get('matched_solution') or None,
            'confidence': str(item.get('confidence') or 'baixa').strip(),
            'observation': str(item.get('observation') or '').strip()
        })
    return {
        'insights': insights,
        'has_unmatched': bool(parsed.get('has_unmatched', False))
    }


def _iata_call_llm(system_msg, user_msg, log_tag):
    """SAI como primário; OpenRouter como fallback."""
    # Tenta SAI primeiro
    question = f"{system_msg}\n\n{user_msg}" if system_msg else user_msg
    raw = _sai_simple_prompt(question)
    if raw and str(raw).strip():
        logger.info(f'[iAta][{log_tag}] SAI respondeu com sucesso')
        return raw, 'SAI'
    logger.info(f'[iAta][{log_tag}] SAI indisponível, tentando OpenRouter.')

    # Fallback: OpenRouter
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if not or_key:
        return None, 'sem_llm'
    or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
    model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
    site_url = (or_settings.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
    app_name = (or_settings.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
    try:
        resp = requests.post(
            'https://openrouter.ai/api/v1/chat/completions',
            json={
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_msg},
                    {'role': 'user', 'content': user_msg}
                ],
                'temperature': 0.2
            },
            headers={
                'Authorization': f'Bearer {or_key}',
                'HTTP-Referer': site_url,
                'X-Title': app_name
            },
            timeout=90
        )
        data = resp.json()
        choices = data.get('choices') or []
        raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
        if raw and str(raw).strip():
            logger.info(f'[iAta][{log_tag}][OpenRouter] Resposta gerada com sucesso')
            return raw, 'OpenRouter'
        logger.warning(f'[iAta][{log_tag}][OpenRouter] Resposta vazia.')
    except Exception as e:
        logger.warning(f'[iAta][{log_tag}][OpenRouter] Falha: {e}.')
    return None, 'sem_llm'


def _iata_generate_ata(raw_text):
    prompt = (
        "Você é um especialista em documentação corporativa. "
        "Analise o texto de reunião e gere uma ATA DE REUNIÃO COMPLETA E DETALHADA em português (Brasil). "
        "A ata deve ser rica em detalhes, capturando o contexto, as discussões e os argumentos apresentados. "
        "Retorne EXCLUSIVAMENTE um objeto JSON válido, sem markdown, sem introdução:\n"
        '{"title":"Título formal da reunião",'
        '"meeting_date":"DD/MM/AAAA ou null se não mencionada",'
        '"meeting_time":"HH:MM ou null se não mencionada",'
        '"location":"Local ou plataforma (ex: Microsoft Teams) ou null",'
        '"topic":"Tema central da reunião em uma frase",'
        '"participants":[{"name":"Nome Completo","role":"Cargo ou empresa se mencionado"}],'
        '"objective":"Objetivo principal da reunião em 2-3 frases completas",'
        '"summary":"Resumo executivo DETALHADO da reunião com 5-8 frases cobrindo contexto, discussões principais e resultados",'
        '"agenda":["Item de pauta 1 — descrição breve","Item de pauta 2 — descrição breve"],'
        '"key_points":["Ponto discutido em detalhe — inclua contexto, argumentos e posições dos participantes. Cada item deve ter 2-4 frases.","Segundo ponto com o mesmo nível de detalhe"],'
        '"decisions":["Decisão formal tomada com contexto — quem decidiu e por quê quando relevante"],'
        '"next_steps":[{"action":"Descrição detalhada da ação a ser realizada","responsible":"Nome completo do responsável","deadline":"DD/MM/AAAA ou null","status":"pendente"}],'
        '"observations":"Observações adicionais relevantes, contexto extra ou informações complementares mencionadas durante a reunião. Use null se não houver."}\n'
        "REGRAS OBRIGATÓRIAS:\n"
        "- Liste TODOS os participantes mencionados com nome e cargo/empresa quando disponível;\n"
        "- key_points: cada item deve ser um parágrafo detalhado (2-4 frases), não apenas uma frase curta;\n"
        "- next_steps: inclua TODAS as pendências, ações e encaminhamentos com responsáveis identificados;\n"
        "- Se não houver responsável claro para uma ação, use 'A definir';\n"
        "- deadline: use formato DD/MM/AAAA quando mencionado, ou null (JSON null, não a string 'null');\n"
        "- Não invente informações que não estejam no texto original;\n"
        "- Preserve nomes próprios, empresas e termos técnicos como aparecem no texto.\n\n"
        f"TEXTO DA REUNIÃO:\n{raw_text[:30000]}"
    )
    raw, source = _iata_call_llm(
        'Você é um especialista em documentação corporativa. Responda SEMPRE e SOMENTE com JSON válido.',
        prompt,
        'Ata'
    )
    if raw is None:
        return None, source
    parsed = _iata_parse_llm_ata(raw)
    if parsed:
        return parsed, source
    logger.warning(f'[iAta] Resposta inválida para parsing de ata (source={source}).')
    return None, 'llm_invalid_response'


def _iata_generate_insights(ata_data, portfolio_offers):
    ata_text = json.dumps(ata_data, ensure_ascii=False)
    offers_text = json.dumps(portfolio_offers, ensure_ascii=False)
    prompt = (
        "Você é um consultor de negócios especialista em soluções tecnológicas. "
        "Analise a ata de reunião e identifique dores/desafios de negócio mencionados. "
        "Cruze cada dor com as soluções disponíveis no portfólio STF. "
        "Retorne EXCLUSIVAMENTE um objeto JSON válido, sem markdown:\n"
        '{"insights":[{"pain":"Dor identificada","matched_offer":"Título da oferta ou null",'
        '"matched_solution":"Solução específica ou null","confidence":"alta/média/baixa",'
        '"observation":"Observação breve"}],"has_unmatched":true}\n'
        "Regras:\n"
        "- Identifique TODAS as dores, desafios e oportunidades mencionados;\n"
        "- matched_offer = null e observation explicando para buscar soluções na Stefanini se não houver match;\n"
        "- confidence: 'alta' se endereça diretamente, 'média' se parcialmente, 'baixa' se tangencialmente;\n"
        "- has_unmatched: true se houver pelo menos uma dor sem match.\n\n"
        f"ATA:\n{ata_text[:15000]}\n\n"
        f"SOLUÇÕES STF:\n{offers_text[:15000]}"
    )
    raw, source = _iata_call_llm(
        'Você é um consultor de negócios. Responda SEMPRE e SOMENTE com JSON válido.',
        prompt,
        'Insights'
    )
    if raw is None:
        return None, source
    parsed = _iata_parse_llm_insights(raw)
    if parsed:
        return parsed, source
    logger.warning(f'[iAta] Resposta inválida para parsing de insights (source={source}).')
    return None, 'llm_invalid_response'


_iata_tasks = {}
_iata_tasks_lock = threading.Lock()


def _iata_task_set(task_id, updates):
    with _iata_tasks_lock:
        task = _iata_tasks.get(task_id, {})
        task.update(updates)
        _iata_tasks[task_id] = task


def _iata_task_get(task_id):
    with _iata_tasks_lock:
        return dict(_iata_tasks.get(task_id) or {})


def _iata_task_cleanup(task_id, delay=300):
    def _do():
        time.sleep(delay)
        with _iata_tasks_lock:
            _iata_tasks.pop(task_id, None)
    threading.Thread(target=_do, daemon=True).start()


def _iata_extract_bytes(file_bytes, filename):
    """Extrai texto de bytes de arquivo (sem file_storage)."""
    class _BytesFS:
        def __init__(self, data, name):
            self.filename = name
            self._data = data
        def read(self):
            return self._data
    return _iata_extract_file_text(_BytesFS(file_bytes, filename))


def _iata_process_async(task_id, file_bytes, filename, raw_text_input):
    try:
        raw_text = raw_text_input or ''
        if file_bytes and filename:
            _iata_task_set(task_id, {'step': 'Extraindo texto do arquivo...', 'progress': 15})
            file_text = _iata_extract_bytes(file_bytes, filename)
            if file_text and raw_text_input:
                raw_text = file_text + '\n\n' + raw_text_input
            else:
                raw_text = file_text or raw_text_input

        if not raw_text:
            raise ValueError('Não foi possível extrair texto do arquivo enviado.')

        _iata_task_set(task_id, {'step': 'Gerando ata com IA...', 'progress': 30})
        ata_data, ata_source = _iata_generate_ata(raw_text)
        if not ata_data:
            if ata_source == 'sem_llm':
                raise RuntimeError('Nenhum serviço de IA configurado (OpenRouter ou SAI).')
            raise RuntimeError('A IA retornou uma resposta inválida. Tente novamente.')

        _iata_task_set(task_id, {'step': 'Buscando soluções STF...', 'progress': 68})
        conn = get_db()
        c = conn.cursor()
        c.execute('SELECT id, title, summary FROM portfolio_offers ORDER BY id DESC')
        offers = [dict_from_row(row) for row in c.fetchall()]
        for offer in offers:
            c.execute('SELECT pain, solution FROM portfolio_offer_items WHERE offer_id = ? ORDER BY sort_order ASC', (offer['id'],))
            offer['items'] = [dict_from_row(row) for row in c.fetchall()]

        if offers:
            _iata_task_set(task_id, {'step': 'Gerando insights de negócio...', 'progress': 78})
            insights_data, _ = _iata_generate_insights(ata_data, offers)
            if not insights_data:
                insights_data = {'insights': [], 'has_unmatched': False}
        else:
            insights_data = {'insights': [], 'has_unmatched': False}

        _iata_task_set(task_id, {'step': 'Salvando ata...', 'progress': 92})
        c.execute(
            '''INSERT INTO iata_records (title, meeting_date, meeting_time, topic, participants, ata_json, insights_json, raw_text)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                ata_data['title'],
                ata_data.get('meeting_date'),
                ata_data.get('meeting_time'),
                ata_data.get('topic'),
                json.dumps([p['name'] if isinstance(p, dict) else str(p) for p in (ata_data.get('participants') or [])], ensure_ascii=False),
                json.dumps(ata_data, ensure_ascii=False),
                json.dumps(insights_data, ensure_ascii=False),
                raw_text[:50000]
            )
        )
        record_id = c.lastrowid
        conn.commit()
        c.execute('SELECT * FROM iata_records WHERE id = ?', (record_id,))
        record = _iata_record_to_dict(c.fetchone())
        conn.close()
        _iata_task_set(task_id, {'step': 'Concluído!', 'progress': 100, 'status': 'done', 'result': record})
    except Exception as e:
        logger.exception(f'[iAta][Task:{task_id}] Erro: {e}')
        _iata_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _iata_task_cleanup(task_id)


def _iata_record_to_dict(record):
    r = dict(record)
    for field in ('participants', 'ata_json', 'insights_json'):
        val = r.get(field)
        if val:
            try:
                r[field] = json.loads(val)
            except Exception:
                r[field] = [] if field == 'participants' else {}
        else:
            r[field] = [] if field == 'participants' else {}
    r['ata'] = r.pop('ata_json', {})
    r['insights'] = r.pop('insights_json', {})
    return r


@app.route('/api/tasks/<task_id>', methods=['GET'])
def get_bg_task_status(task_id):
    """Generic task status poll for background tasks (autofill, linkedin, automapping, iToca)."""
    task = _bg_task_get(task_id)
    if not task:
        return jsonify({'status': 'not_found'}), 404
    return jsonify(task)


def _autofill_process_async(task_id, account_name):
    try:
        _bg_task_set(task_id, {'step': 'Consultando IA sobre a empresa...', 'progress': 20})
        company_data = _account_autofill_via_sai(account_name)

        average_revenue_formatted = ''
        raw_revenue = company_data.get('average_revenue_brl')
        if raw_revenue is not None:
            try:
                cents = int(round(float(raw_revenue) * 100))
                average_revenue_formatted = format_currency_br(cents)
            except Exception as e:
                logger.debug(f'[_autofill_process_async] exceção ignorada: {e}')

        _bg_task_set(task_id, {'step': 'Buscando logo da empresa...', 'progress': 70})
        logo_url = None
        try:
            candidates = _find_image_candidates_on_web(f'{account_name} logo empresa', limit=4)
            if candidates:
                logo_url = candidates[0]
        except Exception as e:
            logger.warning(f'[AccountAutoFill] Falha ao buscar logo: {e}')

        result = {
            'average_revenue': average_revenue_formatted,
            'professionals_count': company_data.get('professionals_count', ''),
            'global_presence': company_data.get('global_presence', ''),
            'logo_url': logo_url,
            'source': company_data.get('source', 'SAI')
        }
        _bg_task_set(task_id, {'step': 'Concluído!', 'progress': 100, 'status': 'done', 'result': result})
    except Exception as e:
        logger.exception(f'[AccountAutoFill][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _bg_task_cleanup(task_id)


# ---------------------------------------------------------------------------
# LinkedIn AutoFill — Preenche ficha de contato via perfil LinkedIn
# ---------------------------------------------------------------------------

def _client_linkedin_autofill_async(task_id, linkedin_url, profile_text, extension_photo_url=''):
    """Extrai dados cadastrais de um contato a partir do texto do perfil LinkedIn via LLM.
    Não tenta scraping — recebe apenas o texto capturado pela extensão ou colado pelo usuário.
    """
    try:
        _bg_task_set(task_id, {'step': 'Extraindo dados cadastrais com IA...', 'progress': 20})

        data_is_rich = bool(profile_text and len(profile_text.strip()) > 100)
        profile_content = (profile_text or '').strip()
        if not profile_content and linkedin_url:
            profile_content = f'URL do perfil LinkedIn: {linkedin_url}'

        quality_note = (
            'Use SOMENTE as informações explicitamente presentes no perfil. '
            if data_is_rich else
            'Dados limitados — use null para campos não encontrados. NÃO invente informações. '
        )

        llm_prompt = (
            f'Analise as informações do perfil LinkedIn abaixo e extraia os dados cadastrais desta pessoa. '
            f'{quality_note}'
            f'PERFIL:\n{profile_content}\n\n'
            'Retorne SOMENTE JSON válido, sem texto adicional: '
            '{"nome": "Nome completo da pessoa", '
            '"empresa": "Nome da empresa atual (somente o nome da empresa)", '
            '"cargo": "Cargo ou título profissional atual", '
            '"area_atuacao": "Área de atuação profissional (ex: Tecnologia, Comercial, RH, Financeiro, Marketing, Jurídico, Operações)", '
            '"email": "email@dominio.com ou null se não encontrado"} '
            'Use null para campos desconhecidos. Responda em português (BR).'
        )

        raw = _sai_simple_prompt(llm_prompt)

        if not raw:
            or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
            if or_key:
                or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
                model = (or_settings.get('openrouter_model') or 'stepfun/step-3.5-flash:free').strip() or 'stepfun/step-3.5-flash:free'
                site_url = (or_settings.get('openrouter_site_url') or 'http://localhost').strip()
                app_name = (or_settings.get('openrouter_app_name') or 'TocaDoCoelho').strip()
                try:
                    or_payload = {
                        'model': model,
                        'messages': [
                            {'role': 'system', 'content': 'Você é um analista de dados. Responda SEMPRE e SOMENTE com JSON válido, sem texto adicional.'},
                            {'role': 'user', 'content': llm_prompt}
                        ],
                        'temperature': 0.1
                    }
                    req = urllib.request.Request(
                        'https://openrouter.ai/api/v1/chat/completions',
                        data=json.dumps(or_payload, ensure_ascii=False).encode('utf-8'),
                        headers={
                            'Content-Type': 'application/json',
                            'Authorization': f'Bearer {or_key}',
                            'HTTP-Referer': site_url,
                            'X-Title': app_name
                        },
                        method='POST'
                    )
                    with urllib.request.urlopen(req, timeout=45) as resp:
                        data = json.loads(resp.read().decode('utf-8'))
                    choices = data.get('choices') or []
                    raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
                except Exception as e:
                    logger.warning(f'[ClientLinkedInAutoFill][OpenRouter] Falha: {e}')

        parsed = {}
        if raw:
            try:
                parsed = json.loads(raw)
            except Exception:
                m = re.search(r'\{[\s\S]*\}', raw)
                if m:
                    try:
                        parsed = json.loads(m.group(0))
                    except Exception as e:
                        logger.debug(f'[_client_linkedin_autofill_async] exceção ignorada: {e}')

        _bg_task_set(task_id, {'step': 'Buscando foto de perfil...', 'progress': 75})

        # Foto: extensão tem prioridade → Bing fallback (sem scraping do LinkedIn)
        photo_url = (extension_photo_url or '').strip() or None
        if not photo_url and parsed.get('nome'):
            try:
                nome = parsed.get('nome') or ''
                empresa = parsed.get('empresa') or ''
                query = f'{nome} {empresa} foto perfil'.strip()
                candidates = _find_image_candidates_on_web(query, limit=3)
                if candidates:
                    photo_url = candidates[0]
            except Exception as e:
                logger.debug(f'[_client_linkedin_autofill_async] exceção ignorada: {e}')

        result = {
            'nome': parsed.get('nome') or '',
            'empresa': parsed.get('empresa') or '',
            'cargo': parsed.get('cargo') or '',
            'area_atuacao': parsed.get('area_atuacao') or '',
            'email': parsed.get('email') or '',
            'photo_url': photo_url or '',
            'data_is_rich': data_is_rich,
        }
        _bg_task_set(task_id, {'step': 'Concluído!', 'progress': 100, 'status': 'done', 'result': result})
    except Exception as e:
        logger.exception(f'[ClientLinkedInAutoFill][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _bg_task_cleanup(task_id)


# ---------------------------------------------------------------------------
# LinkedIn Profile Summarizer
# ---------------------------------------------------------------------------

def _linkedin_try_fetch_public(url):
    """Tenta buscar dados públicos de um perfil LinkedIn. Retorna texto ou None."""
    if not url or 'linkedin.com' not in url:
        return None
    try:
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        req = urllib.request.Request(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'pt-BR,pt;q=0.9,en;q=0.8',
        })
        with urllib.request.urlopen(req, context=ctx, timeout=15) as resp:
            raw_html = resp.read().decode('utf-8', errors='ignore')
        # Extrai texto visível removendo tags HTML
        clean = re.sub(r'<style[^>]*>.*?</style>', ' ', raw_html, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r'<script[^>]*>.*?</script>', ' ', clean, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r'<[^>]+>', ' ', clean)
        clean = html.unescape(clean)
        clean = re.sub(r'\s+', ' ', clean).strip()
        # Se LinkedIn redirecionou para login, o conteúdo tem pouca informação útil
        if len(clean) < 300 or 'Sign in' in clean or 'authwall' in url:
            return None
        return clean[:16000]  # Limita para não exceder contexto do LLM
    except Exception as e:
        logger.debug(f'[LinkedIn] Falha ao buscar URL pública: {e}')
        return None


def _linkedin_generate_summary_via_llm(profile_content, meeting_context='', data_is_rich=True):
    """Gera resumo executivo do perfil LinkedIn via LLM (SAI → OpenRouter)."""
    ctx_part = f'\n\nCONTEXTO DA REUNIÃO: {meeting_context}' if meeting_context.strip() else ''

    if data_is_rich:
        quality_instruction = (
            'IMPORTANTE: Use SOMENTE as informações explicitamente presentes no perfil abaixo. '
            'Seja ESPECÍFICO e DETALHADO: cite nomes reais de empresas, cargos exatos, anos, projetos, '
            'números e realizações concretas mencionados no perfil. '
            'NÃO generalize, NÃO invente informações que não estejam no texto. '
        )
    else:
        quality_instruction = (
            'ATENÇÃO: Os dados disponíveis são LIMITADOS (perfil parcialmente acessível). '
            'Use null para campos que não estiverem explicitamente nos dados fornecidos. '
            'NÃO invente nem generalize — prefira null a informação genérica. '
        )

    llm_prompt = (
        'Analise as informações abaixo de um perfil profissional do LinkedIn e gere um resumo executivo '
        f'para preparar uma reunião de negócios com esta pessoa. {quality_instruction}'
        'IMPORTANTE SOBRE A TRAJETÓRIA: monte a lista "trajetoria" com TODAS as experiências profissionais '
        'mencionadas no perfil (não apenas o cargo atual) — inclua cada empresa/cargo/período encontrado no texto, '
        'da mais recente para a mais antiga, sem limite fixo de itens. Se a pessoa teve mais de um cargo na mesma '
        'empresa (ex.: promoção interna), mencione isso explicitamente (ex.: "CEO desde 2025 (na empresa desde 2012, '
        'anteriormente Diretor Financeiro)"). No campo "cargo_atual", se o perfil indicar uma data de início no cargo '
        'diferente da data de início na empresa, mencione as duas (ex.: "CEO na Dasa desde 2025 — na empresa desde 2012").'
        f'\n\nPERFIL LINKEDIN:\n{profile_content}{ctx_part}\n\n'
        'Retorne SOMENTE um JSON válido, sem texto adicional, com exatamente estes campos: '
        '{"nome": "Nome completo da pessoa", '
        '"cargo_atual": "Cargo exato e empresa atual, com tempo de casa se diferente do tempo no cargo", '
        '"trajetoria": ["experiência específica 1 com empresa, cargo e período", "experiência 2", "experiência 3", "... (todas as experiências do perfil)"], '
        '"formacao": ["curso, instituição e ano se disponível", "outra formação"], '
        '"competencias": ["competência específica 1", "competência 2", "competência 3", "competência 4", "competência 5"], '
        '"pontos_conversa": ["ponto de conversa concreto 1", "ponto 2", "ponto 3", "ponto 4"], '
        '"insights": ["insight específico sobre o perfil 1", "insight 2"], '
        '"tom_sugerido": "tom recomendado com justificativa baseada no perfil"} '
        'Use null para campos não encontrados. Responda em português (BR).'
    )

    raw = _sai_simple_prompt(llm_prompt)
    if raw:
        logger.info('[LinkedIn][SAI] Resumo gerado com sucesso')
        return raw, 'SAI'

    # Fallback: OpenRouter (o template SAI compartilhado pode estourar o limite de uso
    # sob carga concorrente de outras features do app; sem fallback a feature ficava sem resposta)
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
        model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
        site_url = (or_settings.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
        app_name = (or_settings.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
        try:
            or_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'Você é um analista de inteligência executiva. Responda SEMPRE e SOMENTE com um JSON válido, sem texto adicional.'},
                    {'role': 'user', 'content': llm_prompt}
                ],
                'temperature': 0.3
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(or_payload, ensure_ascii=False).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=45) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            choices = data.get('choices') or []
            raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
            if raw:
                logger.info('[LinkedIn][OpenRouter] Resumo gerado com sucesso')
                return raw, 'OpenRouter'
        except Exception as e:
            logger.warning(f'[LinkedIn][OpenRouter] Falha: {e}')

    return None, 'sem_llm'


def _linkedin_parse_summary(raw):
    """Extrai e valida JSON do resumo gerado pelo LLM."""
    if not raw:
        return None
    # Tenta JSON direto
    try:
        return json.loads(raw)
    except Exception as e:
        logger.debug(f'[_linkedin_parse_summary] exceção ignorada: {e}')
    # Tenta extrair JSON do texto
    m = re.search(r'\{[\s\S]*\}', raw)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception as e:
            logger.debug(f'[_linkedin_parse_summary] exceção ignorada: {e}')
    return None


def _normalize_linkedin_url(url):
    """Normaliza URL de LinkedIn para comparação: minúsculas, sem protocolo/www/querystring/trailing slash."""
    if not url:
        return ''
    u = str(url).strip().lower()
    u = u.split('?', 1)[0].split('#', 1)[0]
    u = re.sub(r'^https?://', '', u)
    u = re.sub(r'^www\.', '', u)
    return u.rstrip('/')


def _linkedin_find_contact_and_account(linkedin_url, contact_name):
    """Localiza o contato (por linkedin, depois por nome) e a conta mapeada vinculada
    via clients.company == accounts.name. Retorna (contact|None, account|None)."""
    conn = get_db()
    c = conn.cursor()
    contact = None
    norm = _normalize_linkedin_url(linkedin_url)
    if norm:
        c.execute("SELECT * FROM clients WHERE linkedin IS NOT NULL AND TRIM(linkedin) != ''")
        for row in c.fetchall():
            r = dict_from_row(row)
            if _normalize_linkedin_url(r.get('linkedin')) == norm:
                contact = r
                break
    if not contact and contact_name and contact_name.strip():
        c.execute("SELECT * FROM clients WHERE LOWER(TRIM(name)) = LOWER(TRIM(?)) ORDER BY id LIMIT 1",
                  (contact_name.strip(),))
        row = c.fetchone()
        if row:
            contact = dict_from_row(row)
    account = None
    if contact and (contact.get('company') or '').strip():
        c.execute("SELECT * FROM accounts WHERE LOWER(TRIM(name)) = LOWER(TRIM(?)) ORDER BY id LIMIT 1",
                  (contact['company'].strip(),))
        row = c.fetchone()
        if row:
            account = dict_from_row(row)
    conn.close()
    return contact, account


def _linkedin_resolve_account(account_id):
    """Carrega uma conta por id (usado quando o usuário força a conta no dropdown)."""
    try:
        aid = int(account_id)
    except (TypeError, ValueError):
        return None
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM accounts WHERE id = ? LIMIT 1", (aid,))
    row = c.fetchone()
    conn.close()
    return dict_from_row(row) if row else None


def _linkedin_build_account_context(account, contact):
    """Monta o contexto da conta para o Preparar Reunião.
    Retorna dict com: account, market_moment, relationship_summary,
    stefanini_services, contact_found, contact_photo_url."""
    account_name = (account.get('name') or '').strip()

    # a) Momento de mercado (fresco a cada execução)
    market_moment = None
    try:
        market_moment = _relation_report_fetch_market_context(account_name)
    except Exception as e:
        logger.warning(f'[LinkedIn] Falha ao buscar momento de mercado: {e}')

    # c) Serviços Stefanini mapeados na conta
    stefanini_services = []
    try:
        conn = get_db()
        c = conn.cursor()
        c.execute("""SELECT delivery_name, stf_owner, current_revenue_cents
                     FROM account_presences WHERE account_id = ?
                     ORDER BY delivery_name COLLATE NOCASE""", (account['id'],))
        for r in c.fetchall():
            row = dict_from_row(r)
            stefanini_services.append({
                'delivery_name': row.get('delivery_name') or '',
                'stf_owner': row.get('stf_owner') or '',
                'revenue': format_currency_br(row.get('current_revenue_cents')) if row.get('current_revenue_cents') else None,
            })
        conn.close()
    except Exception as e:
        logger.warning(f'[LinkedIn] Falha ao listar serviços Stefanini: {e}')

    # b) Resumo do relacionamento com o contato (só se o contato existir)
    relationship_summary = None
    contact_found = bool(contact)
    contact_photo_url = (contact.get('photo_url') or '').strip() if contact else ''
    if contact:
        try:
            conn = get_db()
            c = conn.cursor()
            c.execute("""SELECT activity_date, contact_type, description, information
                         FROM activities WHERE client_id = ?
                         ORDER BY datetime(activity_date) DESC, id DESC LIMIT 20""", (contact['id'],))
            acts = [dict_from_row(r) for r in c.fetchall()]
            c.execute("""SELECT activity_date, description FROM account_activities
                         WHERE account_id = ?
                         ORDER BY datetime(activity_date) DESC, created_at DESC LIMIT 15""", (account['id'],))
            acc_acts = [dict_from_row(r) for r in c.fetchall()]
            conn.close()

            contact_lines = []
            for a in acts:
                txt = (a.get('description') or a.get('information') or '').strip().replace('\n', ' ')
                if txt:
                    contact_lines.append(f"- {a.get('activity_date') or 's/data'} [{a.get('contact_type') or 'atividade'}]: {txt[:240]}")
            account_lines = []
            for a in acc_acts:
                txt = (a.get('description') or '').strip().replace('\n', ' ')
                if txt:
                    account_lines.append(f"- {a.get('activity_date') or 's/data'}: {txt[:240]}")

            if contact_lines or account_lines:
                user_msg = (
                    f"Contato: {contact.get('name')} ({contact.get('position') or 'cargo não informado'}) "
                    f"da conta {account_name}.\n\n"
                    "HISTÓRICO DE ATIVIDADES DO CONTATO:\n"
                    + ('\n'.join(contact_lines) if contact_lines else '(sem registros diretos do contato)')
                    + "\n\nCONTEXTO DO RELACIONAMENTO DA CONTA:\n"
                    + ('\n'.join(account_lines) if account_lines else '(sem registros da conta)')
                    + "\n\nEscreva um resumo executivo do relacionamento com este contato em 2-4 frases, "
                    "priorizando o histórico do contato e complementando com o contexto da conta quando o "
                    "histórico direto for escasso. Não invente fatos não listados."
                )
                raw, _src = _iata_call_llm(
                    'Você é um assistente comercial. Responda em português (Brasil), apenas o resumo, sem markdown.',
                    user_msg,
                    'linkedin_relationship'
                )
                if raw and str(raw).strip():
                    relationship_summary = str(raw).strip()
        except Exception as e:
            logger.warning(f'[LinkedIn] Falha ao resumir relacionamento: {e}')

    return {
        'account': {'id': account['id'], 'name': account_name},
        'market_moment': market_moment,
        'relationship_summary': relationship_summary,
        'stefanini_services': stefanini_services,
        'contact_found': contact_found,
        'contact_photo_url': contact_photo_url,
    }


def _linkedin_process_async(task_id, linkedin_url, profile_text, meeting_context, extension_photo_url, forced_account_id=None):
    try:
        _bg_task_set(task_id, {'step': 'Buscando perfil público...', 'progress': 15})
        fetched_text = None
        if linkedin_url and not profile_text:
            fetched_text = _linkedin_try_fetch_public(linkedin_url)

        data_is_rich = bool(profile_text) or (fetched_text and len(fetched_text) > 2000)
        profile_content = profile_text or fetched_text or ''
        if not profile_content and linkedin_url:
            profile_content = f'URL do perfil LinkedIn: {linkedin_url}'

        _bg_task_set(task_id, {'step': 'Gerando resumo executivo com IA...', 'progress': 35})
        raw, source = _linkedin_generate_summary_via_llm(profile_content, meeting_context, data_is_rich=data_is_rich)
        if not raw:
            _bg_task_set(task_id, {'status': 'error', 'error': 'Nenhum serviço de IA configurado ou disponível no momento (SAI e OpenRouter falharam).'})
            return

        parsed = _linkedin_parse_summary(raw)
        contact_name = (parsed or {}).get('nome') or ''

        # Detecção de contato/conta
        contact, detected_account = _linkedin_find_contact_and_account(linkedin_url, contact_name)
        account = _linkedin_resolve_account(forced_account_id) if forced_account_id else detected_account

        _bg_task_set(task_id, {'step': 'Buscando foto do perfil...', 'progress': 70})
        photo_url = None
        photo_source = None

        # Prioridade 1: foto já cadastrada no sistema para este contato — mais confiável
        # que qualquer captura automática do LinkedIn (evita pegar foto de capa/banner errada).
        if contact and (contact.get('photo_url') or '').strip():
            photo_url = contact['photo_url'].strip()
            photo_source = 'system_contact_photo'

        if not photo_url and extension_photo_url:
            try:
                photo_url = _download_remote_image_to_uploads(extension_photo_url, prefix='linkedin-profile-ext')
                photo_source = 'extension_profile_photo'
            except Exception as e:
                logger.debug(f'[LinkedIn] Falha ao baixar foto da extensão: {e}')

        if not photo_url and linkedin_url:
            try:
                og_image = _linkedin_extract_og_image(linkedin_url)
                if og_image:
                    photo_url = _download_remote_image_to_uploads(og_image, prefix='linkedin-profile-og')
                    photo_source = 'linkedin_og_image'
            except Exception as e:
                logger.debug(f'[LinkedIn] Falha ao usar og:image: {e}')

        if not photo_url and parsed and parsed.get('nome'):
            try:
                nome = parsed['nome']
                cargo = parsed.get('cargo_atual', '')
                query = f'{nome} {cargo} foto perfil profissional'.strip()
                candidates = _find_image_candidates_on_web(query, limit=3)
                if candidates:
                    photo_url = _download_remote_image_to_uploads(candidates[0], prefix='linkedin-profile')
                    photo_source = 'web_search_fallback'
            except Exception as e:
                logger.debug(f'[LinkedIn] Falha ao buscar/baixar foto: {e}')

        # Contexto da conta (a/b/c)
        account_context = None
        if account:
            _bg_task_set(task_id, {'step': 'Analisando conta mapeada...', 'progress': 88})
            try:
                account_context = _linkedin_build_account_context(account, contact)
            except Exception as e:
                logger.warning(f'[LinkedIn] Falha ao montar contexto da conta: {e}')

        result = {
            'summary': parsed,
            'raw': raw if not parsed else None,
            'source': source,
            'fetched_from_url': fetched_text is not None,
            'limited_data': not data_is_rich,
            'photo_url': photo_url,
            'photo_source': photo_source,
            'account_context': account_context
        }
        _bg_task_set(task_id, {'step': 'Concluído!', 'progress': 100, 'status': 'done', 'result': result})
    except Exception as e:
        logger.exception(f'[LinkedIn][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _bg_task_cleanup(task_id)


@app.route('/api/linkedin/summarize', methods=['POST'])
def linkedin_summarize():
    """Gera resumo executivo de um perfil LinkedIn para preparação de reunião."""
    try:
        data = request.get_json() or {}
        linkedin_url = (data.get('linkedin_url') or '').strip()
        profile_text = (data.get('profile_text') or '').strip()
        meeting_context = (data.get('meeting_context') or '').strip()
        extension_photo_url = (data.get('extension_photo_url') or '').strip()
        forced_account_id = data.get('account_id')

        if not linkedin_url and not profile_text:
            return jsonify({'error': 'Informe a URL do LinkedIn ou cole o texto do perfil.'}), 400

        task_id = uuid.uuid4().hex
        _bg_task_set(task_id, {'status': 'processing', 'step': 'Iniciando...', 'progress': 5})
        threading.Thread(
            target=_linkedin_process_async,
            args=(task_id, linkedin_url, profile_text, meeting_context, extension_photo_url, forced_account_id),
            daemon=True
        ).start()
        return jsonify({'task_id': task_id}), 202
    except Exception as e:
        logger.exception(f'[LinkedIn] Erro: {e}')
        return jsonify({'error': str(e)}), 500


# ---------------------------------------------------------------------------


def _automapping_process_async(task_id, company, country, industry, force, request_id):
    try:
        _bg_task_set(task_id, {'step': 'Verificando cache...', 'progress': 10})
        query_key = _normalize_automapping_key(company, country, industry)
        conn = get_db()
        c = conn.cursor()
        c.execute('''SELECT id, result_json, created_at FROM automapping_runs
                     WHERE query_key = ?
                       AND datetime(created_at) >= datetime('now', '-20 days')
                     ORDER BY datetime(created_at) DESC
                     LIMIT 1''', (query_key,))
        cached = c.fetchone()

        if cached and not force:
            conn.close()
            result = json.loads(cached['result_json'])
            _bg_task_set(task_id, {
                'step': 'Concluído!', 'progress': 100, 'status': 'done',
                'result': result, 'already_exists': True,
                'run_id': cached['id'], 'created_at': cached['created_at']
            })
            return

        if _is_automapping_cancelled(request_id, consume=True):
            conn.close()
            _bg_task_set(task_id, {'status': 'cancelled', 'error': 'AutoMapping cancelado pelo usuário.'})
            return

        _bg_task_set(task_id, {'step': 'Buscando evidências de mercado...', 'progress': 30})
        try:
            evidence_results, section_errors, execution_meta = _run_tavily_search(company, country, industry)
        except urllib.error.HTTPError as e:
            detail = e.read().decode('utf-8', errors='ignore') if hasattr(e, 'read') else str(e)
            conn.close()
            _bg_task_set(task_id, {'status': 'error', 'error': f'Falha ao consultar Tavily: {detail[:400]}'})
            return
        except Exception as e:
            conn.close()
            _bg_task_set(task_id, {'status': 'error', 'error': f'Falha ao consultar Tavily: {str(e)}'})
            return

        if _is_automapping_cancelled(request_id, consume=True):
            conn.close()
            _bg_task_set(task_id, {'status': 'cancelled', 'error': 'AutoMapping cancelado pelo usuário.'})
            return

        _bg_task_set(task_id, {'step': 'Construindo análise de ambiente...', 'progress': 55})
        result_payload = _build_automapping_payload(company, country, industry, evidence_results, section_errors, execution_meta)
        c.execute('SELECT 1 FROM clients WHERE LOWER(TRIM(company)) = LOWER(TRIM(?)) LIMIT 1', (company,))
        result_payload['client_exists'] = bool(c.fetchone())

        if _is_automapping_cancelled(request_id, consume=True):
            conn.close()
            _bg_task_set(task_id, {'status': 'cancelled', 'error': 'AutoMapping cancelado pelo usuário.'})
            return

        _bg_task_set(task_id, {'step': 'Gerando síntese executiva com IA...', 'progress': 75})
        try:
            llm_summary, llm_meta = _run_openrouter_synthesis(result_payload)
            result_payload['llm_summary'] = llm_summary
            result_payload['llm_meta'] = llm_meta
        except urllib.error.HTTPError as e:
            detail = e.read().decode('utf-8', errors='ignore') if hasattr(e, 'read') else str(e)
            conn.close()
            _bg_task_set(task_id, {'status': 'error', 'error': f'Falha ao consultar OpenRouter: {detail[:400]}'})
            return
        except Exception as e:
            conn.close()
            _bg_task_set(task_id, {'status': 'error', 'error': f'Falha ao consultar OpenRouter: {str(e)}'})
            return

        if _is_automapping_cancelled(request_id, consume=True):
            conn.close()
            _bg_task_set(task_id, {'status': 'cancelled', 'error': 'AutoMapping cancelado pelo usuário.'})
            return

        _bg_task_set(task_id, {'step': 'Salvando resultado...', 'progress': 90})
        c.execute('INSERT INTO automapping_runs (company, country, industry, query_key, result_json) VALUES (?, ?, ?, ?, ?)',
                  (company, country, industry, query_key, json.dumps(result_payload, ensure_ascii=False)))
        run_id = c.lastrowid
        conn.commit()
        conn.close()

        _bg_task_set(task_id, {
            'step': 'Concluído!', 'progress': 100, 'status': 'done',
            'result': result_payload, 'already_exists': False, 'run_id': run_id
        })
    except Exception as e:
        logger.exception(f'[AutoMapping][Task:{task_id}] Erro: {e}')
        _bg_task_set(task_id, {'status': 'error', 'error': str(e)})
    finally:
        _bg_task_cleanup(task_id)


# ─────────────────────────────────────────────────────────────
# WikiToca – Conhecimentos registrados
# ─────────────────────────────────────────────────────────────


# ─────────────────────────────────────────────────────────────
# WikiToca – Documentos
# ─────────────────────────────────────────────────────────────

ALLOWED_WIKI_EXTENSIONS = {'.pdf', '.xls', '.xlsx', '.doc', '.docx'}


# WikiToca - Export/Import XLSX


def _autotoca_parse_account_info_raw(raw):
    """Extrai razao_social, cnpj, endereco e confiavel de uma resposta bruta de LLM."""
    def _coerce(value):
        if value is None:
            return None
        if isinstance(value, dict):
            return value
        if not isinstance(value, str):
            return None
        text = value.strip()
        if text.startswith('```'):
            m = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text, flags=re.IGNORECASE)
            if m:
                text = m.group(1).strip()
        try:
            obj = json.loads(text)
            if isinstance(obj, dict):
                return obj
        except Exception as e:
            logger.debug(f'[_coerce] exceção ignorada: {e}')
        obj = _extract_json_object_from_text(text)
        return obj if isinstance(obj, dict) else None

    parsed = _coerce(raw) or {}
    # Desembrulha respostas aninhadas comuns dos templates SAI.
    for key in ('answer', 'output', 'result', 'text', 'content', 'response', 'data', 'message'):
        if isinstance(parsed, dict) and 'razao_social' not in parsed and 'endereco' not in parsed and key in parsed:
            nested = _coerce(parsed.get(key))
            if isinstance(nested, dict):
                parsed = nested
                break

    if not isinstance(parsed, dict):
        return {}

    def _clean(value):
        if value is None:
            return None
        s = str(value).strip()
        if not s or s.lower() in ('null', 'none', 'n/a', 'não informado', 'nao informado', '-'):
            return None
        return s

    confiavel_raw = parsed.get('confiavel')
    if isinstance(confiavel_raw, str):
        confiavel = confiavel_raw.strip().lower() in ('true', 'sim', 'yes', '1')
    else:
        confiavel = bool(confiavel_raw)

    return {
        'razao_social': _clean(parsed.get('razao_social')),
        'cnpj': _clean(parsed.get('cnpj')),
        'endereco': _clean(parsed.get('endereco')),
        'confiavel': confiavel,
    }


def _autotoca_account_info_via_llm(account_name: str) -> dict:
    """Busca razão social, CNPJ e endereço de sede da empresa via LLM (OpenRouter -> SAI).

    Retorna dict com as chaves razao_social, cnpj e endereco. O endereço é objetivo
    (apenas o endereço). Quando os dados não são confiáveis, o endereço recebe o
    prefixo 'Por favor confirmar: '.
    """
    # Dica de endereço a partir da busca web (mesma lógica do antigo "Buscar Sede").
    heuristic_address = ''
    try:
        heuristic = AccountAddressService().find_headquarter_address(account_name)
        heuristic_address = (heuristic.get('suggested_address') or '').strip()
    except Exception as e:
        logger.warning(f'[AutoToca][AccountInfo] heurística de endereço falhou: {e}')

    prompt = (
        f"Empresa cliente no Brasil: '{account_name}'. "
        + (f"Dica de endereço encontrada na web: {heuristic_address}. " if heuristic_address else "")
        + "Pesquise os dados cadastrais oficiais desta empresa e retorne SOMENTE um JSON válido, "
        "sem markdown e sem texto adicional, no formato exato: "
        '{"razao_social": "...", "cnpj": "...", "endereco": "...", "confiavel": true}. '
        "Regras: 'razao_social' é a razão social completa registrada na Receita Federal; "
        "'cnpj' é o CNPJ da matriz no formato 00.000.000/0000-00; "
        "'endereco' é o endereço completo da sede/matriz no Brasil no formato "
        "'Logradouro, número, bairro, cidade - UF', contendo APENAS o endereço, "
        "sem explicações, sem fontes e sem comentários; "
        "'confiavel' deve ser true somente se você tem certeza dos dados, e false caso contrário. "
        "Use null nos campos que não conseguir determinar."
    )

    raw = None

    # --- Tentativa 1: OpenRouter ---
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
        model = (or_settings.get('openrouter_model') or os.environ.get('OPENROUTER_MODEL', 'stepfun/step-3.5-flash:free')).strip() or 'stepfun/step-3.5-flash:free'
        site_url = (or_settings.get('openrouter_site_url') or os.environ.get('OPENROUTER_SITE_URL', 'http://localhost')).strip() or 'http://localhost'
        app_name = (or_settings.get('openrouter_app_name') or os.environ.get('OPENROUTER_APP_NAME', 'TocaDoCoelho')).strip() or 'TocaDoCoelho'
        try:
            or_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'Você é um analista de dados cadastrais corporativos. Responda SEMPRE e SOMENTE com um JSON válido, sem texto adicional.'},
                    {'role': 'user', 'content': prompt}
                ],
                'temperature': 0.1
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(or_payload, ensure_ascii=False).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=45) as resp:
                or_data = json.loads(resp.read().decode('utf-8'))
            choices = or_data.get('choices') or []
            raw = (choices[0].get('message') or {}).get('content', '') if choices else ''
        except Exception as e:
            logger.warning(f'[AutoToca][AccountInfo][OpenRouter] falha: {e}')
            raw = None

    # --- Tentativa 2: SAI (fallback) ---
    if not raw:
        raw = _sai_simple_prompt(prompt)

    parsed = _autotoca_parse_account_info_raw(raw)

    endereco = parsed.get('endereco')
    confiavel = parsed.get('confiavel')

    # Endereço: usa o do LLM; se vazio, cai para a dica heurística.
    final_address = endereco or heuristic_address or ''
    # Marca como "a confirmar" quando o LLM não classificou o endereço como confiável
    # ou quando o endereço veio apenas da heurística da web.
    address_trustworthy = bool(endereco) and confiavel is True
    if final_address and not address_trustworthy:
        final_address = f'Por favor confirmar: {final_address}'

    return {
        'razao_social': parsed.get('razao_social') or '',
        'cnpj': parsed.get('cnpj') or '',
        'endereco': final_address,
    }


# ─────────────────────────────────────────
#  WhatsApp Update — WAHA (WhatsApp HTTP API)
# ─────────────────────────────────────────

def _phone_to_waha_chatid(phone):
    """Converte telefone normalizado do projeto para chatId do WAHA (5511999999999@c.us)."""
    if not phone:
        return None
    digits = re.sub(r'\D', '', str(phone))
    if len(digits) < 10:
        return None
    if not digits.startswith('55'):
        digits = '55' + digits
    return f'{digits}@c.us'


def _waha_settings():
    s = _load_app_settings_map(['waha_api_url', 'waha_api_key', 'waha_session_name'])
    return (
        (s.get('waha_api_url') or 'http://localhost:3001').strip().rstrip('/'),
        (s.get('waha_api_key') or '').strip(),
        (s.get('waha_session_name') or 'default').strip(),
    )


def _waha_headers(api_key):
    h = {'Content-Type': 'application/json'}
    if api_key:
        h['X-Api-Key'] = api_key
    return h


# ---------------------------------------------------------------------------
# Caixa de respostas pendentes (Bloco 6) — inbound WhatsApp e Outlook.
# Registra a última mensagem recebida de cada cliente sem resposta sua.
# ---------------------------------------------------------------------------

def _inbound_upsert(c, client_id, channel, received_at, preview, source_msg_id):
    """Registra mensagem recebida pendente (dedup por source_msg_id)."""
    c.execute(
        '''INSERT INTO inbound_messages (client_id, channel, received_at, preview, source_msg_id)
           VALUES (?, ?, ?, ?, ?)
           ON CONFLICT(source_msg_id) DO NOTHING''',
        (client_id, channel, received_at, (preview or '')[:280], source_msg_id)
    )


def _inbound_mark_responded(c, client_id, channel, responded_at=None):
    responded_at = responded_at or datetime.now().isoformat(timespec='seconds')
    c.execute(
        '''UPDATE inbound_messages SET responded_at = ?
           WHERE client_id = ? AND channel = ? AND responded_at IS NULL''',
        (responded_at, client_id, channel)
    )


def _inbound_scan_whatsapp():
    """Fase A (polling): consulta o WAHA pelas conversas dos clientes cadastrados
    e registra a última mensagem com fromMe=false posterior à última fromMe=true.
    Comportamento passivo (equivalente a abrir o WhatsApp Web)."""
    api_url, api_key, session = _waha_settings()
    headers = _waha_headers(api_key)
    conn = get_db()
    c = conn.cursor()
    c.execute("""SELECT id, name, phone FROM clients
                 WHERE phone IS NOT NULL AND phone != '' AND COALESCE(is_archived, 0) = 0""")
    clients = c.fetchall()
    since_ts = int((datetime.now() - timedelta(days=14)).timestamp())
    scanned = pending = 0
    for row in clients:
        chat_id = _phone_to_waha_chatid(row['phone'])
        if not chat_id:
            continue
        try:
            resp = requests.get(
                f'{api_url}/api/{session}/chats/{chat_id}/messages',
                headers=headers,
                params={'limit': 50, 'downloadMedia': 'false', 'filter.timestamp.gte': since_ts},
                timeout=20
            )
            if resp.status_code != 200:
                continue
            raw = resp.json()
            messages = raw if isinstance(raw, list) else (raw.get('messages') or raw.get('data') or [])
        except Exception as e:
            logger.debug(f'[Inbound] WAHA indisponível para {row["name"]}: {e}')
            continue
        scanned += 1
        last_in, last_out, last_in_msg = 0, 0, None
        for msg in messages:
            try:
                ts = int(msg.get('timestamp') or 0)
            except (TypeError, ValueError):
                continue
            if msg.get('fromMe'):
                last_out = max(last_out, ts)
            elif ts > last_in and _waha_is_content_message(msg):
                last_in = ts
                last_in_msg = msg
        if last_in and last_in > last_out:
            text, _sender = _waha_extract_text(last_in_msg, row['name'])
            msg_id = last_in_msg.get('id')
            if isinstance(msg_id, dict):
                msg_id = msg_id.get('_serialized') or msg_id.get('id')
            source_id = f'wa:{msg_id or (str(row["id"]) + ":" + str(last_in))}'
            _inbound_upsert(c, row['id'], 'whatsapp',
                            datetime.fromtimestamp(last_in).isoformat(timespec='seconds'),
                            text or '(mensagem sem texto)', source_id)
            pending += 1
            conn.commit()
        elif last_out:
            _inbound_mark_responded(c, row['id'], 'whatsapp',
                                    datetime.fromtimestamp(last_out).isoformat(timespec='seconds'))
            conn.commit()
    conn.close()
    logger.info(f'[Inbound] Scan WhatsApp: {scanned} conversas verificadas, {pending} pendências registradas')
    return {'scanned': scanned, 'pending': pending}


def _inbound_feed_from_outlook(c, emails_data, clients_map):
    """Alimenta o painel de pendências a partir dos e-mails importados do Outlook:
    cliente cujo último e-mail recebido é posterior ao último enviado = pendente."""
    per_client = {}
    for em in emails_data or []:
        direction = em.get('direction', 'received')
        when = (em.get('date') or '').strip()
        if not when:
            continue
        if direction == 'received':
            addr = ((em.get('sender') or {}).get('email') or '').lower()
            entry = clients_map.get(addr)
            if not entry:
                continue
            slot = per_client.setdefault(entry['id'], {'in': None, 'out': ''})
            if slot['in'] is None or when > slot['in']['date']:
                slot['in'] = {'date': when, 'preview': em.get('body_preview') or em.get('subject') or '',
                              'msg_id': em.get('message_id') or ''}
        else:
            for rec in em.get('recipients') or []:
                entry = clients_map.get((rec.get('email') or '').lower())
                if entry:
                    slot = per_client.setdefault(entry['id'], {'in': None, 'out': ''})
                    slot['out'] = max(slot['out'], when)
    for client_id, slot in per_client.items():
        if slot['in'] and slot['in']['date'] > slot['out']:
            source_id = f'em:{slot["in"]["msg_id"] or (str(client_id) + ":" + slot["in"]["date"])}'
            _inbound_upsert(c, client_id, 'email', slot['in']['date'][:19].replace('T', ' '),
                            slot['in']['preview'], source_id)
        elif slot['out']:
            _inbound_mark_responded(c, client_id, 'email', slot['out'][:19].replace('T', ' '))


# ---------------------------------------------------------------------------
# Jobs agendados leves (Blocos 9/13/14): cada job registra sua última execução
# em app_settings e roda quando o intervalo vence. Registrados pelos módulos
# de rotas via _register_scheduled_job().
# ---------------------------------------------------------------------------
_SCHEDULED_JOBS = []
_scheduled_jobs_started = False


def _register_scheduled_job(settings_key, interval_days, fn, only_after_hour=None):
    """only_after_hour: se definido, o job só roda a partir dessa hora local."""
    _SCHEDULED_JOBS.append((settings_key, interval_days, fn, only_after_hour))


def _save_app_setting(key, value):
    conn = get_db()
    conn.execute(
        'INSERT INTO app_settings (key, value) VALUES (?, ?) '
        'ON CONFLICT(key) DO UPDATE SET value = excluded.value, updated_at = CURRENT_TIMESTAMP',
        (key, str(value))
    )
    conn.commit()
    conn.close()


def _start_scheduled_jobs():
    global _scheduled_jobs_started
    if _scheduled_jobs_started or os.environ.get('TOCA_DISABLE_BG_JOBS') == '1':
        return
    _scheduled_jobs_started = True

    def _loop():
        while True:
            time.sleep(30 * 60)  # verifica a cada 30 min o que está vencido
            now = datetime.now()
            for key, interval_days, fn, only_after_hour in list(_SCHEDULED_JOBS):
                try:
                    if only_after_hour is not None and now.hour < only_after_hour:
                        continue
                    last_raw = _resolve_setting(key, '')
                    if last_raw:
                        try:
                            last = datetime.fromisoformat(last_raw)
                            if (now - last) < timedelta(days=interval_days):
                                continue
                        except ValueError:
                            pass
                    logger.info(f'[Jobs] Executando job agendado: {key}')
                    result = fn()
                    if result == 'skip':
                        continue  # fora da janela do job — tenta de novo mais tarde
                    _save_app_setting(key, now.isoformat(timespec='seconds'))
                except Exception as e:
                    logger.warning(f'[Jobs] Job {key} falhou: {e}')

    threading.Thread(target=_loop, daemon=True).start()
    logger.info(f'[Jobs] Agendador iniciado ({len(_SCHEDULED_JOBS)} jobs registrados)')


_inbound_poller_started = False


def _start_inbound_poller():
    """Job em background (Fase A) com intervalo configurável (default 15 min)."""
    global _inbound_poller_started
    if _inbound_poller_started or os.environ.get('TOCA_DISABLE_BG_JOBS') == '1':
        return
    _inbound_poller_started = True

    def _loop():
        while True:
            try:
                minutes = int(_resolve_setting('inbound_poll_minutes', 'INBOUND_POLL_MINUTES') or 15)
            except Exception:
                minutes = 15
            time.sleep(max(minutes, 1) * 60)
            try:
                _inbound_scan_whatsapp()
            except Exception as e:
                logger.debug(f'[Inbound] Poller: scan falhou (WAHA offline?): {e}')

    threading.Thread(target=_loop, daemon=True).start()
    logger.info('[Inbound] Poller de respostas pendentes iniciado')


_waha_last_restart = 0.0

def _waha_runtime_paths():
    """Retorna (node, script) do WAHA-lite. Usa as env vars exportadas pelo launcher;
    se ausentes (ex.: app rodando direto do código-fonte), tenta um fallback razoável."""
    node   = os.environ.get('WAHA_NODE_EXE', '').strip()
    script = os.environ.get('WAHA_SCRIPT', '').strip()
    if not script:
        candidate = Path(__file__).resolve().parent / 'waha-lite' / 'waha-lite.js'
        if candidate.exists():
            script = str(candidate)
    if not node:
        node = shutil.which('node') or shutil.which('node.exe') or ''
    return node, script


def _waha_deps_missing():
    """True se o node_modules do WAHA-lite estiver ausente/vazio. Nesse caso o processo
    Node crasha no primeiro require e nada escuta na porta — reiniciar não adianta."""
    if os.environ.get('WAHA_DEPS_MISSING') == '1':
        return True
    _, script = _waha_runtime_paths()
    if not script:
        return False  # não sabemos onde está o waha-lite; não afirmar que falta
    nm = Path(script).parent / 'node_modules'
    try:
        return (not nm.is_dir()) or (next(nm.iterdir(), None) is None)
    except Exception:
        return False


def _kill_process_on_port(port):
    """Mata o processo que estiver ESCUTANDO na porta informada (WAHA-lite).

    Importante: só mata quem está em LISTENING no endereço local — nunca conexões
    ESTABLISHED, senão poderíamos matar o próprio Toca (que conecta no WAHA-lite)."""
    killed = []
    try:
        import subprocess as _sp
        if sys.platform == 'win32':
            # Colunas do netstat -ano: Proto | Local Address | Foreign Address | State | PID
            #   TCP    127.0.0.1:3001    0.0.0.0:0    LISTENING    12345
            out = _sp.check_output(
                ['netstat', '-ano', '-p', 'TCP'],
                creationflags=_sp.CREATE_NO_WINDOW,
                stderr=_sp.DEVNULL,
                timeout=5,
            ).decode(errors='ignore')
            seen = set()
            for line in out.splitlines():
                parts = line.split()
                if len(parts) < 5:
                    continue
                proto, local, _foreign, state, pid = parts[0], parts[1], parts[2], parts[3], parts[-1]
                if not proto.upper().startswith('TCP'):
                    continue
                if state.upper() != 'LISTENING':
                    continue            # ignora ESTABLISHED/TIME_WAIT etc.
                if not local.endswith(f':{port}'):
                    continue
                if pid.isdigit() and pid != '0' and pid not in seen:
                    seen.add(pid)
                    _sp.call(
                        ['taskkill', '/F', '/T', '/PID', pid],
                        creationflags=_sp.CREATE_NO_WINDOW,
                        stderr=_sp.DEVNULL,
                    )
                    killed.append(pid)
        else:
            # -sTCP:LISTEN garante que só pegamos quem escuta, não quem conecta.
            out = _sp.check_output(
                ['lsof', '-ti', f'TCP:{port}', '-sTCP:LISTEN'],
                stderr=_sp.DEVNULL,
                timeout=5,
            ).decode(errors='ignore').strip()
            import signal as _sig
            for pid in out.splitlines():
                if pid.isdigit():
                    try:
                        os.kill(int(pid), _sig.SIGTERM)
                        killed.append(pid)
                    except Exception as e:
                        logger.debug(f'[_kill_process_on_port] exceção ignorada: {e}')
    except Exception as exc:
        logger.debug(f'[WhatsApp] _kill_process_on_port({port}): {exc}')
    if killed:
        logger.info(f'[WhatsApp] Processo(s) WAHA-lite anterior(es) na porta {port} encerrado(s): {", ".join(killed)}')
    return bool(killed)


def _restart_waha_lite():
    """Reinicia o WAHA-lite. Máximo 1 restart a cada 5 minutos. Não tenta reiniciar quando
    as dependências (node_modules) estão ausentes, pois o Node crasharia de novo."""
    global _waha_last_restart
    now = time.time()
    if now - _waha_last_restart < 300:
        return False
    if _waha_deps_missing():
        logger.warning('[WhatsApp] node_modules do WAHA-lite ausente — restart ignorado (reinstale).')
        return False
    node, script = _waha_runtime_paths()
    if not node or not script:
        return False
    try:
        import subprocess as _sp
        # Encerra o processo atual na porta antes de subir um novo para evitar EADDRINUSE.
        waha_port = int(os.environ.get('WAHA_PORT', '3001'))
        _kill_process_on_port(waha_port)
        import threading as _th
        _th.Event().wait(timeout=1.5)  # aguarda o SO liberar a porta
        env = os.environ.copy()
        # Reaproveita o log do WAHA-lite definido pelo launcher (WAHA_LOG); sem ele, o
        # crash do Node ficaria invisível (DEVNULL), exatamente o que dificultou diagnosticar
        # a falha do WhatsApp Update em produção.
        _waha_log = os.environ.get('WAHA_LOG', '').strip()
        try:
            _out = open(_waha_log, 'a', encoding='utf-8', buffering=1) if _waha_log else _sp.DEVNULL
        except Exception:
            _out = _sp.DEVNULL
        kwargs = {
            'env': env,
            'cwd': str(Path(script).parent),
            'stdout': _out,
            'stderr': _out,
        }
        if sys.platform == 'win32':
            kwargs['creationflags'] = _sp.CREATE_NO_WINDOW
        _sp.Popen([node, script], **kwargs)
        _waha_last_restart = now
        os.environ['WAHA_STARTED_AT'] = str(now)
        logger.info('[WhatsApp] WAHA-lite reiniciado automaticamente.')
        return True
    except Exception as exc:
        logger.warning(f'[WhatsApp] Falha ao reiniciar WAHA-lite: {exc}')
        return False


def _waha_is_content_message(msg):
    """True se a mensagem tem conteúdo real (texto ou mídia) enviado por alguém.
    Filtra avisos de sistema do próprio WhatsApp (ex.: "As mensagens e ligações
    são protegidas com criptografia de ponta a ponta..."), que chegam com
    fromMe=false mas não são mensagens do contato — sem isso, esses avisos
    eram registrados como pendência de resposta com "(mensagem sem texto)"."""
    if (msg.get('body') or '').strip():
        return True
    mtype = (msg.get('type') or '').lower()
    if mtype in ('e2e_notification', 'notification_template', 'notification', 'gp2', 'call_log', 'revoked', 'status'):
        return False
    return bool(msg.get('hasMedia')) or mtype in ('image', 'video', 'ptt', 'audio', 'document', 'sticker', 'location')


def _waha_extract_text(msg, client_name):
    """Extrai texto de uma mensagem do WAHA. Retorna (texto, sender_label)."""
    sender = 'Você' if msg.get('fromMe') else client_name
    body = (msg.get('body') or '').strip()
    if body:
        return f'{sender}: {body}', sender
    # Sem corpo de texto → rotula por tipo de mídia
    mtype = (msg.get('type') or '').lower()
    label_map = {
        'image': '[Imagem]', 'video': '[Vídeo]', 'ptt': '[Áudio]',
        'audio': '[Áudio]', 'document': '[Documento]', 'sticker': '[Figurinha]',
        'location': '[Localização]',
    }
    if msg.get('hasMedia') or mtype in label_map:
        return f'{sender}: {label_map.get(mtype, "[Mídia]")}', sender
    return None, sender


def _whatsapp_sync_async(task_id, period_days):
    try:
        _bg_task_set(task_id, {'step': '📋 Carregando clientes...', 'progress': 5})

        api_url, api_key, session = _waha_settings()
        headers = _waha_headers(api_key)

        db = get_db()
        c = db.cursor()
        c.execute("SELECT id, name, phone FROM clients WHERE phone IS NOT NULL AND phone != '' AND is_archived = 0")
        clients = c.fetchall()

        if not clients:
            _bg_task_set(task_id, {
                'status': 'done', 'progress': 100,
                'step': '✅ Concluído',
                'result': {'imported': 0, 'skipped': 0, 'total_clients': 0,
                           'message': 'Nenhum cliente com telefone cadastrado.'}
            })
            _bg_task_cleanup(task_id)
            return

        # Horário local em todo o app; a conversão para epoch UTC (que a API do
        # WAHA espera) acontece via .timestamp(), que interpreta o datetime
        # naive como horário local — elimina o deslocamento de 3h no corte.
        now_dt = datetime.now()
        since_dt = now_dt - timedelta(days=period_days)
        since_ts = int(since_dt.timestamp())
        now_ts = int(now_dt.timestamp())

        imported = 0
        skipped = 0
        total = len(clients)
        pending_items = []
        period_label = {1: '1 dia', 3: '3 dias', 7: '1 semana', 15: '15 dias', 30: '30 dias'}.get(period_days, f'{period_days} dias')

        for i, (client_id, client_name, phone) in enumerate(clients):
            progress = 10 + int((i / total) * 78)
            _bg_task_set(task_id, {
                'step': f'📱 Verificando {client_name}... ({i + 1}/{total})',
                'progress': progress
            })

            chat_id = _phone_to_waha_chatid(phone)
            if not chat_id:
                skipped += 1
                continue

            # Deduplicação incremental: só resumimos mensagens MAIS NOVAS do que a
            # última já processada para este cliente. Evita re-resumir a mesma conversa.
            c.execute('SELECT MAX(last_message_ts) FROM whatsapp_sync_log WHERE client_id = ?', (client_id,))
            _row = c.fetchone()
            last_processed_ts = int(_row[0]) if _row and _row[0] else 0
            effective_since = max(since_ts, last_processed_ts + 1) if last_processed_ts else since_ts

            try:
                # WAHA: GET /api/{session}/chats/{chatId}/messages com filtro de timestamp server-side
                msg_resp = requests.get(
                    f'{api_url}/api/{session}/chats/{chat_id}/messages',
                    headers=headers,
                    params={
                        'limit': 500,
                        'downloadMedia': 'false',
                        'filter.timestamp.gte': effective_since,
                        'filter.timestamp.lte': now_ts,
                    },
                    timeout=25
                )
                if msg_resp.status_code != 200:
                    skipped += 1
                    continue
                raw = msg_resp.json()
                messages = raw if isinstance(raw, list) else (raw.get('messages') or raw.get('data') or [])
            except Exception:
                skipped += 1
                continue

            if not messages:
                skipped += 1
                continue

            # Ordena por timestamp (segundos). WAHA já filtra pelo período, mas garantimos o range.
            def _msg_ts(m):
                try:
                    return int(m.get('timestamp') or 0)
                except (TypeError, ValueError):
                    return 0
            messages.sort(key=_msg_ts)

            texts = []
            last_ts = 0
            for msg in messages:
                ts = _msg_ts(msg)
                if ts and (ts < effective_since or ts > now_ts):
                    continue
                if ts > last_ts:
                    last_ts = ts
                text, _sender = _waha_extract_text(msg, client_name)
                if text:
                    texts.append(text)

            if not texts:
                skipped += 1
                continue

            content_hash = hashlib.sha256('|'.join(texts).encode('utf-8', errors='replace')).hexdigest()[:40]
            c.execute('SELECT id FROM whatsapp_sync_log WHERE client_id = ? AND content_hash = ?',
                      (client_id, content_hash))
            if c.fetchone():
                skipped += 1
                continue

            conversation_text = '\n'.join(texts[-60:])
            today_str = datetime.now().strftime('%Y-%m-%d')
            prompt = (
                f"Analise a conversa de WhatsApp entre o usuário e '{client_name}' "
                f"(período: {period_label}). A data de hoje é {today_str}. "
                "Retorne SOMENTE um JSON válido (sem markdown, sem texto antes ou depois) no formato:\n"
                '{"resumo": "resumo conciso em 2 a 4 frases em português, formato log de atividade CRM, '
                'mencionando tópicos tratados, decisões e pendências, sem aspas nem markdown", '
                '"followup": {"data": "YYYY-MM-DD", "titulo": "descrição curta do retorno/compromisso combinado"}}\n'
                "Se houver uma data combinada de retorno, reunião, ligação ou follow-up (FUP) na conversa, "
                "preencha o objeto 'followup' resolvendo datas relativas (ex.: 'amanhã', 'semana que vem') a partir de hoje. "
                "Se NÃO houver compromisso de data combinado, use \"followup\": null.\n\n"
                f"Conversa:\n{conversation_text}"
            )

            raw_llm = _sai_simple_prompt(prompt)
            if not raw_llm:
                or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
                if or_key:
                    try:
                        or_s = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
                        or_resp = requests.post(
                            'https://openrouter.ai/api/v1/chat/completions',
                            headers={
                                'Authorization': f'Bearer {or_key}',
                                'Content-Type': 'application/json',
                                'HTTP-Referer': or_s.get('openrouter_site_url') or 'http://localhost',
                                'X-Title': or_s.get('openrouter_app_name') or 'TocaDoCoelho'
                            },
                            json={
                                'model': (or_s.get('openrouter_model') or 'stepfun/step-3.5-flash:free').strip(),
                                'messages': [
                                    {'role': 'system', 'content': 'Você é um assistente de CRM especializado em resumos de conversas de WhatsApp. Responde sempre em JSON válido.'},
                                    {'role': 'user', 'content': prompt}
                                ],
                                'temperature': 0.1
                            },
                            timeout=30
                        )
                        raw_llm = (or_resp.json().get('choices') or [{}])[0].get('message', {}).get('content', '').strip()
                    except Exception as e:
                        logger.debug(f'[_msg_ts] exceção ignorada: {e}')

            # Parse da resposta: JSON {resumo, followup} ou texto puro (fallback retrocompatível)
            summary = ''
            followup_date = ''
            followup_title = ''
            parsed = _extract_json_object_from_text(raw_llm) if raw_llm else None
            if isinstance(parsed, dict):
                summary = (parsed.get('resumo') or parsed.get('summary') or '').strip()
                fu = parsed.get('followup')
                if isinstance(fu, dict):
                    fu_data = (fu.get('data') or fu.get('date') or '').strip()
                    if re.match(r'^\d{4}-\d{2}-\d{2}$', fu_data):
                        followup_date = fu_data
                        followup_title = (fu.get('titulo') or fu.get('title') or '').strip()
            elif raw_llm:
                summary = raw_llm.strip()

            if not summary:
                summary = f'Conversa de WhatsApp com {client_name} no período de {period_label}. {len(texts)} mensagem(ns) trocada(s).'

            # Fallback: se a IA não detectou follow-up, tenta extrair datas do texto via regex
            if not followup_date:
                _dates = extract_future_commitment_dates(conversation_text)
                if _dates:
                    followup_date = _dates[0]
                    followup_title = f'Retorno combinado com {client_name}'

            activity_date = datetime.utcfromtimestamp(last_ts).strftime('%Y-%m-%d %H:%M:%S') if last_ts else now_dt.strftime('%Y-%m-%d %H:%M:%S')
            pending_items.append({
                'client_id': client_id,
                'client_name': client_name,
                'phone': phone,
                'summary': summary,
                'activity_date': activity_date,
                'message_count': len(texts),
                'content_hash': content_hash,
                'last_message_ts': last_ts,
                'period_days': period_days,
                'followup_date': followup_date,
                'followup_title': followup_title or (f'Retorno combinado com {client_name}' if followup_date else ''),
            })
            imported += 1

        _bg_task_set(task_id, {
            'status': 'done', 'progress': 100,
            'step': '✅ Análise concluída — revise os resumos!',
            'result': {
                'pending': imported,
                'skipped': skipped,
                'total_clients': total,
                'items': pending_items,
                'message': f'{imported} conversa(s) para revisar, {skipped} sem novidade.'
            }
        })
        _bg_task_cleanup(task_id)
    except Exception as exc:
        logger.exception(f'[WhatsApp Sync] Erro: {exc}')
        _bg_task_set(task_id, {'status': 'error', 'progress': 0, 'step': f'Erro: {exc}', 'error': str(exc)})
        _bg_task_cleanup(task_id)


# Servir arquivos estaticos


@app.route('/uploads/<filename>')
def serve_upload(filename):
    return send_from_directory(str(UPLOAD_DIR), filename)


# ═══════════════════════════════════════════════════════════════════════════
# CAMPANHA OFENSIVA — planejamento e plano de ação comercial assistido por IA
# Cruza o desafio do vendedor (texto livre) com o Portfólio (Soluções STF) e o
# mapeamento de contatos/contas, classificando em 4 blocos:
#   A = Pronto pra Ataque  (decisor mapeado + solução aderente)
#   B = Aprofundar Relacionamento  (há contatos, faltam decisores na área)
#   C = Iniciar do Zero  (sem mapeamento relevante)
#   D = Sem Solução no Portfólio  (relacionamento forte, mas sem oferta aderente)
# A classificação é determinística (funciona sem LLM); o LLM enriquece o
# entendimento, as áreas/decisores e a aderência semântica das ofertas.
# ═══════════════════════════════════════════════════════════════════════════

def _campaign_llm_text(system_prompt, user_prompt):
    """Chama o LLM seguindo a ordem padrão do projeto: OpenRouter primeiro, SAI fallback.
    Retorna (texto_bruto, fonte) ou (None, 'heuristic') se nenhum LLM responder."""
    or_key = _resolve_setting('openrouter_api_key', 'OPENROUTER_API_KEY')
    if or_key:
        try:
            or_settings = _load_app_settings_map(['openrouter_model', 'openrouter_site_url', 'openrouter_app_name'])
            model = (or_settings.get('openrouter_model') or 'stepfun/step-3.5-flash:free').strip() or 'stepfun/step-3.5-flash:free'
            site_url = (or_settings.get('openrouter_site_url') or 'http://localhost').strip() or 'http://localhost'
            app_name = (or_settings.get('openrouter_app_name') or 'TocaDoCoelho').strip() or 'TocaDoCoelho'
            payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_prompt}
                ],
                'temperature': 0.2
            }
            req = urllib.request.Request(
                'https://openrouter.ai/api/v1/chat/completions',
                data=json.dumps(payload, ensure_ascii=False).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {or_key}',
                    'HTTP-Referer': site_url,
                    'X-Title': app_name
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            choices = data.get('choices') or []
            content = (choices[0].get('message') or {}).get('content', '') if choices else ''
            if content and str(content).strip():
                return content, 'OpenRouter'
        except Exception as e:
            logger.warning(f'[Campaign][OpenRouter] falha: {e}')
    raw = _sai_simple_prompt(f"{system_prompt}\n\n{user_prompt}")
    if raw and str(raw).strip():
        return raw, 'SAI'
    return None, 'heuristic'


def _campaign_llm_json(system_prompt, user_prompt):
    text, source = _campaign_llm_text(system_prompt, user_prompt)
    if not text:
        return None, source
    parsed = _extract_json_object_from_text(text)
    return (parsed if isinstance(parsed, dict) else None), source


def _campaign_load_offers(c):
    c.execute('SELECT * FROM portfolio_offers ORDER BY id')
    offers = [dict_from_row(r) for r in c.fetchall()]
    for o in offers:
        c.execute('SELECT pain, solution FROM portfolio_offer_items WHERE offer_id = ? ORDER BY sort_order, id', (o['id'],))
        o['items'] = [dict_from_row(r) for r in c.fetchall()]
    return offers


def _campaign_offer_blob(offer):
    parts = [offer.get('title') or '', offer.get('summary') or '']
    for it in (offer.get('items') or []):
        parts.append((it.get('pain') or '') + ' ' + (it.get('solution') or ''))
    return ' '.join(parts).lower()


def _campaign_keyword_match_offers(challenge, offers):
    """Fallback determinístico: casa palavras do desafio contra o texto das ofertas."""
    words = set(re.findall(r'[a-zá-úâ-û0-9]{4,}', (challenge or '').lower()))
    stop = {'para', 'como', 'mais', 'nossa', 'nosso', 'clientes', 'cliente', 'setor', 'area', 'área', 'quero', 'aumentar', 'penetração', 'penetracao', 'atuando', 'soluções', 'solucoes', 'solução', 'solucao'}
    words = {w for w in words if w not in stop}
    scored = []
    for off in offers:
        blob = _campaign_offer_blob(off)
        hits = sum(1 for w in words if w in blob)
        if hits >= 2:
            scored.append((hits, off))
    scored.sort(key=lambda x: -x[0])
    return [off for _, off in scored]


def _campaign_seniority(position):
    p = (position or '').lower()
    clevel_keys = [
        'ceo', 'cfo', 'cio', 'cto', 'coo', 'cmo', 'cdo', 'chro', 'cro', 'clo', 'cso', 'cpo',
        'ciso', 'cao', 'cco', 'cxo', 'chief', 'presidente', 'president',
        'vice-pres', 'vice pres', ' vp', 'c-level', 'board',
        'sócio', 'socio', 'founder', 'fundador', 'owner', 'proprietár',
        'managing director', 'managing partner', 'sócio-diretor', 'socio-diretor',
        'diretor executivo', 'diretor-executivo', 'diretor geral', 'diretor-geral',
        'partner', 'sócio gerente', 'socio gerente',
    ]
    if any(k in p for k in clevel_keys):
        return 'clevel'
    if any(k in p for k in ['diretor', 'director', 'head', 'superintend', 'vp']):
        return 'diretor'
    if any(k in p for k in ['gerente', 'manager', 'coordenad', 'supervisor', 'líder', 'lider', 'lead']):
        return 'gerente'
    return 'outro'


def _campaign_seniority_rank(position):
    return {'clevel': 0, 'diretor': 1, 'gerente': 2, 'outro': 3}.get(_campaign_seniority(position), 3)


def _campaign_area_match(contact, areas_lower):
    if not areas_lower:
        return True
    blob = ((contact.get('area_of_activity') or '') + ' ' + (contact.get('position') or '')).lower()
    return any(a and a in blob for a in areas_lower)


def _campaign_account_contacts(c, account_name):
    c.execute("SELECT * FROM clients WHERE COALESCE(is_archived,0)=0 AND lower(trim(company))=lower(trim(?))", (account_name,))
    return [dict_from_row(r) for r in c.fetchall()]


def _campaign_has_mapping(c, account_name):
    c.execute(
        """SELECT COUNT(*) FROM environment_responses er
           JOIN clients cl ON cl.id = er.client_id
           WHERE er.response IS NOT NULL AND trim(er.response) <> ''
             AND lower(trim(cl.company)) = lower(trim(?))""",
        (account_name,)
    )
    return (c.fetchone()[0] or 0) > 0


def _campaign_has_kanban(c, account_id):
    c.execute('SELECT COUNT(*) FROM kanban_cards WHERE account_id = ?', (account_id,))
    return (c.fetchone()[0] or 0) > 0


def _campaign_classify(contacts, areas_lower, has_mapping, has_kanban, has_solution):
    """Classificação determinística por força de mapeamento. Retorna (bloco, score, contato_primário)."""
    non_cold = [c for c in contacts if not c.get('is_cold_contact')]
    in_area = [c for c in contacts if _campaign_area_match(c, areas_lower)]
    targets = [c for c in contacts if c.get('is_target')]
    decisionmakers = [c for c in contacts if _campaign_seniority(c.get('position')) in ('clevel', 'diretor')]
    strong_pool = [
        c for c in contacts
        if (c.get('is_target') or _campaign_seniority(c.get('position')) in ('clevel', 'diretor'))
        and _campaign_area_match(c, areas_lower)
        and not c.get('is_cold_contact')
    ]

    score = 0
    if contacts:
        score += 15
    if in_area:
        score += 20
    if targets:
        score += 25
    if decisionmakers:
        score += 20
    if has_mapping:
        score += 10
    if has_kanban:
        score += 10
    score = min(score, 100)

    if strong_pool:
        block = 'A' if has_solution else 'D'
    elif non_cold:
        block = 'B'
    else:
        block = 'C'

    if strong_pool:
        strong_pool.sort(key=lambda c: (_campaign_seniority_rank(c.get('position')), 0 if c.get('is_target') else 1))
        primary = strong_pool[0]
    elif in_area:
        primary = sorted(in_area, key=lambda c: _campaign_seniority_rank(c.get('position')))[0]
    elif non_cold:
        primary = sorted(non_cold, key=lambda c: _campaign_seniority_rank(c.get('position')))[0]
    elif contacts:
        primary = contacts[0]
    else:
        primary = None
    return block, score, primary


def _campaign_rationale(block, account_name, primary, areas, offer_title):
    pos = (primary or {}).get('position') or ''
    pname = (primary or {}).get('name') or ''
    area_label = (areas[0] if areas else 'área-alvo')
    if block == 'A':
        who = f"{pname} ({pos})" if pname else "decisor mapeado"
        return f"Relacionamento forte: {who}. Solução aderente: {offer_title or 'a definir'}."
    if block == 'B':
        return f"Mapeamento parcial: há contatos, mas é preciso aprofundar decisores em {area_label}."
    if block == 'D':
        who = f"{pname} ({pos})" if pname else "decisores mapeados"
        return f"Relacionamento forte ({who}), porém sem solução aderente no portfólio para este desafio."
    return "Sem mapeamento relevante: iniciar relacionamento do zero."


def _campaign_contacts_in_area(contacts, area_lower):
    """Contatos cuja área de atuação OU cargo casa com a área-alvo, ordenados por senioridade."""
    matched = []
    for ct in contacts:
        blob = ((ct.get('area_of_activity') or '') + ' ' + (ct.get('position') or '')).lower()
        if area_lower and area_lower in blob:
            matched.append(ct)
    matched.sort(key=lambda ct: (_campaign_seniority_rank(ct.get('position')),
                                  0 if ct.get('is_target') else 1,
                                  1 if ct.get('is_cold_contact') else 0))
    return matched


def _campaign_rel_ctx(act_count, last_act_desc, has_map, has_kan):
    rel_parts = []
    if act_count > 0:
        rel_parts.append(f"{act_count} interação(ões) registrada(s)")
    if has_map:
        rel_parts.append("ambiente mapeado")
    if has_kan:
        rel_parts.append("oportunidade(s) no funil")
    rel_ctx = (" Histórico: " + ", ".join(rel_parts) + ".") if rel_parts else ""
    if last_act_desc:
        snip = last_act_desc.strip()[:120]
        if len(last_act_desc.strip()) > 120:
            snip += "..."
        rel_ctx += f" Última interação: \"{snip}\""
    return rel_ctx


def _campaign_build_actions(block, contacts, areas, offer_title, ctx):
    """Gera UMA OU MAIS ações por conta — uma por ângulo (área/decisor) do desafio.

    Para cada área-alvo da campanha, identifica o melhor decisor mapeado naquela área e
    cria uma ação específica (apresentação se houver solução, ou aprofundamento). Áreas sem
    decisor mapeado viram ações de mapeamento daquele ângulo. Assim o desafio multi-área
    (ex.: CIO p/ tecnologia + CMO p/ jornada do paciente) vira ações distintas e justificadas.
    Cada item retornado: {type, title, description, contact, angle_area}.
    """
    rel_ctx = _campaign_rel_ctx(ctx.get('act_count', 0), ctx.get('last_act_desc', ''),
                                ctx.get('has_map', False), ctx.get('has_kan', False))
    areas_eff = [a for a in (areas or []) if a] or ['área-alvo']

    # Conta fria / sem contatos: prospecção única
    if block == 'C' or not contacts:
        area_label = areas_eff[0]
        return [{'type': 'prospeccao', 'angle_area': area_label, 'contact': None,
                 'title': "Iniciar prospecção e mapear decisores",
                 'description': (f"Conta fria: sem mapeamento de contatos em {area_label}. "
                                 f"Iniciar prospecção ativa — identificar organograma, decisores e dores "
                                 f"antes de qualquer abordagem." + rel_ctx)}]

    actions = []
    seen = set()
    for area in areas_eff:
        pool = _campaign_contacts_in_area(contacts, area.lower())
        dm = [ct for ct in pool if _campaign_seniority(ct.get('position')) in ('clevel', 'diretor') and not ct.get('is_cold_contact')]
        warm = [ct for ct in pool if not ct.get('is_cold_contact')]
        chosen = dm[0] if dm else (warm[0] if warm else (pool[0] if pool else None))

        if chosen is None:
            key = ('none', area.lower())
            if key in seen:
                continue
            seen.add(key)
            actions.append({'type': 'mapeamento', 'angle_area': area, 'contact': None,
                            'title': f"Mapear decisor de {area}",
                            'description': (f"Nenhum decisor de {area} mapeado nesta conta. Identificar e qualificar "
                                            f"o responsável por {area} (C-level/Diretor) para abrir esta frente." + rel_ctx)})
            continue

        key = (chosen.get('id'), area.lower())
        if key in seen:
            continue
        seen.add(key)
        cpos = chosen.get('position') or ''
        cname = chosen.get('name') or ''
        who = f"{cname} ({cpos})" if cpos else cname
        sen = _campaign_seniority(chosen.get('position'))
        sen_label = {'clevel': 'C-level', 'diretor': 'Diretoria', 'gerente': 'Gerência'}.get(sen, 'contato')
        why = f"Escolhido por ser {sen_label} responsável por {area}" + (" e contato target" if chosen.get('is_target') else "") + "."

        if block == 'A' and offer_title:
            actions.append({'type': 'apresentacao', 'angle_area': area, 'contact': chosen,
                            'title': f"Apresentar {offer_title} — ângulo {area}",
                            'description': (f"Conduzir apresentação executiva para {who}. {why} "
                                            f"Conectar {offer_title} ao ângulo de {area}, alinhando a solução à dor "
                                            f"específica desta frente." + rel_ctx)})
        else:
            actions.append({'type': 'mapeamento', 'angle_area': area, 'contact': chosen,
                            'title': f"Aprofundar relacionamento em {area}",
                            'description': (f"Frente de {area} com {who}. {why} Aprofundar a dor, validar fit e "
                                            f"preparar terreno para apresentar a solução nesta frente." + rel_ctx)})

    if block == 'D':
        prim = contacts[0] if contacts else None
        d_area = areas_eff[0]
        d_action = {'type': 'portfolio', 'angle_area': d_area, 'contact': prim,
                    'title': "Explorar portfólio para o desafio",
                    'description': (f"Relacionamento maduro, porém sem solução aderente cadastrada para {d_area}. "
                                    f"Explorar portfólio interno ou identificar parceiros que atendam esta demanda." + rel_ctx)}
        actions = [d_action] + actions

    if not actions:
        prim = contacts[0] if contacts else None
        actions = [{'type': 'mapeamento', 'angle_area': areas_eff[0], 'contact': prim,
                    'title': "Aprofundar mapeamento da conta",
                    'description': ("Aprofundar mapeamento de stakeholders e qualificar a dor." + rel_ctx)}]
    return actions


def _campaign_market_insight(account_name, sector, challenge, areas, offer_title):
    """Busca contexto de mercado real via Tavily e sintetiza um 'motivo de abordagem' via LLM.

    Ordem: Tavily (dados reais) → LLM sintetiza com base nas evidências. Sem Tavily, o LLM faz
    uma leitura objetiva do setor. Retorna (insight_text, sources_list)."""
    evidence, sources = [], []
    api_key = _resolve_setting('tavily_api_key', 'TAVILY_API_KEY')
    if api_key:
        try:
            areas_str = ', '.join(areas) if areas else ''
            query = f'"{account_name}" {sector or ""} notícias investimentos estratégia movimentos {areas_str}'.strip()
            results = _run_tavily_request(api_key, query, max_results=5)
            evidence = _extract_tavily_evidence(results, max_items=5)
            sources = [{'title': e.get('title'), 'url': e.get('url')} for e in evidence if e.get('url')][:4]
        except Exception as e:
            logger.warning(f'[Campaign][insight] Tavily falhou p/ {account_name}: {e}')

    evidence_blob = '\n'.join(f"- {e.get('snippet','')[:300]}" for e in evidence if e.get('snippet'))
    sys_p = ('Você é um analista de inteligência comercial B2B. Escreva em português do Brasil um parágrafo '
             'curto (2 a 4 frases), direto e específico, sem markdown, sem títulos e sem citar URLs.')
    base = (f"Conta: {account_name} (setor: {sector or 'n/d'}).\n"
            f"Desafio comercial: {challenge}.\n"
            f"Áreas-alvo: {', '.join(areas) or 'n/d'}. Solução considerada: {offer_title or 'a definir'}.\n\n")
    if evidence_blob:
        user_p = (base + f"Evidências de mercado encontradas:\n{evidence_blob}\n\n"
                  "Com base NAS EVIDÊNCIAS acima, escreva o 'motivo da abordagem': o momento de mercado e os "
                  "movimentos da conta que tornam ESTE plano relevante AGORA. Cite fatos concretos das evidências.")
    else:
        user_p = (base + "Escreva o 'motivo da abordagem': uma leitura objetiva e específica do provável momento "
                  "de mercado do setor que torna este plano relevante agora. Evite generalidades vazias.")
    text, _src = _campaign_llm_text(sys_p, user_p)
    return (text or '').strip(), sources


def _campaign_parse_date(s):
    try:
        return datetime.strptime((s or '').strip(), '%Y-%m-%d')
    except Exception:
        return datetime.now()


def _campaign_add_weeks(start_dt, weeks):
    return (start_dt + timedelta(weeks=weeks)).strftime('%Y-%m-%d')


def _campaign_analyze_core(challenge_text):
    conn = get_db()
    c = conn.cursor()
    offers = _campaign_load_offers(c)
    c.execute("SELECT DISTINCT area_of_activity FROM clients WHERE area_of_activity IS NOT NULL AND trim(area_of_activity) <> ''")
    known_areas = [r[0] for r in c.fetchall()]
    c.execute("SELECT id, name, logo_url, sector FROM accounts WHERE is_target = 1 ORDER BY name")
    target_accounts = [dict_from_row(r) for r in c.fetchall()]
    conn.close()

    offers_brief = '; '.join(f"#{o['id']} {o['title']}" for o in offers[:40]) or '(nenhuma oferta cadastrada)'
    sys_p = 'Você é um consultor comercial estratégico B2B sênior. Responda SEMPRE e SOMENTE com um JSON válido, sem markdown.'
    user_p = (
        f"Desafio comercial do vendedor: \"{challenge_text}\".\n"
        f"Áreas de cliente já conhecidas no sistema: {', '.join(known_areas) or 'nenhuma'}.\n"
        f"Ofertas do portfólio (id e título): {offers_brief}.\n\n"
        "Tarefas:\n"
        "1) Identifique as ÁREAS do cliente a atacar (ex.: TI, Operações, Financeiro, Jurídico, Marketing, Infraestrutura, Manufatura, Sistemas, Segurança da Informação).\n"
        "2) Identifique os NÍVEIS DECISORES relevantes (ex.: C-level, Diretores, Gerentes, e cargos específicos como CIO, CTO).\n"
        "3) Aponte quais OFERTAS do portfólio (pelos ids) são aderentes ao desafio.\n"
        "4) Escreva um ENTENDIMENTO resumido em 2 a 3 frases.\n\n"
        'Retorne EXATAMENTE: {"entendimento":"...","areas":["..."],"decisores":["..."],"oferta_ids":[1,2],"tem_solucao":true}'
    )
    parsed, source = _campaign_llm_json(sys_p, user_p)

    entendimento, areas, decisores, offer_ids = '', [], [], []
    if parsed:
        entendimento = (parsed.get('entendimento') or '').strip()
        areas = [a.strip() for a in (parsed.get('areas') or []) if isinstance(a, str) and a.strip()]
        decisores = [d.strip() for d in (parsed.get('decisores') or []) if isinstance(d, str) and d.strip()]
        valid_ids = {o['id'] for o in offers}
        for i in (parsed.get('oferta_ids') or []):
            try:
                iv = int(i)
                if iv in valid_ids:
                    offer_ids.append(iv)
            except Exception:
                continue

    if not offer_ids:
        offer_ids = [o['id'] for o in _campaign_keyword_match_offers(challenge_text, offers)[:5]]
    has_solution = bool(offer_ids)
    if not areas:
        cl = (challenge_text or '').lower()
        areas = [a for a in known_areas if a and a.lower() in cl][:6]
    if not decisores:
        decisores = ['C-level', 'Diretores', 'Gerentes']
    if not entendimento:
        entendimento = f"Atacar o desafio descrito: {(challenge_text or '').strip()[:240]}"

    portfolio_note = ''
    if not offers:
        portfolio_note = 'Você ainda não cadastrou ofertas no Portfólio (Soluções STF). Cadastre para que o plano recomende soluções específicas.'
    elif not has_solution:
        portfolio_note = 'Nenhuma oferta cadastrada parece aderente a este desafio. Considere aprofundar/registrar soluções da empresa para este segmento.'

    offers_by_id = {o['id']: o for o in offers}
    matched_offers = [{'id': i, 'title': offers_by_id[i]['title']} for i in offer_ids if i in offers_by_id]
    return {
        'entendimento': entendimento,
        'areas': areas,
        'decisores': decisores,
        'offer_ids': offer_ids,
        'matched_offers': matched_offers,
        'has_solution': has_solution,
        'portfolio_note': portfolio_note,
        'target_accounts': target_accounts,
        'target_count': len(target_accounts),
        'llm_source': source,
    }


def _campaign_generate_core(payload, progress_cb=None):
    title = payload['title']
    challenge = payload['challenge_text']
    areas = payload.get('areas') or []
    decisores = payload.get('decisores') or []
    offer_ids = payload.get('offer_ids') or []
    start_date = payload['start_date']
    capacity = max(1, int(payload.get('weekly_capacity') or 5))
    account_ids = payload['account_ids']
    llm_source = payload.get('llm_source', '')

    def _emit(pct, step):
        if progress_cb:
            try:
                progress_cb(pct, step)
            except Exception as e:
                logger.debug(f'[_emit] exceção ignorada: {e}')

    conn = get_db()
    c = conn.cursor()
    offers = _campaign_load_offers(c)
    offers_by_id = {o['id']: o for o in offers}
    matched_offers = [offers_by_id[i] for i in offer_ids if i in offers_by_id]
    has_solution = bool(matched_offers)
    primary_offer = matched_offers[0] if matched_offers else None
    areas_lower = [a.lower() for a in areas if a]

    built = []
    n_accs = len(account_ids)
    for i, aid in enumerate(account_ids):
        c.execute('SELECT * FROM accounts WHERE id = ?', (aid,))
        acc = dict_from_row(c.fetchone())
        if not acc:
            continue
        contacts = _campaign_account_contacts(c, acc['name'])
        has_map = _campaign_has_mapping(c, acc['name'])
        has_kan = _campaign_has_kanban(c, acc['id'])
        block, score, primary = _campaign_classify(contacts, areas_lower, has_map, has_kan, has_solution)
        c.execute(
            "SELECT COUNT(*) FROM activities a JOIN clients cl ON cl.id=a.client_id WHERE lower(trim(cl.company))=lower(trim(?))",
            (acc['name'],)
        )
        act_count = c.fetchone()[0] or 0
        c.execute(
            "SELECT a.description FROM activities a JOIN clients cl ON cl.id=a.client_id WHERE lower(trim(cl.company))=lower(trim(?)) AND (a.description IS NOT NULL AND trim(a.description)<>'') ORDER BY a.activity_date DESC LIMIT 1",
            (acc['name'],)
        )
        _row = c.fetchone()
        last_act_desc = (_row[0] if _row else '') or ''

        offer_title = primary_offer['title'] if (block == 'A' and primary_offer) else None
        _emit(40 + int(45 * i / max(1, n_accs)), f"Pesquisando mercado: {acc['name']} ({i + 1}/{n_accs})")
        insight, sources = _campaign_market_insight(acc['name'], acc.get('sector'), challenge, areas, offer_title)

        built.append({'acc': acc, 'block': block, 'score': score, 'primary': primary,
                      'contacts': contacts, 'act_count': act_count, 'last_act_desc': last_act_desc,
                      'has_map': has_map, 'has_kan': has_kan,
                      'market_insight': insight, 'insight_sources': sources})

    _emit(88, 'Montando blocos e cronograma...')
    order_rank = {'A': 0, 'B': 1, 'C': 2, 'D': 3}
    built.sort(key=lambda b: (order_rank.get(b['block'], 9), -b['score'], (b['acc']['name'] or '').lower()))

    c.execute(
        '''INSERT INTO campaigns
           (title, objective_text, llm_summary, areas_json, decision_levels_json, offer_ids_json,
            has_solution, portfolio_note, start_date, weekly_capacity, llm_source, status, updated_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP)''',
        (title, challenge, payload.get('summary', ''),
         json.dumps(areas, ensure_ascii=False), json.dumps(decisores, ensure_ascii=False), json.dumps(offer_ids),
         1 if has_solution else 0, payload.get('portfolio_note', ''), start_date, capacity, llm_source, 'Ativo')
    )
    campaign_id = c.lastrowid

    start_dt = _campaign_parse_date(start_date)
    idx = 0
    for ord_idx, b in enumerate(built):
        acc = b['acc']
        block = b['block']
        primary = b['primary']
        offer_title = primary_offer['title'] if (block == 'A' and primary_offer) else None
        offer_id = primary_offer['id'] if (block == 'A' and primary_offer) else None
        rationale = _campaign_rationale(block, acc['name'], primary, areas, offer_title)
        c.execute(
            '''INSERT INTO campaign_accounts
               (campaign_id, account_id, account_name, block_type, readiness_score, offer_id, offer_title,
                rationale, sort_order, market_insight, insight_sources)
               VALUES (?,?,?,?,?,?,?,?,?,?,?)''',
            (campaign_id, acc['id'], acc['name'], block, b['score'], offer_id, offer_title,
             rationale, ord_idx, b.get('market_insight', ''),
             json.dumps(b.get('insight_sources') or [], ensure_ascii=False))
        )
        ca_id = c.lastrowid
        ctx = {'act_count': b.get('act_count', 0), 'last_act_desc': b.get('last_act_desc', ''),
               'has_map': b.get('has_map', False), 'has_kan': b.get('has_kan', False)}
        actions = _campaign_build_actions(block, b.get('contacts') or [], areas, offer_title, ctx)
        for a in actions:
            ct = a.get('contact') or {}
            week = idx // capacity
            due = _campaign_add_weeks(start_dt, week)
            c.execute(
                '''INSERT INTO campaign_actions
                   (campaign_account_id, title, description, action_type, contact_id, contact_name,
                    scheduled_week, due_date, status, sort_order, angle_area, updated_at)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP)''',
                (ca_id, a['title'], a['description'], a['type'],
                 ct.get('id'), ct.get('name'),
                 week + 1, due, 'pending', idx, a.get('angle_area', ''))
            )
            idx += 1

    conn.commit()
    conn.close()
    return campaign_id


def _campaign_regenerate_core(cid, progress_cb=None):
    def _emit(pct, step):
        if progress_cb:
            try:
                progress_cb(pct, step)
            except Exception as e:
                logger.debug(f'[_emit] exceção ignorada: {e}')

    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT * FROM campaigns WHERE id = ?', (cid,))
    camp = dict_from_row(c.fetchone())
    if not camp:
        conn.close()
        raise RuntimeError('Campanha não encontrada.')
    areas = json.loads(camp.get('areas_json') or '[]')
    challenge = camp.get('objective_text') or ''
    areas_lower = [a.lower() for a in areas if a]

    offers = _campaign_load_offers(c)
    offers_by_id = {o['id']: o for o in offers}
    stored_ids = set(json.loads(camp.get('offer_ids_json') or '[]'))
    kw_matched = _campaign_keyword_match_offers(challenge, offers)
    matched = [offers_by_id[i] for i in stored_ids if i in offers_by_id]
    for o in kw_matched:
        if o['id'] not in stored_ids:
            matched.append(o)
    has_solution = bool(matched)
    primary_offer = matched[0] if matched else None

    c.execute('SELECT * FROM campaign_accounts WHERE campaign_id = ? ORDER BY sort_order, id', (cid,))
    cas = [dict_from_row(r) for r in c.fetchall()]
    n_cas = len(cas)
    capacity = max(1, int(camp.get('weekly_capacity') or 5))
    start_dt = _campaign_parse_date(camp.get('start_date') or '')
    for ci, ca in enumerate(cas):
        contacts = _campaign_account_contacts(c, ca['account_name'])
        has_map = _campaign_has_mapping(c, ca['account_name'])
        has_kan = _campaign_has_kanban(c, ca['account_id'])
        new_block, score, primary = _campaign_classify(contacts, areas_lower, has_map, has_kan, has_solution)
        c.execute(
            "SELECT COUNT(*) FROM activities a JOIN clients cl ON cl.id=a.client_id WHERE lower(trim(cl.company))=lower(trim(?))",
            (ca['account_name'],)
        )
        act_count = c.fetchone()[0] or 0
        c.execute(
            "SELECT a.description FROM activities a JOIN clients cl ON cl.id=a.client_id WHERE lower(trim(cl.company))=lower(trim(?)) AND (a.description IS NOT NULL AND trim(a.description)<>'') ORDER BY a.activity_date DESC LIMIT 1",
            (ca['account_name'],)
        )
        _row = c.fetchone()
        last_act_desc = (_row[0] if _row else '') or ''
        old_block = ca['block_type']
        offer_title = primary_offer['title'] if (new_block == 'A' and primary_offer) else None
        offer_id = primary_offer['id'] if (new_block == 'A' and primary_offer) else None
        rationale = _campaign_rationale(new_block, ca['account_name'], primary, areas, offer_title)

        _emit(40 + int(50 * ci / max(1, n_cas)), f"Reavaliando mercado: {ca['account_name']} ({ci + 1}/{n_cas})")
        insight, sources = _campaign_market_insight(ca['account_name'], None, challenge, areas, offer_title)
        c.execute(
            'UPDATE campaign_accounts SET block_type=?, readiness_score=?, offer_id=?, offer_title=?, rationale=?, market_insight=?, insight_sources=? WHERE id=?',
            (new_block, score, offer_id, offer_title, rationale, insight,
             json.dumps(sources or [], ensure_ascii=False), ca['id'])
        )

        ctx = {'act_count': act_count, 'last_act_desc': last_act_desc, 'has_map': has_map, 'has_kan': has_kan}
        desired = _campaign_build_actions(new_block, contacts, areas, offer_title, ctx)

        # Ações existentes (preservar status/logs). Matching por (contact_id, angle_area).
        c.execute('SELECT * FROM campaign_actions WHERE campaign_account_id = ? ORDER BY sort_order, id', (ca['id'],))
        existing = [dict_from_row(r) for r in c.fetchall()]
        used_existing = set()

        def _match(want_cid, want_area):
            # 1) match exato contato+ângulo; 2) só ângulo; 3) só contato; 4) qualquer não usada
            for pref in (lambda a: a['id'] not in used_existing and (a.get('contact_id') or None) == (want_cid or None) and (a.get('angle_area') or '').lower() == (want_area or '').lower(),
                         lambda a: a['id'] not in used_existing and (a.get('angle_area') or '').lower() == (want_area or '').lower(),
                         lambda a: a['id'] not in used_existing and (a.get('contact_id') or None) == (want_cid or None),
                         lambda a: a['id'] not in used_existing):
                for a in existing:
                    if pref(a):
                        return a
            return None

        sort_base = ca['sort_order'] * 100
        for k, want in enumerate(desired):
            ct = want.get('contact') or {}
            want_cid = ct.get('id')
            want_area = want.get('angle_area', '')
            match = _match(want_cid, want_area)
            if match:
                used_existing.add(match['id'])
                c.execute(
                    'UPDATE campaign_actions SET title=?, description=?, action_type=?, contact_id=?, contact_name=?, angle_area=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                    (want['title'], want['description'], want['type'], want_cid, ct.get('name'), want_area, match['id'])
                )
                if old_block != new_block:
                    msg = f"Plano atualizado pela IA: bloco {old_block} → {new_block}."
                    if offer_title:
                        msg += f" Solução recomendada: {offer_title}."
                    c.execute("INSERT INTO campaign_action_logs (action_id, log_text, log_type) VALUES (?,?, 'system')", (match['id'], msg))
            else:
                week = (sort_base + k) // capacity
                due = _campaign_add_weeks(start_dt, week)
                c.execute(
                    '''INSERT INTO campaign_actions
                       (campaign_account_id, title, description, action_type, contact_id, contact_name,
                        scheduled_week, due_date, status, sort_order, angle_area, updated_at)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP)''',
                    (ca['id'], want['title'], want['description'], want['type'],
                     want_cid, ct.get('name'), week + 1, due, 'pending', sort_base + k, want_area)
                )

        # Ações antigas não correspondidas: preservar se tiverem trabalho do usuário, senão remover.
        for a in existing:
            if a['id'] in used_existing:
                continue
            c.execute('SELECT COUNT(*) FROM campaign_action_logs WHERE action_id=? AND log_type=?', (a['id'], 'user'))
            has_user_log = (c.fetchone()[0] or 0) > 0
            if has_user_log or a.get('status') != 'pending':
                c.execute("INSERT INTO campaign_action_logs (action_id, log_text, log_type) VALUES (?,?, 'system')",
                          (a['id'], "A IA não priorizou mais esta frente nesta atualização, mas ela foi mantida por já ter execução registrada."))
            else:
                c.execute('DELETE FROM campaign_action_logs WHERE action_id=?', (a['id'],))
                c.execute('DELETE FROM campaign_actions WHERE id=?', (a['id'],))

    new_offer_ids = [o['id'] for o in matched]
    portfolio_note = '' if has_solution else 'Nenhuma oferta aderente no portfólio. Registre soluções da empresa para este segmento.'
    c.execute(
        'UPDATE campaigns SET has_solution=?, offer_ids_json=?, portfolio_note=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
        (1 if has_solution else 0, json.dumps(new_offer_ids), portfolio_note, cid)
    )

    c.execute('SELECT id, block_type, readiness_score, account_name FROM campaign_accounts WHERE campaign_id = ?', (cid,))
    rows = [dict_from_row(r) for r in c.fetchall()]
    order_rank = {'A': 0, 'B': 1, 'C': 2, 'D': 3}
    rows.sort(key=lambda r: (order_rank.get(r['block_type'], 9), -(r['readiness_score'] or 0), (r['account_name'] or '').lower()))
    for i, r in enumerate(rows):
        c.execute('UPDATE campaign_accounts SET sort_order=? WHERE id=?', (i, r['id']))

    conn.commit()
    conn.close()


def _campaign_timeline(camp, accs):
    start_dt = _campaign_parse_date(camp.get('start_date') or '')
    weeks_count = camp['stats']['weeks'] or 1
    weeks = []
    for w in range(weeks_count):
        ws = start_dt + timedelta(weeks=w)
        we = ws + timedelta(days=6)
        weeks.append({'week': w + 1, 'start': ws.strftime('%d/%m'), 'end': we.strftime('%d/%m'), 'counts': {}, 'total': 0})
    for a in accs:
        for act in a['actions']:
            wi = (act.get('scheduled_week') or 1) - 1
            if 0 <= wi < len(weeks):
                t = act.get('action_type') or 'outro'
                weeks[wi]['counts'][t] = weeks[wi]['counts'].get(t, 0) + 1
                weeks[wi]['total'] += 1
    return weeks


def _campaign_serialize(c, campaign_id):
    c.execute('SELECT * FROM campaigns WHERE id = ?', (campaign_id,))
    camp = dict_from_row(c.fetchone())
    if not camp:
        return None
    camp['areas'] = json.loads(camp.get('areas_json') or '[]')
    camp['decisores'] = json.loads(camp.get('decision_levels_json') or '[]')
    camp['offer_ids'] = json.loads(camp.get('offer_ids_json') or '[]')

    c.execute('SELECT * FROM campaign_accounts WHERE campaign_id = ? ORDER BY sort_order, id', (campaign_id,))
    accs = [dict_from_row(r) for r in c.fetchall()]
    total = done = in_prog = 0
    for a in accs:
        c.execute('SELECT logo_url FROM accounts WHERE id = ?', (a['account_id'],))
        row = c.fetchone()
        a['logo_url'] = row['logo_url'] if row else None
        try:
            a['insight_sources'] = json.loads(a.get('insight_sources') or '[]')
        except Exception:
            a['insight_sources'] = []
        c.execute('SELECT * FROM campaign_actions WHERE campaign_account_id = ? ORDER BY sort_order, id', (a['id'],))
        acts = [dict_from_row(r) for r in c.fetchall()]
        for act in acts:
            c.execute('SELECT * FROM campaign_action_logs WHERE action_id = ? ORDER BY datetime(created_at) DESC, id DESC', (act['id'],))
            act['logs'] = [dict_from_row(l) for l in c.fetchall()]
            total += 1
            if act['status'] == 'done':
                done += 1
            elif act['status'] == 'in_progress':
                in_prog += 1
        a['actions'] = acts
        a_total = len(acts)
        a_done = sum(1 for x in acts if x['status'] == 'done')
        a['progress'] = round(100 * a_done / a_total) if a_total else 0

    blocks = {'A': [], 'B': [], 'C': [], 'D': []}
    for a in accs:
        blocks.setdefault(a['block_type'], []).append(a)
    all_weeks = [act['scheduled_week'] for a in accs for act in a['actions']]
    camp['stats'] = {
        'total_actions': total, 'done': done, 'in_progress': in_prog, 'pending': total - done - in_prog,
        'progress': round(100 * done / total) if total else 0,
        'accounts_count': len(accs),
        'weeks': max(all_weeks) if all_weeks else 0,
    }
    camp['blocks'] = blocks
    camp['accounts'] = accs
    camp['timeline'] = _campaign_timeline(camp, accs)
    return camp


@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    if path.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory(app.static_folder, path)


def _home_status_thresholds(c):
    """Carrega thresholds universais + regras por cargo (igual ao home_overview)."""
    c.execute("SELECT key, value FROM app_settings WHERE key IN ('status_green_days', 'status_yellow_days')")
    s_map = {row['key']: row['value'] for row in c.fetchall()}
    try:
        green_days = int(s_map.get('status_green_days') or 7)
    except Exception:
        green_days = 7
    try:
        yellow_days = int(s_map.get('status_yellow_days') or 14)
    except Exception:
        yellow_days = 14
    c.execute('SELECT position, green_days, yellow_days FROM status_rules')
    rules_map = {
        (row['position'] or '').strip().lower(): (int(row['green_days']), int(row['yellow_days']))
        for row in c.fetchall() if row['position']
    }
    def get_thresholds(position):
        return rules_map.get((position or '').strip().lower(), (green_days, yellow_days))
    return get_thresholds


def _home_sort_key(name):
    """Chave de ordenação alfabética insensível a maiúsculas e acentos (pt-BR)."""
    s = unicodedata.normalize('NFKD', str(name or ''))
    return ''.join(ch for ch in s if not unicodedata.combining(ch)).casefold()


def _home_status_for(last_date, green, yellow, now_dt):
    """Retorna ('em_dia'|'atencao'|'atrasado', dias_sem_contato|None)."""
    if not last_date:
        return ('atrasado', None)
    try:
        base = str(last_date).replace('T', ' ').split('.')[0].split('+')[0].strip()
        d = datetime.strptime(base[:19], '%Y-%m-%d %H:%M:%S') if len(base) >= 19 else datetime.strptime(base[:10], '%Y-%m-%d')
        days = (now_dt - d).days
    except Exception:
        return ('atrasado', None)
    if days < green:
        return ('em_dia', days)
    elif days < yellow:
        return ('atencao', days)
    return ('atrasado', days)


@app.errorhandler(Exception)
def handle_unexpected_exception(error):
    if isinstance(error, HTTPException):
        return error
    logger.exception(f'[Unhandled] Erro inesperado: {error}')
    return jsonify({'error': 'Erro interno inesperado. Consulte os logs para suporte.'}), 500


# ---------------------------------------------------------------------------
# Modularização das rotas (Bloco 3): cada arquivo em routes/ contém as rotas
# de um domínio e é executado no namespace deste módulo — comportamento
# idêntico ao arquivo único original (mesmos globals e helpers, mesmas URLs),
# apenas organizado por assunto. Novas rotas devem ir no arquivo do domínio.
# No build PyInstaller, incluir --add-data "routes;routes".
# ---------------------------------------------------------------------------
ROUTE_MODULES = ['clients', 'accounts', 'activities_agenda', 'kanban', 'campaigns',
                 'whatsapp', 'outlook', 'itoca', 'autotoca', 'wikitoca',
                 'portfolio', 'config', 'home']


def _load_route_modules():
    base = Path(__file__).resolve().parent / 'routes'
    for _name in ROUTE_MODULES:
        _path = base / (_name + '.py')
        _code = compile(_path.read_text(encoding='utf-8'), str(_path), 'exec')
        exec(_code, globals())
    logger.info(f'[Rotas] {len(ROUTE_MODULES)} módulos de rotas carregados')


_load_route_modules()
_start_inbound_poller()
_start_scheduled_jobs()


if __name__ == '__main__':
    port = int(os.getenv('PORT', 3000))
    fixed_debug_mode = True
    app.logger.setLevel(logging.DEBUG)
    logging.getLogger().setLevel(logging.DEBUG)
    print('=' * 50)
    print('  TOCA DO COELHO - Gestao de Clientes')
    print('=' * 50)
    logger.info(f'[Database] Banco de dados inicializado')
    print(f'[Server] Iniciando em http://localhost:{port}')
    print(f'[Debug] Modo debug fixo no terminal: {"ATIVO" if fixed_debug_mode else "INATIVO"}')
    print(f'[Server] Pressione CTRL+C para parar')
    print()
    
    # Abrir navegador
    import threading
    def open_browser():
        import time
        time.sleep(2)
        webbrowser.open(f'http://localhost:{port}')
    
    thread = threading.Thread(target=open_browser, daemon=True)
    thread.start()
    
    app.run(host='localhost', port=port, debug=fixed_debug_mode, use_reloader=False)
