"""Validação e extração segura de documentos para o runtime web.

O módulo não importa dependências opcionais no boot. PDF e DOCX são carregados
somente quando a extração é solicitada, preservando o desktop sem esses pacotes.
"""

from __future__ import annotations

import os
import time
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path, PurePosixPath


_PDF_MIMES = {"", "application/octet-stream", "application/pdf"}
_DOCX_MIMES = {
    "",
    "application/octet-stream",
    "application/zip",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_DOCX_SIGNATURES = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")


class DocumentProcessingError(ValueError):
    """Erro controlado de validação, dependência ou extração."""

    def __init__(self, message, *, code, status=422, hint=None):
        super().__init__(message)
        self.code = code
        self.status = int(status)
        self.hint = hint

    def to_payload(self):
        payload = {"error": str(self), "code": self.code}
        if self.hint:
            payload["hint"] = self.hint
        return payload


@dataclass(frozen=True)
class DocumentLimits:
    max_bytes: int = 10 * 1024 * 1024
    max_pages: int = 50
    max_text_chars: int = 250_000
    timeout_seconds: float = 15.0
    max_docx_entries: int = 2_048
    max_docx_uncompressed_bytes: int = 50 * 1024 * 1024
    max_docx_compression_ratio: int = 200


def _positive_env_int(name, default):
    raw = (os.environ.get(name) or "").strip()
    if not raw:
        return default
    try:
        value = int(raw)
    except ValueError:
        return default
    return value if value > 0 else default


def _positive_env_float(name, default):
    raw = (os.environ.get(name) or "").strip()
    if not raw:
        return default
    try:
        value = float(raw)
    except ValueError:
        return default
    return value if value > 0 else default


def document_limits_from_env():
    return DocumentLimits(
        max_bytes=_positive_env_int(
            "TOCA_DOCUMENT_MAX_BYTES", DocumentLimits.max_bytes
        ),
        max_pages=_positive_env_int(
            "TOCA_DOCUMENT_MAX_PAGES", DocumentLimits.max_pages
        ),
        max_text_chars=_positive_env_int(
            "TOCA_DOCUMENT_MAX_TEXT_CHARS", DocumentLimits.max_text_chars
        ),
        timeout_seconds=_positive_env_float(
            "TOCA_DOCUMENT_PARSE_TIMEOUT_SECONDS", DocumentLimits.timeout_seconds
        ),
        max_docx_entries=_positive_env_int(
            "TOCA_DOCX_MAX_ENTRIES", DocumentLimits.max_docx_entries
        ),
        max_docx_uncompressed_bytes=_positive_env_int(
            "TOCA_DOCX_MAX_UNCOMPRESSED_BYTES",
            DocumentLimits.max_docx_uncompressed_bytes,
        ),
        max_docx_compression_ratio=_positive_env_int(
            "TOCA_DOCX_MAX_COMPRESSION_RATIO",
            DocumentLimits.max_docx_compression_ratio,
        ),
    )


def _normalized_mime(value):
    return (value or "").split(";", 1)[0].strip().lower()


def _extension(filename):
    return Path(filename or "").suffix.lower()


def _ensure_size(data, limits):
    if not data:
        raise DocumentProcessingError(
            "O arquivo enviado está vazio.",
            code="DOCUMENT_EMPTY",
            status=400,
        )
    if len(data) > limits.max_bytes:
        raise DocumentProcessingError(
            "O documento excede o limite permitido.",
            code="DOCUMENT_TOO_LARGE",
            status=413,
            hint=f"Limite por documento: {limits.max_bytes} bytes.",
        )


def _validate_pdf_signature(data):
    header = data[:1024].lstrip(b"\xef\xbb\xbf\x00\t\r\n ")
    if not header.startswith(b"%PDF-"):
        raise DocumentProcessingError(
            "O conteúdo não corresponde a um PDF válido.",
            code="DOCUMENT_SIGNATURE_MISMATCH",
            status=415,
        )


def _validate_docx_archive(data, limits):
    if not data.startswith(_DOCX_SIGNATURES):
        raise DocumentProcessingError(
            "O conteúdo não corresponde a um DOCX válido.",
            code="DOCUMENT_SIGNATURE_MISMATCH",
            status=415,
        )
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > limits.max_docx_entries:
                raise DocumentProcessingError(
                    "O DOCX possui arquivos internos demais.",
                    code="DOCUMENT_ARCHIVE_UNSAFE",
                    hint=f"Máximo: {limits.max_docx_entries} entradas.",
                )

            names = set()
            total_uncompressed = 0
            for entry in entries:
                normalized = PurePosixPath(entry.filename.replace("\\", "/"))
                if normalized.is_absolute() or ".." in normalized.parts:
                    raise DocumentProcessingError(
                        "O DOCX contém caminhos internos inseguros.",
                        code="DOCUMENT_ARCHIVE_UNSAFE",
                    )
                if entry.flag_bits & 0x1:
                    raise DocumentProcessingError(
                        "DOCX protegido por senha não é suportado.",
                        code="DOCUMENT_ENCRYPTED",
                    )
                names.add(entry.filename.replace("\\", "/"))
                if entry.is_dir():
                    continue
                total_uncompressed += entry.file_size
                if total_uncompressed > limits.max_docx_uncompressed_bytes:
                    raise DocumentProcessingError(
                        "O conteúdo expandido do DOCX excede o limite permitido.",
                        code="DOCUMENT_ARCHIVE_UNSAFE",
                        hint=(
                            "Máximo descompactado: "
                            f"{limits.max_docx_uncompressed_bytes} bytes."
                        ),
                    )
                ratio = entry.file_size / max(entry.compress_size, 1)
                if (
                    entry.file_size >= 1024 * 1024
                    and ratio > limits.max_docx_compression_ratio
                ):
                    raise DocumentProcessingError(
                        "O DOCX possui taxa de compressão insegura.",
                        code="DOCUMENT_ARCHIVE_UNSAFE",
                    )

            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                raise DocumentProcessingError(
                    "O arquivo ZIP não possui a estrutura obrigatória de um DOCX.",
                    code="DOCUMENT_SIGNATURE_MISMATCH",
                    status=415,
                )
    except DocumentProcessingError:
        raise
    except (zipfile.BadZipFile, OSError, RuntimeError) as exc:
        raise DocumentProcessingError(
            "O DOCX está corrompido ou malformado.",
            code="DOCUMENT_MALFORMED",
        ) from exc


def validate_document_bytes(
    data,
    *,
    filename,
    declared_mime="",
    allowed_kinds=("pdf", "docx"),
    limits=None,
):
    """Valida tamanho, extensão, MIME declarado e assinatura real."""
    limits = limits or document_limits_from_env()
    data = bytes(data or b"")
    _ensure_size(data, limits)

    allowed = {str(kind).lower() for kind in allowed_kinds}
    ext = _extension(filename)
    if ext == ".doc":
        raise DocumentProcessingError(
            "O formato legado .doc não é suportado na versão web. Converta para DOCX.",
            code="DOCUMENT_LEGACY_DOC_UNSUPPORTED",
            status=415,
        )
    kind_by_ext = {".pdf": "pdf", ".docx": "docx"}.get(ext)
    if not kind_by_ext or kind_by_ext not in allowed:
        accepted = ", ".join(sorted(kind.upper() for kind in allowed))
        raise DocumentProcessingError(
            f"Formato de documento não suportado. Aceitos: {accepted}.",
            code="DOCUMENT_UNSUPPORTED_TYPE",
            status=415,
        )

    mime = _normalized_mime(declared_mime)
    accepted_mimes = _PDF_MIMES if kind_by_ext == "pdf" else _DOCX_MIMES
    if mime not in accepted_mimes:
        raise DocumentProcessingError(
            "O tipo MIME declarado não corresponde ao documento enviado.",
            code="DOCUMENT_MIME_MISMATCH",
            status=415,
        )

    if kind_by_ext == "pdf":
        _validate_pdf_signature(data)
    else:
        _validate_docx_archive(data, limits)
    return kind_by_ext


def read_document_upload(
    file_storage,
    *,
    allowed_kinds=("pdf", "docx"),
    limits=None,
):
    """Lê no máximo max_bytes + 1 para limitar uploads sem Content-Length."""
    if not file_storage or not getattr(file_storage, "filename", ""):
        raise DocumentProcessingError(
            "Nenhum documento foi enviado.",
            code="DOCUMENT_MISSING",
            status=400,
        )
    limits = limits or document_limits_from_env()
    data = file_storage.read(limits.max_bytes + 1)
    kind = validate_document_bytes(
        data,
        filename=file_storage.filename,
        declared_mime=getattr(file_storage, "mimetype", ""),
        allowed_kinds=allowed_kinds,
        limits=limits,
    )
    return data, kind


def read_text_upload(
    file_storage,
    *,
    allowed_extensions=(".txt", ".vtt", ".srt", ".csv"),
    limits=None,
):
    """Lê arquivos textuais simples com allowlist, MIME e limite real de bytes."""
    if not file_storage or not getattr(file_storage, "filename", ""):
        raise DocumentProcessingError(
            "Nenhum arquivo de texto foi enviado.",
            code="DOCUMENT_MISSING",
            status=400,
        )
    ext = _extension(file_storage.filename)
    allowed = {str(item).lower() for item in allowed_extensions}
    if ext not in allowed:
        raise DocumentProcessingError(
            "Formato de arquivo não suportado.",
            code="DOCUMENT_UNSUPPORTED_TYPE",
            status=415,
        )
    mime = _normalized_mime(getattr(file_storage, "mimetype", ""))
    if mime and mime != "application/octet-stream" and not mime.startswith("text/"):
        raise DocumentProcessingError(
            "O tipo MIME declarado não corresponde a um arquivo de texto.",
            code="DOCUMENT_MIME_MISMATCH",
            status=415,
        )
    limits = limits or document_limits_from_env()
    data = file_storage.read(limits.max_bytes + 1)
    _ensure_size(data, limits)
    if b"\x00" in data[:4096]:
        raise DocumentProcessingError(
            "O arquivo contém dados binários e não pode ser tratado como texto.",
            code="DOCUMENT_SIGNATURE_MISMATCH",
            status=415,
        )
    return data


def _check_deadline(started_at, limits):
    if time.monotonic() - started_at > limits.timeout_seconds:
        raise DocumentProcessingError(
            "A leitura do documento excedeu o tempo permitido.",
            code="DOCUMENT_PARSE_TIMEOUT",
            hint=f"Limite: {limits.timeout_seconds:g} segundos.",
        )


def _append_text(parts, value, total_chars, limits):
    text = (value or "").strip()
    if not text:
        return total_chars
    total_chars += len(text)
    if total_chars > limits.max_text_chars:
        raise DocumentProcessingError(
            "O texto extraído excede o limite permitido.",
            code="DOCUMENT_TEXT_TOO_LARGE",
            hint=f"Máximo: {limits.max_text_chars} caracteres.",
        )
    parts.append(text)
    return total_chars


def _extract_pdf_text(data, limits, started_at):
    try:
        import pdfplumber
    except ImportError as exc:
        raise DocumentProcessingError(
            "Leitura de PDF indisponível neste ambiente.",
            code="DOCUMENT_PARSER_UNAVAILABLE",
            status=503,
        ) from exc

    parts = []
    total_chars = 0
    try:
        with pdfplumber.open(BytesIO(data)) as pdf:
            pages = pdf.pages or []
            if not pages:
                raise DocumentProcessingError(
                    "O PDF não possui páginas legíveis.",
                    code="DOCUMENT_MALFORMED",
                )
            if len(pages) > limits.max_pages:
                raise DocumentProcessingError(
                    "O PDF excede o limite de páginas.",
                    code="DOCUMENT_TOO_MANY_PAGES",
                    hint=f"Máximo: {limits.max_pages} páginas.",
                )
            for page in pages:
                _check_deadline(started_at, limits)
                total_chars = _append_text(
                    parts, page.extract_text() or "", total_chars, limits
                )
                _check_deadline(started_at, limits)
    except DocumentProcessingError:
        raise
    except Exception as exc:
        raise DocumentProcessingError(
            "O PDF está corrompido, protegido ou malformado.",
            code="DOCUMENT_MALFORMED",
        ) from exc
    return "\n".join(parts).strip()


def _extract_docx_text(data, limits, started_at):
    try:
        import docx
    except ImportError as exc:
        raise DocumentProcessingError(
            "Leitura de DOCX indisponível neste ambiente.",
            code="DOCUMENT_PARSER_UNAVAILABLE",
            status=503,
        ) from exc

    parts = []
    total_chars = 0
    try:
        document = docx.Document(BytesIO(data))
        _check_deadline(started_at, limits)
        for paragraph in document.paragraphs:
            total_chars = _append_text(
                parts, paragraph.text, total_chars, limits
            )
            _check_deadline(started_at, limits)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                total_chars = _append_text(
                    parts, row_text, total_chars, limits
                )
                _check_deadline(started_at, limits)
    except DocumentProcessingError:
        raise
    except Exception as exc:
        raise DocumentProcessingError(
            "O DOCX está corrompido ou malformado.",
            code="DOCUMENT_MALFORMED",
        ) from exc
    return "\n".join(parts).strip()


def extract_document_text(
    data,
    *,
    filename,
    declared_mime="",
    allowed_kinds=("pdf", "docx"),
    limits=None,
    require_text=True,
):
    """Valida e extrai apenas texto digital; OCR não é executado."""
    limits = limits or document_limits_from_env()
    data = bytes(data or b"")
    kind = validate_document_bytes(
        data,
        filename=filename,
        declared_mime=declared_mime,
        allowed_kinds=allowed_kinds,
        limits=limits,
    )
    started_at = time.monotonic()
    if kind == "pdf":
        text = _extract_pdf_text(data, limits, started_at)
    else:
        text = _extract_docx_text(data, limits, started_at)

    if require_text and not text:
        hint = (
            "O arquivo parece ser um PDF escaneado; OCR não está habilitado "
            "no servidor."
            if kind == "pdf"
            else "Confirme se o documento contém texto em parágrafos ou tabelas."
        )
        raise DocumentProcessingError(
            "Não foi encontrado texto digital no documento.",
            code="DOCUMENT_NO_DIGITAL_TEXT",
            hint=hint,
        )
    return text


def extract_document_path(
    file_path,
    *,
    allowed_kinds=("pdf", "docx"),
    limits=None,
    require_text=True,
):
    path = Path(file_path)
    limits = limits or document_limits_from_env()
    if not path.is_file():
        raise DocumentProcessingError(
            "Documento não encontrado.",
            code="DOCUMENT_NOT_FOUND",
            status=404,
        )
    if path.stat().st_size > limits.max_bytes:
        raise DocumentProcessingError(
            "O documento excede o limite permitido.",
            code="DOCUMENT_TOO_LARGE",
            status=413,
        )
    with path.open("rb") as stream:
        data = stream.read(limits.max_bytes + 1)
    return extract_document_text(
        data,
        filename=path.name,
        declared_mime="",
        allowed_kinds=allowed_kinds,
        limits=limits,
        require_text=require_text,
    )
