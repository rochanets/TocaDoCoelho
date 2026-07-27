# -*- coding: utf-8 -*-
"""Fase 7.2: validação e leitura segura de PDF digital/DOCX na web."""

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from reportlab.pdfgen import canvas

import app as toca
import document_processing as documents


@pytest.fixture(autouse=True)
def _isolated_wiki_upload_dir(tmp_path, monkeypatch):
    upload_dir = tmp_path / "wikitoca"
    upload_dir.mkdir()
    monkeypatch.setattr(toca, "WIKI_UPLOAD_DIR", upload_dir)


def _pdf_bytes(*page_texts):
    output = BytesIO()
    pdf = canvas.Canvas(output)
    for text in page_texts:
        if text:
            pdf.drawString(72, 760, text)
        pdf.showPage()
    pdf.save()
    return output.getvalue()


def _docx_bytes(text="Documento DOCX seguro"):
    output = BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Campo"
    table.cell(0, 1).text = "Valor"
    doc.save(output)
    return output.getvalue()


def _auth_on(monkeypatch):
    monkeypatch.setenv("TOCA_AUTH_ENABLED", "1")
    monkeypatch.setitem(toca.app.config, "SESSION_COOKIE_SECURE", False)


def _seed_and_login(client):
    conn = toca.get_db()
    cur = conn.cursor()
    cur.execute("INSERT INTO organizations (name) VALUES ('Org F7.2')")
    org_id = cur.lastrowid
    cur.execute(
        """INSERT INTO users (org_id, email, full_name, role)
           VALUES (?, 'f72@corp.com', 'Usuária F7.2', 'admin')""",
        (org_id,),
    )
    user_id = cur.lastrowid
    conn.commit()
    conn.close()
    with client.session_transaction() as session:
        session["user_id"] = user_id
    return user_id


def test_extracts_valid_pdf_and_docx_text():
    pdf_text = documents.extract_document_text(
        _pdf_bytes("PDF digital F7.2"),
        filename="material.pdf",
        declared_mime="application/pdf",
    )
    docx_text = documents.extract_document_text(
        _docx_bytes(),
        filename="material.docx",
        declared_mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    assert "PDF digital F7.2" in pdf_text
    assert "Documento DOCX seguro" in docx_text
    assert "Campo | Valor" in docx_text


@pytest.mark.parametrize(
    ("data", "filename", "mime", "expected_code"),
    [
        (b"nao e pdf", "fraude.pdf", "application/pdf", "DOCUMENT_SIGNATURE_MISMATCH"),
        (
            _pdf_bytes("PDF"),
            "fraude.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "DOCUMENT_SIGNATURE_MISMATCH",
        ),
        (
            _docx_bytes(),
            "fraude.docx",
            "text/plain",
            "DOCUMENT_MIME_MISMATCH",
        ),
        (b"arquivo legado", "legado.doc", "application/msword", "DOCUMENT_LEGACY_DOC_UNSUPPORTED"),
    ],
    ids=("fake-pdf", "pdf-as-docx", "docx-wrong-mime", "legacy-doc"),
)
def test_rejects_spoofed_mime_signature_and_legacy_doc(
    data, filename, mime, expected_code
):
    with pytest.raises(documents.DocumentProcessingError) as raised:
        documents.validate_document_bytes(
            data,
            filename=filename,
            declared_mime=mime,
        )
    assert raised.value.code == expected_code


def test_enforces_actual_byte_and_page_limits():
    with pytest.raises(documents.DocumentProcessingError) as too_large:
        documents.validate_document_bytes(
            _pdf_bytes("grande"),
            filename="grande.pdf",
            declared_mime="application/pdf",
            limits=documents.DocumentLimits(max_bytes=32),
        )
    assert too_large.value.code == "DOCUMENT_TOO_LARGE"
    assert too_large.value.status == 413

    with pytest.raises(documents.DocumentProcessingError) as too_many_pages:
        documents.extract_document_text(
            _pdf_bytes("página 1", "página 2"),
            filename="paginas.pdf",
            declared_mime="application/pdf",
            limits=documents.DocumentLimits(max_pages=1),
        )
    assert too_many_pages.value.code == "DOCUMENT_TOO_MANY_PAGES"


def test_rejects_pdf_without_digital_text_instead_of_running_ocr():
    with pytest.raises(documents.DocumentProcessingError) as raised:
        documents.extract_document_text(
            _pdf_bytes(""),
            filename="escaneado.pdf",
            declared_mime="application/pdf",
        )
    assert raised.value.code == "DOCUMENT_NO_DIGITAL_TEXT"
    assert "OCR não está habilitado" in raised.value.hint


def test_rejects_docx_with_unsafe_archive_expansion():
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", b"A" * (2 * 1024 * 1024))

    with pytest.raises(documents.DocumentProcessingError) as raised:
        documents.validate_document_bytes(
            output.getvalue(),
            filename="bomba.docx",
            declared_mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            limits=documents.DocumentLimits(max_docx_compression_ratio=10),
        )
    assert raised.value.code == "DOCUMENT_ARCHIVE_UNSAFE"


@pytest.mark.parametrize("unsafe_entry", ("../escape.xml", "/absolute.xml"))
def test_rejects_docx_with_unsafe_internal_paths(unsafe_entry):
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<document/>")
        archive.writestr(unsafe_entry, "conteúdo")

    with pytest.raises(documents.DocumentProcessingError) as raised:
        documents.validate_document_bytes(
            output.getvalue(),
            filename="inseguro.docx",
            declared_mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
    assert raised.value.code == "DOCUMENT_ARCHIVE_UNSAFE"


def test_rejects_zip_without_required_docx_structure():
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("arquivo.txt", "não é DOCX")

    with pytest.raises(documents.DocumentProcessingError) as raised:
        documents.validate_document_bytes(
            output.getvalue(),
            filename="falso.docx",
            declared_mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
    assert raised.value.code == "DOCUMENT_SIGNATURE_MISMATCH"


def test_portfolio_rejects_fake_pdf_before_starting_task(client, monkeypatch):
    _auth_on(monkeypatch)
    _seed_and_login(client)

    response = client.post(
        "/api/portfolio/offers",
        data={
            "upload_file": (
                BytesIO(b"conteudo executavel"),
                "material.pdf",
                "application/pdf",
            )
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 415
    assert response.get_json()["code"] == "DOCUMENT_SIGNATURE_MISMATCH"


def test_iata_rejects_legacy_doc_before_starting_task(client, monkeypatch):
    _auth_on(monkeypatch)
    _seed_and_login(client)

    response = client.post(
        "/api/portfolio/iata",
        data={
            "meeting_file": (
                BytesIO(b"documento legado"),
                "reuniao.doc",
                "application/msword",
            )
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 415
    assert response.get_json()["code"] == "DOCUMENT_LEGACY_DOC_UNSUPPORTED"


def test_wiki_web_rejects_legacy_doc_without_persisting(client, monkeypatch):
    _auth_on(monkeypatch)
    _seed_and_login(client)

    response = client.post(
        "/api/wikitoca/documents",
        data={
            "files": (
                BytesIO(b"documento legado"),
                "legado.doc",
                "application/msword",
            )
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 415
    assert response.get_json()["error_code"] == "DOCUMENT_LEGACY_DOC_UNSUPPORTED"
    conn = toca.get_db()
    count = conn.execute("SELECT COUNT(*) AS n FROM wiki_documents").fetchone()["n"]
    conn.close()
    assert count == 0


def test_wiki_web_accepts_valid_pdf_and_docx(client, monkeypatch):
    _auth_on(monkeypatch)
    _seed_and_login(client)

    response = client.post(
        "/api/wikitoca/documents",
        data={
            "files": [
                (
                    BytesIO(_pdf_bytes("Wiki PDF F7.2")),
                    "material.pdf",
                    "application/pdf",
                ),
                (
                    BytesIO(_docx_bytes("Wiki DOCX F7.2")),
                    "material.docx",
                    (
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                ),
            ]
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 201
    assert {item["file_ext"] for item in response.get_json()} == {".pdf", ".docx"}
    assert {path.suffix for path in toca.WIKI_UPLOAD_DIR.iterdir()} == {
        ".pdf",
        ".docx",
    }


def test_wiki_desktop_keeps_legacy_doc_upload(client, monkeypatch):
    monkeypatch.delenv("TOCA_AUTH_ENABLED", raising=False)

    response = client.post(
        "/api/wikitoca/documents",
        data={
            "files": (
                BytesIO(b"documento legado desktop"),
                "legado.doc",
                "application/msword",
            )
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 201
    assert response.get_json()[0]["file_ext"] == ".doc"
